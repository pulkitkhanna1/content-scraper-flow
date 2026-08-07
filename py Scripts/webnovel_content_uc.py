#!/usr/bin/env python3
"""
webnovel_content_uc.py

Cloudflare-resistant webnovel scraper:
- Mandatory undetected_chromedriver (bypasses most Cloudflare JS challenges)
- 2Captcha solver (hCaptcha / Cloudflare Turnstile) as fallback
- Batch driver restart + threaded timeout watchdog
- Checkpoint / resume support
- Scrapes chapters via Selenium DOM and saves to chunked DOCX files
"""

import argparse
import json
import logging
import os
import re
import sys
import time
import random
import threading
import traceback
from typing import List, Dict, Any, Optional, Tuple
from urllib.parse import urlparse
from datetime import datetime

import requests as http_requests
from bs4 import BeautifulSoup as bs
import undetected_chromedriver as uc
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ──────────────────────────────────────────────
# CONFIG
# ──────────────────────────────────────────────

EMAIL    = "fantasy.team@pocketfm.com"
PASSWORD = "0ForFantasyOnly0"

# EMAIL = "pranjal.trivedi@pocketfm.com"
# PASSWORD = "8bigLHMk3etRskK"

# EMAIL = "commissioning_ops@pocketfm.com"
# PASSWORD = "Usgrowth2026"

# EMAIL = "ops.pocket15@pocketfm.com"
# PASSWORD = "PocketOps1"

BOOK_ID               = "32608793200298405"
START_CHAPTER_ID      = "87599150631287726"
START_CHAPTER_NUMBER  = 1
END_CHAPTER_NUMBER    = 10
OUT_FOLDER            = "Formula 1: The GOAT - trst"

CHUNK_SIZE            = 100
HEADLESS              = False
COOKIE_FILE           = "webnovel_cookies.json"
POLITE_DELAY          = 3
PAGE_LOAD_TIMEOUT     = 60
WAIT_FOR_2FA_SECONDS  = 120
MIN_PAGE_CONTENT_LENGTH = 100

TWOCAPTCHA_API_KEY    = "7c802b56a424742db63aa7d8640377eb"
MAX_CAPTCHA_RETRIES   = 3

CHECKPOINT_FILE       = f"webnovel_checkpoint_{BOOK_ID}.json"

# Driver restart settings
BATCH_SIZE            = 15
MAX_CONSECUTIVE_FAILS = 3
CHAPTER_TIMEOUT       = 90
COOLDOWN_MIN          = 8.0
COOLDOWN_MAX          = 18.0

CHAPTER_DELAY_MIN     = 3.0
CHAPTER_DELAY_MAX     = 7.0

# ──────────────────────────────────────────────
# LOGGING
# ──────────────────────────────────────────────

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler("webnovel_uc_run.log", encoding="utf-8"),
    ],
)
logger = logging.getLogger("webnovel_uc")


# ──────────────────────────────────────────────
# TIMEOUT WATCHDOG
# ──────────────────────────────────────────────

class ChapterTimeoutError(Exception):
    pass


def run_with_timeout(func, args=(), kwargs=None, timeout=CHAPTER_TIMEOUT):
    if kwargs is None:
        kwargs = {}
    result = [None]
    error  = [None]
    done   = threading.Event()

    def wrapper():
        try:
            result[0] = func(*args, **kwargs)
        except Exception as e:
            error[0] = e
        finally:
            done.set()

    t = threading.Thread(target=wrapper, daemon=True)
    t.start()
    done.wait(timeout=timeout)
    if not done.is_set():
        raise ChapterTimeoutError(f"Timed out after {timeout}s")
    if error[0] is not None:
        raise error[0]
    return result[0]


# ──────────────────────────────────────────────
# CAPTCHA DETECTION
# ──────────────────────────────────────────────

def detect_cloudflare_challenge(html: str) -> bool:
    low = html.lower()
    return (
        "just a moment" in low
        or "__cf_chl_opt" in html
        or "cf-challenge" in low
        or "checking your browser" in low
        or ("cloudflare" in low and len(html) < 5000)
    )


def detect_hcaptcha(html: str) -> Tuple[bool, Optional[str]]:
    if "hcaptcha.com" not in html and "h-captcha" not in html:
        return False, None
    m = re.search(r'data-sitekey=["\']([^"\']+)["\']', html)
    return True, (m.group(1) if m else None)


def detect_turnstile(html: str) -> Tuple[bool, Optional[str]]:
    if "challenges.cloudflare.com" not in html and "turnstile" not in html.lower():
        return False, None
    m = re.search(r'data-sitekey=["\']([^"\']+)["\']', html)
    return True, (m.group(1) if m else None)


# ──────────────────────────────────────────────
# 2CAPTCHA SOLVER
# ──────────────────────────────────────────────

def _solve_2captcha(task_payload: dict) -> Optional[dict]:
    """Generic 2captcha task dispatcher. Returns solution dict or None."""
    create_payload = {"clientKey": TWOCAPTCHA_API_KEY, "task": task_payload}
    try:
        resp = http_requests.post(
            "https://api.2captcha.com/createTask",
            json=create_payload,
            timeout=30,
        )
        result = resp.json()
    except Exception as e:
        logger.warning("2captcha createTask failed: %s", e)
        return None

    if result.get("errorId", 1) != 0:
        logger.warning("2captcha error: %s", result.get("errorDescription", result))
        return None

    task_id = result["taskId"]
    logger.info("2captcha task %s submitted — polling...", task_id)

    for poll in range(60):
        time.sleep(3)
        try:
            pr = http_requests.post(
                "https://api.2captcha.com/getTaskResult",
                json={"clientKey": TWOCAPTCHA_API_KEY, "taskId": task_id},
                timeout=30,
            )
            pr_data = pr.json()
        except Exception as e:
            logger.warning("2captcha poll error: %s", e)
            continue

        if pr_data.get("errorId", 0) != 0:
            logger.warning("2captcha poll error: %s", pr_data.get("errorDescription"))
            return None

        if pr_data.get("status") == "ready":
            solution = pr_data.get("solution", {})
            logger.info("2captcha solved: %s", str(solution)[:80])
            return solution

        if poll % 5 == 0:
            logger.info("  still solving... (%ds)", poll * 3)

    logger.warning("2captcha timed out after 180s")
    return None


def solve_hcaptcha(website_url: str, sitekey: str) -> Optional[str]:
    solution = _solve_2captcha({
        "type": "HCaptchaTaskProxyless",
        "websiteURL": website_url,
        "websiteKey": sitekey,
    })
    return solution.get("gRecaptchaResponse") if solution else None


def solve_turnstile(website_url: str, sitekey: str) -> Optional[str]:
    solution = _solve_2captcha({
        "type": "TurnstileTaskProxyless",
        "websiteURL": website_url,
        "websiteKey": sitekey,
    })
    return solution.get("token") if solution else None


def _inject_hcaptcha_token(driver, token: str) -> bool:
    js = f"""
    (function() {{
        var el = document.querySelector('textarea[name="h-captcha-response"]') ||
                 document.querySelector('textarea[id="h-captcha-response"]') ||
                 document.querySelector('[name="h-captcha-response"]');
        if (el) {{ el.value = "{token}"; }}

        if (typeof hcaptcha !== 'undefined') {{
            try {{
                hcaptcha.execute();
            }} catch(e) {{}}
        }}

        var form = document.querySelector('form');
        if (form) {{ form.submit(); }}
        return !!el;
    }})();
    """
    try:
        ok = driver.execute_script(js)
        time.sleep(2)
        driver.refresh()
        time.sleep(2)
        return bool(ok)
    except Exception as e:
        logger.warning("hCaptcha injection error: %s", e)
        return False


def _inject_turnstile_token(driver, token: str) -> bool:
    js = f"""
    (function() {{
        var el = document.querySelector('[name="cf-turnstile-response"]') ||
                 document.querySelector('input[name="cf-turnstile-response"]');
        if (el) {{ el.value = "{token}"; }}
        var form = document.querySelector('form');
        if (form) {{ form.submit(); return true; }}
        return !!el;
    }})();
    """
    try:
        driver.execute_script(js)
        time.sleep(2)
        driver.refresh()
        time.sleep(2)
        return True
    except Exception as e:
        logger.warning("Turnstile injection error: %s", e)
        return False


def wait_for_cloudflare(driver, timeout=30) -> str:
    """Wait for Cloudflare JS challenge to auto-resolve (UC handles it)."""
    logger.info("Cloudflare challenge detected — waiting for auto-resolve...")
    for i in range(timeout // 2):
        time.sleep(2)
        html = driver.page_source
        if not detect_cloudflare_challenge(html):
            logger.info("Cloudflare resolved after ~%ds", (i + 1) * 2)
            return html
        try:
            driver.execute_script(
                "document.dispatchEvent(new MouseEvent('mousemove',"
                " {clientX: Math.random()*800, clientY: Math.random()*600}));"
            )
        except Exception:
            pass
    logger.warning("Cloudflare still showing after %ds", timeout)
    return driver.page_source


def handle_captcha(driver, url: str, max_retries: int = MAX_CAPTCHA_RETRIES) -> bool:
    """
    Detect and resolve Cloudflare / hCaptcha / Turnstile challenges.
    Returns True if page is clear (or best-effort done), False if unresolvable.
    """
    html = driver.page_source

    # Step 1: Cloudflare JS challenge (UC usually handles automatically)
    if detect_cloudflare_challenge(html):
        html = wait_for_cloudflare(driver)
        if detect_cloudflare_challenge(html):
            logger.warning("Cloudflare not auto-resolved — attempting reload")
            try:
                driver.refresh()
                time.sleep(3)
                html = driver.page_source
            except Exception:
                return False

    # Step 2: hCaptcha
    for attempt in range(max_retries):
        is_hcap, sitekey = detect_hcaptcha(html)
        if not is_hcap:
            break
        logger.info("hCaptcha detected (attempt %d/%d, sitekey=%s)", attempt + 1, max_retries, sitekey)
        if not sitekey:
            logger.warning("Could not extract hCaptcha sitekey")
            return False
        token = solve_hcaptcha(url, sitekey)
        if not token:
            time.sleep(3)
            continue
        _inject_hcaptcha_token(driver, token)
        html = driver.page_source
        is_hcap2, _ = detect_hcaptcha(html)
        if not is_hcap2:
            logger.info("hCaptcha bypassed!")
            return True
    else:
        is_hcap, _ = detect_hcaptcha(html)
        if is_hcap:
            logger.warning("Could not bypass hCaptcha after %d attempts", max_retries)
            return False

    # Step 3: Cloudflare Turnstile
    for attempt in range(max_retries):
        is_ts, sitekey = detect_turnstile(html)
        if not is_ts:
            break
        logger.info("Turnstile detected (attempt %d/%d, sitekey=%s)", attempt + 1, max_retries, sitekey)
        if not sitekey:
            logger.warning("Could not extract Turnstile sitekey")
            return False
        token = solve_turnstile(url, sitekey)
        if not token:
            time.sleep(3)
            continue
        _inject_turnstile_token(driver, token)
        html = driver.page_source
        is_ts2, _ = detect_turnstile(html)
        if not is_ts2:
            logger.info("Turnstile bypassed!")
            return True
    else:
        is_ts, _ = detect_turnstile(html)
        if is_ts:
            logger.warning("Could not bypass Turnstile after %d attempts", max_retries)
            return False

    return True


# ──────────────────────────────────────────────
# DRIVER MANAGEMENT
# ──────────────────────────────────────────────

def kill_driver(driver):
    if driver is None:
        return
    try:
        driver.quit()
    except Exception:
        pass
    try:
        import psutil
        current_pid = os.getpid()
        for proc in psutil.process_iter(["pid", "name", "ppid"]):
            try:
                if proc.info["ppid"] == current_pid and "chrome" in proc.info["name"].lower():
                    proc.kill()
            except Exception:
                pass
    except ImportError:
        pass


def _get_local_chrome_major() -> Optional[int]:
    import subprocess
    candidates = [
        "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome",
        "/Applications/Chromium.app/Contents/MacOS/Chromium",
        "google-chrome", "chromium-browser", "chromium", "chrome",
    ]
    for cmd in candidates:
        try:
            out = subprocess.check_output([cmd, "--version"], stderr=subprocess.STDOUT, text=True)
            m = re.search(r"(\d+)\.", out)
            if m:
                ver = int(m.group(1))
                logger.info("Detected local Chrome version: %d", ver)
                return ver
        except Exception:
            continue
    return None


def create_driver(headless: bool = HEADLESS) -> uc.Chrome:
    logger.info("Launching fresh undetected Chrome...")
    
    def get_opts():
        opts = uc.ChromeOptions()
        if headless:
            opts.add_argument("--headless=new")
        opts.add_argument("--start-maximized")
        opts.add_argument("--window-size=1920,1080")
        opts.add_argument("--no-sandbox")
        opts.add_argument("--disable-dev-shm-usage")
        opts.add_argument("--disable-gpu")
        opts.add_argument("--log-level=3")
        opts.add_argument("--dns-prefetch-disable")
        return opts

    chrome_major = _get_local_chrome_major()
    try:
        opts = get_opts()
        if chrome_major:
            driver = uc.Chrome(options=opts, use_subprocess=True, headless=headless, version_main=chrome_major)
        else:
            driver = uc.Chrome(options=opts, use_subprocess=True, headless=headless)
    except TypeError as e:
        if "Binary Location Must be a String" in str(e):
            logger.error("Google Chrome is not installed on this system! "
                         "If you are deploying on Render, you must add the Chrome Buildpack to your settings: "
                         "https://github.com/render-examples/chrome-buildpack.git")
            raise RuntimeError("Google Chrome not found. Please install Chrome or add the Render Chrome Buildpack.") from e
        # older UC versions don't accept version_main or headless as kwarg
        logger.info("Retrying with fallback options...")
        opts = get_opts()
        driver = uc.Chrome(options=opts, use_subprocess=True)

    driver.set_page_load_timeout(PAGE_LOAD_TIMEOUT)
    driver.set_script_timeout(30)
    driver.implicitly_wait(5)
    logger.info("Chrome started")
    return driver


def safe_navigate(driver, url: str, retries: int = 2) -> bool:
    """Navigate to URL, then resolve any captcha challenges. Returns True if page is clean."""
    for attempt in range(1, retries + 1):
        try:
            driver.get(url)
            time.sleep(random.uniform(1.5, 3.0))
            if not handle_captcha(driver, url):
                logger.warning("Captcha unresolved (attempt %d)", attempt)
                if attempt < retries:
                    time.sleep(3)
                    continue
                return False
            return True
        except Exception as e:
            logger.warning("Navigation error (attempt %d): %s", attempt, e)
            if attempt < retries:
                time.sleep(3)
    return False


# ──────────────────────────────────────────────
# COOKIE HELPERS
# ──────────────────────────────────────────────

def dump_cookies(driver, out_path: str = COOKIE_FILE) -> List[Dict[str, Any]]:
    all_cookies = driver.get_cookies()
    logger.info("Total cookies from browser: %d", len(all_cookies))
    wn_cookies = []
    for c in all_cookies:
        dom = c.get("domain", "") or ""
        name = c.get("name", "") or ""
        if "webnovel" in dom or "webnovel" in name or dom.endswith("novel.com"):
            nc = {"name": c.get("name"), "value": c.get("value"), "path": c.get("path", "/")}
            if c.get("domain"):
                nc["domain"] = c["domain"]
            if c.get("expiry"):
                try:
                    nc["expiry"] = int(c["expiry"])
                except Exception:
                    pass
            if c.get("secure"):
                nc["secure"] = bool(c["secure"])
            wn_cookies.append(nc)
    if not wn_cookies:
        for c in all_cookies:
            dom = c.get("domain", "") or ""
            if "novel" in dom:
                nc = {"name": c.get("name"), "value": c.get("value"), "path": c.get("path", "/")}
                if c.get("domain"):
                    nc["domain"] = c["domain"]
                wn_cookies.append(nc)
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(wn_cookies, f, indent=2, ensure_ascii=False)
    logger.info("Saved %d webnovel cookies to %s", len(wn_cookies), out_path)
    return wn_cookies


def load_cookies_from_file(path: str = COOKIE_FILE) -> List[Dict[str, Any]]:
    if not os.path.exists(path):
        logger.warning("Cookie file not found: %s", path)
        return []
    try:
        with open(path, "r", encoding="utf-8") as f:
            cookies = json.load(f)
    except Exception as e:
        logger.exception("Failed to load cookies: %s", e)
        return []
    norm = []
    for c in cookies:
        nc = {"name": c.get("name"), "value": c.get("value"), "path": c.get("path", "/")}
        if c.get("domain"):
            nc["domain"] = c["domain"]
        if c.get("expiry"):
            try:
                nc["expiry"] = int(c["expiry"])
            except Exception:
                pass
        if c.get("secure"):
            nc["secure"] = bool(c["secure"])
        norm.append(nc)
    return norm


def _try_add_cookie_variants(driver, cookie: Dict[str, Any], host_variants: List[str]) -> bool:
    base = cookie.copy()
    base.setdefault("path", "/")
    domains = []
    if base.get("domain"):
        domains.append(base["domain"])
        domains.append(base["domain"].lstrip("."))
    domains.extend(host_variants)
    domains.append(None)
    for d in domains:
        c = base.copy()
        if d is None:
            c.pop("domain", None)
        else:
            c["domain"] = d
        if c.get("expiry") is None:
            c.pop("expiry", None)
        try:
            driver.add_cookie(c)
            return True
        except Exception:
            continue
    return False


def apply_cookies_from_file(driver, cookie_file: str = COOKIE_FILE, target_url: str = "https://www.webnovel.com") -> bool:
    cookies = load_cookies_from_file(cookie_file)
    if not cookies:
        return False
    try:
        driver.get(target_url)
        time.sleep(1.0)
        handle_captcha(driver, target_url)
    except Exception:
        pass
    host = urlparse(target_url).netloc
    host_variants = [host]
    if host.startswith("www."):
        host_variants.append(host.replace("www.", ""))
    else:
        host_variants.append("www." + host)
    any_added = any(_try_add_cookie_variants(driver, c, host_variants) for c in cookies)
    try:
        driver.refresh()
        time.sleep(1.0)
        handle_captcha(driver, target_url)
    except Exception:
        pass
    logger.info("Applied cookies from file (any_added=%s)", any_added)
    return any_added


def create_driver_with_cookies(book_page: str) -> uc.Chrome:
    """Create fresh driver, land on webnovel home (triggering CF check), then inject cookies."""
    driver = create_driver()
    logger.info("Loading webnovel.com to set cookies...")
    try:
        driver.get("https://www.webnovel.com/")
        time.sleep(2)
        handle_captcha(driver, "https://www.webnovel.com/")
    except Exception as e:
        logger.warning("Homepage load issue: %s", e)

    cookies = load_cookies_from_file(COOKIE_FILE)
    if cookies:
        host = urlparse(book_page).netloc
        host_variants = [host, host.lstrip("www."), "www." + host.lstrip("www.")]
        added = sum(1 for c in cookies if _try_add_cookie_variants(driver, c, host_variants))
        logger.info("Injected %d/%d cookies", added, len(cookies))
    return driver


# ──────────────────────────────────────────────
# LOGIN
# ──────────────────────────────────────────────

def open_login_and_submit(driver, email: str, password: str, return_url: str) -> bool:
    from selenium.webdriver.common.by import By
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC

    try:
        return_url_enc = http_requests.utils.quote(return_url, safe="")
        login_url = (
            "https://passport.webnovel.com/emaillogin.html?"
            f"appid=900&areaid=1&source=enweb&format=redirect&returnurl={return_url_enc}&channel=email"
        )
        logger.info("Opening login page")
        if not safe_navigate(driver, login_url):
            logger.warning("Login page may have unresolved captcha")
        WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.ID, "email")))
        try:
            driver.find_element(By.ID, "email").clear()
            driver.find_element(By.ID, "email").send_keys(email)
        except Exception:
            try:
                driver.find_element(By.NAME, "email").send_keys(email)
            except Exception:
                pass
        try:
            driver.find_element(By.NAME, "password").clear()
            driver.find_element(By.NAME, "password").send_keys(password)
        except Exception:
            pass
        time.sleep(0.5)
        submit_clicked = False
        selectors = [
            (By.ID, "submit"),
            (By.CSS_SELECTOR, "button[type='submit']"),
            (By.CSS_SELECTOR, "input[type='submit']"),
            (By.XPATH, "//button[contains(.,'Login') or contains(.,'Log in') or contains(.,'Sign in')]"),
            (By.CSS_SELECTOR, "[class*='login'][class*='btn'], .login-btn, .submit-btn"),
        ]
        for by, sel in selectors:
            try:
                el = driver.find_element(by, sel)
                if el.is_displayed() and el.is_enabled():
                    el.click()
                    submit_clicked = True
                    logger.info("Clicked submit via %s=%s", by, sel)
                    break
            except Exception:
                continue
        if not submit_clicked:
            try:
                form = driver.find_element(By.CSS_SELECTOR, "form")
                driver.execute_script("arguments[0].submit();", form)
                submit_clicked = True
            except Exception:
                pass
        if not submit_clicked:
            logger.warning("Could not find submit button — manual login may be needed")
        return True
    except Exception as e:
        logger.exception("Login submit failed: %s", e)
        return False


def wait_for_login(driver, timeout: int = WAIT_FOR_2FA_SECONDS) -> bool:
    logger.info("Waiting for login / 2FA...")
    deadline = time.time() + timeout
    while time.time() < deadline:
        try:
            page = driver.page_source.lower()
            if "trustcode" in page or "verification code" in page or "enter the code" in page:
                logger.info("2FA detected — waiting for manual entry (%ds left)", int(deadline - time.time()))
                time.sleep(30)
            if "logout" in page or "my library" in page or "reading history" in page:
                logger.info("Login confirmed via UI heuristics")
                return True
            if "webnovel.com/book" in driver.current_url:
                logger.info("Redirected to book page — logged in")
                return True
        except Exception:
            pass
        time.sleep(2)
    logger.warning("Login wait timed out")
    return False


# ──────────────────────────────────────────────
# CATALOG & PARSING
# ──────────────────────────────────────────────

def extract_book_id_from_url(book_url: str) -> Optional[str]:
    if not book_url or not book_url.strip():
        return None
    parsed = urlparse(book_url.strip())
    path = (parsed.path or "").strip("/")
    parts = path.split("/")
    if len(parts) >= 2 and parts[0].lower() == "book":
        candidate = parts[1]
        if candidate.isdigit():
            return candidate
        m = re.search(r"_(\d+)$", candidate)
        if m:
            return m.group(1)
    m = re.search(r"/book/(\d+)", path or book_url, re.I)
    if m:
        return m.group(1)
    m = re.search(r"_(\d+)(?:/|$)", path or book_url)
    if m:
        return m.group(1)
    return None


def get_chapter_links_from_catalog(driver, book_id: str) -> List[Dict[str, str]]:
    catalog_url = f"https://www.webnovel.com/book/{book_id}/catalog"
    if not safe_navigate(driver, catalog_url):
        logger.warning("Catalog page may have unresolved captcha")
    else:
        time.sleep(2.5)
    soup = bs(driver.page_source, "lxml")
    anchors = []
    for a in soup.select(".clearfix.g_row.content-list.mb32 li a"):
        href = a.get("href")
        cid = None
        parent = a.parent
        if parent and parent.has_attr("data-cid"):
            cid = parent.get("data-cid")
        elif a.has_attr("data-cid"):
            cid = a.get("data-cid")
        anchors.append({"href": href, "data-cid": cid})
    logger.info("Found %d chapter links in catalog", len(anchors))
    return anchors


def parse_paragraphs_from_html(html: str) -> Tuple[str, List[str]]:
    soup = bs(html, "lxml")
    title_el = soup.select_one(".cha-tit h1")
    title = title_el.get_text(strip=True) if title_el else ""
    paragraphs = []
    for div in soup.find_all("div", class_=lambda c: c and ("cha-paragraph" in c or "j_paragraph" in c)):
        inner = div.find("div", class_=lambda c: c and "dib" in c and "pr" in c)
        target = inner if inner else div
        for p in target.find_all("p"):
            txt = p.get_text(strip=True)
            if txt:
                paragraphs.append(txt)
    if not paragraphs:
        for p in soup.select(".dib.pr p"):
            txt = p.get_text(strip=True)
            if txt:
                paragraphs.append(txt)
    if not paragraphs:
        for p in soup.find_all("p"):
            txt = p.get_text(strip=True)
            if txt:
                paragraphs.append(txt)
    return title, paragraphs


# ──────────────────────────────────────────────
# CHECKPOINT
# ──────────────────────────────────────────────

def load_checkpoint() -> dict:
    if os.path.exists(CHECKPOINT_FILE):
        try:
            with open(CHECKPOINT_FILE, "r", encoding="utf-8") as f:
                data = json.load(f)
            logger.info("Checkpoint loaded: %d chapters done", len(data.get("completed", [])))
            return data
        except Exception:
            pass
    return {"completed": []}


def save_checkpoint(completed_list: list):
    with open(CHECKPOINT_FILE, "w", encoding="utf-8") as f:
        json.dump({"completed": completed_list, "timestamp": datetime.now().isoformat()}, f, indent=2)


# ──────────────────────────────────────────────
# DOCX BATCH WRITER
# ──────────────────────────────────────────────

class DocxBatchWriter:
    def __init__(self, output_folder: str, book_name: str, chapters_per_doc: int = CHUNK_SIZE):
        self.output_folder   = output_folder
        self.book_name       = book_name
        self.chapters_per_doc = chapters_per_doc
        os.makedirs(output_folder, exist_ok=True)
        self._cache: Dict[tuple, dict] = {}

    def _key(self, chapter_no: int) -> tuple:
        start = ((chapter_no - 1) // self.chapters_per_doc) * self.chapters_per_doc + 1
        return (start, start + self.chapters_per_doc - 1)

    def _get_or_create(self, key: tuple) -> dict:
        if key in self._cache:
            return self._cache[key]
        doc_start, doc_end = key
        safe_name = re.sub(r'[\\/:"*?<>|]+', "_", self.book_name)
        filename  = f"{safe_name}_ch{doc_start}-{doc_end}.docx"
        filepath  = os.path.join(self.output_folder, filename)
        if os.path.exists(filepath):
            doc = Document(filepath)
            logger.info("Loaded existing doc: %s", filename)
        else:
            doc = Document()
            style = doc.styles["Normal"]
            style.font.name = "Calibri"
            style.font.size = Pt(11)
            tp = doc.add_paragraph()
            tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            tr = tp.add_run(self.book_name)
            tr.bold = True
            tr.font.size = Pt(18)
            sp = doc.add_paragraph()
            sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sr = sp.add_run(f"Chapters {doc_start} – {doc_end}")
            sr.font.size = Pt(14)
            sr.font.color.rgb = RGBColor(100, 100, 100)
            doc.add_page_break()
            logger.info("Created new doc: %s", filename)
        entry = {"doc": doc, "path": filepath, "count": 0}
        self._cache[key] = entry
        return entry

    def add_chapter(self, chapter_no: int, chapter_name: str, paragraphs: List[str]):
        entry = self._get_or_create(self._key(chapter_no))
        doc   = entry["doc"]
        doc.add_heading(f"Chapter {chapter_no}: {chapter_name}", level=1)
        for para in paragraphs:
            para = para.strip()
            if para:
                p = doc.add_paragraph(para)
                p.paragraph_format.space_after = Pt(4)
        doc.add_paragraph("")
        entry["count"] += 1

    def save_for_chapter(self, chapter_no: int):
        key = self._key(chapter_no)
        if key in self._cache:
            entry = self._cache[key]
            try:
                entry["doc"].save(entry["path"])
            except Exception as e:
                logger.error("Error saving %s: %s", entry["path"], e)

    def save_all(self):
        for entry in self._cache.values():
            try:
                entry["doc"].save(entry["path"])
                logger.info("Saved %s (%d chapters)", os.path.basename(entry["path"]), entry["count"])
            except Exception as e:
                logger.error("Error saving %s: %s", entry["path"], e)


# ──────────────────────────────────────────────
# CHAPTER SCRAPER (inner)
# ──────────────────────────────────────────────

def _scrape_chapter_inner(driver, url: str) -> Tuple[Optional[List[str]], str]:
    if not safe_navigate(driver, url):
        return None, ""
    page = driver.page_source
    if not page or len(page) < MIN_PAGE_CONTENT_LENGTH:
        return None, ""
    title, paragraphs = parse_paragraphs_from_html(page)
    if not title:
        soup = bs(page, "lxml")
        h = soup.select_one(".cha-tit h1")
        title = h.get_text(strip=True) if h else ""
    return (paragraphs if paragraphs else None), title


def scrape_chapter(driver, url: str) -> Tuple[Optional[List[str]], str]:
    return run_with_timeout(_scrape_chapter_inner, args=(driver, url), timeout=CHAPTER_TIMEOUT)


# ──────────────────────────────────────────────
# ARG PARSING
# ──────────────────────────────────────────────

def parse_args():
    parser = argparse.ArgumentParser(description="Webnovel scraper (undetected-chrome + 2captcha)")
    parser.add_argument("book_url", nargs="?", default=None,
                        help="Webnovel book URL. Falls back to CONFIG BOOK_ID if omitted.")
    parser.add_argument("--out-dir",       "-o", default=None)
    parser.add_argument("--start-chapter", "-s", type=int, default=None)
    parser.add_argument("--end-chapter",   "-e", type=int, default=None)
    parser.add_argument("--confirm-login", "-c", action="store_true",
                        help="Pause and wait for Enter after login before scraping.")
    parser.add_argument("--headless", action="store_true", help="Run Chrome headless.")
    parser.add_argument("--email", default=None, help="Webnovel account email.")
    parser.add_argument("--password", default=None, help="Webnovel account password.")
    return parser.parse_args()


# ──────────────────────────────────────────────
# MAIN
# ──────────────────────────────────────────────

def main():
    args = parse_args()

    book_id = BOOK_ID
    if args.book_url:
        extracted = extract_book_id_from_url(args.book_url)
        if extracted:
            book_id = extracted
            logger.info("Using book_id %s from URL", book_id)
        else:
            logger.warning("Could not extract book_id from URL; using CONFIG BOOK_ID")

    out_folder = (args.out_dir or OUT_FOLDER).strip() or OUT_FOLDER
    book_page  = f"https://www.webnovel.com/book/{book_id}"

    # ── Phase 1: Login & cookie setup ──
    global HEADLESS
    if args.headless:
        HEADLESS = True

    driver = create_driver(headless=HEADLESS)
    email_to_use = args.email or EMAIL
    password_to_use = args.password or PASSWORD
    submitted = open_login_and_submit(driver, email_to_use, password_to_use, book_page)
    if submitted:
        wait_for_login(driver, timeout=WAIT_FOR_2FA_SECONDS)
    else:
        logger.warning("Auto-submit failed — please login manually in the browser")
        try:
            safe_navigate(driver, book_page)
        except Exception:
            pass
        logger.info("Manual login window open for %d seconds", WAIT_FOR_2FA_SECONDS)
        time.sleep(WAIT_FOR_2FA_SECONDS)

    cookies = dump_cookies(driver, COOKIE_FILE)
    if not cookies:
        logger.error("No cookies after login — exiting")
        kill_driver(driver)
        sys.exit(1)

    apply_cookies_from_file(driver, COOKIE_FILE, target_url=book_page)

    # Verify login
    def is_logged_in() -> bool:
        from selenium.webdriver.common.by import By
        try:
            safe_navigate(driver, book_page)
            page = driver.page_source.lower()
            if "logout" in page or "my library" in page or "reading history" in page:
                return True
            if driver.find_elements(By.CSS_SELECTOR, "[class*='avatar'], [class*='user-avatar']"):
                return True
        except Exception:
            pass
        return False

    if not is_logged_in():
        logger.warning("Not confirmed logged in — scraping may fail")
    else:
        logger.info("Login confirmed")

    if args.confirm_login:
        print("\n" + "=" * 60)
        print("  Verify login in browser, then press ENTER to start scraping.")
        print("=" * 60)
        try:
            input("  >>> Press ENTER to begin... ")
        except EOFError:
            pass

    # ── Phase 2: Extract catalog ──
    chapter_links = get_chapter_links_from_catalog(driver, book_id)
    if not chapter_links:
        logger.error("No chapters found in catalog — exiting")
        kill_driver(driver)
        sys.exit(1)

    base = "https://www.webnovel.com"
    chapter_urls = [
        (c["href"] if c["href"].startswith("http") else base + c["href"])
        for c in chapter_links if c.get("href")
    ]
    cids = [c.get("data-cid") for c in chapter_links]
    logger.info("Collected %d chapter URLs", len(chapter_urls))

    # Determine range
    start_index = 0
    if START_CHAPTER_ID and START_CHAPTER_ID in cids:
        start_index = cids.index(START_CHAPTER_ID)

    effective_start = args.start_chapter if args.start_chapter is not None else START_CHAPTER_NUMBER
    effective_end   = args.end_chapter   if args.end_chapter   is not None else END_CHAPTER_NUMBER

    start_offset = start_index
    if args.start_chapter is not None and args.start_chapter > 0:
        start_offset = args.start_chapter - 1

    chapters_to_process = chapter_urls[start_offset:]
    if effective_end is not None:
        count = effective_end - effective_start + 1
        if count <= 0:
            logger.error("END_CHAPTER_NUMBER must be >= START_CHAPTER_NUMBER")
            kill_driver(driver)
            sys.exit(1)
        chapters_to_process = chapters_to_process[:count]

    logger.info("Will scrape %d chapters (ch %s to %s)", len(chapters_to_process), effective_start, effective_end)

    # Done with login driver — kill it; fresh ones created per batch
    kill_driver(driver)
    driver = None

    # ── Phase 3: Scrape with batch restarts ──
    checkpoint   = load_checkpoint()
    completed    = set(checkpoint.get("completed", []))
    doc_writer   = DocxBatchWriter(out_folder, out_folder, CHUNK_SIZE)
    failed       = []
    batch_count  = 0
    consec_fails = 0
    total_scraped = len(completed)
    total_to_scrape = len(chapters_to_process)

    # Build chapter list with numbers
    chapters = [
        {"chapter_no": effective_start + i, "url": url}
        for i, url in enumerate(chapters_to_process)
    ]
    pending = [ch for ch in chapters if ch["chapter_no"] not in completed]
    logger.info("%d chapters remaining", len(pending))

    if not pending:
        logger.info("All chapters already scraped!")
        return

    try:
        for ch in pending:
            ch_no  = ch["chapter_no"]
            ch_url = ch["url"]

            if batch_count >= BATCH_SIZE:
                logger.info("Batch limit (%d) — recycling driver...", BATCH_SIZE)
                doc_writer.save_all()
                kill_driver(driver)
                driver = None

            if consec_fails >= MAX_CONSECUTIVE_FAILS:
                logger.info("%d consecutive failures — recycling driver...", MAX_CONSECUTIVE_FAILS)
                doc_writer.save_all()
                kill_driver(driver)
                driver = None
                consec_fails = 0

            if driver is None:
                cooldown = random.uniform(COOLDOWN_MIN, COOLDOWN_MAX)
                logger.info("Cooling down %.1fs before new driver...", cooldown)
                time.sleep(cooldown)
                driver = create_driver_with_cookies(book_page)
                batch_count = 0

            delay = random.uniform(CHAPTER_DELAY_MIN, CHAPTER_DELAY_MAX)
            logger.info("[Ch %d | %d left] Waiting %.1fs...", ch_no, total_to_scrape - total_scraped, delay)
            time.sleep(delay)

            success = False
            for attempt in range(1, 4):
                if attempt > 1:
                    retry_delay = random.uniform(5, 12)
                    logger.info("Retry %d/3 — restarting driver, waiting %.1fs", attempt, retry_delay)
                    kill_driver(driver)
                    driver = None
                    time.sleep(retry_delay)
                    cooldown = random.uniform(COOLDOWN_MIN, COOLDOWN_MAX)
                    time.sleep(cooldown)
                    driver = create_driver_with_cookies(book_page)
                    batch_count = 0

                try:
                    t0 = time.time()
                    paragraphs, title = scrape_chapter(driver, ch_url)
                    logger.info("  Took %.1fs", time.time() - t0)

                    if paragraphs:
                        # Filter: keep lines with English chars or pure Chinese
                        kept = [
                            p for p in paragraphs
                            if re.search(r"[a-zA-Z0-9]", p)
                            or re.fullmatch(r"[一-鿿\s]+", p)
                        ]
                        final_name = title or f"Chapter {ch_no}"
                        doc_writer.add_chapter(ch_no, final_name, kept)
                        doc_writer.save_for_chapter(ch_no)
                        completed.add(ch_no)
                        save_checkpoint(list(completed))
                        total_scraped += 1
                        batch_count   += 1
                        consec_fails   = 0
                        logger.info("  Saved ch %d — %d paras, %d words", ch_no, len(kept),
                                    len(" ".join(kept).split()))
                        success = True
                        break
                    else:
                        logger.warning("  No content (attempt %d)", attempt)

                except ChapterTimeoutError:
                    logger.warning("  TIMEOUT after %ds — killing driver!", CHAPTER_TIMEOUT)
                    kill_driver(driver)
                    driver = None

                except Exception as e:
                    logger.error("  Error (attempt %d): %s", attempt, e)
                    try:
                        _ = driver.title
                    except Exception:
                        logger.warning("  Driver dead — will recreate")
                        kill_driver(driver)
                        driver = None

            if not success:
                consec_fails += 1
                failed.append(ch_no)
                logger.error("FAILED ch %d after 3 attempts (consec_fails=%d)", ch_no, consec_fails)

    except KeyboardInterrupt:
        logger.warning("Interrupted by user")
    except Exception as e:
        logger.exception("Fatal error: %s", e)
    finally:
        doc_writer.save_all()

        logger.info("=" * 60)
        logger.info("SUMMARY")
        logger.info("Scraped: %d/%d", total_scraped, total_to_scrape)
        logger.info("Failed:  %d — %s", len(failed), failed)
        logger.info("Output:  %s/", out_folder)

        if os.path.exists(out_folder):
            for f in sorted(x for x in os.listdir(out_folder) if x.endswith(".docx")):
                fpath = os.path.join(out_folder, f)
                logger.info("  %s (%.1f KB)", f, os.path.getsize(fpath) / 1024)

        kill_driver(driver)
        logger.info("Browser closed")

        if total_scraped >= total_to_scrape and not failed:
            try:
                os.remove(CHECKPOINT_FILE)
                logger.info("All chapters scraped — checkpoint cleaned up")
            except Exception:
                pass


if __name__ == "__main__":
    main()
