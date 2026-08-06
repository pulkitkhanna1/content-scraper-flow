#!/usr/bin/env python3
"""
WTR-Lab chapter scraper.

Default input:
    content_crawling/wtr-lab/url.txt

url.txt format:
    https://wtr-lab.com/en/novel/49763/book-slug    All
    https://wtr-lab.com/en/novel/49763/book-slug    1-100

Output:
    content_crawling/wtr-lab/<Book_Name>/<Book_Name>_chapters_0001-0100.docx
    content_crawling/wtr-lab/<Book_Name>/chapter_urls.json
    content_crawling/wtr-lab/<Book_Name>/progress.json

Notes:
    The default service is "web" because it can crawl public full books and
    the script can decrypt its encrypted reader payloads.

    WTR-Lab's AI reader allows anonymous preview chapters only. If AI returns
    the registration wall, the script retries through Google Web Translation
    (?service=web) by default.
"""

from __future__ import annotations

import argparse
import base64
import json
import logging
import os
import random
import re
import shutil
import shlex
import sys
import tempfile
import time
from dataclasses import asdict, dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple
from urllib.parse import urljoin, urlparse

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt


SCRIPT_DIR = Path(__file__).parent.resolve()
DEFAULT_URLS_FILE = SCRIPT_DIR / "url.txt"
DEFAULT_COOKIE_FILE = SCRIPT_DIR / "cookie.txt"
WTR_BASE_URL = "https://wtr-lab.com"
DEFAULT_CHROME_USER_DATA_DIR = Path.home() / "Library" / "Application Support" / "Google" / "Chrome"
DEFAULT_CHROME_PROFILE = os.getenv("WTR_CHROME_PROFILE", "Profile 1")
DEFAULT_BATCH_SIZE = 100
DEFAULT_DELAY = 1.0
DEFAULT_JITTER = 0.35
DEFAULT_TURNSTILE_RETRIES = 2
DEFAULT_TURNSTILE_COOLDOWN = 180.0
DEFAULT_TURNSTILE_CHECK_INTERVAL = 15.0
DEFAULT_BROWSER_FETCH_RETRIES = 3
DEFAULT_BROWSER_FETCH_DELAY = 15.0
DEFAULT_CHAPTER_NOT_FOUND_RETRIES = 3
DEFAULT_CHAPTER_NOT_FOUND_DELAY = 5.0
REQUEST_TIMEOUT = 45
MAX_RETRIES = 4
WTR_AES_KEY = b"IJAFUUxjM25hyzL2AZrn0wl7cESED6Ru"[:32]
CJK_RE = re.compile(r"[\u3400-\u9fff]")

HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/126.0.0.0 Safari/537.36"
    ),
    "Accept": "application/json,text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    "Accept-Language": "en-US,en;q=0.9",
    "Accept-Encoding": "gzip, deflate",
    "Origin": WTR_BASE_URL,
}

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[logging.StreamHandler(sys.stdout)],
)
log = logging.getLogger("wtr-lab")


class TurnstileRequired(RuntimeError):
    """Raised when WTR-Lab asks this session to solve browser verification."""


class AiRegistrationRequired(RuntimeError):
    """Raised when WTR-Lab blocks anonymous AI translation after preview chapters."""


class RenderedTranslationNotReady(RuntimeError):
    """Raised when the browser-rendered Google translation did not finish."""


class BrowserFetchFailed(RuntimeError):
    """Raised when Selenium/browser fetch fails before WTR-Lab returns a response."""


class PartialBookStop(RuntimeError):
    """Carries saved chapter count when a book stops before process_book can return."""

    def __init__(self, message: str, *, scraped: int, stop_all: bool = False) -> None:
        super().__init__(message)
        self.scraped = scraped
        self.stop_all = stop_all


class ChromeProfileSession:
    CACHE_DIR_NAMES = {
        "Application Cache",
        "BrowserMetrics",
        "Cache",
        "Code Cache",
        "Crashpad",
        "DawnCache",
        "GPUCache",
        "GrShaderCache",
        "Media Cache",
        "OptimizationGuidePredictionModels",
        "Safe Browsing",
        "ShaderCache",
        "VideoDecodeStats",
    }

    def __init__(
        self,
        source_user_data_dir: Path,
        profile_name: str,
        *,
        headless: bool = False,
        keep_temp_profile: bool = False,
        warmup_seconds: float = 8.0,
    ) -> None:
        self.source_user_data_dir = source_user_data_dir.expanduser()
        self.profile_name = profile_name
        self.headless = headless
        self.keep_temp_profile = keep_temp_profile
        self.warmup_seconds = warmup_seconds
        self.temp_user_data_dir: Optional[Path] = None
        self.driver = None

    def _ignore_profile_files(self, _dir: str, names: List[str]) -> set[str]:
        ignored = {name for name in names if name in self.CACHE_DIR_NAMES}
        ignored.update({name for name in names if name.startswith("Singleton")})
        ignored.add("DevToolsActivePort")
        return ignored

    def copy_profile(self) -> Path:
        source_profile = self.source_user_data_dir / self.profile_name
        if not source_profile.exists():
            raise FileNotFoundError(f"Chrome profile not found: {source_profile}")

        temp_root = Path(tempfile.mkdtemp(prefix="wtr_chrome_profile_"))
        local_state = self.source_user_data_dir / "Local State"
        if local_state.exists():
            shutil.copy2(local_state, temp_root / "Local State")

        shutil.copytree(
            source_profile,
            temp_root / self.profile_name,
            ignore=self._ignore_profile_files,
            dirs_exist_ok=True,
        )
        self.temp_user_data_dir = temp_root
        return temp_root

    def start(self) -> None:
        temp_root = self.copy_profile()
        try:
            from selenium import webdriver
            from selenium.webdriver.chrome.options import Options
        except ImportError as exc:
            raise RuntimeError("Chrome profile mode needs selenium installed.") from exc

        options = Options()
        options.add_argument(f"--user-data-dir={temp_root}")
        options.add_argument(f"--profile-directory={self.profile_name}")
        options.page_load_strategy = "none"
        options.add_argument("--window-size=1400,1000")
        options.add_argument("--disable-blink-features=AutomationControlled")
        options.add_argument("--disable-dev-shm-usage")
        options.add_argument("--no-sandbox")
        if self.headless:
            options.add_argument("--headless=new")

        log.info("Starting Chrome with copied profile: %s", temp_root)
        self.driver = webdriver.Chrome(options=options)
        self.driver.set_page_load_timeout(20)
        self._navigate(f"{WTR_BASE_URL}/en", "Initial WTR-Lab")
        if self.warmup_seconds > 0:
            log.info("Chrome warmup %.1fs. If WTR-Lab shows verification, solve it in this window.", self.warmup_seconds)
            time.sleep(self.warmup_seconds)

    def _stop_loading(self) -> None:
        if not self.driver:
            return
        try:
            self.driver.execute_script("window.stop();")
        except Exception:
            pass

    def _navigate(self, url: str, label: str) -> None:
        if not self.driver:
            raise RuntimeError("Chrome browser is not started")
        try:
            self.driver.get(url)
        except Exception as exc:
            if exc.__class__.__name__ != "TimeoutException":
                raise
            log.warning("%s page load timed out; stopping page load and continuing.", label)
            self._stop_loading()

    def _ensure_wtr_origin(self) -> None:
        if not self.driver:
            raise RuntimeError("Chrome browser is not started")
        current_url = ""
        try:
            current_url = self.driver.current_url or ""
        except Exception:
            pass
        if not current_url.startswith(WTR_BASE_URL):
            self._navigate(f"{WTR_BASE_URL}/en", "WTR-Lab origin")
            time.sleep(2)

    def click_translation_link(self, chapter_url: str, service: str, timeout: float = 25.0) -> str:
        if not self.driver:
            raise RuntimeError("Chrome browser is not started")

        try:
            from selenium.webdriver.common.by import By
        except ImportError as exc:
            raise RuntimeError("Chrome profile mode needs selenium installed.") from exc

        self._navigate(chapter_url, "AI reader")
        time.sleep(2)
        deadline = time.time() + timeout
        service_marker = f"service={service}"
        fallback_url = f"{chapter_url.split('?', 1)[0]}?service={service}"

        while time.time() < deadline:
            anchors = self.driver.find_elements(By.CSS_SELECTOR, f"a[href*='{service_marker}']")
            for anchor in anchors:
                href = anchor.get_attribute("href") or ""
                text = (anchor.text or "").strip()
                if service_marker not in href:
                    continue
                if service == "web" and text and "google web translation" not in text.lower():
                    continue

                clicked_url = urljoin(chapter_url, href)
                self.driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", anchor)
                time.sleep(0.2)
                self.driver.execute_script("arguments[0].click();", anchor)
                time.sleep(2)
                return clicked_url
            time.sleep(0.5)

        log.warning("Could not find %s translation link on page; opening %s directly.", service, fallback_url)
        self._navigate(fallback_url, f"{service} reader")
        time.sleep(2)
        return fallback_url

    def open_verification_page(self, url: str, wait_seconds: float, check_interval: float) -> None:
        if not self.driver:
            raise RuntimeError("Chrome browser is not started")
        self._navigate(url, "Turnstile verification")
        log.warning(
            "Chrome is open on the challenged page. Solve/refresh WTR verification there; retrying every %.0fs for up to %.0fs.",
            check_interval,
            wait_seconds,
        )

    def rendered_chapter_content(
        self,
        url: str,
        chapter_order: int,
        fallback_title: str,
        timeout: float = 120.0,
    ) -> Tuple[str, List[str], Dict[str, Any]]:
        if not self.driver:
            raise RuntimeError("Chrome browser is not started")

        self._navigate(url, "Rendered reader")
        deadline = time.time() + timeout
        last_reload = time.time()
        last_status = 0.0
        data: Dict[str, Any] = {}

        script = """
            const active = document.querySelector('.chapter-tracker.active')
                || document.querySelector('.chapter-tracker')
                || document;
            const body = active.querySelector('.chapter-body') || document.querySelector('.chapter-body');
            const title = active.querySelector('.chapter-title, h1, h2, h3, [class*="chapter-title"]');
            return {
                body: body ? body.innerText : '',
                title: title ? title.innerText : '',
                tracker: active ? active.innerText : '',
                url: location.href
            };
        """
        while time.time() < deadline:
            current_url = ""
            try:
                current_url = self.driver.current_url or ""
            except Exception:
                pass
            if current_url and not same_url_path(current_url, url):
                time.sleep(0.5)
                continue

            data = self.driver.execute_script(script) or {}
            body_text = clean_text(data.get("body"))
            tracker_text = clean_text(data.get("tracker"))
            if is_security_challenge_text(tracker_text):
                raise TurnstileRequired(
                    "WTR-Lab requested Turnstile/browser verification for this rendered chapter page."
                )
            if len(body_text) > 100 and cjk_ratio(body_text) < 0.08:
                break
            if len(body_text) > 100 and time.time() - last_status > 15:
                log.warning(
                    "Rendered Web Translation for chapter %s is still loading/translating (CJK ratio %.2f).",
                    chapter_order,
                    cjk_ratio(body_text),
                )
                last_status = time.time()
            if len(body_text) > 100 and time.time() - last_reload > 35:
                self._navigate(url, "Rendered reader refresh")
                last_reload = time.time()
            time.sleep(1)

        body_text = clean_text(data.get("body"))
        if not body_text:
            raise RuntimeError(f"No rendered chapter body found for chapter {chapter_order}")
        if is_security_challenge_text(body_text):
            raise TurnstileRequired(
                "WTR-Lab requested Turnstile/browser verification for this rendered chapter page."
            )
        if cjk_ratio(body_text) >= 0.2:
            raise RenderedTranslationNotReady(
                f"Rendered Web Translation for chapter {chapter_order} still looks untranslated. "
                "Wait for the browser page to finish translating, then resume."
            )

        lines = [clean_text(line) for line in body_text.splitlines()]
        lines = [line for line in lines if line]
        if not lines:
            raise RuntimeError(f"No rendered chapter paragraphs found for chapter {chapter_order}")

        title = clean_text(data.get("title")) or clean_text(fallback_title) or f"Chapter {chapter_order}"
        if re.search(r"\bchapter\b|\bch\.", lines[0], flags=re.I):
            title = lines.pop(0)
        if not re.search(r"\bchapter\b|\bch\.", title, flags=re.I):
            title = f"Chapter {chapter_order}: {title}"

        return title, lines, {
            "source": "rendered_browser",
            "url": data.get("url") or url,
            "service": "web",
        }

    def get_html(self, url: str) -> str:
        if not self.driver:
            raise RuntimeError("Chrome browser is not started")
        self._navigate(url, "WTR-Lab")
        deadline = time.time() + 20
        html = ""
        while time.time() < deadline:
            current_url = ""
            try:
                current_url = self.driver.current_url or ""
            except Exception:
                pass
            if current_url and not same_url_path(current_url, url):
                time.sleep(0.5)
                continue
            html = self.driver.page_source
            if "__NEXT_DATA__" in html:
                return html
            time.sleep(0.5)
        return html or self.driver.page_source

    def request_json(self, method: str, url: str, payload: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        if not self.driver:
            raise RuntimeError("Chrome browser is not started")
        self._ensure_wtr_origin()

        target = url
        if target.startswith(WTR_BASE_URL):
            target = target[len(WTR_BASE_URL) :]

        script = """
            const method = arguments[0];
            const url = arguments[1];
            const payload = arguments[2];
            const done = arguments[arguments.length - 1];
            const options = {
                method,
                credentials: 'include',
                headers: {
                    'Accept': 'application/json',
                    'Content-Type': 'application/json'
                }
            };
            if (payload !== null && payload !== undefined) {
                options.body = JSON.stringify(payload);
            }
            fetch(url, options)
                .then(async response => {
                    const text = await response.text();
                    done(JSON.stringify({
                        ok: response.ok,
                        status: response.status,
                        text
                    }));
                })
                .catch(error => done(JSON.stringify({
                    ok: false,
                    status: 0,
                    error: String(error)
                })));
        """
        result = self.driver.execute_async_script(script, method, target, payload)
        envelope = json.loads(result)
        if not envelope.get("ok"):
            message = f"Browser fetch failed HTTP {envelope.get('status')}: {envelope.get('error') or envelope.get('text')}"
            if int(envelope.get("status") or 0) == 0:
                raise BrowserFetchFailed(message)
            raise RuntimeError(message)

        text = envelope.get("text") or "{}"
        try:
            return json.loads(text)
        except json.JSONDecodeError as exc:
            raise RuntimeError(f"Browser fetch did not return JSON: {text[:300]}") from exc

    def close(self) -> None:
        if self.driver:
            self.driver.quit()
            self.driver = None
        if self.temp_user_data_dir and not self.keep_temp_profile:
            shutil.rmtree(self.temp_user_data_dir, ignore_errors=True)


@dataclass
class BookSource:
    book_name: Optional[str]
    novel_url: str
    selector: str = "All"


@dataclass
class ChapterLink:
    order: int
    title: str
    name: str
    url: str
    id: Optional[int] = None
    serie_id: Optional[int] = None
    updated_at: Optional[str] = None


@dataclass
class NovelMeta:
    raw_id: int
    title: str
    slug: str
    chapter_count: Optional[int]
    default_service: str
    novel_url: str


def clean_text(value: Any) -> str:
    value = "" if value is None else str(value)
    value = value.replace("\xa0", " ").replace("\u3000", " ")
    value = re.sub(r"[ \t\r\f\v]+", " ", value)
    value = re.sub(r"\n\s+", "\n", value)
    return value.strip()


def cjk_ratio(value: str) -> float:
    text = clean_text(value)
    letters = [char for char in text if char.isalpha()]
    if not letters:
        return 0.0
    return sum(1 for char in letters if CJK_RE.match(char)) / len(letters)


def is_security_challenge_text(value: str) -> bool:
    text = clean_text(value).lower()
    return (
        "security check required" in text
        or "detected unusual reading activity" in text
        or "open challenge in new page" in text
        or "verifying your request" in text
    )


def same_url_path(current_url: str, target_url: str) -> bool:
    try:
        current = urlparse(current_url)
        target = urlparse(target_url)
    except Exception:
        return False
    return current.netloc == target.netloc and current.path.rstrip("/") == target.path.rstrip("/")


def sanitize_filename(value: str, max_len: int = 100) -> str:
    value = clean_text(value)
    value = re.sub(r'[<>:"/\\|?*\x00-\x1f]+', "_", value)
    value = re.sub(r"\s+", "_", value)
    value = value.strip("._ ")
    return (value or "WTR_Lab_Book")[:max_len]


def normalize_novel_url(url: str) -> str:
    url = url.strip().split("#", 1)[0].split("?", 1)[0].rstrip("/")
    return re.sub(r"/chapter-\d+$", "", url, flags=re.I)


def chapter_service_url(chapter: "ChapterLink", service: str) -> str:
    if service == "ai":
        return chapter.url
    return f"{chapter.url}?service={service}"


def extract_raw_id(url: str) -> int:
    match = re.search(r"/novel/(\d+)(?:/|$)", url)
    if not match:
        raise ValueError(f"Could not find WTR-Lab raw_id in URL: {url}")
    return int(match.group(1))


def title_from_slug(url: str) -> str:
    parts = [p for p in urlparse(url).path.split("/") if p]
    slug = parts[3] if len(parts) >= 4 and parts[1] == "novel" else parts[-1]
    return re.sub(r"[-_]+", " ", slug).strip().title() or "WTR Lab Book"


def make_session(cookie: Optional[str] = None) -> requests.Session:
    session = requests.Session()
    session.headers.update(HEADERS)
    if cookie:
        session.headers["Cookie"] = cookie
    return session


def request_with_retries(
    session: requests.Session,
    method: str,
    url: str,
    *,
    expected_json: bool = False,
    **kwargs: Any,
) -> Any:
    last_error: Optional[Exception] = None
    for attempt in range(1, MAX_RETRIES + 1):
        try:
            response = session.request(method, url, timeout=REQUEST_TIMEOUT, **kwargs)
            if response.status_code in {429, 500, 502, 503, 504}:
                raise RuntimeError(f"HTTP {response.status_code}: {response.text[:200]}")
            response.raise_for_status()
            return response.json() if expected_json else response.text
        except Exception as exc:
            last_error = exc
            log.warning("Request failed (%s/%s): %s", attempt, MAX_RETRIES, exc)
            if attempt < MAX_RETRIES:
                time.sleep(2 * attempt)
    raise RuntimeError(f"Request failed for {url}: {last_error}")


def parse_next_data(html: str) -> Dict[str, Any]:
    soup = BeautifulSoup(html, "html.parser")
    node = soup.find("script", id="__NEXT_DATA__")
    if not node:
        return {}
    try:
        return json.loads(node.get_text())
    except json.JSONDecodeError:
        return {}


def fetch_novel_meta(
    session: requests.Session,
    source_url: str,
    browser: Optional[ChromeProfileSession] = None,
) -> NovelMeta:
    novel_url = normalize_novel_url(source_url)
    raw_id = extract_raw_id(novel_url)
    slug = urlparse(novel_url).path.rstrip("/").split("/")[-1]
    title = title_from_slug(novel_url)
    chapter_count: Optional[int] = None
    default_service = "ai"

    if browser:
        html = browser.get_html(novel_url)
    else:
        html = request_with_retries(
            session,
            "GET",
            novel_url,
            headers={"Referer": WTR_BASE_URL + "/en"},
        )
    next_data = parse_next_data(html)
    page_props = next_data.get("props", {}).get("pageProps", {})
    serie = page_props.get("serie") or {}
    serie_data = serie.get("serie_data") or {}

    data = serie_data.get("data") or {}
    if data.get("title"):
        title = clean_text(data["title"])
    if serie_data.get("slug"):
        slug = clean_text(serie_data["slug"])
    if serie_data.get("raw_id"):
        raw_id = int(serie_data["raw_id"])
    if serie_data.get("chapter_count") is not None:
        chapter_count = int(serie_data["chapter_count"])
    if serie.get("default_service"):
        default_service = clean_text(serie["default_service"])

    return NovelMeta(
        raw_id=raw_id,
        title=title,
        slug=slug,
        chapter_count=chapter_count,
        default_service=default_service,
        novel_url=novel_url,
    )


def fetch_chapter_links(
    session: requests.Session,
    meta: NovelMeta,
    browser: Optional[ChromeProfileSession] = None,
) -> List[ChapterLink]:
    endpoint = f"{WTR_BASE_URL}/api/chapters/{meta.raw_id}"
    if browser:
        payload = browser.request_json("GET", endpoint)
    else:
        payload = request_with_retries(
            session,
            "GET",
            endpoint,
            expected_json=True,
            headers={"Referer": meta.novel_url},
        )
    chapters = payload.get("chapters") or []
    links: List[ChapterLink] = []
    for item in chapters:
        order = int(item.get("order") or 0)
        if order <= 0:
            continue
        title = clean_text(item.get("title")) or clean_text(item.get("name")) or f"Chapter {order}"
        links.append(
            ChapterLink(
                order=order,
                title=title,
                name=clean_text(item.get("name")),
                url=f"{meta.novel_url}/chapter-{order}",
                id=item.get("id"),
                serie_id=item.get("serie_id"),
                updated_at=item.get("updated_at"),
            )
        )

    if not links:
        raise RuntimeError(f"No chapters returned for raw_id={meta.raw_id}")

    links.sort(key=lambda chapter: chapter.order)
    return links


def decrypt_wtr_value(value: str) -> Any:
    if not isinstance(value, str):
        return value

    is_array = False
    encrypted = value
    if encrypted.startswith("arr:"):
        is_array = True
        encrypted = encrypted[4:]
    elif encrypted.startswith("str:"):
        encrypted = encrypted[4:]
    else:
        return value

    try:
        from cryptography.hazmat.primitives.ciphers.aead import AESGCM
    except ImportError as exc:
        raise RuntimeError(
            "Encrypted WTR-Lab responses need the 'cryptography' package. "
            "Install it or use --service ai with a logged-in cookie."
        ) from exc

    parts = encrypted.split(":")
    if len(parts) != 3:
        raise RuntimeError("Invalid WTR-Lab encrypted payload format")

    iv = base64.b64decode(parts[0])
    tag = base64.b64decode(parts[1])
    ciphertext = base64.b64decode(parts[2])
    plaintext = AESGCM(WTR_AES_KEY).decrypt(iv, ciphertext + tag, None).decode("utf-8")
    return json.loads(plaintext) if is_array else plaintext


def glossary_terms(data_block: Dict[str, Any]) -> List[List[str]]:
    glossary = data_block.get("glossary_data") or {}
    terms = glossary.get("terms") or []
    return terms if isinstance(terms, list) else []


def apply_glossary(text: str, terms: List[List[str]]) -> str:
    if not terms or "※" not in text:
        return text

    def replace(match: re.Match[str]) -> str:
        index = int(match.group(1))
        if 0 <= index < len(terms) and terms[index]:
            return clean_text(terms[index][0])
        return match.group(0)

    return re.sub(r"※(\d+)[⛬〓]", replace, text)


def body_to_paragraphs(body: Any, terms: List[List[str]]) -> List[str]:
    body = decrypt_wtr_value(body) if isinstance(body, str) else body

    if isinstance(body, list):
        raw_paragraphs = body
    elif isinstance(body, str):
        if "<" in body and ">" in body:
            soup = BeautifulSoup(body, "html.parser")
            for tag in soup.find_all(["script", "style", "noscript"]):
                tag.decompose()
            raw_text = soup.get_text("\n")
        else:
            raw_text = body
        raw_paragraphs = re.split(r"\n\s*\n+|\n", raw_text)
    else:
        raw_paragraphs = [body]

    paragraphs: List[str] = []
    for raw in raw_paragraphs:
        text = apply_glossary(clean_text(raw), terms)
        if text:
            paragraphs.append(text)
    return paragraphs


def fetch_chapter_content(
    session: requests.Session,
    meta: NovelMeta,
    chapter: ChapterLink,
    *,
    service: str,
    language: str,
    browser: Optional[ChromeProfileSession] = None,
) -> Tuple[str, List[str], Dict[str, Any]]:
    payload: Dict[str, Any] = {
        "translate": service,
        "language": language,
        "raw_id": meta.raw_id,
        "chapter_no": chapter.order,
    }
    if chapter.id:
        payload["chapter_id"] = chapter.id

    if browser:
        response = browser.request_json("POST", f"{WTR_BASE_URL}/api/reader/get", payload)
    else:
        referer = chapter_service_url(chapter, service)
        response = request_with_retries(
            session,
            "POST",
            f"{WTR_BASE_URL}/api/reader/get",
            expected_json=True,
            json=payload,
            headers={
                "Accept": "application/json",
                "Content-Type": "application/json",
                "Referer": referer,
            },
        )

    if not response.get("success"):
        code = response.get("code") or response.get("error")
        if code == 1401 or str(code) == "1401":
            raise AiRegistrationRequired(
                "Login required for this service. For full AI chapters, pass "
                "--cookie or set WTR_COOKIE. Anonymous AI preview usually stops "
                "after chapter 10."
            )
        if response.get("requireTurnstile"):
            raise TurnstileRequired(
                "WTR-Lab requested Turnstile/browser verification for this session. "
                "Wait a bit and retry more slowly, or pass a fresh browser Cookie "
                "header with --cookie / WTR_COOKIE after solving the challenge in Chrome."
            )
        raise RuntimeError(response.get("error") or response.get("message") or f"API error: {code}")

    data_wrapper = response.get("data") or {}
    data_block = data_wrapper.get("data") or {}
    if data_block.get("encrypted") and data_block.get("body"):
        data_block["body"] = decrypt_wtr_value(data_block["body"])
        data_block.pop("encrypted", None)

    terms = glossary_terms(data_block)
    paragraphs = body_to_paragraphs(data_block.get("body"), terms)
    if not paragraphs:
        raise RuntimeError(f"No paragraphs found for chapter {chapter.order}")

    api_title = clean_text(data_block.get("title"))
    response_title = clean_text((response.get("chapter") or {}).get("title"))
    title = api_title or response_title or chapter.title or f"Chapter {chapter.order}"
    if not re.search(r"\bchapter\b|\bch\.", title, flags=re.I):
        title = f"Chapter {chapter.order}: {title}"

    return title, paragraphs, response


def parse_chapter_selector(selector: str, total: int) -> Tuple[int, int]:
    selector = clean_text(selector or "All")
    if not selector or selector.lower() == "all":
        return 1, total

    normalized = selector.lower().replace("chapter", "").replace("chapters", "")
    match = re.search(r"(\d+)\s*(?:-|:|to)\s*(\d+)", normalized)
    if match:
        start, end = int(match.group(1)), int(match.group(2))
        return max(1, start), min(total, end)

    match = re.search(r"(\d+)\s*(?:-|:|to)\s*$", normalized)
    if match:
        return max(1, int(match.group(1))), total

    match = re.search(r"^\s*(\d+)\s*$", normalized)
    if match:
        chapter = int(match.group(1))
        return chapter, min(total, chapter)

    raise ValueError(f"Unsupported chapter selector: {selector!r}")


def parse_url_file(urls_file: Path) -> List[BookSource]:
    if not urls_file.exists():
        raise FileNotFoundError(f"URLs file not found: {urls_file}")

    sources: List[BookSource] = []
    for line_number, raw_line in enumerate(urls_file.read_text(encoding="utf-8").splitlines(), start=1):
        line = raw_line.strip()
        if not line or line.startswith("#"):
            continue

        try:
            parts = shlex.split(line)
        except ValueError as exc:
            raise ValueError(f"{urls_file}:{line_number}: could not parse line: {exc}") from exc

        url_index = next((i for i, part in enumerate(parts) if re.match(r"^https?://", part)), None)
        if url_index is None:
            raise ValueError(f"{urls_file}:{line_number}: line must contain a URL")

        book_name = " ".join(parts[:url_index]).strip() or None
        novel_url = parts[url_index]
        selector = " ".join(parts[url_index + 1 :]).strip() or "All"
        sources.append(BookSource(book_name=book_name, novel_url=novel_url, selector=selector))

    if not sources:
        raise ValueError(f"No WTR-Lab URLs found in {urls_file}")
    return sources


def save_chapter_cache(
    output_dir: Path,
    meta: NovelMeta,
    chapters: List[ChapterLink],
    selector: str,
) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)
    payload = {
        "source_url": meta.novel_url,
        "raw_id": meta.raw_id,
        "book_name": meta.title,
        "chapter_count": len(chapters),
        "selector": selector,
        "chapters": [asdict(chapter) for chapter in chapters],
    }
    (output_dir / "chapter_urls.json").write_text(
        json.dumps(payload, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def docx_path_for_chapter(output_dir: Path, book_name: str, order: int, batch_size: int) -> Path:
    start = ((order - 1) // batch_size) * batch_size + 1
    end = start + batch_size - 1
    safe_name = sanitize_filename(book_name)
    return output_dir / f"{safe_name}_chapters_{start:04d}-{end:04d}.docx"


def configure_document(doc: Document) -> None:
    doc.styles["Normal"].font.size = Pt(11)


def add_chapter_to_doc(doc: Document, title: str, paragraphs: Iterable[str]) -> None:
    doc.add_heading(title, level=1)
    for text in paragraphs:
        paragraph = doc.add_paragraph(text)
        paragraph.paragraph_format.space_after = Pt(6)


def load_progress(progress_path: Path, raw_id: int, service: str, language: str) -> Dict[str, Any]:
    if not progress_path.exists():
        return {}
    try:
        progress = json.loads(progress_path.read_text(encoding="utf-8"))
    except Exception as exc:
        log.warning("Could not read progress file: %s", exc)
        return {}

    if progress.get("raw_id") != raw_id:
        return {}
    if progress.get("service") != service or progress.get("language") != language:
        return {}
    return progress


def save_progress(
    progress_path: Path,
    meta: NovelMeta,
    chapter: ChapterLink,
    title: str,
    service: str,
    language: str,
) -> None:
    payload = {
        "source_url": meta.novel_url,
        "raw_id": meta.raw_id,
        "book_name": meta.title,
        "service": service,
        "language": language,
        "last_order": chapter.order,
        "last_title": title,
        "last_url": chapter_service_url(chapter, service),
        "updated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
    }
    progress_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def output_dir_for_book(
    args: argparse.Namespace,
    book_name: str,
    from_url_file: bool,
    multiple_books: bool,
) -> Path:
    if not args.output_dir:
        suffix = "_Chinese" if args.chinese else ""
        return SCRIPT_DIR / f"{sanitize_filename(book_name)}{suffix}"

    output_root = Path(args.output_dir).expanduser().resolve()
    if from_url_file or multiple_books:
        return output_root / sanitize_filename(book_name)
    return output_root


def selected_chapters(
    chapters: List[ChapterLink],
    selector: str,
    limit: Optional[int],
) -> List[ChapterLink]:
    max_order = max(chapter.order for chapter in chapters)
    start, end = parse_chapter_selector(selector, max_order)
    selected = [chapter for chapter in chapters if start <= chapter.order <= end]
    if limit is not None:
        selected = selected[:limit]
    return selected


def polite_sleep(delay: float, jitter: float) -> None:
    if delay <= 0 and jitter <= 0:
        return
    wait = max(0.0, delay + random.uniform(-jitter, jitter))
    time.sleep(wait)


def process_book(
    session: requests.Session,
    source: BookSource,
    args: argparse.Namespace,
    *,
    from_url_file: bool,
    multiple_books: bool,
    browser: Optional[ChromeProfileSession] = None,
) -> int:
    meta = fetch_novel_meta(session, source.novel_url, browser=browser)
    if source.book_name:
        meta.title = source.book_name

    output_dir = output_dir_for_book(args, meta.title, from_url_file, multiple_books)
    output_dir.mkdir(parents=True, exist_ok=True)

    chapters = fetch_chapter_links(session, meta, browser=browser)
    selector = args.chapter_range or source.selector
    todo = selected_chapters(chapters, selector, args.limit)
    if not todo:
        log.warning("No chapters selected for %s (%s)", meta.title, selector)
        return 0

    save_chapter_cache(output_dir, meta, chapters, selector)
    log.info("Saved %s chapter URLs -> %s", len(chapters), output_dir / "chapter_urls.json")
    progress_path = output_dir / "progress.json"
    active_service = args.service
    progress = load_progress(progress_path, meta.raw_id, active_service, args.language) if args.resume else {}
    if (
        args.resume
        and not progress
        and args.service == "ai"
        and args.ai_fallback_service != "none"
    ):
        fallback_progress = load_progress(progress_path, meta.raw_id, args.ai_fallback_service, args.language)
        if fallback_progress:
            progress = fallback_progress
            active_service = args.ai_fallback_service
    last_order = int(progress.get("last_order", 0) or 0)
    append_existing = bool(progress and args.resume)
    docs: Dict[Path, Document] = {}
    scraped = 0
    skipped = 0

    log.info("Book: %s", meta.title)
    log.info("Raw ID: %s | Service: %s | Chapters selected: %s", meta.raw_id, active_service, len(todo))
    log.info("Output: %s", output_dir)
    if last_order:
        log.info("Resuming after chapter %s", last_order)
    if active_service != args.service:
        log.info("Continuing with %s because the previous run fell back from %s.", active_service, args.service)

    for chapter in todo:
        if args.resume and chapter.order <= last_order:
            skipped += 1
            continue

        docx_path = docx_path_for_chapter(output_dir, meta.title, chapter.order, args.batch_size)
        if docx_path not in docs:
            if append_existing and docx_path.exists():
                doc = Document(docx_path)
            else:
                doc = Document()
                configure_document(doc)
            docs[docx_path] = doc
        doc = docs[docx_path]

        try:
            log.info("Fetching chapter %s: %s", chapter.order, chapter.title)

            def fetch_current_chapter() -> Tuple[str, List[str], Dict[str, Any]]:
                nonlocal active_service
                if browser and active_service in {"web", "webplus"} and not args.chinese:
                    return browser.rendered_chapter_content(
                        chapter_service_url(chapter, active_service),
                        chapter.order,
                        chapter.title,
                    )

                try:
                    return fetch_chapter_content(
                        session,
                        meta,
                        chapter,
                        service=active_service,
                        language=args.language,
                        browser=browser,
                    )
                except AiRegistrationRequired:
                    if active_service != "ai" or args.ai_fallback_service == "none":
                        raise

                    active_service = args.ai_fallback_service
                    fallback_label = "Google Web Translation" if active_service == "web" else active_service
                    log.warning(
                        "AI requires registration at chapter %s; retrying with %s (%s).",
                        chapter.order,
                        fallback_label,
                        chapter_service_url(chapter, active_service),
                    )
                    if browser and args.click_web_link_fallback:
                        clicked_url = browser.click_translation_link(chapter.url, active_service)
                        log.info("Clicked fallback translation link: %s", clicked_url)
                        return browser.rendered_chapter_content(clicked_url, chapter.order, chapter.title)
                    return fetch_chapter_content(
                        session,
                        meta,
                        chapter,
                        service=active_service,
                        language=args.language,
                        browser=browser,
                    )

            turnstile_attempt = 0
            browser_fetch_attempt = 0
            chapter_not_found_attempt = 0
            while True:
                try:
                    title, paragraphs, _ = fetch_current_chapter()
                    break
                except TurnstileRequired:
                    if turnstile_attempt >= args.turnstile_retries:
                        raise
                    turnstile_attempt += 1
                    wait = args.turnstile_cooldown * turnstile_attempt
                    if browser:
                        log.warning(
                            "Turnstile at chapter %s; opening browser page before retry %s/%s.",
                            chapter.order,
                            turnstile_attempt,
                            args.turnstile_retries,
                        )
                        browser.open_verification_page(
                            chapter_service_url(chapter, active_service),
                            wait,
                            args.turnstile_check_interval,
                        )
                        deadline = time.time() + wait
                        fetched_after_verification = False
                        while time.time() < deadline:
                            sleep_for = min(args.turnstile_check_interval, max(0.0, deadline - time.time()))
                            if sleep_for > 0:
                                time.sleep(sleep_for)
                            try:
                                title, paragraphs, _ = fetch_current_chapter()
                                fetched_after_verification = True
                                break
                            except TurnstileRequired:
                                remaining = max(0.0, deadline - time.time())
                                log.warning(
                                    "Still challenged at chapter %s; retrying again in %.0fs (%.0fs left in window).",
                                    chapter.order,
                                    min(args.turnstile_check_interval, remaining),
                                    remaining,
                                )
                                continue
                        if fetched_after_verification:
                            break
                    else:
                        log.warning(
                            "Turnstile at chapter %s; waiting %.0fs before retry %s/%s.",
                            chapter.order,
                            wait,
                            turnstile_attempt,
                            args.turnstile_retries,
                        )
                        time.sleep(wait)
                    continue
                except BrowserFetchFailed:
                    if browser_fetch_attempt >= args.browser_fetch_retries:
                        raise
                    browser_fetch_attempt += 1
                    wait = args.browser_fetch_delay * browser_fetch_attempt
                    log.warning(
                        "Browser fetch failed at chapter %s; refreshing browser context and retrying in %.0fs (%s/%s).",
                        chapter.order,
                        wait,
                        browser_fetch_attempt,
                        args.browser_fetch_retries,
                    )
                    if browser:
                        browser.open_verification_page(
                            chapter_service_url(chapter, active_service),
                            wait,
                            args.turnstile_check_interval,
                        )
                    else:
                        time.sleep(wait)
                    continue
                except RuntimeError as exc:
                    if isinstance(exc, (AiRegistrationRequired, RenderedTranslationNotReady, PartialBookStop)):
                        raise
                    if chapter_not_found_attempt >= DEFAULT_CHAPTER_NOT_FOUND_RETRIES:
                        raise
                    chapter_not_found_attempt += 1
                    wait = DEFAULT_CHAPTER_NOT_FOUND_DELAY * chapter_not_found_attempt
                    log.warning(
                        "Chapter %s fetch failed (%s); retrying in %.0fs (%s/%s).",
                        chapter.order,
                        exc,
                        wait,
                        chapter_not_found_attempt,
                        DEFAULT_CHAPTER_NOT_FOUND_RETRIES,
                    )
                    time.sleep(wait)
                    continue

            add_chapter_to_doc(doc, title, paragraphs)
            doc.save(docx_path)
            save_progress(progress_path, meta, chapter, title, active_service, args.language)
            scraped += 1
            log.info("Saved chapter %s (%s paragraphs) -> %s", chapter.order, len(paragraphs), docx_path.name)
        except Exception as exc:
            log.error("Failed chapter %s: %s", chapter.order, exc)
            if isinstance(exc, (TurnstileRequired, RenderedTranslationNotReady, BrowserFetchFailed)):
                raise PartialBookStop(str(exc), scraped=scraped, stop_all=True) from exc
            if args.skip_failed_chapters:
                continue
            if args.stop_on_error:
                raise
            break

        polite_sleep(args.delay, args.jitter)

    log.info("Done %s: scraped=%s skipped=%s", meta.title, scraped, skipped)
    return scraped


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Scrape WTR-Lab chapters into DOCX files.")
    parser.add_argument("--urls-file", type=Path, default=DEFAULT_URLS_FILE, help="WTR-Lab URL list file")
    parser.add_argument("--index-url", default=None, help="Scrape one WTR-Lab novel URL instead of url.txt")
    parser.add_argument("--book-name", default=None, help="Override detected book name with --index-url")
    parser.add_argument("--chapter-range", default=None, help='Override URL-file range, e.g. "1-100" or "All"')
    parser.add_argument(
        "--chinese",
        action="store_true",
        help="Raw Chinese mode: web API body, separate *_Chinese output folder",
    )
    parser.add_argument("--service", default="web", choices=["ai", "web", "webplus"], help="Reader service")
    parser.add_argument(
        "--ai-fallback-service",
        default="web",
        choices=["none", "web", "webplus"],
        help="Service to retry with when anonymous AI translation requires registration",
    )
    parser.add_argument(
        "--no-click-web-link-fallback",
        dest="click_web_link_fallback",
        action="store_false",
        help="Do not open the chapter page and click the Web Translation link before browser fallback",
    )
    parser.add_argument("--language", default="en", help="Reader language")
    parser.add_argument("--cookie", default=os.getenv("WTR_COOKIE"), help="Raw Cookie header, or set WTR_COOKIE")
    parser.add_argument(
        "--cookie-file",
        type=Path,
        default=DEFAULT_COOKIE_FILE if DEFAULT_COOKIE_FILE.exists() else None,
        help="File containing a raw Cookie header. Defaults to ./cookie.txt if it exists.",
    )
    parser.add_argument(
        "--chrome-profile",
        default=DEFAULT_CHROME_PROFILE,
        help="Chrome profile directory name to copy into temp, e.g. 'Profile 6'",
    )
    parser.add_argument(
        "--chrome-user-data-dir",
        type=Path,
        default=DEFAULT_CHROME_USER_DATA_DIR,
        help="Chrome user data root containing Local State and profile folders",
    )
    parser.add_argument("--no-chrome-profile", action="store_true", help="Use direct requests instead of copied Chrome profile")
    parser.add_argument(
        "--use-chrome-profile",
        action="store_true",
        help="Force copied Chrome profile mode even when a cookie is provided",
    )
    parser.add_argument("--headless", action="store_true", help="Run copied-profile Chrome headless")
    parser.add_argument("--keep-temp-profile", action="store_true", help="Do not delete the copied temp Chrome profile")
    parser.add_argument("--browser-warmup-seconds", type=float, default=8.0, help="Seconds to wait after opening WTR-Lab in Chrome")
    parser.add_argument("--output-dir", default=None, help="Output directory. Defaults to ./<book name>")
    parser.add_argument("--batch-size", type=int, default=DEFAULT_BATCH_SIZE, help="Chapters per DOCX file")
    parser.add_argument("--delay", type=float, default=DEFAULT_DELAY, help="Base delay between chapter requests")
    parser.add_argument("--jitter", type=float, default=DEFAULT_JITTER, help="Random +/- jitter added to delay")
    parser.add_argument(
        "--turnstile-retries",
        type=int,
        default=DEFAULT_TURNSTILE_RETRIES,
        help="Cooldown retries for the same chapter when WTR-Lab requests Turnstile",
    )
    parser.add_argument(
        "--turnstile-cooldown",
        type=float,
        default=DEFAULT_TURNSTILE_COOLDOWN,
        help="Max seconds to keep retrying within each Turnstile retry window",
    )
    parser.add_argument(
        "--turnstile-check-interval",
        type=float,
        default=DEFAULT_TURNSTILE_CHECK_INTERVAL,
        help="Seconds between browser-mode Turnstile retry checks",
    )
    parser.add_argument(
        "--browser-fetch-retries",
        type=int,
        default=DEFAULT_BROWSER_FETCH_RETRIES,
        help="Retries for transient Selenium/browser fetch failures",
    )
    parser.add_argument(
        "--browser-fetch-delay",
        type=float,
        default=DEFAULT_BROWSER_FETCH_DELAY,
        help="Base seconds to wait before retrying transient browser fetch failures",
    )
    parser.add_argument("--limit", type=int, default=None, help="Limit selected chapters per book for testing")
    parser.add_argument("--no-resume", dest="resume", action="store_false", help="Ignore progress.json")
    parser.add_argument("--skip-failed-chapters", action="store_true", help="Continue within a book after failures")
    parser.add_argument("--stop-on-error", action="store_true", help="Stop all processing on first failed book/chapter")
    parser.set_defaults(resume=True, click_web_link_fallback=True)
    return parser.parse_args()


def apply_mode_defaults(args: argparse.Namespace) -> None:
    if args.chinese:
        args.service = "web"
        args.ai_fallback_service = "none"
        args.click_web_link_fallback = False
        if not args.use_chrome_profile:
            args.no_chrome_profile = True


def resolve_cookie(args: argparse.Namespace) -> Optional[str]:
    if args.cookie:
        return args.cookie.strip()
    if not args.cookie_file:
        return None

    cookie_path = args.cookie_file.expanduser()
    if not cookie_path.exists():
        raise FileNotFoundError(f"Cookie file not found: {cookie_path}")
    cookie = cookie_path.read_text(encoding="utf-8").strip()
    return cookie or None


def validate_args(args: argparse.Namespace) -> None:
    if args.batch_size < 1:
        raise ValueError("--batch-size must be >= 1")
    if args.limit is not None and args.limit < 1:
        raise ValueError("--limit must be >= 1")
    if args.delay < 0 or args.jitter < 0:
        raise ValueError("--delay and --jitter must be >= 0")
    if args.turnstile_retries < 0 or args.turnstile_cooldown < 0:
        raise ValueError("--turnstile-retries and --turnstile-cooldown must be >= 0")
    if args.turnstile_check_interval <= 0:
        raise ValueError("--turnstile-check-interval must be > 0")
    if args.browser_fetch_retries < 0 or args.browser_fetch_delay < 0:
        raise ValueError("--browser-fetch-retries and --browser-fetch-delay must be >= 0")
    if args.browser_warmup_seconds < 0:
        raise ValueError("--browser-warmup-seconds must be >= 0")
    if args.book_name and not args.index_url:
        log.warning("--book-name is only used with --index-url")
    if args.no_chrome_profile and args.service == "ai" and not args.cookie:
        log.warning(
            "Using --service ai without a cookie usually only works for preview chapters. "
            "Use the default --service web for public full-book crawling."
        )


def load_sources(args: argparse.Namespace) -> Tuple[List[BookSource], bool]:
    if args.index_url:
        return [BookSource(book_name=args.book_name, novel_url=args.index_url, selector=args.chapter_range or "All")], False
    return parse_url_file(args.urls_file.expanduser()), True


def main() -> None:
    args = parse_args()
    apply_mode_defaults(args)
    args.cookie = resolve_cookie(args)
    validate_args(args)
    sources, from_url_file = load_sources(args)
    multiple_books = len(sources) > 1

    session = make_session(args.cookie)
    total = 0
    browser: Optional[ChromeProfileSession] = None

    try:
        use_browser = bool(
            args.chrome_profile
            and not args.no_chrome_profile
            and (args.use_chrome_profile or not args.cookie)
        )
        if args.cookie and not use_browser:
            log.info("Using provided Cookie header with direct API requests.")
        elif args.cookie and use_browser:
            log.info("Using copied Chrome profile. Provided Cookie header will only be used by fallback requests.")

        if use_browser:
            browser = ChromeProfileSession(
                args.chrome_user_data_dir,
                args.chrome_profile,
                headless=args.headless,
                keep_temp_profile=args.keep_temp_profile,
                warmup_seconds=args.browser_warmup_seconds,
            )
            browser.start()

        for source in sources:
            try:
                total += process_book(
                    session,
                    source,
                    args,
                    from_url_file=from_url_file,
                    multiple_books=multiple_books,
                    browser=browser,
                )
            except PartialBookStop as exc:
                total += exc.scraped
                log.error("%s", exc)
                if exc.stop_all:
                    log.error("Stopping all books because the current WTR-Lab session is challenged.")
                    break
            except Exception as exc:
                log.error("Book failed: %s", exc)
                if args.stop_on_error:
                    raise
    finally:
        if browser:
            browser.close()

    log.info("All done. Total chapters scraped: %s", total)


if __name__ == "__main__":
    main()