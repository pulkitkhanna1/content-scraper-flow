#!/usr/bin/env python3
"""
scraper_server.py — HTML GUI for Novel Scraper
Starts a local web server on http://localhost:7799 and opens it in your browser.

Run:
    python scraper_server.py
"""

import sys
import os
import re
import time
import json
import uuid
import queue
import threading
import subprocess
import webbrowser
from pathlib import Path

SCRIPT_DIR = Path(__file__).parent.resolve()
sys.path.insert(0, str(SCRIPT_DIR))

try:
    from flask import Flask, request, jsonify, Response, stream_with_context
    FLASK_OK = True
except ImportError:
    FLASK_OK = False
    print("Flask not installed. Run: pip install flask")
    sys.exit(1)

try:
    import requests as _requests
    from bs4 import BeautifulSoup
    HTTP_OK = True
except ImportError:
    HTTP_OK = False

try:
    from docx import Document
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

try:
    from layer1_router import route as layer1_route
    ROUTER_OK = True
except Exception as _e:
    ROUTER_OK = False
    _ROUTER_ERR = str(_e)

PORT = int(os.environ.get("PORT", 7799))
IS_SERVER = os.environ.get("RAILWAY_ENVIRONMENT") or os.environ.get("RENDER") or os.environ.get("SERVER_MODE")
DEFAULT_OUT_DIR = os.environ.get("OUT_DIR", "/tmp/novels" if IS_SERVER else str(Path.home() / "Downloads" / "Novels"))
app = Flask(__name__)

HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/124.0.0.0 Safari/537.36"
    ),
    "Referer": "https://www.google.com/",
    "Accept-Language": "en-US,en;q=0.9",
}
CHAPTERS_PER_FILE = 100

GDRIVE_FOLDER_ID   = "1UyCUOcPTQLGSkII4DoEPd-_gKGZwLO9E"
# Apps Script web app URL — set via env var or paste directly
APPS_SCRIPT_URL    = os.environ.get("APPS_SCRIPT_URL", "https://script.google.com/macros/s/AKfycbyHP8IS2ZTAabP0EVycUAYoXmYws0RFwRnOok1EEc5ZxcW1rp_tUZcbFdVtXaPIE9f2/exec")


def upload_to_drive(file_path: str, folder_id: str, log_fn, subfolder_name: str = ""):
    """Upload a .docx via Apps Script web app — no API keys needed."""
    if not APPS_SCRIPT_URL:
        log_fn("Drive upload skipped — set APPS_SCRIPT_URL env var first.")
        return
    import base64
    try:
        name = Path(file_path).name
        with open(file_path, "rb") as f:
            encoded = base64.b64encode(f.read()).decode("utf-8")
        resp = _requests.post(
            APPS_SCRIPT_URL,
            json={"filename": name, "content": encoded, "folder_id": folder_id, "subfolder_name": subfolder_name},
            timeout=60,
        )
        result = resp.json()
        if result.get("status") == "ok":
            log_fn(f"  Drive: uploaded {name}")
        else:
            log_fn(f"  Drive upload failed: {result.get('message', resp.text)}")
    except Exception as exc:
        log_fn(f"  Drive upload failed: {exc}")


# ── Job store ──────────────────────────────────────────────────────────────────

_jobs: dict = {}


def _new_job() -> str:
    jid = uuid.uuid4().hex[:8]
    _jobs[jid] = {
        "log_queue": queue.Queue(),
        "stop_event": threading.Event(),
        "status": "running",
    }
    return jid


# ── Inline scrapers ────────────────────────────────────────────────────────────

def _safe_name(name: str) -> str:
    return re.sub(r'[<>:"/\\|?*\x00-\x1f]', "", name).strip() or "Novel"


def _file_path(book_dir, safe_name, file_idx):
    s = (file_idx - 1) * CHAPTERS_PER_FILE + 1
    e = s + CHAPTERS_PER_FILE - 1
    return os.path.join(book_dir, f"{safe_name} {s}-{e}.docx")


def _write_metadata_docx(book_dir, result, book_name, platform, start_ch, end_ch):
    """Write a metadata.docx to the book folder with show info."""
    from datetime import datetime
    path = os.path.join(book_dir, "metadata.docx")
    doc = Document()
    doc.add_heading(book_name, level=1)
    fields = [
        ("Platform",        platform),
        ("Source URL",      result.get("canonical_url") or result.get("url") or ""),
        ("Author",          result.get("author") or ""),
        ("Total Chapters",  str(result.get("chapter_count") or "")),
        ("Scraped",         f"Ch {start_ch} – Ch {end_ch}"),
        ("Date",            datetime.now().strftime("%Y-%m-%d %H:%M")),
    ]
    for label, value in fields:
        if value:
            p = doc.add_paragraph()
            p.add_run(f"{label}: ").bold = True
            p.add_run(value)
    doc.save(path)
    return path


def scrape_allnovel(start_url, book_name, start_ch, end_ch, out_dir, log, stop, on_file_complete=None):
    """Scrape allnovel.org — works with plain requests (no Cloudflare)."""
    safe = _safe_name(book_name)
    book_dir = os.path.join(out_dir, "scraped_novels", safe)
    os.makedirs(book_dir, exist_ok=True)
    log(f"Output: {book_dir}")

    ch_num = start_ch
    file_idx = (ch_num - 1) // CHAPTERS_PER_FILE + 1
    ch_in_file = (ch_num - 1) % CHAPTERS_PER_FILE
    out_path = _file_path(book_dir, safe, file_idx)
    doc = Document(out_path) if os.path.exists(out_path) else Document()
    url = start_url

    while url and ch_num <= end_ch:
        if stop.is_set():
            break
        for attempt in range(1, 4):
            try:
                resp = _requests.get(url, headers=HEADERS, timeout=15)
                if resp.status_code == 404:
                    log(f"  Chapter {ch_num}: 404 — no more chapters")
                    url = None
                    break
                resp.raise_for_status()
                soup = BeautifulSoup(resp.text, "html.parser")

                # Title — h2 with chapter name
                title_el = soup.find("h2") or soup.find("h1")
                title = title_el.get_text(strip=True) if title_el else f"Chapter {ch_num}"

                # Content in div.chapter.container or div.chapter-content
                content = soup.select_one("div.chapter.container, div.chapter-content, div#chapter-content")
                paragraphs = []
                if content:
                    for p in content.find_all("p"):
                        text = p.get_text(strip=True)
                        if text and len(text) > 10:
                            paragraphs.append(text)

                if not paragraphs:
                    log(f"  Chapter {ch_num}: no content, skipping")
                    ch_num += 1
                    url = None
                    break

                doc.add_heading(title, level=1)
                count = 0
                for para in paragraphs:
                    doc.add_paragraph(para)
                    count += 1

                doc.save(out_path)
                log(f"  Saved: {title} ({count} paragraphs)")

                # Next chapter link — find "next chapter" anchor
                next_href = None
                for a in soup.find_all("a", href=True):
                    if "next" in a.get_text(strip=True).lower() and "chapter" in a.get("href", ""):
                        next_href = a.get("href")
                        break

                if next_href:
                    url = next_href if next_href.startswith("http") else f"https://allnovel.org{next_href}"
                else:
                    url = None

                ch_num += 1
                ch_in_file += 1

                if ch_num <= end_ch and ch_in_file >= CHAPTERS_PER_FILE:
                    completed = out_path
                    log("100 chapters done — starting new file...")
                    file_idx += 1
                    ch_in_file = 0
                    out_path = _file_path(book_dir, safe, file_idx)
                    doc = Document()
                    if on_file_complete:
                        on_file_complete(completed)

                time.sleep(1)
                break

            except Exception as exc:
                if attempt < 3:
                    log(f"  Error (attempt {attempt}/3): {exc}, retrying...")
                    time.sleep(5 * attempt)
                else:
                    log(f"  Skipping chapter {ch_num}: {exc}")
                    ch_num += 1
                    break

    doc.save(out_path)
    if on_file_complete:
        on_file_complete(out_path)
    log(f"Done. Files saved in: {book_dir}")


def scrape_novelbin(start_url, book_name, start_ch, end_ch, out_dir, log, stop, on_file_complete=None):
    safe = _safe_name(book_name)
    book_dir = os.path.join(out_dir, "scraped_novels", safe)
    os.makedirs(book_dir, exist_ok=True)
    log(f"Output: {book_dir}")

    ch_num = start_ch
    file_idx = (ch_num - 1) // CHAPTERS_PER_FILE + 1
    ch_in_file = (ch_num - 1) % CHAPTERS_PER_FILE
    out_path = _file_path(book_dir, safe, file_idx)
    doc = Document(out_path) if os.path.exists(out_path) else Document()
    url = start_url

    while url and ch_num <= end_ch:
        if stop.is_set():
            doc.save(out_path)
            log("Stopped.")
            return

        log(f"Chapter {ch_num}: {url}")
        for attempt in range(1, 4):
            try:
                resp = _requests.get(url, headers=HEADERS, timeout=15)
                if resp.status_code != 200:
                    if attempt < 3:
                        log(f"  HTTP {resp.status_code}, retry {attempt}/3...")
                        time.sleep(3 * attempt)
                        continue
                    log(f"  HTTP {resp.status_code}, skipping.")
                    ch_num += 1
                    break

                soup = BeautifulSoup(resp.content, "html.parser")
                title_tag = soup.find("span", class_="chr-text")
                title = title_tag.get_text(strip=True) if title_tag else f"Chapter {ch_num}"
                doc.add_heading(title, level=1)

                content_div = soup.find("div", id="chr-content")
                if not content_div:
                    if attempt < 3:
                        log(f"  No content, retry {attempt}/3...")
                        time.sleep(3 * attempt)
                        continue
                    log("  No content found, skipping.")
                else:
                    count = 0
                    for p in content_div.find_all("p"):
                        for br in p.find_all("br"):
                            br.replace_with("\n")
                        text = p.get_text(separator="", strip=True)
                        if text:
                            doc.add_paragraph(text)
                            count += 1
                    log(f"  Saved: {title} ({count} paragraphs)")

                doc.save(out_path)
                ch_num += 1
                ch_in_file += 1

                next_link = soup.find("a", class_="js-chapter-nav",
                                      attrs={"data-chapter-nav": "next"})
                if (not next_link
                        or next_link.get("disabled") is not None
                        or "disabled" in (next_link.get("class") or [])):
                    log("No more chapters found.")
                    url = None
                    break

                url = next_link.get("data-chapter-url") or next_link.get("href")
                time.sleep(1)

                if ch_in_file >= CHAPTERS_PER_FILE:
                    completed = out_path
                    log("100 chapters done — starting new file...")
                    file_idx += 1
                    ch_in_file = 0
                    out_path = _file_path(book_dir, safe, file_idx)
                    doc = Document(out_path) if os.path.exists(out_path) else Document()
                    if on_file_complete:
                        on_file_complete(completed)
                break

            except Exception as exc:
                if attempt < 3:
                    log(f"  Error (attempt {attempt}/3): {exc}, retrying...")
                    time.sleep(5 * attempt)
                else:
                    log(f"  Skipping chapter {ch_num}: {exc}")
                    ch_num += 1
                    break

    log(f"Done. Files saved in: {book_dir}")
    if on_file_complete:
        on_file_complete(out_path)


def scrape_royalroad(start_url, book_name, start_ch, end_ch, out_dir, log, stop, on_file_complete=None):
    safe = _safe_name(book_name)
    book_dir = os.path.join(out_dir, "scraped_novels", safe)
    os.makedirs(book_dir, exist_ok=True)
    log(f"Output: {book_dir}")

    ch_num = start_ch
    file_idx = (ch_num - 1) // CHAPTERS_PER_FILE + 1
    ch_in_file = (ch_num - 1) % CHAPTERS_PER_FILE
    out_path = _file_path(book_dir, safe, file_idx)
    doc = Document(out_path) if os.path.exists(out_path) else Document()
    url = start_url

    while url and ch_num <= end_ch:
        if stop.is_set():
            doc.save(out_path)
            log("Stopped.")
            return

        log(f"Chapter {ch_num}: {url}")
        for attempt in range(1, 4):
            try:
                resp = _requests.get(url, headers=HEADERS, timeout=15)
                if resp.status_code != 200:
                    if attempt < 3:
                        log(f"  HTTP {resp.status_code}, retry {attempt}/3...")
                        time.sleep(3 * attempt)
                        continue
                    log(f"  HTTP {resp.status_code}, skipping.")
                    ch_num += 1
                    break

                soup = BeautifulSoup(resp.content, "html.parser")
                # RoyalRoad chapter title is in <h1> or the chapter-title element
                title_el = (soup.select_one(".chapter-title")
                            or soup.select_one("h1")
                            or soup.select_one("h2"))
                title = title_el.get_text(strip=True) if title_el else f"Chapter {ch_num}"
                doc.add_heading(title, level=1)

                # Content: div.chapter-inner.chapter-content
                content_div = soup.find("div", class_=lambda c: c and "chapter-inner" in c)
                if not content_div:
                    content_div = soup.find("div", class_="chapter-content")
                if not content_div:
                    if attempt < 3:
                        log(f"  No content, retry {attempt}/3...")
                        time.sleep(3 * attempt)
                        continue
                    log("  No content found, skipping.")
                else:
                    count = 0
                    for p in content_div.find_all("p"):
                        text = p.get_text(strip=True)
                        if text:
                            doc.add_paragraph(text)
                            count += 1
                    log(f"  Saved: {title} ({count} paragraphs)")

                doc.save(out_path)
                ch_num += 1
                ch_in_file += 1

                # RoyalRoad's Next button has child <i> elements so string= won't work.
                # Match by class + get_text() instead.
                next_btn = None
                for a in soup.find_all("a"):
                    cls = a.get("class") or []
                    href_candidate = a.get("href", "")
                    text = a.get_text(strip=True)
                    if ("btn-primary" in cls
                            and re.search(r"next", text, re.I)
                            and "/chapter/" in href_candidate):
                        next_btn = a
                        break
                # Fallback: any link with "Next Chapter" text
                if not next_btn:
                    for a in soup.find_all("a"):
                        href_candidate = a.get("href", "")
                        text = a.get_text(strip=True)
                        if re.search(r"next\s*chapter", text, re.I) and "/chapter/" in href_candidate:
                            next_btn = a
                            break

                if not next_btn:
                    log("No more chapters found.")
                    url = None
                    break

                href = next_btn.get("href", "")
                url = href if href.startswith("http") else f"https://www.royalroad.com{href}"
                time.sleep(1)

                if ch_in_file >= CHAPTERS_PER_FILE:
                    completed = out_path
                    log("100 chapters done — starting new file...")
                    file_idx += 1
                    ch_in_file = 0
                    out_path = _file_path(book_dir, safe, file_idx)
                    doc = Document(out_path) if os.path.exists(out_path) else Document()
                    if on_file_complete:
                        on_file_complete(completed)
                break

            except Exception as exc:
                if attempt < 3:
                    log(f"  Error (attempt {attempt}/3): {exc}, retrying...")
                    time.sleep(5 * attempt)
                else:
                    log(f"  Skipping chapter {ch_num}: {exc}")
                    ch_num += 1
                    break

    log(f"Done. Files saved in: {book_dir}")
    if on_file_complete:
        on_file_complete(out_path)


def scrape_wuxiaworld(start_url, book_name, start_ch, end_ch, out_dir, log, stop, on_file_complete=None, cookies_str=""):
    safe = _safe_name(book_name)
    book_dir = os.path.join(out_dir, "scraped_novels", safe)
    os.makedirs(book_dir, exist_ok=True)
    log(f"Output: {book_dir}")

    cookies = _parse_cookie_str(cookies_str) if cookies_str else {}

    ch_num = start_ch
    file_idx = (ch_num - 1) // CHAPTERS_PER_FILE + 1
    ch_in_file = (ch_num - 1) % CHAPTERS_PER_FILE
    out_path = _file_path(book_dir, safe, file_idx)
    doc = Document(out_path) if os.path.exists(out_path) else Document()
    url = start_url

    while ch_num <= end_ch and url:
        if stop.is_set():
            break
        for attempt in range(1, 4):
            try:
                resp = _requests.get(url, headers=HEADERS, cookies=cookies, timeout=15)
                if resp.status_code == 404:
                    log(f"  Chapter {ch_num}: 404 — no more chapters")
                    url = None
                    break
                resp.raise_for_status()
                soup = BeautifulSoup(resp.text, "html.parser")

                # Title: h4 with "Chapter" text, fallback to <title>
                title = ""
                for h in soup.find_all(["h4", "h3", "h2"]):
                    t = h.get_text(strip=True)
                    if "Chapter" in t:
                        title = t
                        break
                if not title:
                    t_tag = soup.find("title")
                    if t_tag:
                        parts = t_tag.get_text().split(" - ", 1)
                        title = parts[1].strip() if len(parts) > 1 else f"Chapter {ch_num}"

                # Content
                cc = soup.select_one("div.chapter-content")
                paragraphs = []
                if cc:
                    for p in cc.find_all("p"):
                        text = p.get_text(strip=True)
                        if text:
                            paragraphs.append(text)

                if not paragraphs:
                    log(f"  Chapter {ch_num}: no content, skipping")
                    ch_num += 1
                    url = None
                    break

                doc.add_heading(title or f"Chapter {ch_num}", level=1)
                count = 0
                for para in paragraphs:
                    doc.add_paragraph(para)
                    count += 1

                doc.save(out_path)
                log(f"  Saved: {title} ({count} paragraphs)")

                # Next: find any link to next chapter number on the page
                next_ch = ch_num + 1
                slug_m = re.search(r"/novel/([^/]+)/", url)
                slug = slug_m.group(1) if slug_m else None
                if slug:
                    url = f"https://www.wuxiaworld.com/novel/{slug}/{slug}-chapter-{next_ch}"
                else:
                    url = None

                ch_num += 1
                ch_in_file += 1

                if ch_num <= end_ch and ch_in_file >= CHAPTERS_PER_FILE:
                    completed = out_path
                    log("100 chapters done — starting new file...")
                    file_idx += 1
                    ch_in_file = 0
                    out_path = _file_path(book_dir, safe, file_idx)
                    doc = Document()
                    if on_file_complete:
                        on_file_complete(completed)

                time.sleep(1)
                break

            except Exception as exc:
                if attempt < 3:
                    log(f"  Error (attempt {attempt}/3): {exc}, retrying...")
                    time.sleep(5 * attempt)
                else:
                    log(f"  Skipping chapter {ch_num}: {exc}")
                    ch_num += 1
                    break

    doc.save(out_path)
    if on_file_complete:
        on_file_complete(out_path)
    log(f"Done. Files saved in: {book_dir}")


def _parse_cookie_str(cookie_str: str) -> dict:
    """Parse a document.cookie-style string into a dict."""
    cookies = {}
    for part in cookie_str.split(";"):
        part = part.strip()
        if "=" in part:
            k, v = part.split("=", 1)
            cookies[k.strip()] = v.strip()
    return cookies


def scrape_webnovel(book_id, book_name, start_ch, end_ch, out_dir, cookies_str, log, stop, on_file_complete=None):
    """Scrape webnovel.com using session cookies + internal JSON API — no browser needed."""
    safe = _safe_name(book_name)
    book_dir = os.path.join(out_dir, "scraped_novels", safe)
    os.makedirs(book_dir, exist_ok=True)

    cookies = _parse_cookie_str(cookies_str)
    headers = {**HEADERS, "Referer": "https://www.webnovel.com/", "Accept": "application/json, */*"}

    # 1. Get chapter list
    log("Fetching chapter list from webnovel API...")
    try:
        resp = _requests.get(
            f"https://www.webnovel.com/go/pcm/chapter/getChapterList?bookId={book_id}",
            headers=headers, cookies=cookies, timeout=20
        )
        data = resp.json()
        all_chapters = data.get("data", {}).get("chapterItems", [])
        if not all_chapters:
            log(f"  API returned no chapters (status {resp.status_code}). Check your cookies.")
            return
    except Exception as e:
        log(f"  Failed to fetch chapter list: {e}")
        return

    log(f"  Total chapters available: {len(all_chapters)}")
    chapters_to_scrape = all_chapters[start_ch - 1 : end_ch]
    log(f"  Scraping chapters {start_ch}–{start_ch + len(chapters_to_scrape) - 1}")

    # 2. Write chapters
    file_idx = (start_ch - 1) // CHAPTERS_PER_FILE + 1
    ch_in_file = (start_ch - 1) % CHAPTERS_PER_FILE
    out_path = _file_path(book_dir, safe, file_idx)
    doc = Document(out_path) if os.path.exists(out_path) else Document()

    for i, ch in enumerate(chapters_to_scrape):
        if stop.is_set():
            log("Stopped.")
            break

        ch_no = start_ch + i
        ch_id = ch.get("chapterId") or ch.get("id", "")
        ch_name = ch.get("chapterName") or ch.get("name") or f"Chapter {ch_no}"
        log(f"  [{ch_no}] {ch_name}")

        try:
            resp = _requests.get(
                f"https://www.webnovel.com/go/pcm/chapter/getContent?type=9"
                f"&bookId={book_id}&chapterId={ch_id}",
                headers=headers, cookies=cookies, timeout=20
            )
            cdata = resp.json()
            ch_info = cdata.get("data", {}).get("chapterInfo", {})
            html_content = ch_info.get("content", "")
            if not html_content:
                log(f"    (empty — VIP/locked chapter)")
                paragraphs = ["[VIP chapter — unlock required]"]
            else:
                soup = BeautifulSoup(html_content, "html.parser")
                paragraphs = [p.get_text() for p in soup.find_all("p") if p.get_text(strip=True)]
                if not paragraphs:
                    paragraphs = [t.strip() for t in soup.get_text("\n").split("\n") if t.strip()]
        except Exception as e:
            log(f"    Error: {e}")
            paragraphs = [f"[Error fetching chapter: {e}]"]

        doc.add_heading(ch_name, level=1)
        for para in paragraphs:
            doc.add_paragraph(para)

        ch_in_file += 1
        if ch_in_file >= CHAPTERS_PER_FILE:
            doc.save(out_path)
            if on_file_complete:
                on_file_complete(out_path)
            file_idx += 1
            ch_in_file = 0
            out_path = _file_path(book_dir, safe, file_idx)
            doc = Document()

        time.sleep(0.8)

    doc.save(out_path)
    if on_file_complete:
        on_file_complete(out_path)
    log(f"Done. Files saved in: {book_dir}")


def scrape_freewebnovel(book_url, book_name, start_ch, end_ch, out_dir, log, stop, on_file_complete=None):
    safe = _safe_name(book_name)
    book_dir = os.path.join(out_dir, "scraped_novels", safe)
    os.makedirs(book_dir, exist_ok=True)
    log(f"Output: {book_dir}")

    m = re.search(r"/novel/([^/?#]+)", book_url)
    slug = m.group(1) if m else None
    if not slug:
        log("ERROR: Could not extract slug from book URL")
        return

    ch_num = start_ch
    file_idx = (ch_num - 1) // CHAPTERS_PER_FILE + 1
    ch_in_file = (ch_num - 1) % CHAPTERS_PER_FILE
    out_path = _file_path(book_dir, safe, file_idx)
    doc = Document(out_path) if os.path.exists(out_path) else Document()

    while ch_num <= end_ch:
        if stop.is_set():
            break

        ch_url = f"https://freewebnovel.com/novel/{slug}/chapter-{ch_num}"
        try:
            resp = _requests.get(ch_url, headers=HEADERS, timeout=15)
            if resp.status_code == 404:
                log(f"  Chapter {ch_num}: 404 — no more chapters")
                break
            if resp.status_code != 200:
                log(f"  Chapter {ch_num}: HTTP {resp.status_code}, skipping")
                ch_num += 1
                continue

            html = resp.text

            # Title from <title> tag: "BookName - Chapter N - ChapterTitle"
            title_m = re.search(r"<title>([^<|]+)", html)
            raw_title = title_m.group(1).strip() if title_m else ""
            parts = raw_title.split(" - ", 1)
            title = parts[1].strip() if len(parts) > 1 else f"Chapter {ch_num}"

            # Content between chapter-start and chapter-end markers
            body_m = re.search(r'class="chapter-start">(.*?)class="chapter-end"', html, re.S)
            paragraphs = []
            if body_m:
                inner = re.sub(r"<script[^>]*>.*?</script>", "", body_m.group(1), flags=re.S)
                for chunk in re.split(r"</?p[^>]*>", inner):
                    text = re.sub(r"<[^>]+>", "", chunk).strip()
                    if len(text) > 20:
                        paragraphs.append(text)

            if not paragraphs:
                log(f"  Chapter {ch_num}: no content, skipping")
                ch_num += 1
                continue

            doc.add_heading(title, level=1)
            count = 0
            for para in paragraphs:
                doc.add_paragraph(para)
                count += 1

            doc.save(out_path)
            log(f"  Saved: {title} ({count} paragraphs)")

            ch_num += 1
            ch_in_file += 1

            if ch_num <= end_ch and ch_in_file >= CHAPTERS_PER_FILE:
                completed = out_path
                log("100 chapters done — starting new file...")
                file_idx += 1
                ch_in_file = 0
                out_path = _file_path(book_dir, safe, file_idx)
                doc = Document()
                if on_file_complete:
                    on_file_complete(completed)

            time.sleep(1)

        except Exception as exc:
            log(f"  Chapter {ch_num}: error — {exc}")
            ch_num += 1

    doc.save(out_path)
    if on_file_complete:
        on_file_complete(out_path)
    log(f"Done. Files saved in: {book_dir}")


def run_subprocess(cmd, log, stop, cwd=None):
    log(f"Running: {' '.join(str(c) for c in cmd)}")
    try:
        proc = subprocess.Popen(
            [str(c) for c in cmd],
            stdout=subprocess.PIPE, stderr=subprocess.STDOUT,
            text=True, cwd=str(cwd or SCRIPT_DIR), bufsize=1,
        )
        for line in proc.stdout:
            if stop.is_set():
                proc.terminate()
                log("Stopped.")
                return -1
            log(line.rstrip())
        proc.wait()
        return proc.returncode
    except Exception as exc:
        log(f"Failed to launch script: {exc}")
        return -1


PLATFORM_MODE = {
    "novelbin":     "inline",
    "allnovel":     "inline",
    "royalroad":    "inline",
    "freewebnovel": "hardcoded",   # Cloudflare-blocked; keep as manual
    "webnovel":     "inline",
    "wuxiaworld":   "inline",
    "69shuba":      "subprocess",
    "babelnovel":   "subprocess",
    "tapas":        "subprocess",
    "wattpad":      "subprocess",
    "kakao":        "hardcoded",   # OCR + Korean auth; manual only
    "qidian":       "hardcoded",   # 2captcha + ZH auth; manual only
    "qdmm":         "hardcoded",   # 2captcha + ZH auth; manual only
    "hengyan":      "subprocess",
}


def _upload_dir_docx(book_dir: str, log_fn, subfolder_name: str = ""):
    """After a subprocess scraper finishes, upload all .docx files in book_dir to Drive."""
    files = sorted(Path(book_dir).glob("*.docx"))
    if not files:
        log_fn("No .docx files found to upload.")
        return
    for f in files:
        log_fn(f"Uploading to Google Drive: {f.name} ...")
        upload_to_drive(str(f), GDRIVE_FOLDER_ID, log_fn, subfolder_name=subfolder_name)


def _run_scrape_job(jid, platform, result, start_ch, end_ch, out_dir, use_drive=False, login_email="", login_password="", cookies_str=""):
    job = _jobs[jid]
    log = lambda msg: job["log_queue"].put(str(msg))
    stop = job["stop_event"]
    book_name = result.get("book_name") or "Novel"
    scraper_args = result.get("scraper_args", {})

    safe = _safe_name(book_name)

    def on_file_complete(path):
        if use_drive:
            log(f"Uploading to Google Drive: {Path(path).name} ...")
            upload_to_drive(path, GDRIVE_FOLDER_ID, log, subfolder_name=safe)

    try:
        book_dir = os.path.join(out_dir, "scraped_novels", safe)
        os.makedirs(book_dir, exist_ok=True)
        _write_metadata_docx(book_dir, result, book_name, platform, start_ch, end_ch)

        log(f"{'='*50}")
        log(f"Scraping: {book_name}")
        log(f"Platform: {platform} | Chapters: {start_ch}–{end_ch}")
        log(f"Output:   {book_dir}")
        log(f"{'='*50}")

        mode = PLATFORM_MODE.get(platform, "hardcoded")

        if mode == "inline" and platform == "novelbin":
            url = scraper_args.get("start_url") or result.get("canonical_url")
            scrape_novelbin(url, book_name, start_ch, end_ch, out_dir, log, stop, on_file_complete)

        elif mode == "inline" and platform == "allnovel":
            url = scraper_args.get("start_url") or result.get("canonical_url")
            scrape_allnovel(url, book_name, start_ch, end_ch, out_dir, log, stop, on_file_complete)

        elif mode == "inline" and platform == "webnovel":
            book_id = scraper_args.get("book_id") or result.get("book_id", "")
            if not book_id:
                # Extract from canonical URL
                import re as _re
                m = _re.search(r"_(\d+)$", result.get("canonical_url", "").rstrip("/").split("/")[-1])
                book_id = m.group(1) if m else ""
            if not cookies_str:
                log("Error: webnovel requires cookies. Paste them in the login modal.")
            else:
                scrape_webnovel(book_id, book_name, start_ch, end_ch, out_dir, cookies_str, log, stop, on_file_complete)

        elif mode == "inline" and platform == "royalroad":
            url = scraper_args.get("url") or result.get("canonical_url")
            scrape_royalroad(url, book_name, start_ch, end_ch, out_dir, log, stop, on_file_complete)

        elif mode == "inline" and platform == "freewebnovel":
            book_url = scraper_args.get("book_url") or result.get("canonical_url")
            scrape_freewebnovel(book_url, book_name, start_ch, end_ch, out_dir, log, stop, on_file_complete)

        elif mode == "inline" and platform == "wuxiaworld":
            slug = scraper_args.get("slug")
            first_ch_url = scraper_args.get("first_chapter_url", "")
            if slug:
                start_url = f"https://www.wuxiaworld.com/novel/{slug}/{slug}-chapter-{start_ch}"
            else:
                start_url = first_ch_url or result.get("canonical_url")
            scrape_wuxiaworld(start_url, book_name, start_ch, end_ch, out_dir, log, stop, on_file_complete, cookies_str=cookies_str)

        elif mode == "subprocess" and platform == "webnovel":
            canonical = result.get("canonical_url", "")
            cmd = [sys.executable, str(SCRIPT_DIR / "webnovel_content_uc.py"),
                   canonical,
                   "--out-dir", book_dir,
                   "--start-chapter", str(start_ch),
                   "--end-chapter", str(end_ch)]
            if login_email:
                cmd += ["--email", login_email]
            if login_password:
                cmd += ["--password", login_password]
            rc = run_subprocess(cmd, log, stop)
            if use_drive and rc == 0:
                _upload_dir_docx(book_dir, log, safe)

        elif mode == "subprocess" and platform == "tapas":
            start_url = scraper_args.get("start_url") or result.get("canonical_url", "")
            cmd = [sys.executable, str(SCRIPT_DIR / "tapas_content_fixed.py"),
                   "--start-url", start_url,
                   "--book-name", book_name,
                   "--start-chapter", str(start_ch),
                   "--max-chapters", str(end_ch - start_ch + 1),
                   "--out", book_dir]
            rc = run_subprocess(cmd, log, stop)
            if use_drive and rc == 0:
                _upload_dir_docx(book_dir, log, safe)

        elif mode == "subprocess" and platform == "babelnovel":
            canonical = result.get("canonical_url", "")
            log("Note: Babelnovel requires manual login in the browser window that opens.")
            cmd = [sys.executable, str(SCRIPT_DIR / "babelnovel_content.py"),
                   "--url", canonical,
                   "--start", str(start_ch),
                   "--end", str(end_ch),
                   "--output", book_dir]
            rc = run_subprocess(cmd, log, stop)
            if use_drive and rc == 0:
                _upload_dir_docx(book_dir, log, safe)

        elif mode == "subprocess" and platform == "hengyan":
            canonical = result.get("canonical_url", "")
            cmd = [sys.executable, str(SCRIPT_DIR / "hengyan/s.py"),
                   "--book-url", canonical,
                   "--book-name", book_name,
                   "--start-chapter", str(start_ch),
                   "--end-chapter", str(end_ch),
                   "--output-dir", book_dir]
            rc = run_subprocess(cmd, log, stop)
            if use_drive and rc == 0:
                _upload_dir_docx(book_dir, log, safe)

        elif mode == "subprocess" and platform == "69shuba":
            start_url = scraper_args.get("start_url") or result.get("canonical_url", "")
            cmd = [sys.executable, str(SCRIPT_DIR / "69shuba_new.py"),
                   "--start-url", start_url,
                   "--book-name", book_name,
                   "--start-chapter", str(start_ch),
                   "--end-chapter", str(end_ch),
                   "--output-dir", book_dir]
            rc = run_subprocess(cmd, log, stop)
            if use_drive and rc == 0:
                _upload_dir_docx(book_dir, log, safe)

        elif mode == "subprocess" and platform == "wattpad":
            canonical = result.get("canonical_url", "")
            cmd = [sys.executable, str(SCRIPT_DIR / "wattpad_content.py"),
                   "--url", canonical,
                   "--book-name", book_name,
                   "--start-chapter", str(start_ch),
                   "--end-chapter", str(end_ch),
                   "--output-dir", book_dir]
            if login_email:
                cmd += ["--email", login_email]
            if login_password:
                cmd += ["--password", login_password]
            rc = run_subprocess(cmd, log, stop)
            if use_drive and rc == 0:
                _upload_dir_docx(book_dir, log, safe)

        else:
            # Platforms that require special manual setup (kakao, qdmm, freewebnovel, etc.)
            scraper = result.get("scraper_script", "")
            canonical = result.get("canonical_url", "")

            log(f"\nPlatform '{platform}' requires manual setup.")
            log(f"Script: py Scripts/{scraper}")
            log(f"\nOpen the script and set these values at the top:")
            log("-" * 44)

            if platform in ("qidian", "qdmm"):
                log(f'  book_url      = "{canonical}"')
                log(f'  book_name     = "{book_name}"')
                log(f'  start_chapter = {start_ch}')
                log(f'  end_chapter   = {end_ch}')
            elif platform == "kakao":
                log(f'  # Requires Korean account login.')
                log(f'  book_url = "{canonical}"')
            elif platform == "freewebnovel":
                log(f'  book_url = "{canonical}"')
            else:
                for k, v in scraper_args.items():
                    log(f"  {k} = {repr(v)}")

            log("-" * 44)
            log(f"\nThen run from py Scripts/:")
            log(f"  python {scraper}")

    except Exception as exc:
        log(f"\nError: {exc}")
    finally:
        job["status"] = "done"
        job["log_queue"].put(None)  # sentinel


# ── Flask routes ───────────────────────────────────────────────────────────────

@app.route("/")
def index():
    return HTML


@app.route("/api/detect", methods=["POST"])
def api_detect():
    if not ROUTER_OK:
        return jsonify({"status": "error", "error": f"layer1_router unavailable: {_ROUTER_ERR}"})
    data = request.get_json(force=True)
    url = (data.get("url") or "").strip()
    if not url:
        return jsonify({"status": "error", "error": "No URL provided"})
    try:
        result = layer1_route(url, search_translations=True)
        return jsonify(result)
    except Exception as exc:
        return jsonify({"status": "error", "error": str(exc)})


@app.route("/api/scrape", methods=["POST"])
def api_scrape():
    data = request.get_json(force=True)
    result     = data.get("detection_result", {})
    platform   = result.get("platform", "")
    start_ch   = int(data.get("start_ch", 1))
    end_ch     = int(data.get("end_ch", 100))
    out_dir    = os.path.expanduser(data.get("out_dir") or DEFAULT_OUT_DIR)
    use_drive  = bool(data.get("upload_to_drive", False))
    login_email    = data.get("email", "").strip()
    login_password = data.get("password", "").strip()
    cookies_str    = data.get("cookies", "").strip()

    if not platform:
        return jsonify({"error": "No detection result. Run Detect first."}), 400

    os.makedirs(out_dir, exist_ok=True)
    jid = _new_job()
    threading.Thread(
        target=_run_scrape_job,
        args=(jid, platform, result, start_ch, end_ch, out_dir, use_drive, login_email, login_password, cookies_str),
        daemon=True,
    ).start()
    return jsonify({"job_id": jid})


@app.route("/api/drive-status")
def api_drive_status():
    ready = bool(APPS_SCRIPT_URL)
    return jsonify({
        "ready":     ready,
        "url_set":   ready,
        "folder_id": GDRIVE_FOLDER_ID,
    })


PLATFORM_CREDS = {
    "webnovel": {"email": "fantasy.team@pocketfm.com",      "password": "0ForFantasyOnly0"},
    "wattpad":  {"email": "shohtaj_singh@pocketfm.com",     "password": "PocketContent@123"},
}

@app.route("/api/platform-creds/<platform>")
def api_platform_creds(platform):
    creds = PLATFORM_CREDS.get(platform, {})
    return jsonify(creds)


@app.route("/api/drive-auth", methods=["POST"])
def api_drive_auth():
    """Test the Apps Script URL is reachable."""
    if not APPS_SCRIPT_URL:
        return jsonify({"ok": False, "error": "APPS_SCRIPT_URL not set."})
    try:
        resp = _requests.get(APPS_SCRIPT_URL, timeout=10)
        result = resp.json()
        if result.get("status") == "ok":
            return jsonify({"ok": True, "message": "Apps Script reachable."})
        return jsonify({"ok": False, "error": str(result)})
    except Exception as exc:
        return jsonify({"ok": False, "error": str(exc)})


@app.route("/api/log/<job_id>")
def api_log(job_id):
    def generate():
        job = _jobs.get(job_id)
        if not job:
            yield "data: {\"line\": \"Job not found\"}\n\n"
            yield "data: \"[DONE]\"\n\n"
            return
        q = job["log_queue"]
        while True:
            try:
                msg = q.get(timeout=30)
                if msg is None:
                    yield "data: \"[DONE]\"\n\n"
                    return
                payload = json.dumps({"line": msg})
                yield f"data: {payload}\n\n"
            except queue.Empty:
                yield ": keepalive\n\n"

    return Response(
        stream_with_context(generate()),
        mimetype="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.route("/api/stop/<job_id>", methods=["POST"])
def api_stop(job_id):
    job = _jobs.get(job_id)
    if job:
        job["stop_event"].set()
        return jsonify({"ok": True})
    return jsonify({"ok": False, "error": "Job not found"}), 404


@app.route("/api/open-folder", methods=["POST"])
def api_open_folder():
    path = os.path.expanduser(request.get_json(force=True).get("path", ""))
    if os.path.isdir(path):
        subprocess.Popen(["open", path])
        return jsonify({"ok": True})
    return jsonify({"ok": False, "error": f"Directory not found: {path}"})


@app.route("/api/status")
def api_status():
    return jsonify({
        "router": ROUTER_OK,
        "http_deps": HTTP_OK,
        "docx": DOCX_OK,
    })


# ── HTML ───────────────────────────────────────────────────────────────────────

HTML = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Novel Scraper</title>
<link rel="icon" type="image/svg+xml" href="data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 32 32'%3E%3Crect width='32' height='32' rx='5' fill='%23E8001D'/%3E%3Ctext x='16' y='23' font-family='Arial,sans-serif' font-size='17' font-weight='bold' text-anchor='middle' fill='white'%3ELS%3C/text%3E%3C/svg%3E">
<style>
  :root {
    --bg:      #0f1117;
    --surface: #1a1d27;
    --card:    #21252f;
    --border:  #2e3244;
    --fg:      #cdd6f4;
    --muted:   #6c7086;
    --blue:    #89b4fa;
    --green:   #a6e3a1;
    --red:     #f38ba8;
    --yellow:  #f9e2af;
    --teal:    #94e2d5;
    --r:       8px;
  }

  *, *::before, *::after { box-sizing: border-box; margin: 0; padding: 0; }

  body {
    background: var(--bg);
    color: var(--fg);
    font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif;
    font-size: 14px;
    min-height: 100vh;
    display: flex;
    flex-direction: column;
  }

  /* ── Header ── */
  header {
    display: flex;
    align-items: center;
    gap: 12px;
    padding: 14px 20px;
    border-bottom: 1px solid var(--border);
    background: var(--surface);
  }
  header h1 {
    font-size: 18px;
    font-weight: 700;
    letter-spacing: -0.3px;
    color: var(--blue);
  }
  .header-status {
    margin-left: auto;
    display: flex;
    gap: 8px;
    align-items: center;
  }
  .dot {
    width: 8px; height: 8px;
    border-radius: 50%;
    background: var(--muted);
    display: inline-block;
  }
  .dot.ok  { background: var(--green); }
  .dot.err { background: var(--red); }
  .dot-label { font-size: 11px; color: var(--muted); }

  /* ── URL bar ── */
  .url-bar {
    display: flex;
    gap: 8px;
    padding: 14px 20px;
    background: var(--surface);
    border-bottom: 1px solid var(--border);
  }
  .url-bar input {
    flex: 1;
    background: var(--card);
    border: 1px solid var(--border);
    border-radius: var(--r);
    color: var(--fg);
    font-size: 14px;
    padding: 9px 14px;
    outline: none;
    transition: border-color .15s;
  }
  .url-bar input:focus { border-color: var(--blue); }
  .url-bar input::placeholder { color: var(--muted); }

  /* ── Buttons ── */
  button {
    border: none;
    border-radius: var(--r);
    cursor: pointer;
    font-size: 13px;
    font-weight: 600;
    padding: 9px 16px;
    transition: opacity .15s, transform .1s;
    white-space: nowrap;
  }
  button:active { transform: scale(.97); }
  button:disabled { opacity: .4; cursor: not-allowed; transform: none; }

  .btn-primary  { background: var(--blue);  color: #1e1e2e; }
  .btn-danger   { background: var(--red);   color: #1e1e2e; }
  .btn-ghost    { background: var(--card);  color: var(--fg); border: 1px solid var(--border); }
  .btn-green    { background: var(--green); color: #1e1e2e; }
  .btn-sm { padding: 5px 10px; font-size: 12px; }

  /* ── Main grid ── */
  .main {
    display: grid;
    grid-template-columns: 1fr 280px;
    gap: 14px;
    padding: 14px 20px;
    flex: 1;
  }

  /* ── Cards ── */
  .card {
    background: var(--card);
    border: 1px solid var(--border);
    border-radius: var(--r);
    padding: 16px;
  }
  .card-title {
    font-size: 11px;
    font-weight: 700;
    letter-spacing: .8px;
    text-transform: uppercase;
    color: var(--muted);
    margin-bottom: 12px;
  }

  /* ── Info card ── */
  .info-card { display: flex; flex-direction: column; gap: 10px; }
  .badge {
    display: inline-flex;
    gap: 8px;
    background: var(--surface);
    border: 1px solid var(--border);
    border-radius: 4px;
    padding: 4px 10px;
    font-size: 12px;
    font-weight: 600;
    color: var(--blue);
    align-items: center;
    align-self: flex-start;
  }
  .badge .lang-tag {
    background: var(--blue);
    color: #1e1e2e;
    border-radius: 3px;
    padding: 1px 5px;
    font-size: 10px;
  }
  .badge .warn-tag {
    background: var(--yellow);
    color: #1e1e2e;
    border-radius: 3px;
    padding: 1px 5px;
    font-size: 10px;
  }
  .book-title {
    font-size: 20px;
    font-weight: 700;
    line-height: 1.3;
    color: var(--fg);
  }
  .book-author { color: var(--green); font-size: 14px; }
  .book-meta   { color: var(--muted); font-size: 12px; }
  .book-desc   { color: #bac2de; font-size: 13px; line-height: 1.5; }

  /* Translation candidates */
  .trans-section { margin-top: 8px; }
  .trans-title { font-size: 11px; color: var(--muted); font-weight: 600;
                 text-transform: uppercase; letter-spacing: .6px; margin-bottom: 6px; }
  .trans-item {
    display: flex;
    align-items: center;
    justify-content: space-between;
    gap: 10px;
    padding: 7px 10px;
    background: var(--surface);
    border: 1px solid var(--border);
    border-radius: 6px;
    margin-bottom: 4px;
  }
  .trans-item-info { flex: 1; min-width: 0; }
  .trans-source { font-size: 10px; color: var(--muted); font-weight: 600;
                  text-transform: uppercase; }
  .trans-name { font-size: 13px; color: var(--teal); overflow: hidden;
                text-overflow: ellipsis; white-space: nowrap; }

  /* ── Config card ── */
  .config-card { display: flex; flex-direction: column; gap: 12px; }
  .field-label {
    font-size: 11px; color: var(--muted); font-weight: 600;
    text-transform: uppercase; letter-spacing: .5px;
    margin-bottom: 4px;
  }
  .field-row { display: grid; grid-template-columns: 1fr 1fr; gap: 8px; }
  input[type=number], input[type=text].path-input {
    width: 100%;
    background: var(--surface);
    border: 1px solid var(--border);
    border-radius: 6px;
    color: var(--fg);
    font-size: 13px;
    padding: 7px 10px;
    outline: none;
    transition: border-color .15s;
  }
  input[type=number]:focus, input[type=text].path-input:focus {
    border-color: var(--blue);
  }

  .scraper-info {
    background: var(--surface);
    border: 1px solid var(--border);
    border-radius: 6px;
    padding: 8px 10px;
    font-size: 11px;
    color: var(--muted);
    line-height: 1.6;
  }

  .btn-group { display: flex; flex-direction: column; gap: 6px; }
  .btn-group button { width: 100%; }

  /* ── Log ── */
  .log-section {
    padding: 0 20px 14px;
  }
  .log-header {
    display: flex;
    align-items: center;
    justify-content: space-between;
    margin-bottom: 8px;
  }
  .log-label {
    font-size: 11px; font-weight: 700; letter-spacing: .8px;
    text-transform: uppercase; color: var(--muted);
  }
  .log-box {
    background: #0a0c10;
    border: 1px solid var(--border);
    border-radius: var(--r);
    font-family: "Menlo", "Fira Code", monospace;
    font-size: 12px;
    line-height: 1.6;
    height: 240px;
    overflow-y: auto;
    padding: 10px 14px;
    color: #a6adc8;
  }
  .log-line { animation: fadeIn .15s ease; }
  .log-line.ok   { color: var(--green); }
  .log-line.err  { color: var(--red); }
  .log-line.info { color: var(--blue); }
  .log-line.sep  { color: var(--border); }
  @keyframes fadeIn { from { opacity: 0; } to { opacity: 1; } }

  /* ── Progress bar ── */
  .progress-wrap {
    height: 3px;
    background: var(--border);
    border-radius: 2px;
    margin-bottom: 8px;
    overflow: hidden;
    display: none;
  }
  .progress-wrap.active { display: block; }
  .progress-bar {
    height: 100%;
    background: var(--blue);
    border-radius: 2px;
    transition: width .3s;
    width: 0%;
  }
  .progress-bar.indeterminate {
    width: 30%;
    animation: slide 1.2s infinite ease-in-out;
  }
  @keyframes slide {
    0%   { transform: translateX(-100%); }
    100% { transform: translateX(400%); }
  }

  /* ── Spinner ── */
  .spinner {
    width: 14px; height: 14px;
    border: 2px solid var(--border);
    border-top-color: var(--blue);
    border-radius: 50%;
    animation: spin .7s linear infinite;
    display: none;
  }
  .spinner.active { display: inline-block; }
  @keyframes spin { to { transform: rotate(360deg); } }

  /* ── Placeholder state ── */
  .placeholder { color: var(--muted); font-style: italic; }

  /* ── Scrollbar ── */
  ::-webkit-scrollbar { width: 6px; }
  ::-webkit-scrollbar-track { background: transparent; }
  ::-webkit-scrollbar-thumb { background: var(--border); border-radius: 3px; }
  ::-webkit-scrollbar-thumb:hover { background: var(--muted); }
</style>
</head>
<body>

<!-- Header -->
<header>
  <svg width="22" height="22" viewBox="0 0 24 24" fill="none" stroke="var(--blue)" stroke-width="2">
    <path d="M4 19.5A2.5 2.5 0 0 1 6.5 17H20"/>
    <path d="M6.5 2H20v20H6.5A2.5 2.5 0 0 1 4 19.5v-15A2.5 2.5 0 0 1 6.5 2z"/>
  </svg>
  <h1>Novel Scraper</h1>
  <div class="header-status">
    <span class="dot" id="dot-router"></span>
    <span class="dot-label">router</span>
    <span class="dot" id="dot-http"></span>
    <span class="dot-label">http</span>
    <span class="dot" id="dot-docx"></span>
    <span class="dot-label">docx</span>
    <span class="dot" id="dot-drive"></span>
    <span class="dot-label">drive</span>
    <div class="spinner" id="spinner"></div>
  </div>
</header>

<!-- URL bar -->
<div class="url-bar">
  <input id="url-input" type="text" placeholder="Paste a novel URL (novelbin.com, royalroad.com, webnovel.com, ...)" />
  <button class="btn-primary" id="detect-btn" onclick="detect()">Detect</button>
</div>

<!-- Progress bar -->
<div class="progress-wrap" id="progress-wrap">
  <div class="progress-bar indeterminate" id="progress-bar"></div>
</div>

<!-- Main grid -->
<div class="main">

  <!-- Info card -->
  <div class="card info-card" id="info-card">
    <div class="card-title">Book Info</div>
    <div id="info-placeholder" class="placeholder">
      Paste a URL above and click Detect to identify the platform and extract book metadata.
    </div>
    <div id="info-content" style="display:none; display:flex; flex-direction:column; gap:10px;">
      <div id="badge" class="badge"></div>
      <div id="book-title" class="book-title"></div>
      <div id="book-author" class="book-author"></div>
      <div id="book-meta" class="book-meta"></div>
      <div id="book-desc" class="book-desc"></div>
      <div id="trans-section" class="trans-section" style="display:none">
        <div class="trans-title">English Translation Candidates</div>
        <div id="trans-list"></div>
      </div>
    </div>
  </div>

  <!-- Config card -->
  <div class="card config-card">
    <div class="card-title">Scrape Config</div>

    <div>
      <div class="field-label">Chapter Range</div>
      <div class="field-row">
        <div>
          <div style="font-size:11px;color:var(--muted);margin-bottom:3px">Start</div>
          <input type="number" id="start-ch" value="1" min="1" />
        </div>
        <div>
          <div style="font-size:11px;color:var(--muted);margin-bottom:3px">End</div>
          <input type="number" id="end-ch" value="100" min="1" />
        </div>
      </div>
    </div>

    <div class="scraper-info" id="scraper-info">
      Run Detect first to see scraper details.
    </div>

    <div class="btn-group">
      <button class="btn-green" id="scrape-btn" onclick="startScrape()" disabled>
        &#9654;&nbsp; Scrape
      </button>
      <button class="btn-danger" id="stop-btn" onclick="stopScrape()" disabled>
        &#9632;&nbsp; Stop
      </button>
      <button class="btn-ghost" onclick="openFolder()">
        &#128193;&nbsp; Open Drive Folder
      </button>
    </div>

    <div style="border-top:1px solid var(--border);padding-top:12px;margin-top:2px">
      <div class="field-label" style="margin-bottom:8px">Google Drive</div>
      <div style="font-size:12px;color:var(--muted);line-height:1.5" id="drive-hint">
        Files upload automatically after each 100 chapters.<br>
        Folder: <a href="https://drive.google.com/drive/folders/1UyCUOcPTQLGSkII4DoEPd-_gKGZwLO9E"
                   target="_blank" style="color:var(--blue)">Open in Drive</a>
      </div>
    </div>
    <div style="display:none">
    </div>
  </div>
</div>

<!-- Log -->
<div class="log-section">
  <div class="log-header">
    <span class="log-label">Log</span>
    <button class="btn-ghost btn-sm" onclick="clearLog()">Clear</button>
  </div>
  <div class="log-box" id="log-box"></div>
</div>

<script>
// ── State ─────────────────────────────────────────────────────────────────────
let detectionResult = null;
let currentJobId    = null;
let eventSource     = null;

// ── Init ──────────────────────────────────────────────────────────────────────
window.onload = async () => {
  // Check backend status
  try {
    const s = await fetch('/api/status').then(r => r.json());
    setDot('dot-router', s.router);
    setDot('dot-http',   s.http_deps);
    setDot('dot-docx',   s.docx);
    if (!s.router)    log('layer1_router not available. Check py Scripts/layer1_router.py', 'err');
    if (!s.http_deps) log('Missing: requests / beautifulsoup4. Run: pip install requests beautifulsoup4', 'err');
    if (!s.docx)      log('Missing: python-docx. Run: pip install python-docx', 'err');
    if (s.router && s.http_deps && s.docx) log('Ready. Paste a URL above and click Detect.', 'info');

    // Drive status
    try {
      const d = await fetch('/api/drive-status').then(r => r.json());
      setDot('dot-drive', d.ready);
    } catch(e) { setDot('dot-drive', false); }
  } catch(e) {
    log('Could not reach server: ' + e, 'err');
  }

  // Enter key on URL input
  document.getElementById('url-input').addEventListener('keydown', e => {
    if (e.key === 'Enter') detect();
  });
};

function setDot(id, ok) {
  const el = document.getElementById(id);
  el.classList.remove('ok', 'err');
  el.classList.add(ok ? 'ok' : 'err');
}

// ── Detection ─────────────────────────────────────────────────────────────────
async function detect() {
  const url = document.getElementById('url-input').value.trim();
  if (!url) { log('Enter a URL first.', 'err'); return; }

  setLoading(true);
  log('');
  log('Detecting: ' + url, 'info');

  try {
    const res = await fetch('/api/detect', {
      method: 'POST',
      headers: {'Content-Type': 'application/json'},
      body: JSON.stringify({ url })
    });
    const data = await res.json();
    setLoading(false);
    handleDetection(data);
  } catch(e) {
    setLoading(false);
    log('Detection failed: ' + e, 'err');
  }
}

function handleDetection(data) {
  if (data.status !== 'ok') {
    log('Error: ' + (data.error || 'Unknown error'), 'err');
    return;
  }

  detectionResult = data;

  const platform    = (data.platform || 'unknown').toUpperCase();
  const lang        = (data.source_language || '?').toUpperCase();
  const needsTrans  = data.needs_translation || false;
  const needsBr     = data.needs_browser !== false;
  const needsLogin  = data.needs_login || false;
  const chCount     = data.chapter_count || 0;
  const bookName    = data.book_name || '';
  const author      = data.author || '';
  const desc        = (data.description || '').slice(0, 280);
  const scrScript   = data.scraper_script || '';

  // Show info content
  document.getElementById('info-placeholder').style.display = 'none';
  const ic = document.getElementById('info-content');
  ic.style.display = 'flex';

  // Badge
  const badge = document.getElementById('badge');
  badge.innerHTML = `${platform}
    <span class="lang-tag">${lang}</span>
    ${needsTrans ? '<span class="warn-tag">needs EN translation</span>' : ''}`;

  document.getElementById('book-title').textContent  = bookName || '(title not found)';
  document.getElementById('book-author').textContent = author;
  document.getElementById('book-meta').textContent   = chCount
    ? `${chCount.toLocaleString()} chapters` : 'chapter count unknown';
  document.getElementById('book-desc').textContent   = desc
    ? (desc.length < (data.description||'').length ? desc + '…' : desc) : '';

  // Auto-fill end chapter
  if (chCount > 0) document.getElementById('end-ch').value = chCount;

  // Translation candidates
  const candidates = data.translation_candidates || [];
  if (candidates.length > 0) {
    document.getElementById('trans-section').style.display = 'block';
    const list = document.getElementById('trans-list');
    list.innerHTML = '';
    candidates.slice(0, 6).forEach(c => {
      const item = document.createElement('div');
      item.className = 'trans-item';
      item.innerHTML = `
        <div class="trans-item-info">
          <div class="trans-source">${c.source || ''}</div>
          <div class="trans-name" title="${c.url||''}">${c.title || c.url || ''}</div>
        </div>
        <button class="btn-ghost btn-sm" onclick="useUrl(${JSON.stringify(c.url)})">Use</button>`;
      list.appendChild(item);
    });
  } else {
    document.getElementById('trans-section').style.display = 'none';
  }

  // Scraper info
  const modeMap = {
    novelbin: 'Built-in (requests, fast)',
    royalroad: 'Built-in (requests, fast)',
    freewebnovel: 'Subprocess (CLI args)',
  };
  const modeName = modeMap[data.platform] || 'Browser script (Chrome required)';
  let infoLines = [`Script: ${scrScript || 'none'}`, `Method: ${modeName}`];
  if (needsLogin) infoLines.push('Requires login (see script)');
  document.getElementById('scraper-info').innerHTML = infoLines
    .map(l => `<div>${l}</div>`).join('');

  document.getElementById('scrape-btn').disabled = false;

  log(`Platform: ${platform} | Lang: ${lang} | Chapters: ${chCount}`, 'ok');
  if (needsTrans && candidates.length > 0)
    log(`Found ${candidates.length} EN translation candidate(s) — shown above.`, 'info');
  else if (needsTrans)
    log('No EN translations found. May need direct scrape + translation.', 'info');
}

function useUrl(url) {
  document.getElementById('url-input').value = url;
  log('');
  log('Switching to: ' + url, 'info');
  detect();
}

// ── Scraping ──────────────────────────────────────────────────────────────────
const LOGIN_PLATFORMS  = ['wattpad'];
const COOKIE_PLATFORMS = ['webnovel', 'wuxiaworld'];

function startScrape() {
  if (!detectionResult) { log('Run Detect first.', 'err'); return; }
  const platform = detectionResult.platform || '';
  if (COOKIE_PLATFORMS.includes(platform)) {
    // Show cookie-paste modal
    document.getElementById('login-modal').style.display = 'flex';
    document.getElementById('login-modal-title').textContent =
      'Paste Cookies — ' + platform.charAt(0).toUpperCase() + platform.slice(1);
    document.getElementById('login-fields').style.display = 'none';
    document.getElementById('cookie-field').style.display = 'block';
    document.getElementById('login-cookies').value = '';
    const hint = document.getElementById('cookie-creds-hint');
    if (hint) hint.style.display = 'none';
    fetch('/api/platform-creds/' + platform).then(r => r.json()).then(c => {
      if (hint && c.email) {
        hint.textContent = 'Log in as: ' + c.email + ' / ' + c.password;
        hint.style.display = 'block';
      }
    }).catch(() => {});
  } else if (LOGIN_PLATFORMS.includes(platform)) {
    // Show email/password modal
    document.getElementById('login-modal').style.display = 'flex';
    document.getElementById('login-modal-title').textContent =
      'Login — ' + platform.charAt(0).toUpperCase() + platform.slice(1);
    document.getElementById('login-fields').style.display = 'block';
    document.getElementById('cookie-field').style.display = 'none';
    document.getElementById('login-email').value = '';
    document.getElementById('login-password').value = '';
    fetch('/api/platform-creds/' + platform).then(r => r.json()).then(c => {
      if (c.email)    document.getElementById('login-email').value = c.email;
      if (c.password) document.getElementById('login-password').value = c.password;
    }).catch(() => {});
  } else {
    _doScrape('', '', '');
  }
}

function submitLogin() {
  const email    = document.getElementById('login-email').value.trim();
  const password = document.getElementById('login-password').value.trim();
  const cookies  = document.getElementById('login-cookies').value.trim();
  document.getElementById('login-modal').style.display = 'none';
  _doScrape(email, password, cookies);
}

function cancelLogin() {
  document.getElementById('login-modal').style.display = 'none';
}

async function _doScrape(email, password, cookies) {
  const startCh = parseInt(document.getElementById('start-ch').value) || 1;
  const endCh   = parseInt(document.getElementById('end-ch').value)   || 100;
  const outDir  = '';

  if (startCh > endCh) { log('Start chapter must be <= end chapter.', 'err'); return; }

  document.getElementById('scrape-btn').disabled = true;
  document.getElementById('stop-btn').disabled   = false;
  setProgress(true);

  try {
    const res = await fetch('/api/scrape', {
      method: 'POST',
      headers: {'Content-Type': 'application/json'},
      body: JSON.stringify({
        detection_result: detectionResult,
        start_ch: startCh,
        end_ch:   endCh,
        out_dir:  outDir,
        upload_to_drive: true,
        email:    email,
        password: password,
        cookies:  cookies,
      })
    });
    const data = await res.json();
    if (data.error) { log('Error: ' + data.error, 'err'); scrapeDone(); return; }

    currentJobId = data.job_id;
    log(`Job started [${currentJobId}]`, 'info');
    startLogStream(currentJobId);
  } catch(e) {
    log('Failed to start scrape: ' + e, 'err');
    scrapeDone();
  }
}

function startLogStream(jobId) {
  if (eventSource) eventSource.close();

  eventSource = new EventSource(`/api/log/${jobId}`);
  eventSource.onmessage = e => {
    const payload = JSON.parse(e.data);
    if (payload === '[DONE]') {
      eventSource.close();
      scrapeDone();
      return;
    }
    const line = payload.line || '';
    const cls  = line.startsWith('Done') || line.includes('Saved') ? 'ok'
               : line.includes('Error') || line.includes('Failed') ? 'err'
               : line.startsWith('=') ? 'sep'
               : '';
    log(line, cls);
  };
  eventSource.onerror = () => {
    eventSource.close();
    scrapeDone();
  };
}


async function stopScrape() {
  if (!currentJobId) return;
  document.getElementById('stop-btn').disabled = true;
  await fetch(`/api/stop/${currentJobId}`, { method: 'POST' });
  log('Stop requested...', 'info');
}

function scrapeDone() {
  setProgress(false);
  document.getElementById('scrape-btn').disabled = false;
  document.getElementById('stop-btn').disabled   = true;
  currentJobId = null;
}

function openFolder() {
  window.open('https://drive.google.com/drive/folders/1UyCUOcPTQLGSkII4DoEPd-_gKGZwLO9E', '_blank');
}

// ── Log helpers ───────────────────────────────────────────────────────────────
function log(msg, cls) {
  const box = document.getElementById('log-box');
  if (!msg && msg !== '') return;
  if (msg === '') { box.appendChild(document.createElement('br')); box.scrollTop = 9e9; return; }
  const div = document.createElement('div');
  div.className = 'log-line' + (cls ? ' ' + cls : '');
  div.textContent = msg;
  box.appendChild(div);
  box.scrollTop = 9e9;
}

function clearLog() {
  document.getElementById('log-box').innerHTML = '';
}

// ── UI helpers ────────────────────────────────────────────────────────────────
function setLoading(on) {
  document.getElementById('spinner').classList.toggle('active', on);
  document.getElementById('detect-btn').disabled = on;
  setProgress(on);
}

function setProgress(on) {
  document.getElementById('progress-wrap').classList.toggle('active', on);
}
</script>

<!-- Login Modal -->
<div id="login-modal" style="display:none;position:fixed;inset:0;background:rgba(0,0,0,.65);z-index:9999;align-items:center;justify-content:center;">
  <div style="background:var(--surface);border:1px solid var(--border);border-radius:10px;padding:28px 32px;width:380px;max-width:90vw;">
    <div style="font-size:15px;font-weight:600;margin-bottom:18px;color:var(--text)" id="login-modal-title">Login</div>

    <!-- Email/password (wattpad etc.) -->
    <div id="login-fields">
      <div style="margin-bottom:12px">
        <label style="font-size:11px;color:var(--muted);display:block;margin-bottom:4px">Email</label>
        <input id="login-email" type="email" placeholder="your@email.com"
          style="width:100%;box-sizing:border-box;background:var(--bg);border:1px solid var(--border);color:var(--text);border-radius:6px;padding:8px 10px;font-size:13px;outline:none;">
      </div>
      <div style="margin-bottom:20px">
        <label style="font-size:11px;color:var(--muted);display:block;margin-bottom:4px">Password</label>
        <input id="login-password" type="password" placeholder="••••••••"
          style="width:100%;box-sizing:border-box;background:var(--bg);border:1px solid var(--border);color:var(--text);border-radius:6px;padding:8px 10px;font-size:13px;outline:none;">
      </div>
    </div>

    <!-- Cookie paste (webnovel etc.) -->
    <div id="cookie-field" style="display:none;margin-bottom:20px">
      <label style="font-size:11px;color:var(--muted);display:block;margin-bottom:4px">
        Paste cookies from browser
        <span style="color:var(--blue);cursor:pointer;margin-left:6px" title="1. Go to webnovel.com and log in&#10;2. Open DevTools (F12) → Console&#10;3. Run: copy(document.cookie)&#10;4. Paste here">ⓘ how?</span>
      </label>
      <textarea id="login-cookies" rows="4" placeholder="Paste cookie string here..."
        style="width:100%;box-sizing:border-box;background:var(--bg);border:1px solid var(--border);color:var(--text);border-radius:6px;padding:8px 10px;font-size:12px;outline:none;resize:vertical;font-family:monospace"></textarea>
      <div id="cookie-creds-hint" style="display:none;font-size:11px;color:var(--blue);margin-top:6px;padding:6px 8px;background:rgba(74,144,226,.08);border-radius:5px;font-family:monospace"></div>
      <div style="font-size:11px;color:var(--muted);margin-top:4px">Open the site → DevTools Console → type <code style="background:var(--bg);padding:1px 4px;border-radius:3px">copy(document.cookie)</code> → paste above</div>
    </div>

    <div style="display:flex;gap:10px;justify-content:flex-end">
      <button onclick="cancelLogin()" style="padding:7px 16px;border-radius:6px;border:1px solid var(--border);background:transparent;color:var(--muted);cursor:pointer;font-size:13px">Cancel</button>
      <button onclick="submitLogin()" style="padding:7px 18px;border-radius:6px;border:none;background:var(--green);color:#fff;cursor:pointer;font-size:13px;font-weight:600">Start Scraping</button>
    </div>
  </div>
</div>
</body>
</html>
"""


# ── Entry point ────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    url = f"http://localhost:{PORT}"
    print(f"Starting Novel Scraper at {url}")
    print("Press Ctrl+C to stop.\n")
    if not IS_SERVER:
        threading.Timer(1.2, lambda: webbrowser.open(url)).start()
    app.run(host="0.0.0.0", port=PORT, debug=False, threaded=True)
