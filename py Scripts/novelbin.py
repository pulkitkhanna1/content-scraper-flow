#!/usr/bin/env python3
"""
NovelBin Scraper with Streaming Output and File Splitting
---------------------------------------------------------
Scrapes chapters continuously and writes them into .docx files,
with each file containing at most 100 chapters, stored in a
dedicated folder named after the book.
"""

import requests
from bs4 import BeautifulSoup
from docx import Document
import time
import os

# ===============================
# CONFIGURATION
# ===============================
BOOK_NAME = "Abandoned Luna: Now Untouchable"
START_URL = "https://novelbin.com/b/abandoned-luna-now-untouchable/chapter-1-deceptive"
START_CHAPTER_NUMBER = 1  # chapter number corresponding to START_URL
END_CHAPTER_NUMBER = 432  # stop after this chapter
CHAPTERS_PER_FILE = 100
MAX_RETRIES = 3  # Number of retries per chapter

# ===============================
# PATH SETUP
# ===============================
OUTPUT_DIR = BOOK_NAME.strip().replace(" ", "_")
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ===============================
# HTTP HEADERS
# ===============================
HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/114.0.0.0 Safari/537.36"
    ),
    "Referer": "https://www.google.com/",
    "Accept-Language": "en-US,en;q=0.9",
}

# ===============================
# HELPER FUNCTIONS
# ===============================
def get_output_filename(file_index):
    """Generate a file name like 'Reincarnated_with_the_Strongest_System/1-100.docx'"""
    start_ch = (file_index - 1) * CHAPTERS_PER_FILE + 1
    end_ch = start_ch + CHAPTERS_PER_FILE - 1
    filename = f"{BOOK_NAME} {start_ch}-{end_ch}.docx"
    return os.path.join(OUTPUT_DIR, filename)


def create_new_document(file_index):
    """Create or load an existing .docx document for this file index."""
    filename = get_output_filename(file_index)
    if os.path.exists(filename):
        doc = Document(filename)
    else:
        doc = Document()
    return doc, filename


# ===============================
# MAIN SCRAPER LOGIC
# ===============================
def scrape_novel():
    url = START_URL
    chapter_number = START_CHAPTER_NUMBER
    file_index = (chapter_number - 1) // CHAPTERS_PER_FILE + 1
    chapter_count_in_file = (chapter_number - 1) % CHAPTERS_PER_FILE

    doc, output_file = create_new_document(file_index)
    print(f"📁 Using output file: {output_file}")

    while url and chapter_number <= END_CHAPTER_NUMBER:
        print(f"📖 Scraping Chapter {chapter_number}: {url}")

        success = False
        for attempt in range(1, MAX_RETRIES + 1):
            try:
                response = requests.get(url, headers=HEADERS, timeout=15)
                if response.status_code != 200:
                    if attempt < MAX_RETRIES:
                        print(f"⚠️ Attempt {attempt}/{MAX_RETRIES} failed to load page ({response.status_code}). Retrying...")
                        time.sleep(3 * attempt)  # Exponential backoff
                        continue
                    else:
                        print(f"❌ Failed to load page ({response.status_code}) after {MAX_RETRIES} attempts. Skipping chapter.")
                        break

                soup = BeautifulSoup(response.content, "html.parser")

                # ---- Extract title ----
                title_tag = soup.find("span", class_="chr-text")
                title = title_tag.get_text(strip=True) if title_tag else f"Chapter {chapter_number}"
                doc.add_heading(title, level=1)

                # ---- Extract content ----
                content_div = soup.find("div", id="chr-content")
                if not content_div:
                    if attempt < MAX_RETRIES:
                        print(f"⚠️ Attempt {attempt}/{MAX_RETRIES}: No content found. Retrying...")
                        time.sleep(3 * attempt)  # Exponential backoff
                        continue
                    else:
                        print("⚠️ No content found in this chapter after all retries.")
                else:
                    for p in content_div.find_all("p"):
                        for br in p.find_all("br"):
                            br.replace_with("\n")
                        text = p.get_text(separator="", strip=True)
                        if text:
                            doc.add_paragraph(text)

                # ---- Save progress continuously ----
                doc.save(output_file)
                print(f"💾 Saved: {title}")
                success = True

                chapter_number += 1
                chapter_count_in_file += 1

                # ---- Move to next chapter ----
                next_link = soup.find("a", class_="js-chapter-nav", attrs={"data-chapter-nav": "next"})
                if not next_link or next_link.get("disabled") is not None or "disabled" in (next_link.get("class") or []):
                    print("✅ No more chapters available.")
                    break

                url = next_link.get("data-chapter-url") or next_link.get("href")
                time.sleep(1)

                # ---- Start new DOCX file after 100 chapters ----
                if chapter_count_in_file >= CHAPTERS_PER_FILE:
                    print(f"📚 Finished 100 chapters. Starting new file...")
                    doc.save(output_file)
                    file_index += 1
                    chapter_count_in_file = 0
                    doc, output_file = create_new_document(file_index)
                    print(f"📁 New output file: {output_file}")

                break  # Successfully scraped, exit retry loop

            except Exception as e:
                if attempt < MAX_RETRIES:
                    print(f"⚠️ Attempt {attempt}/{MAX_RETRIES} failed for chapter {chapter_number}: {e}. Retrying...")
                    time.sleep(5 * attempt)  # Exponential backoff
                else:
                    print(f"❌ Error scraping chapter {chapter_number} after {MAX_RETRIES} attempts: {e}. Skipping chapter.")
                    chapter_number += 1
                    # Try to get next URL if possible
                    try:
                        response = requests.get(url, headers=HEADERS, timeout=15)
                        if response.status_code == 200:
                            soup = BeautifulSoup(response.content, "html.parser")
                            next_link = soup.find("a", class_="js-chapter-nav", attrs={"data-chapter-nav": "next"})
                            if next_link and next_link.get("disabled") is None and "disabled" not in (next_link.get("class") or []):
                                url = next_link.get("data-chapter-url") or next_link.get("href")
                    except:
                        pass
                    break
        
        if not success:
            continue

    print(f"\n✅ Scraping complete. Last file saved: {output_file}")


if __name__ == "__main__":
    scrape_novel()
