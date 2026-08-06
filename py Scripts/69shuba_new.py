import json
import os
import re
import subprocess
import time
import undetected_chromedriver as uc

from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import (
    WebDriverException,
    NoSuchElementException,
    TimeoutException,
)
from docx import Document
from bs4 import BeautifulSoup

PAGE_SUFFIX_RE = re.compile(r"\s*\(\s*\d+\s*/\s*\d+\s*\)\s*$")

_CHROME_PATHS = [
    "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome",
    "/Applications/Chromium.app/Contents/MacOS/Chromium",
]

def _detect_chrome_major_version():
    for path in _CHROME_PATHS:
        try:
            out = subprocess.check_output([path, "--version"], stderr=subprocess.DEVNULL).decode().strip()
            return int(out.split()[-1].split(".")[0])
        except Exception:
            continue
    return None
CHAPTER_PREFIX_RE = re.compile(r"^第\s*[\d一二三四五六七八九十百千万零〇两]+\s*章\s*")

MAX_RETRIES = 10
PAGE_TIMEOUT = 12
BATCH_SIZE = 100  # chapters per saved docx; aligned to multiples (1-100, 101-200, ...)
CHECKPOINT_FILENAME = ".checkpoint.json"


def load_checkpoint(output_dir, expected_end_chapter):
    """Read checkpoint if present and matching this end_chapter target."""
    path = os.path.join(output_dir, CHECKPOINT_FILENAME)
    if not os.path.exists(path):
        return None
    try:
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
    except (OSError, json.JSONDecodeError) as e:
        print(f"⚠️ Could not read checkpoint at {path}: {e}")
        return None
    if int(data.get("end_chapter", -1)) != int(expected_end_chapter):
        print(
            "⚠️ Checkpoint end_chapter "
            f"({data.get('end_chapter')}) does not match requested end_chapter "
            f"({expected_end_chapter}); ignoring checkpoint."
        )
        return None
    return data


def save_checkpoint(output_dir, next_chapter, next_url, end_chapter):
    """Atomically persist progress so a crash/restart can resume cleanly."""
    path = os.path.join(output_dir, CHECKPOINT_FILENAME)
    tmp = path + ".tmp"
    payload = {
        "next_chapter": int(next_chapter),
        "next_url": next_url,
        "end_chapter": int(end_chapter),
        "timestamp": time.time(),
    }
    try:
        with open(tmp, "w", encoding="utf-8") as f:
            json.dump(payload, f, ensure_ascii=False, indent=2)
        os.replace(tmp, path)
    except OSError as e:
        print(f"⚠️ Could not write checkpoint: {e}")


def clear_checkpoint(output_dir):
    path = os.path.join(output_dir, CHECKPOINT_FILENAME)
    try:
        os.remove(path)
    except FileNotFoundError:
        pass
    except OSError as e:
        print(f"⚠️ Could not delete checkpoint {path}: {e}")


def count_words(text):
    """Non-whitespace character count — the standard 'word count' for Chinese text."""
    return sum(1 for ch in text if not ch.isspace())


def append_word_count_summary(doc, chapter_stats):
    """Append a summary table of per-chapter word counts at the end of the doc."""
    if not chapter_stats:
        return
    doc.add_page_break()
    doc.add_heading("Chapter Word Counts", level=1)
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = "Chapter"
    hdr[1].text = "Title"
    hdr[2].text = "Word count"
    total = 0
    for n, t, c in chapter_stats:
        row = table.add_row().cells
        row[0].text = f"Chapter {n}"
        row[1].text = t
        row[2].text = f"{c:,}"
        total += c
    doc.add_paragraph(f"Total: {total:,} words across {len(chapter_stats)} chapter(s)")


def _next_chapter_url(driver):
    """Resolve the next chapter URL via #pb_next (preferred) or link-text fallback."""
    try:
        link = driver.find_element(By.ID, "pb_next")
        href = link.get_attribute("href")
        if href:
            return href
    except NoSuchElementException:
        pass
    try:
        link = driver.find_element(By.LINK_TEXT, "下一章")
        return link.get_attribute("href")
    except NoSuchElementException:
        return None

def sanitize_text(text):
    """Remove NULL bytes and control characters that aren't XML compatible."""
    if not text:
        return ""
    # Remove NULL bytes
    text = text.replace('\x00', '')
    # Remove other control characters except common whitespace (tab, newline, carriage return)
    # Keep: \t (0x09), \n (0x0A), \r (0x0D)
    # Remove others in range 0x00-0x1F except those above
    cleaned = []
    for char in text:
        code = ord(char)
        # Keep printable characters, common whitespace, and valid Unicode
        if code >= 32 or code in [9, 10, 13]:  # 9=tab, 10=newline, 13=carriage return
            cleaned.append(char)
        # Replace other control characters with space
        elif code < 32:
            cleaned.append(' ')
    return ''.join(cleaned)

def scrape_chapters(start_url, start_chapter, end_chapter, book_name):
    output_dir = book_name
    os.makedirs(output_dir, exist_ok=True)

    # --------------------------
    # UNDDetected Chrome (Headful)
    # --------------------------
    options = uc.ChromeOptions()
    options.add_argument("--no-sandbox")
    options.add_argument("--disable-blink-features=AutomationControlled")
    options.add_argument("--disable-dev-shm-usage")
    # Force headful mode
    options.headless = False
    # Don't block on slow ad scripts/images.
    options.page_load_strategy = "eager"
    prefs = {
        "profile.managed_default_content_settings.images": 2,
        "profile.default_content_setting_values.notifications": 2,
    }
    try:
        options.add_experimental_option("prefs", prefs)
    except Exception:
        pass

    # Resume from checkpoint if one exists for this end_chapter target.
    cp = load_checkpoint(output_dir, end_chapter)
    if cp:
        resumed_chapter = int(cp["next_chapter"])
        resumed_url = cp["next_url"]
        if resumed_chapter > end_chapter:
            print(
                f"✅ Checkpoint says next_chapter={resumed_chapter} > end_chapter="
                f"{end_chapter}; nothing to do. Delete {os.path.join(output_dir, CHECKPOINT_FILENAME)} to rerun."
            )
            return
        print(f"↩️  Resuming from chapter {resumed_chapter} ({resumed_url})")
        start_chapter = resumed_chapter
        start_url = resumed_url

    chrome_version = _detect_chrome_major_version()
    if chrome_version:
        print(f"Detected Chrome version: {chrome_version}")
    driver = uc.Chrome(options=options, version_main=chrome_version)
    wait = WebDriverWait(driver, PAGE_TIMEOUT)

    current_url = start_url
    current_chapter = start_chapter
    chapters_in_file = 0
    doc = Document()
    chapter_stats = []

    def save_doc(end_chapter_num, suffix=""):
        """Save the current doc with the word-count summary appended. Returns path or None."""
        if chapters_in_file <= 0:
            return None
        file_start = end_chapter_num - chapters_in_file + 1
        file_end = end_chapter_num
        filename = f"{book_name}_{file_start}-{file_end}{suffix}.docx"
        filepath = os.path.join(output_dir, filename)
        try:
            append_word_count_summary(doc, chapter_stats)
            doc.save(filepath)
            print(f"✅ Saved: {filepath}")
            return filepath
        except Exception as save_err:
            print(f"⚠️ Failed to save {filepath}: {save_err}")
            return None

    try:
        while current_chapter <= end_chapter:
            retries = 0
            success = False

            while retries < MAX_RETRIES and not success:
                try:
                    print(f"[{current_chapter}] Attempt {retries + 1}: {current_url}")
                    driver.get(current_url)
                    wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "div.txtnav")))

                    page_html = driver.page_source
                    soup = BeautifulSoup(page_html, "html.parser")

                    # Extract chapter title from <div class="txtnav"> > <h1>
                    h1 = soup.select_one("div.txtnav > h1")
                    if h1 is None:
                        raise NoSuchElementException("Chapter title not found")
                    full_title = h1.get_text(strip=True)
                    full_title = PAGE_SUFFIX_RE.sub("", full_title)
                    title = CHAPTER_PREFIX_RE.sub("", full_title).strip() or full_title
                    title = sanitize_text(title)

                    # Locate chapter body
                    content = soup.select_one("div.txtnav")
                    if content is None:
                        raise NoSuchElementException("txtnav content container not found")

                    # Strip metadata, ads, and scripts inside the body
                    for cls in ["reader-ad", "contentadv", "bottom-ad", "txtinfo"]:
                        for tag in content.find_all("div", class_=cls):
                            tag.decompose()

                    # Remove the top-right advertisement container if present.
                    txtright = content.find("div", id="txtright")
                    if txtright:
                        txtright.decompose()

                    for s in content.find_all("script"):
                        s.decompose()

                    # Add heading
                    doc.add_heading(f"Chapter {current_chapter}", level=1)
                    doc.add_heading(title, level=2)

                    # Remove the heading from the body so it is not duplicated.
                    body_title = content.find("h1")
                    if body_title:
                        body_title.decompose()

                    # New 69shuba pages separate paragraphs with <br> tags.
                    # Convert only those tags into newlines, then extract text deeply.
                    for br in content.find_all("br"):
                        br.replace_with("\n")

                    raw = content.get_text("", strip=False)
                    paragraphs_text = [
                        line.strip()
                        for line in raw.splitlines()
                        if line.strip()
                    ]

                    chapter_word_count = 0
                    for line in paragraphs_text:
                        cleaned_line = sanitize_text(line).strip()
                        if cleaned_line:
                            doc.add_paragraph(cleaned_line)
                            chapter_word_count += count_words(cleaned_line)

                    chapter_stats.append((current_chapter, title, chapter_word_count))
                    success = True
                    chapters_in_file += 1

                    # Save at chapter-number boundaries (100, 200, ...) or at end_chapter,
                    # so batches stay aligned (1-100, 101-200, ...) across resumes.
                    if current_chapter % BATCH_SIZE == 0 or current_chapter == end_chapter:
                        save_doc(current_chapter)
                        doc = Document()
                        chapters_in_file = 0
                        chapter_stats = []

                    # Move to next chapter
                    next_url = _next_chapter_url(driver)
                    next_chapter = current_chapter + 1

                    # Persist progress before navigating, so resume picks the right next URL.
                    if next_url and next_chapter <= end_chapter:
                        save_checkpoint(output_dir, next_chapter, next_url, end_chapter)

                    if not next_url:
                        print("No next-chapter link found. Ending early.")
                        current_url = None
                        break
                    current_url = next_url
                    current_chapter = next_chapter

                except (WebDriverException, NoSuchElementException, TimeoutException) as e:
                    print(f"⚠️ Error scraping chapter {current_chapter}: {e}")
                    retries += 1
                    time.sleep(0.8)

            if not success:
                print(f"❌ Skipping Chapter {current_chapter} after {MAX_RETRIES} failed attempts.")
                try:
                    nxt = _next_chapter_url(driver)
                except Exception as nav_err:
                    print(f"⚠️ Could not query next-chapter link: {nav_err}")
                    break
                if not nxt:
                    print("No next-chapter link found. Ending early.")
                    break
                # Move past the skipped chapter and checkpoint so a rerun doesn't retry it forever.
                current_url = nxt
                current_chapter += 1
                if current_chapter <= end_chapter:
                    save_checkpoint(output_dir, current_chapter, current_url, end_chapter)

        # Clean completion: clear the checkpoint so the next run won't think it's mid-job.
        clear_checkpoint(output_dir)
    finally:
        # Flush whatever is still in the in-memory doc — covers normal completion,
        # KeyboardInterrupt, urllib3/Selenium crashes, and any other unhandled exception.
        if chapters_in_file > 0:
            last_done = chapter_stats[-1][0] if chapter_stats else current_chapter
            save_doc(last_done, suffix="_partial")
        try:
            driver.quit()
        except Exception:
            pass


# === Configuration ===
start_url = "https://www.69shuba.com/txt/84861/39386250"
book_name = "The head chef in the kitchen has won another"
start_chapter = 1
end_chapter = 10

scrape_chapters(start_url, start_chapter, end_chapter, book_name)
