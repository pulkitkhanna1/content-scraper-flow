#!/usr/bin/env python3
"""
Kakao Webtoon OCR Scraper - Full script (with post-login cookie wait)
- Waits after login to collect/save/sync cookies before fetching viewer content
- Calls ensure_cookies_synced_and_wait() after programmatic login and popup login flows
- Uses new GET API: tries GET /viewer/data first, falls back to existing GraphQL POST
- Uses Kakao content product-list API to collect chapter product IDs before scraping
- Other behavior unchanged from previous version
"""

import os
import time
import random
import ssl
import certifi
import cv2
import numpy as np
import requests
from requests.adapters import HTTPAdapter
from requests.utils import cookiejar_from_dict
from urllib3.util.retry import Retry
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import re
import shutil
import tempfile
import pickle
import subprocess
import logging
import base64
import csv
import html as _html
from urllib.parse import urljoin
import getpass
from pathlib import Path
from datetime import datetime
from typing import Optional

# Selenium + undetected imports
try:
    import undetected_chromedriver as uc
    UNDETECTED_AVAILABLE = True
except Exception:
    UNDETECTED_AVAILABLE = False

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager

# SSL fix for requests
import urllib.request
ssl_context = ssl.create_default_context(cafile=certifi.where())
urllib.request.install_opener(urllib.request.build_opener(urllib.request.HTTPSHandler(context=ssl_context)))

# =======================
# Configuration
# =======================
MAX_CHAPTERS = 553
START_CHAPTER = 1
START_PRODUCT_ID = None
SERIES_ID = 55657093
ALLOW_IMAGE_OCR_FALLBACK = False
PRODUCT_LIST_API_URL = "https://bff-page.kakao.com/api/gateway/api/v2/content/product/list"

# Output / folder
try:
    CODEBASE_DIR = Path(__file__).parent.absolute()
except Exception:
    CODEBASE_DIR = Path.cwd()
DEFAULT_OUTPUT_FOLDER = "Output(27557-27559)"
NOVEL_TITLE = "The Regressed Extra Becomes a Genius"

# Credentials (change env or replace)
EMAIL = os.getenv('KAKAO_USER') or "100thpageonly@gmail.com"
PASSWORD = os.getenv('KAKAO_PASS') or "pfm2024*"
if not EMAIL:
    EMAIL = input('Kakao email: ').strip()
if not PASSWORD:
    try:
        PASSWORD = getpass.getpass('Kakao password (hidden): ')
    except Exception:
        PASSWORD = None

DEFAULT_CHROME_PROFILE = os.path.expanduser("~/Library/Application Support/Google/Chrome/Profile 1")
USE_CHROME_PROFILE = True
KEEP_BROWSER_OPEN = False

LAST_LOGIN_ATTEMPT = 0
LOGIN_COOLDOWN = 25
TEMP_PROFILE_DIR = None

# Logging
log = logging.getLogger('kakao_scraper')
log.setLevel(logging.INFO)
ch = logging.StreamHandler()
ch.setLevel(logging.INFO)
formatter = logging.Formatter("%(asctime)s [%(levelname)s] %(message)s")
ch.setFormatter(formatter)
log.addHandler(ch)

# ============================================================================
# Folder & Docx management
# ============================================================================

class FolderManager:
    def __init__(self, folder_name: str = DEFAULT_OUTPUT_FOLDER):
        self.base_dir = CODEBASE_DIR / folder_name
        self.base_dir.mkdir(exist_ok=True)
        log.info(f"📁 Output folder: {self.base_dir}")

    def get_novel_folder(self, novel_title: str) -> Path:
        safe_name = "".join(c for c in novel_title if c.isalnum() or c in (' ', '-', '_'))[:50]
        novel_folder = self.base_dir / safe_name
        novel_folder.mkdir(exist_ok=True)
        return novel_folder


class DocxManager:
    CHAPTERS_PER_FILE = 100

    def __init__(self, novel_folder: Path, novel_title: str, author: str = "Unknown"):
        self.novel_folder = novel_folder
        self.novel_title = novel_title
        self.novel_author = author
        self.current_doc = None
        self.current_chapter_start = None
        self.current_chapter_end = None
        self.chapter_count = 0
        self.total_chapters_added = 0
        self._create_new_doc()

    def _create_new_doc(self):
        self.current_doc = Document()
        title = self.current_doc.add_heading(self.novel_title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author = self.current_doc.add_paragraph(f"Author: {self.novel_author}")
        author.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            author.runs[0].italic = True
        except Exception:
            pass
        date_para = self.current_doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.current_doc.add_paragraph()
        self.chapter_count = 0

    def add_chapter(self, chapter_number: int, chapter_name: str, content: str) -> Optional[str]:
        if self.chapter_count == 0:
            self.current_chapter_start = chapter_number

        chapter_name = clean_docx_heading_text(chapter_name)
        content = clean_extracted_chapter_text(content)
        heading = self.current_doc.add_heading(f"Chapter {chapter_number}: {chapter_name}", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        content_para = self.current_doc.add_paragraph(content)
        try:
            content_para.paragraph_format.line_spacing = 1.5
        except Exception:
            pass

        self.current_doc.add_paragraph()
        self.chapter_count += 1
        self.total_chapters_added += 1
        self.current_chapter_end = chapter_number

        if self.chapter_count >= self.CHAPTERS_PER_FILE:
            return self._save_current_doc()
        return None

    def _save_current_doc(self) -> str:
        if self.chapter_count == 0:
            return None
        file_path = self.novel_folder / f"Chapter_{self.current_chapter_start}to{self.current_chapter_end}.docx"
        self.current_doc.save(str(file_path))
        log.info(f"✅ Saved: {file_path.name} (Chapters {self.current_chapter_start}-{self.current_chapter_end})")
        saved_path = str(file_path)
        self._create_new_doc()
        return saved_path

    def finalize(self) -> Optional[str]:
        if self.chapter_count > 0:
            return self._save_current_doc()
        return None

# ============================================================================
# HTTP session, helpers, OCR
# ============================================================================

def make_session():
    s = requests.Session()
    retries = Retry(total=5, backoff_factor=1.5,
                    status_forcelist=[429, 500, 502, 503, 504],
                    allowed_methods=["GET", "POST", "HEAD"])
    s.mount("https://", HTTPAdapter(max_retries=retries))
    s.headers.update({
        "User-Agent": "Mozilla/5.0 (Macintosh) OCR",
        "Accept": "*/*",
        "Referer": "https://page.kakao.com/"
    })
    return s

SESSION = make_session()

def _safe_filename(s: str) -> str:
    return re.sub(r"[^0-9A-Za-z._-]", "_", s)[:200]

def print_api_response(label, url, resp=None, content_bytes=None, exception=None, truncate=5000):
    try:
        print('\n' + '='*40)
        print(f'API RESPONSE - {label} - {url}')
        if exception is not None:
            print('Exception:', repr(exception)); print('='*40 + '\n'); return
        if resp is not None:
            status = getattr(resp, 'status_code', None)
            print('status:', status)
            headers = {}
            try:
                headers = dict(resp.headers or {})
            except Exception:
                pass
            safe_headers = {k:v for k,v in headers.items() if k.lower() not in ('set-cookie','cookie','authorization')}
            print('headers:', safe_headers)
            ctype = headers.get('Content-Type','') if headers else ''
            try:
                is_json_like = ('json' in ctype.lower()) or ('textviewer' in (resp.url or '').lower()) or (resp.text.strip().startswith('{') or resp.text.strip().startswith('['))
            except Exception:
                is_json_like = False
            if is_json_like:
                try:
                    parsed = resp.json()
                except Exception:
                    try:
                        parsed = json.loads(resp.content.decode('utf-8', errors='replace'))
                    except Exception:
                        parsed = None
                if parsed is not None:
                    try:
                        text, _ = parse_textviewer_content_json(parsed, ats_base='')
                        if not text:
                            text = extract_text_from_json(parsed)
                    except Exception:
                        text = extract_text_from_json(parsed)
                    if text:
                        if len(text) > truncate:
                            print(text[:truncate]); print('...<truncated>...')
                        else:
                            print(text)
                        print('='*40 + '\n'); return
                    else:
                        s = json.dumps(parsed, ensure_ascii=False)
                        if len(s) > truncate:
                            print(s[:truncate]); print('...<truncated>...')
                        else:
                            print(s)
                        print('='*40 + '\n'); return
            try:
                if isinstance(resp.text, str) and (('json' in ctype) or ('text' in ctype) or len(resp.text) < 200000):
                    data_text = resp.text
                    if len(data_text) > truncate:
                        print(data_text[:truncate]); print('...<truncated>...')
                    else:
                        print(data_text)
                else:
                    try:
                        blen = len(resp.content)
                    except Exception:
                        blen = 'unknown'
                    print(f'binary content length: {blen}, content-type: {ctype}')
            except Exception as e:
                print('failed to print resp text:', e)
        elif content_bytes is not None:
            try:
                sample_text = content_bytes[:3000].decode('utf-8', errors='replace').strip()
                if sample_text.startswith('{') or sample_text.startswith('[') or 'paragraphList' in sample_text or '<div' in sample_text:
                    try:
                        parsed = json.loads(content_bytes.decode('utf-8', errors='replace'))
                        text, _ = parse_textviewer_content_json(parsed, ats_base='')
                        if not text:
                            text = extract_text_from_json(parsed)
                        if text:
                            if len(text) > truncate:
                                print(text[:truncate]); print('...<truncated>...')
                            else:
                                print(text)
                            print('='*40 + '\n'); return
                    except Exception:
                        pass
            except Exception:
                pass
            print('binary content length:', len(content_bytes))
            try:
                sample = content_bytes[:2000].decode('utf-8', errors='replace')
                if len(sample) > truncate:
                    print(sample[:truncate]); print('...<truncated>...')
                else:
                    print(sample)
            except Exception:
                try:
                    preview = base64.b64encode(content_bytes[:256]).decode('ascii', errors='replace')
                    print('base64 preview (first 256 bytes):', preview)
                except Exception:
                    print('[binary data preview unavailable]')
        print('='*40 + '\n')
    except Exception:
        try:
            print('api-response-print-failed for', label, url)
        except Exception:
            pass

_OCR_READER = None
def get_ocr_reader():
    global _OCR_READER
    if _OCR_READER is None:
        import easyocr
        _OCR_READER = easyocr.Reader(['ko'], gpu=False, verbose=False)
        log.info('Initialized EasyOCR')
    return _OCR_READER

def fetch_graphql(session, graphql_url, payload: dict, referer=None, timeout=30):
    headers = {
        "Content-Type": "application/json",
        "Accept": "application/json, text/plain, */*",
        "X-Requested-With": "XMLHttpRequest",
    }
    if referer:
        headers["Referer"] = referer
    headers.setdefault("Origin", "https://page.kakao.com")
    try:
        resp = session.post(graphql_url, json=payload, headers=headers, timeout=timeout)
    except Exception as e:
        print_api_response('GraphQL POST', graphql_url, exception=e)
        return None
    print_api_response('GraphQL POST', graphql_url, resp=resp)
    try:
        j = resp.json()
    except Exception:
        try:
            text = resp.content.decode('utf-8', errors='replace')
            j = json.loads(text)
        except Exception:
            j = None
    return j

# -----------------------
# New: GET viewer/data helper (tries the new BFF GET endpoint)
# -----------------------
def fetch_viewer_data_via_get(session, series_id, product_id, referer=None, timeout=20):
    """
    Try the new GET API:
      GET https://bff-page.kakao.com/api/gateway/api/v1/viewer/data
      params: series_id, product_id

    Returns parsed JSON on success, or None on failure.
    """
    url = "https://bff-page.kakao.com/api/gateway/api/v1/viewer/data"
    headers = {
        "Accept": "application/json, text/plain, */*",
        "X-Requested-With": "XMLHttpRequest",
    }
    if referer:
        headers["Referer"] = referer
    # keep Origin header like other requests
    headers.setdefault("Origin", "https://page.kakao.com")
    try:
        resp = session.get(url, params={"series_id": series_id, "product_id": product_id},
                           headers=headers, timeout=timeout)
    except Exception as e:
        print_api_response("Viewer GET (network error)", url, exception=e)
        return None

    # Print/log the response for debugging
    print_api_response("Viewer GET", resp.url, resp=resp)

    # If 403/401 -> likely missing auth/cookies
    if getattr(resp, "status_code", None) in (401, 403):
        log.warning("GET viewer/data returned %s — ensure cookies/auth synced.", resp.status_code)
        return None

    try:
        j = resp.json()
    except Exception:
        try:
            j = json.loads(resp.content.decode("utf-8", errors="replace"))
        except Exception:
            j = None

    return j

def fetch_image_via_selenium(driver, url, timeout=20):
    driver.set_script_timeout(timeout)
    script = """
    var url = arguments[0];
    var cb = arguments[1];
    fetch(url, { credentials: 'include' }).then(function(resp){
        if(!resp.ok){ cb(JSON.stringify({error: 'http:'+resp.status})); return; }
        return resp.arrayBuffer();
    }).then(function(buf){
        if(!buf){ return; }
        var bytes = new Uint8Array(buf);
        var binary = '';
        var chunkSize = 0x8000;
        for (var i = 0; i < bytes.length; i += chunkSize) {
            var chunk = bytes.subarray(i, i + chunkSize);
            var str = '';
            for (var j = 0; j < chunk.length; j++) { str += String.fromCharCode(chunk[j]); }
            binary += str;
        }
        var b64 = btoa(binary);
        cb(JSON.stringify({b64: b64}));
    }).catch(function(e){ cb(JSON.stringify({error: String(e)})); });
    """
    raw = driver.execute_async_script(script, url)
    try:
        print('selenium raw fetch preview (first 2000 chars):', (raw[:2000] if isinstance(raw, str) else str(raw)) )
    except Exception:
        pass
    data = None
    try:
        data = json.loads(raw) if isinstance(raw, str) else raw
    except Exception:
        data = raw
    if isinstance(data, dict) and data.get('b64'):
        b = base64.b64decode(data['b64'])
        print_api_response('Selenium Fetch (decoded)', url, content_bytes=b)
        return b
    raise Exception(f"selenium-fetch-error: {data}")

def download_resource_with_fallback(driver, session, url, timeout=20):
    headers = dict(session.headers or {})
    try:
        cur = getattr(driver, 'current_url', None)
        if cur:
            headers['Referer'] = cur
    except Exception:
        pass
    try:
        resp = session.get(url, headers=headers, timeout=timeout)
        print_api_response('Resource GET', url, resp=resp)
        resp.raise_for_status()
        ctype = resp.headers.get('Content-Type', '') or ''
        return resp.content, ctype
    except Exception as req_err:
        print_api_response('Resource GET (failed)', url, exception=req_err)
        try:
            data = fetch_image_via_selenium(driver, url, timeout=timeout+10)
            return data, ''
        except Exception as sel_err:
            print_api_response('Selenium fallback failed', url, exception=sel_err)
            raise Exception(f"resource-fetch-failed: requests_error={req_err}, selenium_error={sel_err}")

# JSON parsing helpers (already in your code)
def _collect_strings_from_json(obj, out):
    if obj is None:
        return
    if isinstance(obj, str):
        s = obj.strip()
        if s:
            out.append(s)
        return
    if isinstance(obj, (int, float, bool)):
        out.append(str(obj)); return
    if isinstance(obj, dict):
        for key in ('text', 'paragraph', 'paragraphs', 'body', 'lines', 'content', 'contents', 'items'):
            if key in obj:
                _collect_strings_from_json(obj[key], out)
        for v in obj.values():
            _collect_strings_from_json(v, out)
        return
    if isinstance(obj, (list, tuple)):
        for it in obj:
            _collect_strings_from_json(it, out)
        return

def extract_text_from_json(parsed_json):
    pieces = []
    _collect_strings_from_json(parsed_json, pieces)
    clean = []
    seen = set()
    for p in pieces:
        if not p: continue
        if p in seen: continue
        if len(p) <= 1 and p.isdigit(): continue
        seen.add(p)
        clean.append(p)
    return "\n\n".join(clean) if clean else ''

def resolve_token_url(token: str, ats_base: str = "") -> str:
    if not token: return token
    token = token.strip()
    if token.startswith('http://') or token.startswith('https://'): return token
    if ats_base:
        if ats_base.endswith('=') or '?' in ats_base:
            return ats_base + token
        return urljoin(ats_base, token)
    return token

def parse_textviewer_content_json(payload: dict, ats_base: str = ""):
    if not payload: return "", []
    content_info = payload.get('contentInfo') or payload
    paragraph_list = content_info.get('paragraphList') or []
    out_paragraphs = []
    images = []
    def extract_from_paragraph(para, depth=0):
        if not isinstance(para, dict): return [], []
        local_text_pieces = []
        local_images = []
        top_text = para.get('text') or ''
        if top_text and top_text.strip(): local_text_pieces.append(top_text.strip())
        img = para.get('image')
        if img:
            token = None
            if isinstance(img, str):
                token = img
            elif isinstance(img, dict):
                token = img.get('src') or img.get('token') or img.get('fileName') or None
            if token:
                resolved = resolve_token_url(token, ats_base)
                pid = para.get('id') or None
                local_images.append({'paragraph_id': pid, 'image_token': token, 'resolved_url': resolved})
                local_text_pieces.append(f'[IMAGE: {resolved}]')
        child_list = para.get('childParagraphList') or []
        if child_list:
            for child in child_list:
                if not isinstance(child, dict): continue
                child_text = child.get('text') or ''
                if child_text and child_text.strip(): local_text_pieces.append(child_text.strip())
                child_img = child.get('image')
                if child_img:
                    token = None
                    if isinstance(child_img, str):
                        token = child_img
                    elif isinstance(child_img, dict):
                        token = child_img.get('src') or child_img.get('token') or child_img.get('fileName') or None
                    if token:
                        resolved = resolve_token_url(token, ats_base)
                        pid = child.get('id') or None
                        local_images.append({'paragraph_id': pid, 'image_token': token, 'resolved_url': resolved})
                        local_text_pieces.append(f'[IMAGE: {resolved}]')
                nested_child_list = child.get('childParagraphList')
                if nested_child_list:
                    nested_texts, nested_images = extract_from_paragraph(child, depth + 1)
                    local_text_pieces.extend(nested_texts)
                    local_images.extend(nested_images)
        return local_text_pieces, local_images

    for para in paragraph_list:
        para_texts, para_images = extract_from_paragraph(para)
        if para_texts:
            joined = "".join([t for t in para_texts if t])
            if joined: out_paragraphs.append(joined)
        images.extend(para_images)
    final_text = "\n\n".join(out_paragraphs)
    return final_text, images

def _iter_json_values(obj):
    if isinstance(obj, dict):
        for value in obj.values():
            yield from _iter_json_values(value)
    elif isinstance(obj, (list, tuple)):
        for value in obj:
            yield from _iter_json_values(value)
    else:
        yield obj

def response_says_not_purchased(payload) -> bool:
    for value in _iter_json_values(payload):
        if isinstance(value, str) and value.strip().lower() == 'not_purchased':
            return True
    return False

def normalize_resource_url(url: str) -> str:
    if not url:
        return ''
    return _html.unescape(str(url)).replace('\\u0026', '&').strip().strip('\'",')

def is_text_resource_url(url: str) -> bool:
    normalized = normalize_resource_url(url)
    low = normalized.lower()
    if not normalized or 'resource?kid=' not in low:
        return False
    if 'textviewercontentmeta.json' in low:
        return False
    if 'filename=th' in low or 'thumbnail' in low:
        return False
    return '.json' in low or 'textviewer' in low

def extract_text_resource_urls_from_api_response(payload):
    """
    Kakao's GET viewer/data can return an ats base URL plus signed JSON tokens
    instead of the older viewerData.contentsList shape. Rebuild those text URLs.
    """
    strings = [normalize_resource_url(v) for v in _iter_json_values(payload) if isinstance(v, str)]
    bases = []
    direct_urls = []
    tokens = []

    for value in strings:
        low = value.lower()
        if 'resource?kid=' in low and value.startswith(('http://', 'https://')):
            if is_text_resource_url(value):
                direct_urls.append(value)
            if low.endswith('kid=') or re.search(r'resource\?kid=$', low):
                bases.append(value)
        elif 'filename=' in low and '.json' in low and ('signature=' in low or 'credential=' in low):
            tokens.append(value)

    urls = []
    for url in direct_urls:
        if is_text_resource_url(url):
            urls.append(url)

    for base in bases:
        for token in tokens:
            rebuilt = base + token
            if is_text_resource_url(rebuilt):
                urls.append(rebuilt)

    deduped = []
    seen = set()
    for url in urls:
        if url not in seen:
            deduped.append(url)
            seen.add(url)
    return deduped

def _looks_like_useful_chapter_text(text: str) -> bool:
    if not text:
        return False
    clean = text.strip()
    if len(clean) < 20:
        return False
    low = clean.lower()
    if 'resource?kid=' in low or ('signature=' in low and 'credential=' in low):
        return False
    return True

def strip_leading_layout_metadata(text: str) -> str:
    if not text:
        return text
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if len(lines) < 8:
        return text

    first_part = lines[:min(len(lines), 30)]
    has_cover_meta = (
        'IMG' in first_part
        and 'True' in first_part
        and (
            'cover_f.jpg' in first_part
            or 'cover' in first_part
            or 'book' in first_part
            or 'KAGE' in first_part
        )
    )
    if not has_cover_meta:
        return text

    true_idx = first_part.index('True')
    if 'Start' in first_part:
        start_idx = first_part.index('Start')
        prefix = lines[:start_idx]
        cleaned = prefix + lines[true_idx + 1:]
    elif lines[0] == 'DIV' or 'KAGE' in first_part:
        cleaned = lines[true_idx + 1:]
    else:
        return text

    return "\n\n".join(cleaned)

# ============================================================================
# Cookies, driver helpers, login, popup handling, waits
# ============================================================================

def cookies_file_path():
    base = TEMP_PROFILE_DIR if TEMP_PROFILE_DIR else '.'
    return os.path.join(base, 'kakao_cookies.pkl')

def save_cookies(driver, path=None):
    path = path or cookies_file_path()
    try:
        cookies = driver.get_cookies()
        with open(path, 'wb') as f:
            pickle.dump(cookies, f)
        log.info('Saved %d cookies to %s', len(cookies), path)
    except Exception:
        log.exception('Failed to save cookies')

def load_cookies(driver, path=None, base_url='https://page.kakao.com'):
    path = path or cookies_file_path()
    if not os.path.exists(path):
        return False
    try:
        with open(path, 'rb') as f:
            cookies = pickle.load(f)
        driver.get(base_url)
        time.sleep(1.0)
        added = 0
        for c in cookies:
            cookie = {k: v for k, v in c.items() if k not in ('sameSite',)}
            try:
                driver.add_cookie(cookie)
                added += 1
            except Exception:
                pass
        driver.refresh()
        time.sleep(1.0)
        log.info('Loaded %d cookies', added)
        return True
    except Exception:
        log.exception('Failed to load cookies')
        return False

def sync_cookies_to_requests(driver, session):
    try:
        selenium_cookies = driver.get_cookies()
        jar = {}
        for c in selenium_cookies:
            jar[c['name']] = c.get('value', '')
        session.cookies = cookiejar_from_dict(jar)
        try:
            current = driver.current_url
            if current:
                session.headers.update({'Referer': current})
        except Exception:
            pass
        if 'User-Agent' not in session.headers:
            session.headers.update({'User-Agent': 'Mozilla/5.0 (Macintosh) OCR'})
        log.debug('Synced cookies to requests')
    except Exception:
        log.exception('Failed to sync cookies')

def ensure_cookies_synced_and_wait(driver, session, timeout=15, min_cookies=3):
    """
    After a login action, call this to wait until cookies appear in the browser,
    save them, sync to requests.Session, and give servers a small grace period.
    """
    deadline = time.time() + timeout
    last_count = 0
    while time.time() < deadline:
        try:
            ck = driver.get_cookies()
            count = len(ck) if ck else 0
            if count != last_count:
                log.info("Browser cookies count: %d", count)
                last_count = count
            # prefer is_logged_in if possible
            if count >= min_cookies or is_logged_in(driver):
                try:
                    save_cookies(driver)
                    sync_cookies_to_requests(driver, session)
                    # small randomized sleep to let server-side session propagate
                    t = random.uniform(1.2, 3.0)
                    log.info("Cookies appear synced — sleeping %.2fs to let session propagate", t)
                    time.sleep(t)
                    return True
                except Exception:
                    log.exception("Failed to save/sync cookies; retrying")
            else:
                # still waiting for more cookies or login indicator
                time.sleep(0.5)
        except Exception:
            time.sleep(0.5)
    log.warning("Timeout waiting for cookies to be created/synced (timeout=%ds)", timeout)
    return False

def is_logged_in(driver) -> bool:
    try:
        elems = driver.find_elements(By.CSS_SELECTOR, "[data-tiara-layer='gnb_profile']")
        if elems and any(e.is_displayed() for e in elems): return True
        if 'accounts.kakao.com' in (driver.current_url or ''): return False
    except Exception:
        pass
    return False

def wait_for_load(driver, timeout=15):
    try:
        WebDriverWait(driver, timeout).until(lambda d: d.execute_script('return document.readyState') == 'complete')
        time.sleep(0.35 + random.random() * 0.5)
        return True
    except Exception:
        return False

def wait_for_viewer_ready(driver, timeout=20):
    """
    Wait for viewer shell to be hydrated OR for signed CDN URLs to appear in page.
    Returns True if viewer looks ready (JS hydrated), False if timeout.
    """
    deadline = time.time() + timeout
    log.info("Waiting up to %ds for viewer to be ready...", timeout)
    while time.time() < deadline:
        try:
            # 1) readyState complete
            ready = driver.execute_script("return document.readyState") == 'complete'
            # 2) page contains dn-img-page token = signed url present
            page_src = (driver.page_source or '').lower()
            if 'dn-img-page.kakao.com' in page_src or '.json?signature=' in page_src:
                log.info("Found signed CDN url in page source -> viewer ready")
                wait_for_load(driver, timeout=2)
                return True
            # 3) DOM viewer container heuristics (common attributes)
            selectors = [
                "div.viewer", "div[data-test='viewer']", "div[data-testid='viewer']",
                "[data-test='viewer-container']", "[data-tiara-layer='viewer']",
                "div.DC2CN", "div[class*='textviewer']"
            ]
            for s in selectors:
                try:
                    els = driver.find_elements(By.CSS_SELECTOR, s)
                    if els and any(e.is_displayed() for e in els):
                        log.info("Detected viewer DOM element (%s)", s)
                        wait_for_load(driver, timeout=2)
                        return True
                except Exception:
                    continue
            # 4) check some inline JS variable that may contain viewerData (best-effort)
            try:
                js_check = driver.execute_script("""
                    try {
                        var found = false;
                        if (window.__PRELOADED_STATE__ && JSON.stringify(window.__PRELOADED_STATE__).length > 20) { found = true; }
                        if (!found && window.viewerData) found = true;
                        return found;
                    } catch(e) { return false; }
                """)
                if js_check:
                    log.info("Detected page global viewer data (via JS)")
                    wait_for_load(driver, timeout=2)
                    return True
            except Exception:
                pass
        except Exception:
            pass
        time.sleep(0.6)
    log.warning("Timeout waiting for viewer readiness")
    return False

def dump_auth_state(driver, label="auth_state"):
    try:
        cookies = driver.get_cookies()
        cookie_names = [c.get('name') for c in cookies]
        log.info("[%s] Selenium cookies: %s", label, cookie_names)
    except Exception:
        log.exception("Failed to read cookies for %s", label)
    try:
        script = """
            var out = { keys: [], kv: {} };
            try {
                for (var i=0;i<localStorage.length;i++){
                    var k = localStorage.key(i);
                    out.keys.push(k);
                    try { out.kv[k] = localStorage.getItem(k); } catch(e) { out.kv[k] = '[error]'; }
                }
            } catch(e) { out.err = String(e); }
            return out;
        """
        ls = driver.execute_script(script)
        keys = ls.get('keys') if isinstance(ls, dict) else []
        log.info("[%s] localStorage keys: %s", label, keys[:30])
        kv = ls.get('kv', {}) if isinstance(ls, dict) else {}
        sample = {k: (v[:80] + '...') if isinstance(v, str) and len(v) > 80 else v for k, v in list(kv.items())[:8]}
        if sample:
            log.info("[%s] localStorage sample: %s", label, sample)
    except Exception:
        log.exception("Failed to dump localStorage for %s", label)

def wait_for_post_login_redirect(driver, original_viewer_page=None, timeout=30):
    deadline = time.time() + timeout
    log.info("Waiting up to %ds for post-login redirect to complete...", timeout)
    while time.time() < deadline:
        try:
            cur = (driver.current_url or '').lower()
            if '/content/' in cur and '/viewer/' in cur:
                log.info("Detected viewer page after login: %s", cur)
                wait_for_load(driver, timeout=5)
                return True
            if 'page.kakao.com/relay/login' in cur:
                log.info("On relay/login page: %s", cur)
            if is_logged_in(driver):
                log.info("is_logged_in() returned True after login; continuing.")
                return True
        except Exception:
            pass
        time.sleep(0.7)
    log.warning("Timeout waiting for post-login redirect (timeout %ds)", timeout)
    return False

def handle_possible_login_popup(driver, viewer_page=None):
    """
    Detect login-required popup and attempt to log in programmatically (or click to let user).
    """
    global LAST_LOGIN_ATTEMPT
    try:
        # close trivial close buttons
        close_selectors = ["button[aria-label='닫기']", "button.btn-close"]
        for s in close_selectors:
            try:
                els = driver.find_elements(By.CSS_SELECTOR, s)
                for el in els:
                    try:
                        if el.is_displayed():
                            driver.execute_script('arguments[0].click();', el)
                            time.sleep(0.15)
                    except Exception:
                        pass
            except Exception:
                pass

        popup_xpaths = [
            "//*[contains(normalize-space(.), '해당 콘텐츠를 열람하시려면 로그인이 필요합니다.')]",
            "//*[contains(normalize-space(.), '로그인이 필요합니다')]",
            "//*[contains(normalize-space(.), '로그인 하기')]",
            "//*[contains(normalize-space(.), 'You must log in to view this content.')]",
            "//*[contains(@class, 'bg-sp-bg-popup') or contains(@class, 'z-100')]",
        ]

        popup_elem = None
        for xp in popup_xpaths:
            try:
                els = driver.find_elements(By.XPATH, xp)
                if els:
                    for e in els:
                        try:
                            if e.is_displayed():
                                popup_elem = e
                                break
                        except Exception:
                            continue
                if popup_elem:
                    break
            except Exception:
                continue

        if not popup_elem:
            try:
                # fallback scanning common modal containers
                candidates = driver.find_elements(By.CSS_SELECTOR, "div[class*='popup'], div[class*='modal'], div[class*='bg-sp-bg-popup'], div[class*='z-100']")
                for c in candidates:
                    try:
                        if not c.is_displayed(): continue
                        text = (c.text or "").strip()
                        if '로그인' in text or '로그인이 필요' in text or 'You must log in' in text:
                            popup_elem = c
                            break
                    except Exception:
                        continue
            except Exception:
                pass

        if not popup_elem:
            return False

        log.info('Login-required popup detected')

        btn_candidates = []
        try:
            btn_candidates.extend(popup_elem.find_elements(By.TAG_NAME, "button"))
            btn_candidates.extend(popup_elem.find_elements(By.TAG_NAME, "a"))
        except Exception:
            pass

        login_btn = None
        for b in btn_candidates:
            try:
                if not b.is_displayed(): continue
                txt = (b.text or '').strip()
                txt_low = txt.lower()
                if any(k in txt_low for k in ('로그인', '로그인 하기', 'login', 'log in')):
                    login_btn = b
                    break
            except Exception:
                continue

        if not login_btn and btn_candidates:
            for b in reversed(btn_candidates):
                try:
                    if b.is_displayed():
                        login_btn = b
                        break
                except Exception:
                    continue

        if not login_btn:
            log.info('No login button found inside popup; aborting popup handler')
            return False

        now = time.time()
        if now - LAST_LOGIN_ATTEMPT < LOGIN_COOLDOWN:
            log.info('Skipping programmatic login due to cooldown; clicking login to let user manually finish.')
            try:
                driver.execute_script('arguments[0].click();', login_btn)
            except Exception:
                try:
                    login_btn.click()
                except Exception:
                    pass
            # after clicking, give more time for the user flow
            time.sleep(1.2)
            ensure_cookies_synced_and_wait(driver, SESSION, timeout=10)
            return True

        # click login
        try:
            driver.execute_script('arguments[0].click();', login_btn)
        except Exception:
            try:
                login_btn.click()
            except Exception:
                log.warning('Could not click login button programmatically')

        LAST_LOGIN_ATTEMPT = now

        # wait for any login form or redirect
        try:
            WebDriverWait(driver, 6).until(
                EC.presence_of_element_located((By.CSS_SELECTOR, "input[type='password'], input[type='email'], input[type='text']"))
            )
            # try auto-fill if possible
            try:
                el_email = None
                el_pwd = None
                try:
                    el_email = driver.find_element(By.CSS_SELECTOR, "input[type='email']")
                except Exception:
                    try:
                        el_email = driver.find_element(By.CSS_SELECTOR, "input[type='text']")
                    except Exception:
                        el_email = None
                try:
                    el_pwd = driver.find_element(By.CSS_SELECTOR, "input[type='password']")
                except Exception:
                    el_pwd = None

                if el_email and el_pwd:
                    if EMAIL:
                        try:
                            el_email.clear(); el_email.send_keys(EMAIL)
                        except Exception:
                            pass
                    if PASSWORD:
                        try:
                            el_pwd.clear(); el_pwd.send_keys(PASSWORD)
                        except Exception:
                            pass

                    submit_btn = None
                    try:
                        submit_btn = driver.find_element(By.CSS_SELECTOR, "button[type='submit']")
                    except Exception:
                        pass
                    if not submit_btn:
                        try:
                            btns = driver.find_elements(By.TAG_NAME, "button")
                            for b in btns:
                                try:
                                    if not b.is_displayed(): continue
                                    t = (b.text or '').strip()
                                    if '로그인' in t or 'Login' in t or 'log in' in t.lower():
                                        submit_btn = b; break
                                except Exception:
                                    continue
                        except Exception:
                            pass
                    if submit_btn:
                        try:
                            driver.execute_script('arguments[0].click();', submit_btn)
                        except Exception:
                            try:
                                submit_btn.click()
                            except Exception:
                                pass
                        # wait for redirect / login success then ensure cookies
                        wait_for_post_login_redirect(driver, original_viewer_page=viewer_page, timeout=12)
                        ensure_cookies_synced_and_wait(driver, SESSION, timeout=12)
                        log.info('Login attempt via popup-click + form-fill completed')
                        return True
                    else:
                        log.info('Could not find submit button; falling back to login_kakao()')
                        try:
                            login_kakao(driver, EMAIL, PASSWORD)
                            ensure_cookies_synced_and_wait(driver, SESSION, timeout=12)
                        except Exception:
                            log.exception('Fallback programmatic login failed')
                        return True
                else:
                    log.info('Login inputs missing; falling back to login_kakao()')
                    try:
                        login_kakao(driver, EMAIL, PASSWORD)
                        ensure_cookies_synced_and_wait(driver, SESSION, timeout=12)
                    except Exception:
                        log.exception('Fallback programmatic login failed')
                    return True
            except Exception:
                log.exception('Error auto-filling login form; falling back to login_kakao()')
                try:
                    login_kakao(driver, EMAIL, PASSWORD)
                    ensure_cookies_synced_and_wait(driver, SESSION, timeout=12)
                except Exception:
                    log.exception('Fallback programmatic login failed')
                return True
        except Exception:
            log.info('No login form detected immediately; calling login_kakao() fallback')
            try:
                login_kakao(driver, EMAIL, PASSWORD)
                ensure_cookies_synced_and_wait(driver, SESSION, timeout=12)
            except Exception:
                log.exception('Fallback programmatic login failed')
            return True
    except Exception:
        log.exception('handle_possible_login_popup failed')
        return False

def detect_local_chrome_major_version():
    candidates = [
        "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome",
        shutil.which('google-chrome'), shutil.which('google-chrome-stable'), shutil.which('chromium-browser'), shutil.which('chromium'), shutil.which('chrome')
    ]
    for cmd in candidates:
        if not cmd: continue
        try:
            out = subprocess.check_output([cmd, '--version'], stderr=subprocess.STDOUT)
            m = re.search(r"(\d+)\.", out.decode('utf-8', errors='replace'))
            if m:
                return int(m.group(1))
        except Exception:
            continue
    return None

def get_chrome_driver():
    global TEMP_PROFILE_DIR
    if USE_CHROME_PROFILE and DEFAULT_CHROME_PROFILE and os.path.exists(DEFAULT_CHROME_PROFILE):
        try:
            log.info('Copying Chrome profile...')
            TEMP_PROFILE_DIR = tempfile.mkdtemp(prefix='kakao_profile_')
            profile_name = os.path.basename(DEFAULT_CHROME_PROFILE)
            temp_profile_path = os.path.join(TEMP_PROFILE_DIR, profile_name)
            ignore_list = shutil.ignore_patterns('RunningChromeVersion', 'Singleton*', 'Lock*', '*.lock')
            shutil.copytree(DEFAULT_CHROME_PROFILE, temp_profile_path, ignore=ignore_list, dirs_exist_ok=True)
            log.info('Profile copied to: %s', TEMP_PROFILE_DIR)
        except Exception:
            log.exception('Profile copy failed'); TEMP_PROFILE_DIR = None
    if UNDETECTED_AVAILABLE:
        try:
            options = uc.ChromeOptions()
            if TEMP_PROFILE_DIR:
                options.add_argument(f"--user-data-dir={TEMP_PROFILE_DIR}")
                options.add_argument(f"--profile-directory={os.path.basename(DEFAULT_CHROME_PROFILE)}")
            options.add_argument('--start-maximized')
            major = detect_local_chrome_major_version()
            return uc.Chrome(options=options, version_main=major if major else None, use_subprocess=True, headless=False)
        except Exception:
            log.warning('undetected-chromedriver failed, falling back')
    options = webdriver.ChromeOptions()
    if TEMP_PROFILE_DIR:
        options.add_argument(f"--user-data-dir={TEMP_PROFILE_DIR}")
        options.add_argument(f"--profile-directory={os.path.basename(DEFAULT_CHROME_PROFILE)}")
    options.add_argument('--start-maximized')
    service = Service(ChromeDriverManager().install())
    driver = webdriver.Chrome(service=service, options=options)
    try:
        driver.execute_script("Object.defineProperty(navigator, 'webdriver', {get: () => undefined})")
    except Exception:
        pass
    return driver

def login_kakao(driver, email, password):
    driver.get(
        "https://accounts.kakao.com/login/?continue=https%3A%2F%2Fkauth.kakao.com%2Foauth%2Fauthorize%3Fclient_id%3D49bbb48c5fdb0199e5da1b89de359484%26state%3Dhttps%25253A%25252F%25252Fpage.kakao.com%25252Fcontent%25252F46928984%26redirect_uri%3Dhttps%253A%252F%252Fpage.kakao.com%252Frelay%252Flogin%26response_type%3Dcode%26auth_tran_id%3Dvzszdy51upc49bbb48c5fdb0199e5da1b89de359484mhrcxwtu%26ka%3Dsdk%252F1.38.0%2520os%252Fjavascript%2520lang%252Fen-GB%2520device%252FMacIntel%2520origin%252Fhttps%25253A%25252F%25252Fpage.kakao.com%2526is_popup%3Dfalse%26through_account%3Dtrue&talk_login=hidden&login_type=simple#login"
    )
    wait_for_load(driver, timeout=12)
    try:
        el_email = WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.CSS_SELECTOR, "input[type='email'], input[type='text']")))
        el_pwd = WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.CSS_SELECTOR, "input[type='password']")))
        el_email.clear(); el_email.send_keys(email)
        el_pwd.clear(); el_pwd.send_keys(password)
        btn = driver.find_element(By.CSS_SELECTOR, "button[type='submit']")
        try:
            btn.click()
        except Exception:
            driver.execute_script('arguments[0].click();', btn)
        WebDriverWait(driver, 30).until(lambda d: 'page.kakao.com' in (d.current_url or '') or 'accounts.kakao.com' in (d.current_url or ''))
        # wait for redirect/login to complete and cookies to be set
        wait_for_post_login_redirect(driver, timeout=12)
        ensure_cookies_synced_and_wait(driver, SESSION, timeout=12)
        log.info('Login successful')
    except Exception:
        log.exception('Login failed'); raise

def extract_product_id_from_data_t_obj(data_t_obj_value):
    if not data_t_obj_value: return None
    try:
        dec = _html.unescape(data_t_obj_value)
        if dec.strip().startswith('{') and "'" in dec and '"' not in dec:
            dec = dec.replace("'", '"')
        parsed = json.loads(dec)
        evt = parsed.get('eventMeta') or {}
        nid = evt.get('id') or evt.get('productId') or evt.get('product_id')
        if nid: return int(nid)
    except Exception:
        pass
    m = re.search(r'(\d{5,})', data_t_obj_value)
    if m: return int(m.group(1))
    return None

def clean_chapter_title(title, product_id=None):
    title = (title or "").strip()
    if title:
        return title
    if product_id:
        return f"Product ID {product_id}"
    return "Unknown"

def build_empty_chapter_note(product_id, reason=None):
    reason = reason or "No readable content extracted, or Kakao requires a buy ticket/unlock for this product."
    lines = ["[EMPTY]", f"Product ID: {product_id}", f"Reason: {reason}"]
    return "\n".join(lines)

def viewer_url_for_product(series_id, product_id):
    return f"https://page.kakao.com/content/{int(series_id)}/viewer/{int(product_id)}"

LOGIN_PAGE_MARKERS = [
    "Enter Account Information",
    "Enter Password",
    "Save Login Information",
    "View help",
    "Log In",
    "Log in with QR code",
    "Sign Up",
    "Find Account",
    "Reset Password",
]

def _compact_text(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip().lower()

def clean_extracted_chapter_text(text: str) -> str:
    if not text:
        return ""
    cleaned_lines = []
    for raw_line in str(text).splitlines():
        line = _html.unescape(raw_line).replace("\xa0", " ")
        line = re.sub(r"[ \t\r\f\v]+", " ", line).strip()
        if not line:
            if cleaned_lines and cleaned_lines[-1] != "":
                cleaned_lines.append("")
            continue
        cleaned_lines.append(line)
    return "\n".join(cleaned_lines).strip()

def clean_docx_heading_text(text: str) -> str:
    cleaned = clean_extracted_chapter_text(text)
    return re.sub(r"\s+", " ", cleaned).strip() or "Unknown"

def count_login_page_markers(text: str) -> int:
    body = _compact_text(text)
    return sum(1 for marker in LOGIN_PAGE_MARKERS if _compact_text(marker) in body)

def split_docx_chapter_heading(heading_text: str):
    m = re.match(r"^Chapter\s+(\d+)\s*:\s*(.+)$", (heading_text or "").strip(), re.I)
    if not m:
        return None, (heading_text or "").strip()
    return int(m.group(1)), m.group(2).strip()

def iter_docx_chapters(docx_path: Path):
    doc = Document(str(docx_path))
    current = None
    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip()
        if not text:
            continue
        style = paragraph.style.name if paragraph.style is not None else ""
        is_heading = style.startswith("Heading") and re.match(r"^Chapter\s+\d+\s*:", text, re.I)
        if is_heading:
            if current:
                yield current
            chapter_no, chapter_name = split_docx_chapter_heading(text)
            current = {
                "heading": text,
                "chapter_no": chapter_no,
                "chapter_name": chapter_name,
                "body": [],
            }
        elif current:
            current["body"].append(text)
    if current:
        yield current

def not_populated_reason(chapter_body: str):
    body = clean_extracted_chapter_text(chapter_body)
    login_marker_count = count_login_page_markers(body)
    if login_marker_count >= 4:
        return "login_page_text", login_marker_count
    if not body:
        return "empty_body", login_marker_count
    if body.startswith("[EMPTY]") or "\n[EMPTY]" in body:
        return "empty_placeholder", login_marker_count
    return None, login_marker_count

def create_not_populated_chapters_csv(novel_folder: Path, series_id, chapter_metadata_by_no=None, chapter_catalog=None):
    chapter_metadata_by_no = chapter_metadata_by_no or {}
    chapter_catalog = chapter_catalog or []
    title_catalog = {}
    for row in chapter_catalog:
        title = (row.get("title") or "").strip()
        pid = row.get("product_id")
        if title and pid:
            title_catalog[title] = {
                "product_id": str(pid),
                "chapter_url": row.get("viewer_url") or viewer_url_for_product(series_id, pid),
            }

    rows = []
    for docx_path in sorted(novel_folder.glob("*.docx")):
        if docx_path.name.startswith("~$") or ".bak_" in docx_path.name:
            continue
        if not re.match(r"^Chapter_\d+to\d+\.docx$", docx_path.name):
            continue
        try:
            chapters = list(iter_docx_chapters(docx_path))
        except Exception as exc:
            log.warning("Could not scan %s for not-populated chapters: %s", docx_path.name, exc)
            continue

        for chapter in chapters:
            body = "\n".join(chapter.get("body") or [])
            reason, login_marker_count = not_populated_reason(body)
            if not reason:
                continue

            chapter_no = chapter.get("chapter_no")
            chapter_name = chapter.get("chapter_name") or ""
            meta = chapter_metadata_by_no.get(chapter_no, {}) if chapter_no is not None else {}
            if not meta:
                meta = title_catalog.get(chapter_name, {})
            product_id = meta.get("product_id") or ""
            chapter_url = meta.get("chapter_url") or ""
            if product_id and not chapter_url:
                chapter_url = viewer_url_for_product(series_id, product_id)

            rows.append({
                "file": docx_path.name,
                "chapter_no": chapter_no if chapter_no is not None else "",
                "chapter_name": chapter_name,
                "product_id": product_id,
                "chapter_url": chapter_url,
                "populated": "no",
                "matched_login_markers": login_marker_count,
                "reason": reason,
            })

    rows.sort(key=lambda row: int(row["chapter_no"]) if str(row["chapter_no"]).isdigit() else 999999)
    csv_path = novel_folder / "not_populated_chapters.csv"
    fieldnames = [
        "file",
        "chapter_no",
        "chapter_name",
        "product_id",
        "chapter_url",
        "populated",
        "matched_login_markers",
        "reason",
    ]
    with csv_path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)
    log.info("🧾 Not-populated CSV created: %s (%d row(s))", csv_path, len(rows))
    return csv_path, rows

def fetch_chapter_catalog_via_api(
    session,
    series_id,
    start_product_id=None,
    max_count=None,
    sort_type="asc",
    window_size=100,
    timeout=20,
):
    """
    Fetch chapter/product rows from Kakao's content-home product-list API.
    Returns rows with product_id/title/viewer_url/cursor metadata.
    """
    catalog = []
    seen = set()
    cursor_index = 0
    # Kakao's frontend uses NEXT when a sort_type is supplied.
    cursor_direction = "NEXT" if sort_type else "INIT"
    total_count = None

    for page_no in range(1, 200):
        params = {
            "series_id": int(series_id),
            "cursor_index": int(cursor_index or 0),
            "cursor_direction": cursor_direction,
            "window_size": int(window_size),
        }
        if sort_type:
            params["sort_type"] = sort_type

        headers = {
            "Accept": "application/json, text/plain, */*",
            "Referer": f"https://page.kakao.com/content/{series_id}",
            "Origin": "https://page.kakao.com",
        }
        try:
            resp = session.get(PRODUCT_LIST_API_URL, params=params, headers=headers, timeout=timeout)
            if resp.status_code != 200:
                log.warning("Product-list API returned HTTP %s", resp.status_code)
                print_api_response("content product/list", resp.url, resp=resp, truncate=1200)
                break
            data = resp.json()
        except Exception as exc:
            log.warning("Product-list API request failed: %s", exc)
            break

        result = data.get("result") if isinstance(data, dict) else None
        rows = result.get("list") if isinstance(result, dict) else None
        if not isinstance(rows, list):
            log.warning("Product-list API response did not contain result.list")
            break

        total_count = result.get("total_count", total_count)
        added = 0
        for row in rows:
            item = row.get("item") if isinstance(row, dict) else None
            if not isinstance(item, dict):
                continue
            try:
                pid = int(item.get("product_id"))
            except Exception:
                continue
            if pid in seen:
                continue
            seen.add(pid)
            catalog.append({
                "product_id": pid,
                "title": item.get("title") or f"Product ID {pid}",
                "viewer_url": viewer_url_for_product(series_id, pid),
                "cursor_index": row.get("cursor_index"),
                "order_value": item.get("order_value"),
                "slide_type": item.get("slide_type"),
                "is_text_viewer": bool((item.get("operator_property") or {}).get("is_text_viewer")),
                "is_free": bool(item.get("is_free")),
            })
            added += 1

        log.info(
            "Product-list API page %d: fetched %d row(s), total catalog=%d%s",
            page_no,
            added,
            len(catalog),
            f"/{total_count}" if total_count else "",
        )

        if start_product_id and max_count:
            ids = [row["product_id"] for row in catalog]
            if int(start_product_id) in ids:
                start_index = ids.index(int(start_product_id))
                if len(catalog[start_index:]) >= int(max_count):
                    break
        elif max_count and len(catalog) >= int(max_count):
            break

        has_next = bool(result.get("has_next"))
        if not has_next or not rows:
            break
        last_cursor = rows[-1].get("cursor_index")
        if last_cursor is None or int(last_cursor) == int(cursor_index or 0):
            break
        cursor_index = int(last_cursor)
        cursor_direction = "NEXT"

    if catalog:
        log.info("Product-list API collected %d product row(s)", len(catalog))
    return catalog

def collect_product_ids_from_series_dom(driver):
    product_ids = []
    seen = set()

    elems = driver.find_elements(By.XPATH, "//*[@data-t-obj]")
    for e in elems:
        try:
            v = e.get_attribute('data-t-obj') or ''
            if not v:
                continue
            pid = extract_product_id_from_data_t_obj(v)
            if pid and pid not in seen:
                product_ids.append(pid)
                seen.add(pid)
        except Exception:
            continue

    anchors = driver.find_elements(By.XPATH, "//a[contains(@href,'/viewer/')]")
    for a in anchors:
        try:
            href = a.get_attribute('href') or ''
            m = re.search(r'/viewer/(\d+)', href)
            if m:
                pid = int(m.group(1))
                if pid not in seen:
                    product_ids.append(pid)
                    seen.add(pid)
        except Exception:
            pass

    try:
        html = driver.page_source or ''
        for m in re.finditer(r'/viewer/(\d+)', html):
            pid = int(m.group(1))
            if pid not in seen:
                product_ids.append(pid)
                seen.add(pid)
    except Exception:
        pass

    return product_ids

def prepare_product_sequence(product_ids, start_product_id=None, max_count=None, preserve_order=False):
    sequence = []
    seen = set()
    for pid in product_ids or []:
        try:
            pid = int(pid)
        except Exception:
            continue
        if pid not in seen:
            sequence.append(pid)
            seen.add(pid)

    if len(sequence) > 1 and not preserve_order:
        increasing = sum(1 for a, b in zip(sequence, sequence[1:]) if a < b)
        decreasing = sum(1 for a, b in zip(sequence, sequence[1:]) if a > b)
        if decreasing > increasing:
            sequence.reverse()

    if start_product_id and start_product_id in sequence:
        sequence = sequence[sequence.index(start_product_id):]

    if max_count:
        sequence = sequence[:max_count]
    return sequence

def build_product_list_from_series_page(driver, series_id, expected_count=None, max_scroll_rounds=90):
    series_url = f'https://page.kakao.com/content/{series_id}'
    try:
        driver.get(series_url)
        wait_for_load(driver, timeout=8)
    except Exception:
        pass
    time.sleep(1.0 + random.random() * 1.0)

    product_ids = []
    seen = set()
    stable_rounds = 0
    last_count = 0
    last_scroll_state = None

    for round_no in range(max_scroll_rounds):
        current_ids = collect_product_ids_from_series_dom(driver)
        for pid in current_ids:
            if pid not in seen:
                product_ids.append(pid)
                seen.add(pid)

        if len(product_ids) > last_count:
            log.info("Series page product scan: found %d product id(s)", len(product_ids))
            last_count = len(product_ids)
            stable_rounds = 0
        elif expected_count and len(product_ids) >= expected_count:
            break
        else:
            stable_rounds += 1

        if expected_count and len(product_ids) >= expected_count:
            break

        try:
            scroll_state = driver.execute_script("""
                const before = {
                    y: window.scrollY || document.documentElement.scrollTop || 0,
                    h: document.documentElement.scrollHeight || document.body.scrollHeight || 0
                };
                const step = Math.max(700, Math.floor((window.innerHeight || 900) * 0.85));
                window.scrollBy(0, step);
                const scrollables = Array.from(document.querySelectorAll('div, section, main, ul'))
                    .filter(el => el.scrollHeight > el.clientHeight + 120);
                for (const el of scrollables.slice(0, 12)) {
                    el.scrollTop = Math.min(el.scrollTop + Math.max(500, Math.floor(el.clientHeight * 0.85)), el.scrollHeight);
                }
                const clickables = Array.from(document.querySelectorAll('button, a'))
                    .filter(el => {
                        const text = (el.innerText || el.textContent || '').trim().toLowerCase();
                        const rect = el.getBoundingClientRect();
                        return rect.width > 0 && rect.height > 0 && (text.includes('더보기') || text.includes('more'));
                    });
                if (clickables[0]) clickables[0].click();
                return {
                    y: window.scrollY || document.documentElement.scrollTop || 0,
                    h: document.documentElement.scrollHeight || document.body.scrollHeight || 0,
                    clickedMore: Boolean(clickables[0])
                };
            """)
            if scroll_state == last_scroll_state and stable_rounds >= 8:
                break
            last_scroll_state = scroll_state
        except Exception:
            if stable_rounds >= 8:
                break

        time.sleep(0.7 + random.random() * 0.5)

    current_ids = collect_product_ids_from_series_dom(driver)
    for pid in current_ids:
        if pid not in seen:
            product_ids.append(pid)
            seen.add(pid)

    return product_ids

# ============================================================================
# Fallback extraction of signed URLs from page
# ============================================================================

def extract_signed_urls_from_page(driver):
    """
    Returns a list of unique signed CDN URLs found in script tags or HTML.
    Looks for dn-img-page.kakao.com urls and .json?signature= tokens.
    Also logs all matches found.
    """
    js = r"""
    var out = new Set();
    try {
        // script texts
        var scripts = document.querySelectorAll('script');
        for (var i=0;i<scripts.length;i++){
            var s = scripts[i].textContent || '';
            if(!s) continue;
            var re = /https?:\/\/dn-img-page\.kakao\.com\/[^\s'"]+?(\.json)?[^\s'"]*/g;
            var m;
            while((m = re.exec(s)) !== null) { out.add(m[0]); }
            var re2 = /https?:\/\/[^'"\s]+dn-img-page\.kakao\.com[^'"\s]*/g;
            while((m = re2.exec(s)) !== null) { out.add(m[0]); }
        }
        // inline HTML
        var html = document.documentElement.innerHTML || '';
        var re3 = /https?:\/\/dn-img-page\.kakao\.com\/[^\s'"]+?(\.json)?[^\s'"]*/g;
        var mm;
        while((mm = re3.exec(html)) !== null) { out.add(mm[0]); }
        // also look for common token patterns
        var tokenRe = /download\/resource\?kid=[^'"\s<>]*/g;
        while((mm = tokenRe.exec(html)) !== null) { out.add(window.location.protocol + '//' + window.location.host + '/' + mm[0]); }
    } catch(e) {}
    return Array.from(out);
    """
    try:
        urls = driver.execute_script(js) or []
        urls = list(dict.fromkeys(urls))  # preserve order, unique
        if urls:
            log.info("extract_signed_urls_from_page: found %d signed url(s) on page:", len(urls))
            for u in urls:
                try:
                    log.info("  - %s", u)
                except Exception:
                    log.info("  - (unprintable url)")
        else:
            log.info("extract_signed_urls_from_page: no signed urls found on page")
        return list(urls)
    except Exception:
        log.exception("extract_signed_urls_from_page failed")
        return []

def extract_inline_text_from_page(driver):
    """
    Attempts to extract rendered text from common content containers (e.g. div.DC2CN).
    Returns (text, selector_used) or (None, None).
    """
    try:
        js = r"""
        try {
            function visible(el){
                if(!el) return false;
                var style = window.getComputedStyle(el);
                return style && style.visibility !== 'hidden' && style.display !== 'none' && el.offsetParent !== null;
            }
            var selectors = [
                'div.DC2CN',
                'div[class*=\"DC\"]',
                'div[class*=\"textviewer\"]',
                '[data-p-id]',
                'article',
                'div.viewer'
            ];
            for (var i=0;i<selectors.length;i++){
                var s = selectors[i];
                var els = document.querySelectorAll(s);
                for (var j=0;j<els.length;j++){
                    var el = els[j];
                    if(!visible(el)) continue;
                    // prefer elements with paragraph children
                    var ps = el.querySelectorAll('p');
                    if(ps && ps.length>0){
                        var arr = [];
                        for(var k=0;k<ps.length;k++){
                            var t = ps[k].innerText || ps[k].textContent || '';
                            t = t.trim();
                            if(t) arr.push(t);
                        }
                        if(arr.length>0) return {text: arr.join('\\n\\n'), selector: s};
                    }
                    // fallback: innerText
                    var txt = el.innerText || el.textContent || '';
                    txt = txt.trim();
                    if(txt && txt.length>100) return {text: txt, selector: s};
                }
            }
            return {text: null, selector: null};
        } catch(e) { return {text: null, selector: null}; }
        """
        res = driver.execute_script(js)
        if isinstance(res, dict) and res.get('text'):
            log.info("extract_inline_text_from_page: got inline text using selector: %s", res.get('selector'))
            return res.get('text'), res.get('selector')
        return None, None
    except Exception:
        log.exception("extract_inline_text_from_page failed")
        return None, None

def extract_text_from_html_bytes(content_bytes):
    """
    Fallback parser for HTML bytes when a resource returns HTML (not JSON).
    Extracts <p> blocks and returns joined text.
    Very conservative: strips tags with regex and unescapes HTML entities.
    """
    try:
        s = content_bytes.decode('utf-8', errors='replace')
        # find paragraph contents
        paras = re.findall(r'<p[^>]*>(.*?)</p>', s, flags=re.I|re.S)
        cleaned = []
        for p in paras:
            # remove tags inside paragraph
            t = re.sub(r'<[^>]+>', '', p)
            t = _html.unescape(t).strip()
            if t:
                cleaned.append(t)
        if cleaned:
            return '\n\n'.join(cleaned)
        # fallback: try to extract visible text from common container
        m = re.search(r'<div[^>]*class=["\']?DC2CN[^>]*>(.*?)</div>', s, flags=re.I|re.S)
        if m:
            inner = m.group(1)
            paras = re.findall(r'<p[^>]*>(.*?)</p>', inner, flags=re.I|re.S)
            cleaned = []
            for p in paras:
                t = re.sub(r'<[^>]+>', '', p)
                t = _html.unescape(t).strip()
                if t:
                    cleaned.append(t)
            if cleaned:
                return '\n\n'.join(cleaned)
        # last resort: strip tags and return large text chunk
        text = re.sub(r'<[^>]+>', '', s)
        text = _html.unescape(text).strip()
        if len(text) > 50:
            return text
    except Exception:
        log.exception("extract_text_from_html_bytes failed")
    return None

# Helper to process a signed json file content (raw bytes) and add to docx
def process_signed_json_bytes_and_save(content_bytes, docx_manager, chapter_no, meta_url='', chapter_name=None):
    try:
        parsed = json.loads(content_bytes.decode('utf-8', errors='replace'))
    except Exception as e:
        log.exception('Failed to parse signed JSON: %s', e)
        return None
    text, images = parse_textviewer_content_json(parsed, ats_base=meta_url)
    # --- NEW: log any image tokens / resolved urls discovered inside the JSON ---
    try:
        if images:
            log.info("process_signed_json: found %d embedded image token(s) in signed JSON:", len(images))
            for im in images:
                try:
                    pid = im.get('paragraph_id')
                    token = im.get('image_token')
                    resolved = im.get('resolved_url')
                    log.info("  - paragraph_id=%s token=%s resolved=%s", pid, (token[:120] + '...') if token and len(token) > 120 else token, resolved)
                except Exception:
                    log.info("  - (image entry unprintable)")
        else:
            log.info("process_signed_json: no embedded images found in signed JSON (meta_url=%s)", meta_url)
    except Exception:
        log.exception("Failed to log embedded images from signed JSON")
    if not text:
        text = extract_text_from_json(parsed) or '[No text extracted from signed JSON]'
    text = strip_leading_layout_metadata(text)
    title = clean_chapter_title(chapter_name) if chapter_name else f"Signed JSON {meta_url or 'unknown'}"
    saved = docx_manager.add_chapter(chapter_no, title, text)
    if saved:
        log.info("Saved DOCX from signed JSON: %s", saved)
    return saved

def process_signed_text_resource_urls_and_save(urls, driver, session, docx_manager, chapter_no, chapter_name=None):
    text_parts = []
    for url in urls:
        if not is_text_resource_url(url):
            log.info("Skipping non-text resource URL: %s", url)
            continue
        try:
            content_bytes, ctype = download_resource_with_fallback(driver, session, url, timeout=20)
        except Exception as e:
            log.exception("Failed to fetch text JSON resource %s: %s", url, e)
            continue
        if not content_bytes:
            continue

        looks_json = False
        if ctype and 'json' in ctype.lower():
            looks_json = True
        if '.json' in normalize_resource_url(url).lower():
            looks_json = True
        try:
            sample = content_bytes[:500].decode('utf-8', errors='ignore').strip()
            if sample.startswith('{') or sample.startswith('['):
                looks_json = True
        except Exception:
            pass
        if not looks_json:
            log.info("Skipping non-JSON resource returned from %s (content-type=%s)", url, ctype)
            continue

        try:
            parsed = json.loads(content_bytes.decode('utf-8', errors='replace'))
        except Exception as e:
            log.exception("Failed to parse text JSON resource %s: %s", url, e)
            continue

        text, _images = parse_textviewer_content_json(parsed, ats_base='')
        if not text:
            text = extract_text_from_json(parsed)
        text = strip_leading_layout_metadata(text)
        if _looks_like_useful_chapter_text(text):
            text_parts.append(text.strip())
        else:
            log.info("No readable chapter text found in JSON resource: %s", url)

    if not text_parts:
        return None

    title = chapter_name or f"Chapter {chapter_no}"
    full_text = '\n\n'.join(text_parts)
    saved = docx_manager.add_chapter(chapter_no, title, full_text)
    if saved:
        log.info("Saved DOCX after signed text JSON resources: %s", saved)
    return saved

# Main processing of viewerData (GraphQL path)
def process_viewer_contents_and_ocr(viewerData, driver, session, docx_manager, chapter_no, chapter_name=None):
    if not viewerData:
        return None
    contents = viewerData.get('contentsList') or []
    if not contents:
        return None
    ats_base = viewerData.get('atsServerUrl') or ''
    meta_secure = viewerData.get('metaSecureUrl')
    chapter_texts = []
    for idx, entry in enumerate(contents, start=1):
        chapter_id = entry.get('chapterId') or f'item_{idx}'
        content_id = entry.get('contentId') or ''
        secure_url = entry.get('secureUrl') or ''
        text_parts = [f'[Chapter ID: {chapter_id}, Content ID: {content_id}]']
        if not secure_url:
            if meta_secure:
                secure_url = meta_secure
            else:
                text_parts.append('[No secureUrl found]'); continue
        try:
            if secure_url.startswith('http://') or secure_url.startswith('https://'):
                final_url = secure_url
            else:
                if ats_base:
                    if ats_base.endswith('=') or '?' in ats_base:
                        final_url = ats_base + secure_url
                    else:
                        final_url = urljoin(ats_base, secure_url)
                else:
                    final_url = secure_url
        except Exception:
            final_url = secure_url
        try:
            resource_bytes, content_type = download_resource_with_fallback(driver, session, final_url)
        except Exception as e:
            text_parts.append(f'[Download failed: {e}]'); chapter_texts.append('\n'.join(text_parts)); continue
        is_json = False
        if content_type:
            try:
                ct_lower = content_type.lower()
                if 'json' in ct_lower: is_json = True
                if 'html' in ct_lower: is_html = True
                else: is_html = False
            except Exception:
                is_json = False; is_html = False
        else:
            is_html = False
        if not is_json:
            if final_url.lower().endswith('.json') or 'textviewer' in final_url.lower():
                is_json = True
        # detect HTML
        if not is_json and not is_html:
            try:
                sample = resource_bytes[:1000].decode('utf-8', errors='ignore').strip()
                if (sample.startswith('{') or sample.startswith('[')) and ('paragraphList' in sample or 'contentInfo' in sample or 'KAGE' in sample):
                    is_json = True
                elif sample.startswith('<'):
                    is_html = True
            except Exception:
                pass
        if is_json:
            try:
                try:
                    parsed = json.loads(resource_bytes.decode('utf-8'))
                except Exception:
                    parsed = json.loads(resource_bytes.decode('utf-8', errors='replace'))
                text, image_list = parse_textviewer_content_json(parsed, ats_base=ats_base)
                text = strip_leading_layout_metadata(text)
                if text:
                    text_parts.append(text)
                else:
                    text_parts.append('[JSON content - no readable text extracted]')
            except Exception as e:
                text_parts.append(f'[JSON parse/extract failed: {e}]')
            chapter_texts.append('\n\n'.join(text_parts)); continue
        if is_html:
            # parse html paragraphs
            html_text = extract_text_from_html_bytes(resource_bytes)
            if html_text:
                title = clean_chapter_title(chapter_name, viewerData.get('productId')) if chapter_name else f"HTML content {final_url}"
                saved = docx_manager.add_chapter(chapter_no, title, html_text)
                if saved:
                    log.info(f"📄 Created DOCX from inline HTML: {Path(saved).name}")
                chapter_texts.append(f'[Saved HTML content from {final_url}]')
                continue
        if not ALLOW_IMAGE_OCR_FALLBACK:
            text_parts.append(f'[Skipped non-text resource: {final_url}]')
            chapter_texts.append('\n'.join(text_parts))
            continue
        # binary -> image -> OCR
        try:
            nparr = np.frombuffer(resource_bytes, np.uint8)
            img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
            if img is None:
                raise ValueError('cv2.imdecode returned None')
        except Exception as e:
            text_parts.append(f'[Image decode failed: {e}]'); chapter_texts.append('\n'.join(text_parts)); continue
        try:
            ocr_reader = get_ocr_reader()
            ocr_result_lines = ocr_reader.readtext(img, detail=0, paragraph=True)
            if isinstance(ocr_result_lines, str):
                ocr_text = ocr_result_lines.strip()
            else:
                ocr_text = '\n'.join([ln.strip() for ln in ocr_result_lines if ln and ln.strip()])
            if not ocr_text:
                ocr_text = '[No text detected]'
            text_parts.append(ocr_text)
        except Exception as e:
            text_parts.append(f'[OCR failed: {e}]')
        chapter_texts.append('\n\n'.join(text_parts))
        time.sleep(random.uniform(0.25, 0.8))
    full_content = '\n\n---\n\n'.join(chapter_texts)
    title = clean_chapter_title(chapter_name, viewerData.get('productId'))
    saved_path = docx_manager.add_chapter(chapter_no, title, full_content)
    if saved_path: log.info(f"📄 Created new DOCX file: {Path(saved_path).name}")
    return saved_path

# ============================================================================
# MAIN RUN
# ============================================================================
if __name__ == '__main__':
    log.info('🚀 Starting Kakao Scraper (Fixed with cookie wait & GET-first)')

    try:
        folder_name = input(f"📁 Output folder name [{DEFAULT_OUTPUT_FOLDER}]: ").strip()
    except Exception:
        folder_name = DEFAULT_OUTPUT_FOLDER
    if not folder_name:
        folder_name = DEFAULT_OUTPUT_FOLDER

    folder_manager = FolderManager(folder_name)
    novel_folder = folder_manager.get_novel_folder(NOVEL_TITLE)
    docx_manager = DocxManager(novel_folder, NOVEL_TITLE, "Unknown Author")

    driver = None
    try:
        driver = get_chrome_driver()
        try:
            driver.get('https://page.kakao.com/main')
            wait_for_load(driver, timeout=8)
            try:
                driver.find_element(By.CSS_SELECTOR, "[data-tiara-layer='gnb_profile']")
                sync_cookies_to_requests(driver, SESSION)
            except Exception:
                log.info('Not logged in, attempting programmatic login')
                try:
                    login_kakao(driver, EMAIL, PASSWORD)
                    # login_kakao already calls ensure_cookies_synced_and_wait
                except Exception:
                    log.exception('Initial programmatic login failed')
        except Exception:
            log.warning('Login check failed; continuing')

        graphql_url = 'https://bff-page.kakao.com/api/gateway/api/v1/viewer/data'
        chapter_no = START_CHAPTER
        product_id = START_PRODUCT_ID
        visited = 0
        product_list = []
        product_title_by_id = {}
        chapter_catalog = []
        chapter_metadata_by_no = {}
        product_cursor = 0
        product_start_offset = max(int(START_CHAPTER or 1) - 1, 0)
        try:
            catalog_max_count = MAX_CHAPTERS
            if START_PRODUCT_ID is None and product_start_offset:
                catalog_max_count = product_start_offset + MAX_CHAPTERS + 25
            elif START_PRODUCT_ID is None and not ALLOW_IMAGE_OCR_FALLBACK:
                catalog_max_count += 25
            chapter_catalog = fetch_chapter_catalog_via_api(
                SESSION,
                SERIES_ID,
                start_product_id=START_PRODUCT_ID,
                max_count=catalog_max_count,
            )
            if not ALLOW_IMAGE_OCR_FALLBACK:
                original_catalog_len = len(chapter_catalog)
                chapter_catalog = [row for row in chapter_catalog if row.get("is_text_viewer")]
                skipped_non_text = original_catalog_len - len(chapter_catalog)
                if skipped_non_text:
                    log.info(
                        'Skipping %d non-text product(s) from product-list API because image OCR fallback is disabled.',
                        skipped_non_text
                    )
            product_title_by_id = {
                int(row["product_id"]): row.get("title") or ""
                for row in chapter_catalog
                if row.get("product_id")
            }
            raw_product_list = [int(row["product_id"]) for row in chapter_catalog if row.get("product_id")]
            ordered_product_list = prepare_product_sequence(raw_product_list, START_PRODUCT_ID, None, preserve_order=True)
            if START_PRODUCT_ID is None and product_start_offset:
                if product_start_offset >= len(ordered_product_list):
                    log.warning(
                        "START_CHAPTER=%d requires offset %d, but product-list API only returned %d item(s).",
                        START_CHAPTER,
                        product_start_offset,
                        len(ordered_product_list),
                    )
                    product_list = []
                else:
                    product_list = ordered_product_list[product_start_offset:product_start_offset + MAX_CHAPTERS]
                    log.info(
                        "Using product-list API offset for START_CHAPTER=%d: skipped %d product row(s).",
                        START_CHAPTER,
                        product_start_offset,
                    )
            else:
                product_list = ordered_product_list[:MAX_CHAPTERS]
            if product_list:
                product_id = int(product_list[0])
                log.info(
                    'Built ordered product list from product-list API (%d item(s)); first=%s last=%s',
                    len(product_list),
                    product_list[0],
                    product_list[-1]
                )
                first_title = product_title_by_id.get(product_list[0])
                if first_title:
                    log.info('First API chapter title: %s', first_title)
                if len(product_list) < MAX_CHAPTERS:
                    log.warning(
                        'Product list has only %d item(s), less than MAX_CHAPTERS=%d. Will process the collected list only.',
                        len(product_list),
                        MAX_CHAPTERS
                    )
            else:
                log.warning('Product-list API did not produce product IDs; trying series-page DOM scan.')
        except Exception:
            log.exception('Failed to build product list from product-list API')
            product_list = []

        if not product_list:
            try:
                dom_expected_count = product_start_offset + MAX_CHAPTERS if START_PRODUCT_ID is None else MAX_CHAPTERS
                raw_product_list = build_product_list_from_series_page(driver, SERIES_ID, expected_count=dom_expected_count)
                ordered_product_list = prepare_product_sequence(raw_product_list, START_PRODUCT_ID, None)
                if START_PRODUCT_ID is None and product_start_offset:
                    product_list = ordered_product_list[product_start_offset:product_start_offset + MAX_CHAPTERS]
                else:
                    product_list = ordered_product_list[:MAX_CHAPTERS]
                if product_list:
                    product_id = int(product_list[0])
                    log.info(
                        'Built ordered product list from series page (%d item(s)); first=%s last=%s',
                        len(product_list),
                        product_list[0],
                        product_list[-1]
                    )
                    if len(product_list) < MAX_CHAPTERS:
                        log.warning(
                            'Product list has only %d item(s), less than MAX_CHAPTERS=%d. Will process the collected list only.',
                            len(product_list),
                            MAX_CHAPTERS
                        )
                else:
                    log.warning('Could not build product list from API or series page; will use numeric increment fallback.')
            except Exception:
                log.exception('Failed to build product list from series page')
                product_list = []

        if not product_list and not product_id:
            log.error('No START_PRODUCT_ID and no product list could be built; stopping.')
            raise SystemExit(1)

        seen_products = set()
        while visited < MAX_CHAPTERS:
            try:
                chapter_title = product_title_by_id.get(int(product_id)) if product_id else None
                docx_chapter_title = clean_chapter_title(chapter_title, product_id)
                chapters_added_before = docx_manager.total_chapters_added
                empty_reason = "No readable content extracted, or Kakao requires a buy ticket/unlock for this product."
                if chapter_title:
                    log.info('📖 Processing chapter %s -> product_id=%s - %s', chapter_no, product_id, chapter_title)
                else:
                    log.info('📖 Processing chapter %s -> product_id=%s', chapter_no, product_id)
                if product_id in seen_products:
                    log.warning('Product %s already seen; stopping to avoid loop', product_id)
                    break
                seen_products.add(product_id)

                viewer_page = f'https://page.kakao.com/content/{SERIES_ID}/viewer/{product_id}'
                chapter_metadata_by_no[int(chapter_no)] = {
                    "chapter_name": docx_chapter_title,
                    "product_id": int(product_id) if product_id else "",
                    "chapter_url": viewer_page,
                }
                try:
                    driver.get(viewer_page)
                    wait_for_load(driver, timeout=10)
                    wait_for_viewer_ready(driver, timeout=12)
                except Exception:
                    pass

                # Try to handle a popup if present (this may call login flows)
                handle_possible_login_popup(driver, viewer_page=viewer_page)

                try:
                    if is_logged_in(driver):
                        # ensure cookies synced before GraphQL/GET
                        ensure_cookies_synced_and_wait(driver, SESSION, timeout=12)
                    else:
                        # still attempt to sync whatever is present
                        sync_cookies_to_requests(driver, SESSION)
                except Exception:
                    pass

                # --- NEW: Try GET endpoint first, then fallback to GraphQL POST ---
                g = None
                viewerData = None
                got_from_graphql = False
                api_text_urls = []
                api_not_purchased = False
                try:
                    try:
                        # ensure cookies synced from Selenium first (important)
                        ensure_cookies_synced_and_wait(driver, SESSION, timeout=8)
                    except Exception:
                        sync_cookies_to_requests(driver, SESSION)

                    # 1) Try the new GET endpoint
                    g_get = fetch_viewer_data_via_get(SESSION, SERIES_ID, product_id, referer=driver.current_url)
                    if isinstance(g_get, dict):
                        api_not_purchased = response_says_not_purchased(g_get)
                        api_text_urls = extract_text_resource_urls_from_api_response(g_get)
                        if api_not_purchased and api_text_urls:
                            log.info(
                                "GET viewer/data says product_id=%s is not_purchased, but text JSON resources are present; treating it as readable/free.",
                                product_id
                            )
                        elif api_not_purchased:
                            log.warning(
                                "GET viewer/data says product_id=%s is not_purchased; not attempting gated resources.",
                                product_id
                            )
                        elif api_text_urls:
                            log.info("GET viewer/data exposed %d signed text JSON resource(s)", len(api_text_urls))
                        # a) direct viewerData at top-level
                        if g_get.get('viewerData'):
                            viewerData = g_get.get('viewerData')
                            got_from_graphql = True
                        # b) nested under data.viewerInfo.viewerData
                        elif g_get.get('data') and g_get['data'].get('viewerInfo') and g_get['data']['viewerInfo'].get('viewerData'):
                            viewerData = g_get['data']['viewerInfo']['viewerData']
                            got_from_graphql = True
                        # c) maybe under result or payload
                        elif g_get.get('result') and isinstance(g_get['result'], dict) and g_get['result'].get('viewerData'):
                            viewerData = g_get['result']['viewerData']
                            got_from_graphql = True
                        else:
                            # try some heuristics: sometimes the GET returns viewerInfo
                            if g_get.get('viewerInfo') and isinstance(g_get.get('viewerInfo'), dict) and g_get['viewerInfo'].get('viewerData'):
                                viewerData = g_get['viewerInfo']['viewerData']
                                got_from_graphql = True
                    g = g_get
                    # 3) Fallback: if GET failed or returned 403 or didn't have viewerData, try the existing GraphQL POST
                    if not got_from_graphql and not api_text_urls and not api_not_purchased:
                        log.info("GET endpoint didn't return viewerData — falling back to GraphQL POST")
                        payload = {
                            'query': "\n    query viewerInfo($seriesId: Long!, $productId: Long!) {\n  viewerInfo(seriesId: $seriesId, productId: $productId) {\n    item { id productId seriesId title }\n    nextItem { productId }\n    viewerData { ...TextViewerData }\n  }\n}\n    \n    fragment TextViewerData on TextViewerData {\n  type\n  atsServerUrl\n  metaSecureUrl\n  contentsList { chapterId contentId secureUrl }\n}\n",
                            'variables': {'seriesId': SERIES_ID, 'productId': product_id}
                        }
                        try:
                            SESSION.headers.update({'Referer': driver.current_url})
                        except Exception:
                            pass
                        g_post = fetch_graphql(SESSION, graphql_url, payload, referer=driver.current_url)
                        if isinstance(g_post, dict) and g_post.get('data') and g_post['data'].get('viewerInfo'):
                            vi = g_post['data']['viewerInfo']
                            if not chapter_title:
                                docx_chapter_title = clean_chapter_title((vi.get('item') or {}).get('title'), product_id)
                                if int(chapter_no) in chapter_metadata_by_no:
                                    chapter_metadata_by_no[int(chapter_no)]["chapter_name"] = docx_chapter_title
                            viewerData = vi.get('viewerData')
                            got_from_graphql = True
                        else:
                            g = g_post  # keep the last response in g variable for debug/logging
                except Exception:
                    log.exception("Error while attempting GET/GraphQL fetch")
                    viewerData = None

                # --- Log GraphQL-derived resolved URLs and continue if available ---
                if got_from_graphql and viewerData:
                    try:
                        ats_base = viewerData.get('atsServerUrl') or ''
                        meta_secure = viewerData.get('metaSecureUrl') or ''
                        pid = viewerData.get('productId') or (g.get('data', {}) .get('viewerInfo', {}) .get('item', {}).get('productId', product_id) if isinstance(g, dict) else product_id)
                        log.info("viewerData present for product_id=%s; atsServerUrl=%s metaSecureUrl=%s", pid, bool(ats_base), bool(meta_secure))
                        if meta_secure:
                            resolved_meta = resolve_token_url(meta_secure, ats_base)
                            log.info("  metaSecureUrl resolved -> %s", resolved_meta)
                        contents = viewerData.get('contentsList') or []
                        if not contents:
                            log.info("  viewerData.contentsList is empty")
                        for idx, entry in enumerate(contents, start=1):
                            sec = entry.get('secureUrl') or ''
                            content_id = entry.get('contentId') or ''
                            chapter_id = entry.get('chapterId') or ''
                            resolved = resolve_token_url(sec, ats_base)
                            log.info("  content[%d] chapterId=%s contentId=%s secureUrl=%s -> resolved=%s",
                                     idx, chapter_id, content_id, (sec[:200] if sec else ''), resolved)
                    except Exception:
                        log.exception("Failed to enumerate/print viewerData URLs")
                    process_viewer_contents_and_ocr(
                        viewerData,
                        driver,
                        SESSION,
                        docx_manager,
                        chapter_no,
                        chapter_name=docx_chapter_title,
                    )
                elif api_text_urls:
                    log.info("Processing signed text JSON resource(s) from GET viewer/data for product_id=%s", product_id)
                    saved = process_signed_text_resource_urls_and_save(
                        api_text_urls,
                        driver,
                        SESSION,
                        docx_manager,
                        chapter_no,
                        chapter_name=docx_chapter_title,
                    )
                    if not saved:
                        empty_reason = "Signed text JSON resources were present but did not yield readable text."
                        log.warning("Signed text JSON resources did not yield readable text for product_id=%s", product_id)
                elif api_not_purchased:
                    empty_reason = "Kakao marked this product as not_purchased / buy ticket required."
                    log.warning(
                        "Skipping product_id=%s because Kakao marks it not_purchased. Unlock/purchase it in the browser, then rerun.",
                        product_id
                    )
                else:
                    log.warning('viewerData missing or gated for product_id=%s — attempting page fallback', product_id)
                    # fallback: try to extract signed urls from the loaded page HTML
                    raw_signed_urls = extract_signed_urls_from_page(driver)
                    signed_urls = [u for u in raw_signed_urls if is_text_resource_url(u)]
                    skipped_urls = len(raw_signed_urls) - len(signed_urls)
                    if skipped_urls:
                        log.info("Ignored %d non-text signed URL(s), such as thumbnails", skipped_urls)
                    if signed_urls:
                        log.info('Found %d signed urls on page; attempting to fetch them', len(signed_urls))
                        fetched_any = False
                        for su in signed_urls:
                            try:
                                content_bytes, ctype = download_resource_with_fallback(driver, SESSION, su, timeout=20)
                                if content_bytes:
                                    # if JSON, parse it and save
                                    if (ctype and 'json' in ctype.lower()) or su.lower().endswith('.json') or (content_bytes[:2] in (b'{', b'[')):
                                        fetched_any = True
                                        process_signed_json_bytes_and_save(
                                            content_bytes,
                                            docx_manager,
                                            chapter_no,
                                            meta_url=su,
                                            chapter_name=docx_chapter_title,
                                        )
                                        break
                                    # if HTML (page-like), try to extract paragraphs
                                    try:
                                        sample = content_bytes[:1000].decode('utf-8', errors='ignore').strip()
                                        if sample.startswith('<'):
                                            html_text = extract_text_from_html_bytes(content_bytes)
                                            if html_text:
                                                saved = docx_manager.add_chapter(chapter_no, docx_chapter_title, html_text)
                                                if saved: log.info("Saved DOCX from HTML resource: %s", Path(saved).name)
                                                fetched_any = True
                                                break
                                    except Exception:
                                        pass
                                    else:
                                        if not ALLOW_IMAGE_OCR_FALLBACK:
                                            log.info("Skipping non-text signed resource because image OCR fallback is disabled: %s", su)
                                            continue
                                        # if image, OCR it directly
                                        try:
                                            nparr = np.frombuffer(content_bytes, np.uint8)
                                            img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
                                            if img is not None:
                                                ocr_reader = get_ocr_reader()
                                                ocr_lines = ocr_reader.readtext(img, detail=0, paragraph=True)
                                                ocr_text = ocr_lines if isinstance(ocr_lines, str) else '\n'.join([ln.strip() for ln in ocr_lines if ln and ln.strip()])
                                                docx_manager.add_chapter(chapter_no, docx_chapter_title, ocr_text or '[No text]')
                                                fetched_any = True
                                                break
                                        except Exception:
                                            pass
                            except Exception as e:
                                log.exception('Failed to fetch signed url %s: %s', su, e)
                        if not fetched_any:
                            log.warning('Fallback signed-URL approach failed for product_id=%s', product_id)
                            inline_text, selector = extract_inline_text_from_page(driver)
                            if inline_text:
                                saved = docx_manager.add_chapter(chapter_no, docx_chapter_title, inline_text)
                                if saved:
                                    log.info("Saved DOCX from inline page content: %s", Path(saved).name)
                            else:
                                empty_reason = "Fallback signed URLs failed and no inline page text was found."
                                log.warning("No inline page text found for product_id=%s", product_id)
                    else:
                        log.info('No signed CDN URLs found on page for product_id=%s; trying to extract inline HTML text', product_id)
                        inline_text, selector = extract_inline_text_from_page(driver)
                        if inline_text:
                            saved = docx_manager.add_chapter(chapter_no, docx_chapter_title, inline_text)
                            if saved:
                                log.info("Saved DOCX from inline page content: %s", Path(saved).name)
                        else:
                            empty_reason = "No signed CDN URLs or inline HTML text were found."
                            log.warning('No inline HTML content found on page for product_id=%s', product_id)

                if docx_manager.total_chapters_added == chapters_added_before:
                    saved = docx_manager.add_chapter(
                        chapter_no,
                        docx_chapter_title,
                        build_empty_chapter_note(product_id, empty_reason),
                    )
                    if saved:
                        log.info("Saved DOCX after empty placeholder chapter: %s", Path(saved).name)
                    log.warning(
                        "Added EMPTY DOCX chapter for product_id=%s with title: %s",
                        product_id,
                        docx_chapter_title,
                    )

                visited += 1
                chapter_no += 1
                if product_list:
                    product_cursor += 1
                    if product_cursor >= len(product_list):
                        log.info('Reached end of prebuilt product list (%d item(s)); stopping.', len(product_list))
                        break
                    product_id = int(product_list[product_cursor])
                else:
                    log.info('No prebuilt product list available — falling back to numeric increment.')
                    product_id = int(product_id) + 1
                time.sleep(random.uniform(0.8, 1.6))
            except KeyboardInterrupt:
                log.info('Interrupted by user'); break
            except Exception:
                log.exception('Error processing chapter %s', chapter_no)
                try:
                    fallback_title = locals().get('docx_chapter_title') or clean_chapter_title(None, product_id)
                    chapters_added_before_error = locals().get('chapters_added_before')
                    if (
                        chapters_added_before_error is not None
                        and docx_manager.total_chapters_added == chapters_added_before_error
                    ):
                        saved = docx_manager.add_chapter(
                            chapter_no,
                            fallback_title,
                            build_empty_chapter_note(product_id, "Error while processing this product; continuing to next chapter."),
                        )
                        if saved:
                            log.info("Saved DOCX after error placeholder chapter: %s", Path(saved).name)
                    visited += 1
                    chapter_no += 1
                    if product_list:
                        product_cursor += 1
                        if product_cursor >= len(product_list):
                            break
                        product_id = int(product_list[product_cursor])
                    else:
                        product_id = product_id + 1
                except Exception:
                    break
                continue

        final_path = docx_manager.finalize()
        if final_path:
            log.info(f"📄 Final DOCX saved: {Path(final_path).name}")
        not_populated_csv, not_populated_rows = create_not_populated_chapters_csv(
            novel_folder,
            SERIES_ID,
            chapter_metadata_by_no=chapter_metadata_by_no,
            chapter_catalog=chapter_catalog,
        )

        print(f"\n{'='*60}")
        print(f"✅ SCRAPING COMPLETE!")
        print(f"{'='*60}")
        print(f"📂 Location: {novel_folder}")
        print(f"📚 Files created in format: Chapter_1to100.docx, etc.")
        print(f"🧾 Not populated CSV: {not_populated_csv} ({len(not_populated_rows)} row(s))")
        print(f"{'='*60}\n")

    finally:
        try:
            if driver and KEEP_BROWSER_OPEN:
                log.info('Leaving browser open for inspection')
            else:
                if driver:
                    driver.quit()
        except Exception:
            pass
        try:
            if TEMP_PROFILE_DIR and os.path.exists(TEMP_PROFILE_DIR):
                shutil.rmtree(TEMP_PROFILE_DIR)
        except Exception:
            pass
        log.info('Finished')
