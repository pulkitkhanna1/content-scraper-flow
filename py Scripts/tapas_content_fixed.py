import argparse
import logging
import os
import time
from typing import List

from docx import Document
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from webdriver_manager.chrome import ChromeDriverManager


# -------------------- Logging --------------------

def setup_logger():
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s | %(levelname)s | %(message)s",
        handlers=[logging.StreamHandler()]
    )


# -------------------- Driver --------------------

def create_driver(headless: bool):
    options = Options()
    options.add_argument("--disable-blink-features=AutomationControlled")
    options.add_argument("--start-maximized")
    options.add_argument("--lang=en-US")

    if headless:
        options.add_argument("--headless=new")

    service = Service(ChromeDriverManager().install())
    return webdriver.Chrome(service=service, options=options)


# -------------------- Helpers --------------------

def wait_for_user_login():
    input("\n✅ Log in to Tapas in the opened browser.\n"
          "👉 Once logged in and the page is fully loaded, press ENTER to continue...\n")


def safe_find(driver, by, value, timeout=15):
    return WebDriverWait(driver, timeout).until(
        EC.presence_of_element_located((by, value))
    )


def navigate_to_series_page(driver):
    """If on an episode page, find and navigate to the series info page."""
    current_url = driver.current_url
    if "/episode/" not in current_url:
        return

    logging.info("📖 Detected episode page — finding series info page...")

    # Tapas episode pages have breadcrumb/header links back to the series
    series_links = driver.find_elements(By.CSS_SELECTOR, 'a[href*="/series/"]')
    for link in series_links:
        href = link.get_attribute("href") or ""
        if "/series/" in href and "/episode/" not in href:
            # Normalise to the /info variant
            slug_part = href.split("/series/")[1].split("/")[0]
            base = href.split("/series/")[0]
            info_url = f"{base}/series/{slug_part}/info"
            logging.info(f"🔗 Navigating to series page: {info_url}")
            driver.get(info_url)
            time.sleep(3)
            return

    raise RuntimeError(
        "Could not find series link on this episode page. "
        "Pass the series info page URL (e.g. https://tapas.io/series/SLUG/info) as --start-url."
    )


# Candidate selectors for the episode list, tried in order
_EPISODE_LINK_SELECTORS = [
    'ul.episode-list.js-episode-list li a[href^="/episode/"]',
    'ul.episode-list li a[href^="/episode/"]',
    '.js-episode-list li a[href^="/episode/"]',
    'li.episode__item a[href^="/episode/"]',
    'a.episode__title[href^="/episode/"]',
    'a[href^="/episode/"]',
]


def _find_episode_links(driver):
    for selector in _EPISODE_LINK_SELECTORS:
        items = driver.find_elements(By.CSS_SELECTOR, selector)
        if items:
            return items
    return []


def scroll_until_chapters_loaded(driver, max_chapters: int) -> List[str]:
    logging.info("🔄 Scrolling to load chapters...")

    seen_hrefs = set()
    chapter_hrefs = []
    last_count = 0
    retries = 0

    while len(chapter_hrefs) < max_chapters and retries < 10:
        items = _find_episode_links(driver)

        for a in items:
            href = a.get_attribute("href")
            if href and href not in seen_hrefs:
                seen_hrefs.add(href)
                chapter_hrefs.append(href)

        logging.info(f"📚 Chapters discovered: {len(chapter_hrefs)}")

        if len(chapter_hrefs) == last_count:
            retries += 1
        else:
            retries = 0

        last_count = len(chapter_hrefs)

        driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
        time.sleep(2)

    if len(chapter_hrefs) < max_chapters:
        logging.warning(f"⚠️ Only {len(chapter_hrefs)} chapters found")

    return chapter_hrefs[:max_chapters]


def extract_chapter_content(driver, chapter_href: str, retries=3) -> str:
    for attempt in range(1, retries + 1):
        try:
            logging.info(f"➡️ Navigating to chapter: {chapter_href}")
            driver.get(chapter_href)

            viewer = safe_find(driver, By.CLASS_NAME, "viewer", timeout=20)
            text = viewer.text.strip()

            if not text:
                raise ValueError("Empty chapter content")

            return text

        except Exception as e:
            logging.warning(f"❌ Failed attempt {attempt} for {chapter_href}: {e}")
            time.sleep(3)

    logging.error(f"🚫 Skipping chapter after {retries} retries: {chapter_href}")
    return ""


# -------------------- Main --------------------

def scrape(args):
    os.makedirs(args.out, exist_ok=True)
    doc = Document()
    doc.add_heading(args.book_name, level=0)

    driver = create_driver(args.headless)
    driver.get(args.start_url)

    wait_for_user_login()

    navigate_to_series_page(driver)

    chapter_hrefs = scroll_until_chapters_loaded(driver, args.max_chapters)

    start_idx = args.start_chapter - 1  # convert to 0-based
    chapter_hrefs = chapter_hrefs[start_idx:]

    logging.info(f"✅ Total chapters to scrape: {len(chapter_hrefs)} (starting from chapter {args.start_chapter})")

    for idx, chapter_href in enumerate(chapter_hrefs, start=args.start_chapter):
        content = extract_chapter_content(driver, chapter_href)

        if not content:
            continue

        doc.add_heading(f"Chapter {idx}", level=1)
        for para in content.split("\n"):
            if para.strip():
                doc.add_paragraph(para.strip())

        logging.info(f"✍️ Saved Chapter {idx}")

    out_file = os.path.join(
        args.out,
        f"{args.book_name.replace(' ', '_')}.docx"
    )
    doc.save(out_file)
    logging.info(f"📄 DOCX saved at: {out_file}")

    driver.quit()


# -------------------- CLI --------------------

if __name__ == "__main__":
    setup_logger()

    parser = argparse.ArgumentParser(description="Tapas Novel Scraper")
    parser.add_argument("--start-url", required=True, help="Series info page URL")
    parser.add_argument("--book-name", required=True, help="Book name for DOCX")
    parser.add_argument("--max-chapters", type=int, default=150)
    parser.add_argument("--start-chapter", type=int, default=1, help="First chapter to scrape (1-based)")
    parser.add_argument("--out", default="./output")
    parser.add_argument("--headless", action="store_true")

    args = parser.parse_args()
    scrape(args)
