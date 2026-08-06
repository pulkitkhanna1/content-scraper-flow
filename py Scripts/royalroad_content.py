import os
import time
import logging
import argparse
from docx import Document
import undetected_chromedriver as uc
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException, WebDriverException


# ================= CONFIG =================
MAX_CHAPTERS_PER_DOC = 100
RETRY_COUNT = 3
RETRY_DELAY = 3  # seconds
WAIT_TIME = 15
# =========================================


def setup_logger():
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s [%(levelname)s] %(message)s",
        handlers=[
            logging.FileHandler("royalroad_scraper.log"),
            logging.StreamHandler()
        ]
    )


def setup_driver():
    options = uc.ChromeOptions()
    options.add_argument("--start-maximized")
    # Explicitly specify Chrome version 144 to match installed browser
    # This forces undetected_chromedriver to download/use the correct ChromeDriver
    try:
        driver = uc.Chrome(options=options, version_main=144, use_subprocess=True)
        logging.info("ChromeDriver initialized successfully with version 144")
        return driver
    except Exception as e:
        logging.warning(f"Initialization with version_main=144 failed: {e}")
        logging.info("Attempting without version specification...")
        try:
            # Try with use_subprocess but let it auto-detect
            driver = uc.Chrome(options=options, use_subprocess=True)
            logging.info("ChromeDriver initialized with auto-detection")
            return driver
        except Exception as e2:
            logging.warning(f"Auto-detection failed: {e2}")
            logging.info("Attempting basic initialization...")
            # Last resort: basic initialization
            return uc.Chrome(options=options)


def safe_get_text(driver, by, value):
    return WebDriverWait(driver, WAIT_TIME).until(
        EC.presence_of_element_located((by, value))
    ).text.strip()


def scrape_chapter(driver):
    title = safe_get_text(driver, By.TAG_NAME, "h1")

    content_div = WebDriverWait(driver, WAIT_TIME).until(
        EC.presence_of_element_located(
            (By.CSS_SELECTOR, "div.chapter-inner.chapter-content")
        )
    )

    paragraphs = content_div.find_elements(By.TAG_NAME, "p")
    content = "\n\n".join(p.text for p in paragraphs if p.text.strip())

    return title, content


def get_next_chapter_url(driver):
    try:
        next_btn = driver.find_element(
            By.XPATH,
            "//a[contains(@class,'btn-primary') and contains(.,'Next')]"
        )
        return next_btn.get_attribute("href")
    except WebDriverException:
        return None


def save_doc(doc, directory, book_title, part):
    filename = f"{book_title}_part_{part}.docx"
    path = os.path.join(directory, filename)
    doc.save(path)
    logging.info(f"Saved: {path}")


def scrape_book(start_url, book_title):
    directory = book_title.replace(" ", "_")
    os.makedirs(directory, exist_ok=True)

    driver = setup_driver()
    driver.get(start_url)

    chapter_count = 0
    file_part = 1
    doc = Document()

    try:
        while True:
            for attempt in range(RETRY_COUNT):
                try:
                    title, content = scrape_chapter(driver)
                    break
                except Exception as e:
                    logging.warning(
                        f"Retry {attempt + 1}/{RETRY_COUNT} failed: {e}"
                    )
                    time.sleep(RETRY_DELAY)
            else:
                logging.error("Skipping chapter after retries")
                break

            doc.add_heading(title, level=1)
            doc.add_paragraph(content)

            chapter_count += 1
            logging.info(f"Scraped chapter {chapter_count}: {title}")

            if chapter_count % MAX_CHAPTERS_PER_DOC == 0:
                save_doc(doc, directory, book_title, file_part)
                file_part += 1
                doc = Document()

            next_url = get_next_chapter_url(driver)
            if not next_url:
                logging.info("No next chapter found. Stopping.")
                break

            driver.get(next_url)
            time.sleep(1)

        if len(doc.paragraphs) > 0:
            save_doc(doc, directory, book_title, file_part)

    finally:
        driver.quit()


if __name__ == "__main__":
    setup_logger()

    parser = argparse.ArgumentParser(description="RoyalRoad Chapter Scraper")
    parser.add_argument("--url", required=True, help="URL of starting chapter")
    parser.add_argument("--title", required=True, help="Book title")

    args = parser.parse_args()

    scrape_book(args.url, args.title)
