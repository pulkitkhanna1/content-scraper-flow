import time
from bs4 import BeautifulSoup as bs
import undetected_chromedriver as uc
import re 
from docx import Document
from urllib.parse import urljoin
from selenium.webdriver.common.by import By
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException, NoSuchElementException
import os
from pathlib import Path
from datetime import datetime

print("Starting Wattpad Scraper with Login and Unlock Support...\n")

# Configuration
BASE_OUTPUT_FOLDER = "wattpad_stories_output"
EMAIL = "shohtaj_singh@pocketfm.com"
PASSWORD = "PocketContent@123"

# List of URLs to process
# Each entry: (url, story_name, start_chapter, end_chapter)
URLS_TO_PROCESS = [
    ("https://www.wattpad.com/927889770-lost-luna-a-werewolf-romance-new", "Lost Luna: A Werewolf Romance", 1, 500),
]

def create_folder_structure(base_folder):
    """Create base output folder if it doesn't exist"""
    Path(base_folder).mkdir(parents=True, exist_ok=True)
    return base_folder

def get_story_folder(base_folder, story_name):
    """Create and return story-specific subfolder"""
    story_folder = os.path.join(base_folder, story_name)
    Path(story_folder).mkdir(parents=True, exist_ok=True)
    return story_folder

def setup_driver():
    """Set up and return Chrome driver"""
    options = uc.ChromeOptions()
    options.page_load_strategy = 'eager'
    # options.add_argument("--headless=new")  # optional for faster run
    options.add_argument("--no-sandbox")
    options.add_argument("--disable-dev-shm-usage")
    options.add_argument("--disable-blink-features=AutomationControlled")
    options.add_argument("--disable-gpu")
    options.add_argument("--window-size=1920,1080")
    options.add_argument("--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/92.0.4515.159 Safari/537.36")
    
    # Match your installed Chrome major version (148). Update if Chrome updates.
    print("   Finding/downloading matching ChromeDriver (may take 1–2 min on first run)...")
    driver = uc.Chrome(options=options, version_main=150)
    return driver

# Driver will be initialized in main
driver = None

def login_to_wattpad():
    """Login to Wattpad using provided credentials"""
    global driver
    print("🔐 Step 1: Navigating to login page...")
    driver.get("https://www.wattpad.com/login")
    time.sleep(1)
    
    try:
        # Click the initial submit button to show the login form
        print("⏳ Step 2: Clicking initial login button...")
        submit_button = driver.find_element(By.CSS_SELECTOR, ".btn-block.btn-primary.submit-btn-new")
        driver.execute_script("arguments[0].scrollIntoView(true);", submit_button)
        time.sleep(0.3)
        submit_button.click()
        time.sleep(1)
        
        # Fill in the login form
        print("📝 Step 3: Filling in credentials...")
        
        # Fill username/email
        username_field = driver.find_element(By.ID, "login-username")
        username_field.clear()
        username_field.send_keys(EMAIL)
        time.sleep(0.2)
        
        # Fill password
        password_field = driver.find_element(By.ID, "login-password")
        password_field.clear()
        password_field.send_keys(PASSWORD)
        time.sleep(0.2)
        
        # Click the log in button in the form
        print("🔑 Step 4: Submitting login form...")
        login_button = driver.find_element(By.CSS_SELECTOR, "button.footer-button-margin.button__Y70Pw.primary-variant__NO4pv")
        driver.execute_script("arguments[0].scrollIntoView(true);", login_button)
        time.sleep(0.3)
        login_button.click()
        
        # Wait for login to complete and redirect
        time.sleep(2)
        print("✅ Login successful!\n")
        
    except Exception as e:
        print(f"❌ Error during login: {e}")
        raise

def unlock_chapter():
    """Check and click the unlock chapter button if available"""
    global driver
    try:
        unlock_button = driver.find_element(By.CSS_SELECTOR, ".purchase-buttons-container .unlock-part")
        print("   🔓 Unlock button found! Clicking to unlock chapter...")
        driver.execute_script("arguments[0].scrollIntoView(true);", unlock_button)
        time.sleep(0.3)
        unlock_button.click()
        time.sleep(2)
        print("   ✅ Chapter unlocked!")
        return True
    except:
        # Chapter is free, no unlock needed
        return False

def process_story(url, story_name, start_chapter, end_chapter, output_folder):
    """Process a single story and save to the specified folder"""
    global driver
    print(f"\n{'='*70}")
    print(f"📖 Processing: {story_name}")
    print(f"{'='*70}\n")
    
    try:
        # Ensure full URL (driver.get() requires a valid URL with scheme)
        if not url.startswith(("http://", "https://")):
            url = "https://www." + url.lstrip("/")
        
        # Create story-specific folder
        story_folder = get_story_folder(output_folder, story_name)
        
        # Load the main story page
        print(f"📖 Loading story page...")
        driver.get(url)
        time.sleep(2)

        # Click the "Parts" tab to reveal the chapter list
        # Try multiple selectors in case Wattpad's markup changes
        PARTS_TAB_SELECTORS = [
            (By.ID, "tab-parts"),
            (By.CSS_SELECTOR, "[data-tab='parts']"),
            (By.XPATH, "//button[normalize-space()='Parts']"),
            (By.XPATH, "//a[normalize-space()='Parts']"),
            (By.XPATH, "//*[contains(@class,'tab') and normalize-space()='Parts']"),
        ]
        wait = WebDriverWait(driver, 10)
        tab_clicked = False
        for sel_by, sel_val in PARTS_TAB_SELECTORS:
            try:
                parts_tab = wait.until(EC.presence_of_element_located((sel_by, sel_val)))
                driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", parts_tab)
                time.sleep(0.3)
                try:
                    parts_tab.click()
                except Exception:
                    driver.execute_script("arguments[0].click();", parts_tab)
                print(f"   ✓ Clicked 'Parts' tab ({sel_by}={sel_val!r})")
                time.sleep(1.5)
                tab_clicked = True
                break
            except Exception:
                continue
        if not tab_clicked:
            print("   ⚠ Could not find 'Parts' tab with any known selector")

        soup = bs(driver.page_source, "lxml")

        # Extract chapter links — try multiple selectors
        CHAPTER_LINK_SELECTORS = [
            'ul[aria-label="story-parts"] li a',
            '.story-parts li a',
            '[class*="story-parts"] li a',
            '.table-of-contents li a',
            'ol.parts li a',
        ]
        chapter_links = []
        for sel in CHAPTER_LINK_SELECTORS:
            chapter_links = soup.select(sel)
            if chapter_links:
                print(f"   ✓ Found chapters using selector: {sel!r}")
                break

        if not chapter_links:
            dump_path = os.path.join(output_folder, f"{story_name}_page_source_debug.html")
            with open(dump_path, "w", encoding="utf-8") as f:
                f.write(driver.page_source)
            print(f"   ⚠ No chapter links found with any selector.")
            print(f"   ⚠ Page source saved to: {dump_path}")
            print("   ⚠ Open that file in a browser and inspect the chapter list markup to find the right selector.")
        document = Document()
        
        print(f"📚 Found {len(chapter_links)} chapters")
        print(f"📝 Processing chapters {start_chapter} to {end_chapter}\n")
        
        # Loop over chapters
        for idx, link in enumerate(chapter_links[start_chapter - 1:end_chapter], start=start_chapter):
            href = urljoin("https://www.wattpad.com", link['href'])
            print(f"[Chapter {idx}/{end_chapter}] Fetching: {href}")
            
            try:
                try:
                    driver.get(href)
                except TimeoutException:
                    pass  # eager strategy loaded the DOM; remaining resources can be ignored
                time.sleep(1)

                # Check and unlock chapter if needed
                unlock_chapter()
                
                # Scroll to load all content - improved method with better waiting
                print("   📜 Loading content...")
                
                # Initial wait for page to stabilize
                time.sleep(0.5)
                
                # Get initial page height
                last_height = driver.execute_script("return document.body.scrollHeight")
                print(f"   📏 Initial page height: {last_height}px")
                
                # Scroll down gradually to trigger lazy loading
                scroll_attempts = 0
                max_scroll_attempts = 300  # Increased limit for very long chapters
                no_change_count = 0
                stable_count = 0  # Count how many times height stayed the same
                
                while scroll_attempts < max_scroll_attempts:
                    # Scroll down
                    driver.execute_script("window.scrollBy(0, 400)")
                    time.sleep(0.2)  # Reduced wait time for faster processing
                    
                    # Check if new content has loaded
                    new_height = driver.execute_script("return document.body.scrollHeight")
                    
                    # Get current scroll position
                    current_position = driver.execute_script("return window.pageYOffset + window.innerHeight")
                    
                    # If page height increased, reset counters
                    if new_height > last_height:
                        last_height = new_height
                        no_change_count = 0
                        stable_count = 0
                        if scroll_attempts % 30 == 0:  # Print progress every 30 scrolls
                            print(f"   📏 Page height increased to {new_height}px...")
                    else:
                        # Height didn't change
                        stable_count += 1
                        
                        # If we're at the bottom and height is stable
                        if current_position >= new_height - 50:  # 50px threshold
                            no_change_count += 1
                            # Need more consecutive stable checks to be sure
                            if no_change_count >= 5 and stable_count >= 10:
                                print(f"   ✓ Reached bottom after {scroll_attempts} scrolls (final height: {new_height}px)")
                                break
                    
                    scroll_attempts += 1
                
                if scroll_attempts >= max_scroll_attempts:
                    print(f"   ⚠ Reached max scroll attempts ({max_scroll_attempts}), final height: {new_height}px")
                
                # Enhanced final loading sequence
                print("   📜 Finalizing content load...")
                
                # Multiple scroll-to-bottom cycles with increasing wait times
                for cycle in range(3):
                    driver.execute_script("window.scrollTo(0, document.body.scrollHeight)")
                    time.sleep(0.3 + cycle * 0.2)  # Reduced wait: 0.3s, 0.5s, 0.7s
                    
                    # Check if height changed
                    check_height = driver.execute_script("return document.body.scrollHeight")
                    if check_height > new_height:
                        new_height = check_height
                        print(f"   📏 Additional content loaded (height: {new_height}px)")
                
                # Scroll back up and down one more time to trigger any remaining lazy loads
                driver.execute_script("window.scrollTo(0, document.body.scrollHeight - 1500)")
                time.sleep(0.5)
                driver.execute_script("window.scrollTo(0, document.body.scrollHeight)")
                time.sleep(1)  # Reduced final wait
                
                # One more verification scroll
                final_height = driver.execute_script("return document.body.scrollHeight")
                driver.execute_script("window.scrollTo(0, document.body.scrollHeight)")
                time.sleep(0.5)  # Reduced final wait before extraction
                
                # Verify height is stable
                verify_height = driver.execute_script("return document.body.scrollHeight")
                if verify_height > final_height:
                    print(f"   ⚠ Content still loading, waiting more...")
                    time.sleep(1)
                    final_height = verify_height
                
                print(f"   ✓ Content loading complete (final page height: {final_height}px)")

                soup = bs(driver.page_source, "lxml")
                title_tag = soup.find("h1", class_="h2")
                title = title_tag.text.strip() if title_tag else f"Chapter {idx}"
                document.add_heading(title, level=1)

                content_paragraphs = soup.select('pre p')
                paragraph_count = 0
                for c in content_paragraphs:
                    raw_html = str(c)
                    clean_text = re.sub(r'<br\s*/?>', '\n', raw_html)
                    plain_text = bs(clean_text, "lxml").get_text()
                    for para in plain_text.split('\n'):
                        if para.strip():
                            document.add_paragraph(para.strip())
                            paragraph_count += 1
                
                document.add_page_break()
                print(f"   ✓ {title} extracted ({paragraph_count} paragraphs)\n")

            except Exception as e:
                print(f"   ❌ Error fetching chapter: {e}\n")

            time.sleep(1)  # Reduced delay between chapters
        
        # Save the document to story folder
        output_file = os.path.join(story_folder, f"{story_name}_{start_chapter}_{end_chapter}.docx")
        document.save(output_file)
        print(f"✅ Document saved: {output_file}\n")
        
        return True
        
    except Exception as e:
        print(f"❌ Error processing story '{story_name}': {e}\n")
        return False

# Main execution
if __name__ == "__main__":
    import argparse as _argparse
    _parser = _argparse.ArgumentParser(description="Wattpad story scraper", add_help=True)
    _parser.add_argument("--url", default=None, help="Story URL")
    _parser.add_argument("--book-name", default=None, help="Story name")
    _parser.add_argument("--start-chapter", type=int, default=1)
    _parser.add_argument("--end-chapter", type=int, default=500)
    _parser.add_argument("--output-dir", default=None, help="Output folder path")
    _args = _parser.parse_args()

    try:
        print("🚀 Initializing Chrome driver...")
        driver = setup_driver()
        print("✅ Driver initialized successfully\n")

        print("🔐 Logging in to Wattpad...\n")
        login_to_wattpad()

        if _args.url and _args.book_name:
            # Single story mode (launched by scraper_server)
            out_folder = _args.output_dir or BASE_OUTPUT_FOLDER
            os.makedirs(out_folder, exist_ok=True)
            print(f"📁 Output folder: {os.path.abspath(out_folder)}\n")
            process_story(_args.url, _args.book_name, _args.start_chapter, _args.end_chapter, out_folder)
        else:
            # Legacy batch mode
            base_folder = create_folder_structure(BASE_OUTPUT_FOLDER)
            print(f"📁 Output folder: {os.path.abspath(base_folder)}\n")
            total_stories = len(URLS_TO_PROCESS)
            successful = 0
            failed = 0
            for idx, (url, name, start, end) in enumerate(URLS_TO_PROCESS, 1):
                print(f"\n⏳ Processing story {idx}/{total_stories}...")
                if process_story(url, name, start, end, base_folder):
                    successful += 1
                else:
                    failed += 1
            print(f"\n{'='*70}")
            print(f"📊 SUMMARY")
            print(f"{'='*70}")
            print(f"✅ Successful: {successful}/{total_stories}")
            print(f"❌ Failed: {failed}/{total_stories}")
            print(f"📁 Output folder: {os.path.abspath(base_folder)}\n")

    except Exception as e:
        print(f"❌ Fatal error: {e}")
        import traceback
        traceback.print_exc()
    finally:
        # Close the driver
        if driver:
            try:
                driver.quit()
                print("✅ Driver closed successfully")
            except:
                pass
        print("🏁 Script completed!")