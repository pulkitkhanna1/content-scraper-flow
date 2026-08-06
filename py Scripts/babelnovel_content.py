import time
import logging
import sys
import os
import argparse
from typing import List

import undetected_chromedriver as uc
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException, WebDriverException

from docx import Document

# ---------------- CONFIG ----------------

MAX_CHAPTERS_PER_DOCX = 100
MAX_RETRIES = 3
WAIT_TIMEOUT = 20

# -------------------------------------- #

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s | %(levelname)s | %(message)s",
    handlers=[
        logging.StreamHandler(sys.stdout),
        logging.FileHandler("babelnovel_scraper.log")
    ]
)

logger = logging.getLogger(__name__)


def detect_chrome_major_version() -> int:
    """Best-effort detection of the installed Chrome major version on macOS/Linux."""
    import re
    import shutil
    import subprocess

    candidates = [
        "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome",
        shutil.which("google-chrome"),
        shutil.which("google-chrome-stable"),
        shutil.which("chromium"),
        shutil.which("chrome"),
    ]
    for path in candidates:
        if not path or not os.path.exists(path):
            continue
        try:
            out = subprocess.check_output([path, "--version"], stderr=subprocess.STDOUT, timeout=5).decode()
            match = re.search(r"(\d+)\.\d+\.\d+", out)
            if match:
                return int(match.group(1))
        except (subprocess.SubprocessError, OSError):
            continue
    return 0


def init_driver(version_main: int = 0):
    """Initialize undetected ChromeDriver in headful mode.

    If version_main <= 0, auto-detect the installed Chrome major version.
    """
    options = uc.ChromeOptions()
    # Headful mode - no headless argument
    options.add_argument("--start-maximized")
    options.add_argument("--no-sandbox")
    options.add_argument("--disable-dev-shm-usage")

    if version_main and version_main > 0:
        detected = version_main
        logger.info(f"Using user-specified Chrome major version: {detected}")
    else:
        detected = detect_chrome_major_version()
        if detected:
            logger.info(f"Detected Chrome major version: {detected}")
        else:
            logger.warning(
                "Could not detect Chrome version; letting undetected_chromedriver pick a default."
            )

    kwargs = {"options": options}
    if detected:
        kwargs["version_main"] = detected

    driver = uc.Chrome(**kwargs)
    return driver


def wait_for_login_confirmation(driver, login_url: str = "https://babelnovel.com"):
    """Navigate to login page and wait for the user to log in manually."""
    try:
        logger.info(f"Opening login page: {login_url}")
        driver.get(login_url)
    except WebDriverException as e:
        logger.warning(f"Could not open login URL ({e}); continuing on current page.")

    print("\n===================================================")
    print("MANUAL LOGIN REQUIRED")
    print(f"  1. Log in to BabelNovel in the opened browser ({login_url}).")
    print("  2. Make sure you can access paid/locked chapters.")
    print("  3. Once login is complete, return here and press ENTER.")
    print("===================================================\n")
    input("Press ENTER after you have finished logging in... ")
    logger.info("Login confirmed by user. Starting chapter scraping.")


def build_chapter_url(base_url: str, chapter_no: int) -> str:
    return base_url.rsplit("c", 1)[0] + f"c{chapter_no}"


def scrape_chapter(driver, url: str, wait_timeout: int = WAIT_TIMEOUT) -> (str, List[str]):
    for attempt in range(1, MAX_RETRIES + 1):
        try:
            logger.info(f"Fetching: {url} (Attempt {attempt})")
            driver.get(url)
            
            # Wait 4 seconds for all content to load
            time.sleep(4)

            wait = WebDriverWait(driver, wait_timeout)

            title_el = wait.until(
                EC.presence_of_element_located(
                    (By.CSS_SELECTOR, "h3.chapter_title__3Dp-H")
                )
            )

            p_tags = wait.until(
                EC.presence_of_all_elements_located((By.TAG_NAME, "p"))
            )

            title = title_el.text.strip()
            paragraphs = [p.text.strip() for p in p_tags if p.text.strip()]

            if not paragraphs:
                raise ValueError("No chapter content found")

            return title, paragraphs

        except (TimeoutException, WebDriverException, ValueError) as e:
            logger.warning(f"Error scraping chapter: {e}")
            if attempt == MAX_RETRIES:
                raise
            time.sleep(3)


def save_chapter_to_docx(
    document: Document,
    chapter_title: str,
    paragraphs: List[str]
):
    document.add_heading(chapter_title, level=1)
    for para in paragraphs:
        document.add_paragraph(para)


def count_words(paragraphs: List[str]) -> int:
    return sum(len(p.split()) for p in paragraphs)


def append_word_count_summary(document: Document, chapter_stats: List[tuple]):
    """Append a chapter-wise word count summary table to the document."""
    document.add_page_break()
    document.add_heading("Chapter-wise Word Count Summary", level=1)

    if not chapter_stats:
        document.add_paragraph("No chapters were scraped.")
        return

    table = document.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    header_cells = table.rows[0].cells
    header_cells[0].text = "Chapter #"
    header_cells[1].text = "Title"
    header_cells[2].text = "Word Count"

    for chapter_no, title, word_count in chapter_stats:
        row = table.add_row().cells
        row[0].text = str(chapter_no)
        row[1].text = title
        row[2].text = str(word_count)

    total_words = sum(wc for _, _, wc in chapter_stats)
    document.add_paragraph("")
    document.add_paragraph(f"Total chapters scraped: {len(chapter_stats)}")
    document.add_paragraph(f"Total word count: {total_words}")
    document.add_paragraph(
        f"Average words per chapter: {total_words // len(chapter_stats)}"
    )


def parse_arguments():
    """Parse command-line arguments"""
    parser = argparse.ArgumentParser(
        description="Scrape chapters from BabelNovel and save to DOCX files",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python babelnovel_content.py --url "https://babelnovel.com/books/book-name/chapters/c1" --start 1 --end 100
  python babelnovel_content.py --url "https://babelnovel.com/books/book-name/chapters/c1" --start 1 --end 50 --output "my_book"
        """
    )
    parser.add_argument(
        "--url",
        required=True,
        help="Base URL of the first chapter (e.g., https://babelnovel.com/books/book-name/chapters/c1)"
    )
    parser.add_argument(
        "--start",
        type=int,
        required=True,
        help="Starting chapter number"
    )
    parser.add_argument(
        "--end",
        type=int,
        required=True,
        help="Ending chapter number (inclusive)"
    )
    parser.add_argument(
        "--output",
        default="babelnovel_chapters",
        help="Output folder name (default: babelnovel_chapters)"
    )
    parser.add_argument(
        "--max-chapters-per-docx",
        type=int,
        default=MAX_CHAPTERS_PER_DOCX,
        help=f"Maximum chapters per DOCX file (default: {MAX_CHAPTERS_PER_DOCX})"
    )
    parser.add_argument(
        "--wait-timeout",
        type=int,
        default=WAIT_TIMEOUT,
        help=f"Wait timeout in seconds (default: {WAIT_TIMEOUT})"
    )
    parser.add_argument(
        "--delay",
        type=float,
        default=1.5,
        help="Delay between chapter requests in seconds (default: 1.5)"
    )
    parser.add_argument(
        "--login-url",
        default="https://babelnovel.com",
        help="URL to open for manual login (default: https://babelnovel.com)"
    )
    parser.add_argument(
        "--chrome-version",
        type=int,
        default=0,
        help="Override Chrome major version for ChromeDriver (default: auto-detect)"
    )

    return parser.parse_args()


def main():
    args = parse_arguments()
    
    # Validate arguments
    if args.start < 1:
        logger.error("Start chapter must be >= 1")
        sys.exit(1)
    if args.end < args.start:
        logger.error("End chapter must be >= start chapter")
        sys.exit(1)
    
    logger.info(f"Starting scraper with URL: {args.url}")
    logger.info(f"Chapters: {args.start} to {args.end}")
    logger.info(f"Output folder: {args.output}")
    
    # Create output folder
    os.makedirs(args.output, exist_ok=True)
    logger.info(f"Created output folder: {args.output}")
    
    driver = init_driver(version_main=args.chrome_version)

    wait_for_login_confirmation(driver, args.login_url)

    current_doc = Document()
    doc_index = 1
    chapters_in_current_doc = 0
    chapter_stats: List[tuple] = []

    for chapter_no in range(args.start, args.end + 1):
        chapter_url = build_chapter_url(args.url, chapter_no)

        try:
            title, content = scrape_chapter(driver, chapter_url, args.wait_timeout)

            word_count = count_words(content)
            chapter_stats.append((chapter_no, title, word_count))

            logger.info(
                f"Scraped Chapter {chapter_no}: {title} ({word_count} words)"
            )

            save_chapter_to_docx(current_doc, title, content)
            chapters_in_current_doc += 1

            is_last_chapter = chapter_no == args.end
            if (
                chapters_in_current_doc >= args.max_chapters_per_docx
                and not is_last_chapter
            ):
                output_file = os.path.join(args.output, f"part_{doc_index}.docx")
                current_doc.save(output_file)
                logger.info(f"Saved {output_file}")

                doc_index += 1
                chapters_in_current_doc = 0
                current_doc = Document()

            time.sleep(args.delay)

        except Exception as e:
            logger.error(f"Failed chapter {chapter_no}: {e}")

    # Append chapter-wise word count summary at the end of the final doc.
    if chapters_in_current_doc == 0:
        # All chapters were already flushed; create a final doc just for the summary.
        current_doc = Document()
        current_doc.add_heading("Final Summary", level=0)

    append_word_count_summary(current_doc, chapter_stats)

    output_file = os.path.join(args.output, f"part_{doc_index}.docx")
    current_doc.save(output_file)
    logger.info(f"Saved {output_file} (includes word count summary)")

    driver.quit()
    logger.info("Scraping completed.")


if __name__ == "__main__":
    main()
