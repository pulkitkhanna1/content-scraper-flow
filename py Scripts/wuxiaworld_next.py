import os
import time
import argparse
import logging
from docx import Document
from docx.shared import Pt
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException, StaleElementReferenceException
from webdriver_manager.chrome import ChromeDriverManager

# --------------------------------------
# Setup logging
# --------------------------------------
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
)

# --------------------------------------
# Credentials (hardcoded)
# --------------------------------------
LOGIN_EMAIL = "fantasy.team@pocketfm.com"
LOGIN_PASSWORD = "mH8tF_u.hFj#8dX"
HOME_URL = "https://www.wuxiaworld.com/"

# Selector that targets the bottom-bar "Next" button by its label text. The
# Previous button uses the same MUI classes, so we MUST match on the label
# text "Next" — otherwise the script will navigate backwards.
NEXT_BUTTON_XPATH = (
    '//button[contains(@class, "MuiBottomNavigationAction-root")]'
    '[.//span[contains(@class, "MuiBottomNavigationAction-label")'
    ' and normalize-space(text())="Next"]]'
)


def login_to_wuxiaworld(driver):
    """Authenticate via the SPA's OAuth flow.

    The SPA generates fresh OAuth params (state + PKCE code_challenge) when it
    redirects to identity.wuxiaworld.com — those values must come from the SPA's
    own session, so we cannot navigate directly to a hand-rolled login URL.
    """
    logging.info("Navigating to WuxiaWorld homepage...")
    driver.get(HOME_URL)

    wait = WebDriverWait(driver, 30)
    time.sleep(3)  # let the SPA hydrate its header

    # The LOG IN button lives inside a dropdown that's opened by the profile-nav icon.
    # NOTE: the site removed its data-cy attributes, so we now match the profile
    # icon by its aria-label and the LOG IN entry by its button text.
    logging.info("Opening profile-nav menu...")
    profile_btn = wait.until(
        EC.element_to_be_clickable(
            (By.CSS_SELECTOR, 'button[aria-label="profile nav"]')
        )
    )
    driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", profile_btn)
    profile_btn.click()

    logging.info("Clicking LOG IN button in profile menu...")
    login_btn = wait.until(
        EC.element_to_be_clickable(
            (By.XPATH, '//button[normalize-space(text())="LOG IN"]')
        )
    )
    driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", login_btn)
    login_btn.click()

    logging.info("Waiting for redirect to identity.wuxiaworld.com login page...")
    wait.until(EC.url_contains("identity.wuxiaworld.com/Account/Login"))
    logging.info(f"On identity login page: {driver.current_url}")

    logging.info("Filling in credentials...")
    email_input = wait.until(EC.presence_of_element_located((By.ID, "Username")))
    email_input.clear()
    email_input.send_keys(LOGIN_EMAIL)

    password_input = driver.find_element(By.ID, "Password")
    password_input.clear()
    password_input.send_keys(LOGIN_PASSWORD)

    logging.info("Submitting login form...")
    submit_btn = wait.until(
        EC.element_to_be_clickable(
            (By.CSS_SELECTOR, "button[type='submit'][name='button'][value='login']")
        )
    )
    driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", submit_btn)
    # Native click — guarantees the form's submitter is set so `button=login` is
    # included in the POST body. JS .click() on submit buttons does not always do this.
    try:
        submit_btn.click()
    except Exception as e:
        logging.warning(f"Native click failed ({e}); falling back to JS click.")
        driver.execute_script("arguments[0].click();", submit_btn)

    logging.info("Waiting for OAuth flow to complete (back on www.wuxiaworld.com)...")
    try:
        WebDriverWait(driver, 60).until(
            lambda d: "identity.wuxiaworld.com" not in d.current_url
                      and "www.wuxiaworld.com" in d.current_url
                      and "/auth/callback" not in d.current_url
        )
    except TimeoutException:
        error_text = ""
        try:
            error_text = driver.find_element(
                By.CSS_SELECTOR,
                ".validation-summary-errors, .field-validation-error, [class*='error']",
            ).text
        except Exception:
            pass
        raise RuntimeError(
            f"Login did not complete within 60s. Current URL: {driver.current_url}. "
            f"Error message on page: {error_text or 'none'}"
        )

    logging.info(f"Login successful — landed on: {driver.current_url}")
    time.sleep(3)  # allow SPA to finish auth state setup

# --------------------------------------
# Helper Functions
# --------------------------------------

def create_docx(file_index, start_chapter, output_folder, filename_pattern):
    """Create a new DOCX file whose name matches real chapter numbers"""
    os.makedirs(output_folder, exist_ok=True)
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # compute real start and end chapter for this batch
    batch_start = start_chapter + file_index * 50
    batch_end = batch_start + 49

    # Format filename pattern with batch_start and batch_end
    filename = os.path.join(output_folder, filename_pattern.format(
        batch_start=batch_start,
        batch_end=batch_end,
        file_index=file_index
    ))

    return doc, filename


def save_docx(doc, filename):
    """Save the DOCX file immediately (streaming-style)"""
    doc.save(filename)
    logging.info(f"Saved progress to: {filename}")


def read_chapter(driver):
    """Read the current chapter's title and body. Returns (title, body)."""
    wait = WebDriverWait(driver, 15)

    title_el = wait.until(
        EC.presence_of_element_located((By.CSS_SELECTOR, "h4.font-set-b18.flex.items-start span"))
    )
    chapter_title = title_el.text.strip()

    body_el = wait.until(
        EC.presence_of_element_located((By.CSS_SELECTOR, "div.chapter-content"))
    )

    # Wait for content to populate — the SPA hydrates body text after mount.
    max_content_wait = 10
    content_wait_start = time.time()
    chapter_body = ""
    while time.time() - content_wait_start < max_content_wait:
        try:
            chapter_body = body_el.text.strip()
        except StaleElementReferenceException:
            body_el = driver.find_element(By.CSS_SELECTOR, "div.chapter-content")
            chapter_body = body_el.text.strip()
        if len(chapter_body) > 50:
            break
        time.sleep(1)
        body_el = driver.find_element(By.CSS_SELECTOR, "div.chapter-content")

    return chapter_title, chapter_body


def go_to_next_chapter(driver, current_title, current_url):
    """Click the bottom-bar Next button and wait until navigation completes.

    Returns True if the chapter changed, False otherwise.
    """
    wait = WebDriverWait(driver, 15)

    next_button = wait.until(
        EC.element_to_be_clickable((By.XPATH, NEXT_BUTTON_XPATH))
    )
    driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", next_button)
    try:
        next_button.click()
    except Exception:
        driver.execute_script("arguments[0].click();", next_button)

    # Wait for navigation: URL or title must change.
    deadline = time.time() + 30
    while time.time() < deadline:
        try:
            new_url = driver.current_url
            if new_url != current_url:
                # URL changed — wait for new title to render so the next read is clean.
                try:
                    WebDriverWait(driver, 10).until(
                        EC.presence_of_element_located(
                            (By.CSS_SELECTOR, "h4.font-set-b18.flex.items-start span")
                        )
                    )
                except TimeoutException:
                    pass
                return True
            new_title_el = driver.find_element(
                By.CSS_SELECTOR, "h4.font-set-b18.flex.items-start span"
            )
            new_title = new_title_el.text.strip()
            if new_title and new_title != current_title:
                return True
        except StaleElementReferenceException:
            pass
        time.sleep(0.5)

    return False

# --------------------------------------
# Main Scraper
# --------------------------------------

def scrape_wuxiaworld(start_chapter, end_chapter, start_url, output_folder, filename_pattern, email=None, password=None):
    global LOGIN_EMAIL, LOGIN_PASSWORD
    if email:
        LOGIN_EMAIL = email
    if password:
        LOGIN_PASSWORD = password

    # --- Chrome setup with headers ---
    chrome_options = Options()
    chrome_options.add_argument("--disable-gpu")
    chrome_options.add_argument("--no-sandbox")

    headers = {
        "accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8",
        "accept-language": "en-GB,en;q=0.6",
        "cache-control": "max-age=0",
        "user-agent": 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/144.0.0.0 Safari/537.36',
        "sec-ch-ua": '"Not)A;Brand";v="8", "Chromium";v="144", "Brave";v="144"',
        "sec-ch-ua-mobile": "?0",
        "sec-ch-ua-platform": '"macOS"',
    }

    for key, value in headers.items():
        chrome_options.add_argument(f'--header="{key}: {value}"')

    service = Service(ChromeDriverManager().install())
    driver = webdriver.Chrome(service=service, options=chrome_options)

    # ----------------------------
    # Step 1: Automated login (via SPA-initiated OAuth flow)
    # ----------------------------
    login_to_wuxiaworld(driver)

    # ----------------------------
    # Step 2: Navigate to the starting chapter
    # ----------------------------
    logging.info(f"Loading starting chapter: {start_url}")
    driver.get(start_url)
    time.sleep(5)

    total_chapters = end_chapter - start_chapter + 1
    logging.info(
        f"Will scrape {total_chapters} chapter(s) by clicking Next: "
        f"{start_chapter} → {end_chapter}"
    )

    chapter_count = 50
    file_index = 0
    doc, filename = create_docx(file_index, start_chapter, output_folder, filename_pattern)

    try:
        for offset in range(total_chapters):
            chapter_num = start_chapter + offset
            logging.info(f"Scraping chapter {chapter_num}: {driver.current_url}")

            success = False
            chapter_title = ""
            for attempt in range(10):
                try:
                    chapter_title, chapter_body = read_chapter(driver)
                    if len(chapter_body) <= 50:
                        logging.warning(
                            f"Chapter {chapter_num} may have incomplete content "
                            f"(length: {len(chapter_body)})"
                        )

                    doc.add_heading(chapter_title, level=2)
                    doc.add_paragraph(chapter_body)
                    doc.add_page_break()

                    chapter_count += 1
                    save_docx(doc, filename)

                    if chapter_count % 50 == 0:
                        file_index += 1
                        doc, filename = create_docx(file_index, start_chapter, output_folder, filename_pattern)

                    success = True
                    break

                except Exception as e:
                    logging.error(f"Attempt {attempt+1}/10 failed for chapter {chapter_num}: {e}")
                    if attempt < 9:
                        time.sleep(5)
                        # Reload the page to recover from transient failures.
                        try:
                            driver.refresh()
                            time.sleep(3)
                        except Exception:
                            pass

            if not success:
                logging.error(f"Skipping chapter {chapter_num} after 10 failed attempts.")

            # Click Next to move to the following chapter (skip after the last one).
            if offset < total_chapters - 1:
                current_url = driver.current_url
                navigated = False
                for nav_attempt in range(5):
                    try:
                        if go_to_next_chapter(driver, chapter_title, current_url):
                            navigated = True
                            time.sleep(2)
                            break
                        logging.warning(
                            f"Next click did not change page (attempt {nav_attempt+1}/5)"
                        )
                        time.sleep(2)
                    except Exception as e:
                        logging.error(
                            f"Next-button navigation attempt {nav_attempt+1}/5 failed: {e}"
                        )
                        time.sleep(2)

                if not navigated:
                    logging.error(
                        f"Could not advance past chapter {chapter_num} via Next button — aborting."
                    )
                    break

        logging.info("✅ Scraping completed successfully.")

    finally:
        driver.quit()
        logging.info("Browser closed.")


# --------------------------------------
# CLI Entry
# --------------------------------------
if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Scrape chapters from WuxiaWorld via Next-button navigation.")
    parser.add_argument("--start", type=int, required=True, help="Starting chapter number (used for output filenames)")
    parser.add_argument("--end", type=int, required=True, help="Ending chapter number (inclusive)")
    parser.add_argument("--start-url", type=str, required=True,
                       help="Full URL of the first chapter to scrape (e.g., 'https://www.wuxiaworld.com/novel/foo/foo-chapter-1')")
    parser.add_argument("--output-folder", type=str,
                       default="Netherworld_Investigator",
                       help="Output folder name for saved DOCX files")
    parser.add_argument("--filename-pattern", type=str,
                       default="Chapters_{batch_start}_{batch_end}.docx",
                       help="Filename pattern (use {batch_start}, {batch_end}, {file_index} as placeholders)")
    parser.add_argument("--email", type=str, default=None, help="WuxiaWorld account email")
    parser.add_argument("--password", type=str, default=None, help="WuxiaWorld account password")
    args = parser.parse_args()

    scrape_wuxiaworld(args.start, args.end, args.start_url, args.output_folder, args.filename_pattern, email=args.email, password=args.password)
