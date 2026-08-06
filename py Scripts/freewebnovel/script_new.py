#!/usr/bin/env python3
"""
script_new.py

Scraper for freewebnovel.com that:
  1. Navigates to the book's main page
  2. Reads the full Table of Contents to collect all chapter URLs
  3. Scrapes chapters in the requested range and saves as .docx parts

Usage:
    python script_new.py \
        --book-url "https://freewebnovel.com/novel/goblin-dependency/" \
        --start-chapter 1 \
        --end-chapter 50 \
        --book-name "Goblin Dependency"

    python script_new.py --url-file url.txt

url.txt format:
    "Book Name" https://freewebnovel.com/novel/book-slug 1 1000
"""

import re
import os
import json
import shlex
import subprocess
import time
import argparse
from docx import Document
import undetected_chromedriver as uc
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException, NoSuchElementException

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
DEFAULT_URL_FILE = "url.txt"
DEFAULT_OUTPUT_ROOT = os.path.join(SCRIPT_DIR, "ouput")
PROGRESS_FILE_NAME = "progress.json"
CHAPTERS_PER_PART = 100
DELAY_BETWEEN_CHAPTERS = 2
TIMEOUT = 15
# Restart the browser every N chapters to keep Chrome's memory from building up
# (long-lived sessions eventually die with "tab crashed" = renderer OOM).
RESTART_EVERY = 150
# How many times to recreate the browser and retry a single chapter on a crash.
MAX_CHAPTER_RETRIES = 3


def sanitize_filename(name):
    name = re.sub(r'[<>:"/\\|?*]', '', name).strip()
    return re.sub(r'\s+', '_', name)


def resolve_path(path):
    """Resolve paths relative to cwd first, then this script directory."""
    if os.path.isabs(path) or os.path.exists(path):
        return path
    script_path = os.path.join(SCRIPT_DIR, path)
    if os.path.exists(script_path):
        return script_path
    return path


def _parse_chapter_number(value, label, line_no):
    if value.lower() in {"none", "all", "-"}:
        return None
    try:
        return int(value)
    except ValueError:
        raise ValueError(f"Line {line_no}: invalid {label} chapter '{value}'")


def parse_url_file(path):
    """
    Read book entries from a text file.

    Expected line format:
        "Book Name" https://freewebnovel.com/novel/book-slug 1 1000

    Start/end chapters are optional. End can be "all", "none", or "-".
    """
    entries = []
    path = resolve_path(path)

    with open(path, "r", encoding="utf-8") as f:
        for line_no, raw_line in enumerate(f, start=1):
            line = raw_line.strip()
            if not line or line.startswith("#"):
                continue

            try:
                parts = shlex.split(line, comments=True)
            except ValueError as e:
                raise ValueError(f"Line {line_no}: could not parse quoted text: {e}")

            url_index = next(
                (i for i, part in enumerate(parts)
                 if part.startswith(("http://", "https://"))),
                None,
            )
            if url_index is None:
                raise ValueError(f"Line {line_no}: missing book URL")
            if url_index == 0:
                raise ValueError(f"Line {line_no}: missing book name before URL")

            book_name = " ".join(parts[:url_index]).strip()
            book_url = parts[url_index]
            range_parts = parts[url_index + 1:]

            if len(range_parts) > 2:
                raise ValueError(
                    f"Line {line_no}: expected at most start and end chapters after URL"
                )

            start_chapter = (
                _parse_chapter_number(range_parts[0], "start", line_no)
                if range_parts else 1
            )
            end_chapter = (
                _parse_chapter_number(range_parts[1], "end", line_no)
                if len(range_parts) > 1 else None
            )

            if start_chapter is None:
                raise ValueError(f"Line {line_no}: start chapter cannot be all/none")

            entries.append({
                "book_name": book_name,
                "book_url": book_url,
                "start_chapter": start_chapter,
                "end_chapter": end_chapter,
            })

    if not entries:
        raise ValueError(f"No book entries found in {path}")

    return entries


# ---------------------------------------------------------------------------
# Driver setup (same as existing script)
# ---------------------------------------------------------------------------

def _chrome_major_version():
    """Major version from locally installed Chrome/Chromium (--version output)."""
    candidates = (
        "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome",
        "/Applications/Chromium.app/Contents/MacOS/Chromium",
        "google-chrome",
        "google-chrome-stable",
        "chromium",
        "chrome",
    )
    for cmd in candidates:
        try:
            out = subprocess.check_output(
                [cmd, "--version"],
                stderr=subprocess.STDOUT,
                text=True,
                timeout=15,
            )
            m = re.search(r"(\d+)\.", out)
            if m:
                return int(m.group(1))
        except Exception:
            continue
    return None


def setup_driver():
    opts = uc.ChromeOptions()
    opts.add_argument("--start-maximized")
    opts.add_argument("--ignore-certificate-errors")
    opts.add_argument("--no-sandbox")
    opts.add_argument("--disable-dev-shm-usage")
    # Reduce per-tab memory growth over long runs
    opts.add_argument("--disable-application-cache")
    opts.add_argument("--disk-cache-size=1")
    opts.add_argument("--disable-gpu")

    major = _chrome_major_version()
    if major is not None:
        print(f"Matching Chromedriver to Chrome major version {major}")
    driver = uc.Chrome(options=opts, version_main=major, use_subprocess=True)
    driver.set_page_load_timeout(30)
    driver.implicitly_wait(5)
    return driver


# ---------------------------------------------------------------------------
# TOC scraping
# ---------------------------------------------------------------------------

def _format_number_ranges(numbers):
    """Compact [1, 2, 3, 7, 8] into '1-3, 7-8' for logs."""
    if not numbers:
        return ""

    ranges = []
    start = prev = numbers[0]
    for num in numbers[1:]:
        if num == prev + 1:
            prev = num
            continue
        ranges.append((start, prev))
        start = prev = num
    ranges.append((start, prev))

    parts = []
    for start, end in ranges:
        parts.append(str(start) if start == end else f"{start}-{end}")
    return ", ".join(parts)


def _infer_missing_chapter_urls(chapter_urls):
    """
    Some freewebnovel TOCs expose early chapters plus the latest chapters, but
    leave the middle hidden from the DOM. Chapter URLs are still sequential, so
    fill numeric gaps by reusing the observed URL pattern.
    """
    if len(chapter_urls) < 2:
        return []

    nums = sorted(chapter_urls)
    missing = [num for num in range(nums[0], nums[-1] + 1)
               if num not in chapter_urls]
    if not missing:
        return []

    template = None
    for num in nums:
        href = chapter_urls[num]
        m = re.match(r"^(.*chapter-?)(\d+)(.*)$", href, re.IGNORECASE)
        if m and int(m.group(2)) == num:
            template = (m.group(1), m.group(3))
            break

    if template is None:
        print("  Warning: TOC has missing chapter numbers, but no URL pattern "
              "could be inferred")
        return []

    prefix, suffix = template
    for num in missing:
        chapter_urls[num] = f"{prefix}{num}{suffix}"

    return missing


def collect_toc_urls(driver, book_url):
    """
    Navigate to the book page and collect all chapter URLs from the
    Table of Contents. Handles multi-page TOC by clicking through pages.
    Returns a list of (chapter_number, url) tuples sorted by chapter number.
    """
    print(f"Loading book page: {book_url}")
    driver.get(book_url)
    time.sleep(2)

    # Detect redirect to 404 page
    current = driver.current_url
    if "404" in current or "not-found" in current.lower():
        raise RuntimeError(
            f"Book URL redirected to: {current}\n"
            "The book page URL is likely wrong. Open the book in a browser and copy the exact URL "
            "of the main page that shows the chapter list (e.g. https://freewebnovel.com/goblin-dependency.html)"
        )

    chapter_urls = {}  # chapter_number -> url

    def extract_chapter_num(url):
        m = re.search(r'chapter-?(\d+)', url, re.IGNORECASE)
        return int(m.group(1)) if m else None

    def collect_visible_links():
        """Collect all chapter hrefs currently visible in the DOM."""
        found = []
        elements = driver.find_elements(By.CSS_SELECTOR, "a[href*='/chapter-']")
        for el in elements:
            href = el.get_attribute("href") or ""
            if "chapter" in href.lower():
                found.append(href)
        return found

    # --- Expand the full chapter list ---
    # freewebnovel shows only the latest 6 chapters by default.
    # A "See all chapters" / "Show all" link (href="javascript:;") toggles the full list.
    expand_xpaths = [
        "//a[contains(translate(text(),'ABCDEFGHIJKLMNOPQRSTUVWXYZ','abcdefghijklmnopqrstuvwxyz'),'see all')]",
        "//a[contains(translate(text(),'ABCDEFGHIJKLMNOPQRSTUVWXYZ','abcdefghijklmnopqrstuvwxyz'),'show all')]",
        "//a[contains(translate(text(),'ABCDEFGHIJKLMNOPQRSTUVWXYZ','abcdefghijklmnopqrstuvwxyz'),'all chapters')]",
        "//a[contains(translate(text(),'ABCDEFGHIJKLMNOPQRSTUVWXYZ','abcdefghijklmnopqrstuvwxyz'),'view all')]",
        "//a[@href='javascript:;' and contains(@onclick,'show')]",
    ]
    for xpath in expand_xpaths:
        try:
            btn = driver.find_element(By.XPATH, xpath)
            print(f"  Clicking expand button: '{btn.text.strip()}'")
            driver.execute_script("arguments[0].click();", btn)
            time.sleep(2)
            break
        except NoSuchElementException:
            continue
    else:
        # No explicit expand button — try clicking any javascript:; link near the chapter list
        try:
            js_links = driver.find_elements(By.XPATH, "//a[@href='javascript:;']")
            for link in js_links:
                text = link.text.strip().lower()
                if any(kw in text for kw in ["chapter", "all", "list", "more"]):
                    print(f"  Clicking JS link: '{link.text.strip()}'")
                    driver.execute_script("arguments[0].click();", link)
                    time.sleep(2)
                    break
        except Exception:
            pass

    # Collect all links now visible
    links = collect_visible_links()
    for href in links:
        n = extract_chapter_num(href)
        if n and n not in chapter_urls:
            chapter_urls[n] = href

    print(f"  Found {len(chapter_urls)} chapters in TOC")

    if not chapter_urls:
        raise RuntimeError(
            "Could not find any chapter links on the book page. "
            "The site structure may have changed."
        )

    missing = _infer_missing_chapter_urls(chapter_urls)
    if missing:
        print(f"  Inferred {len(missing)} missing chapter URL(s): "
              f"{_format_number_ranges(missing)}")

    sorted_chapters = sorted(chapter_urls.items())  # [(num, url), ...]
    print(f"Total chapters found in TOC: {len(sorted_chapters)}")
    return sorted_chapters


# ---------------------------------------------------------------------------
# Chapter content extraction (same selectors as existing script)
# ---------------------------------------------------------------------------

def extract_chapter_title(driver, chapter_num):
    try:
        el = driver.find_element(By.CSS_SELECTOR, "div.top span.chapter")
        return el.text.strip()
    except NoSuchElementException:
        pass

    for sel in ["h4", "h1", "div.top h1.tit", ".chapter-title"]:
        try:
            el = driver.find_element(By.CSS_SELECTOR, sel)
            t = el.text.strip()
            if t:
                return t
        except NoSuchElementException:
            continue

    return f"Chapter {chapter_num}"


def _is_ad_content(text):
    ad_keywords = ["advertisement", "sponsored", "click here", "subscribe", "follow us"]
    text_lower = text.lower()
    return any(k in text_lower for k in ad_keywords)


def extract_chapter_content(driver, chapter_num):
    try:
        WebDriverWait(driver, TIMEOUT).until(
            EC.presence_of_element_located((By.ID, "article"))
        )
    except TimeoutException:
        print(f"  Warning: timeout waiting for content on chapter {chapter_num}")

    paragraphs = []
    for sel in [
        "#article p",
        "div.chapter-content p",
        "div.article-content p",
        "div.entry-content p",
        "article p",
    ]:
        elements = driver.find_elements(By.CSS_SELECTOR, sel)
        if elements:
            for el in elements:
                text = el.text.strip()
                if text and len(text) > 10 and not _is_ad_content(text):
                    paragraphs.append(text)
            if paragraphs:
                break

    return paragraphs


# ---------------------------------------------------------------------------
# Main scraper
# ---------------------------------------------------------------------------

class BookScraper:
    def __init__(self, book_url, book_name, start_chapter, end_chapter,
                 chapters_per_part=CHAPTERS_PER_PART, output_dir=None):
        self.book_url = book_url
        self.book_name = self._sanitize(book_name)
        self.start_chapter = start_chapter
        self.end_chapter = end_chapter
        self.chapters_per_part = chapters_per_part

        self.output_dir = output_dir or os.path.join(DEFAULT_OUTPUT_ROOT, self.book_name)
        os.makedirs(self.output_dir, exist_ok=True)

        self.part = 1
        self.doc = Document()
        self.chapters_in_part = 0
        self.part_start_chapter = None
        self.part_end_chapter = None
        self.total_scraped = 0
        self.current_part_checkpoint_path = None

        self.progress_path = os.path.join(self.output_dir, PROGRESS_FILE_NAME)
        self.completed_chapters = set()
        self.failed_chapters = {}
        self.last_completed_url = None
        self._load_progress()

        self.driver = setup_driver()

    def _sanitize(self, name):
        return sanitize_filename(name)

    def _load_progress(self):
        if not os.path.exists(self.progress_path):
            return

        try:
            with open(self.progress_path, "r", encoding="utf-8") as f:
                progress = json.load(f)
        except Exception as e:
            print(f"  Warning: could not read {self.progress_path}: {e}")
            return

        saved_url = progress.get("book_url")
        if saved_url and str(saved_url).rstrip("/") != self.book_url.rstrip("/"):
            print("  Warning: progress.json belongs to a different book URL; "
                  "starting this book from the requested range.")
            return

        completed = set()
        for value in progress.get("completed_chapters", []):
            try:
                completed.add(int(value))
            except (TypeError, ValueError):
                continue
        self.completed_chapters = completed

        failed = {}
        raw_failed = progress.get("failed_chapters", {})
        if isinstance(raw_failed, dict):
            for key, value in raw_failed.items():
                try:
                    failed[int(key)] = str(value)
                except (TypeError, ValueError):
                    continue
        self.failed_chapters = failed
        self.last_completed_url = progress.get("last_completed_url")

        if self.completed_chapters:
            last = max(self.completed_chapters)
            print(f"  Loaded progress.json: {len(self.completed_chapters)} "
                  f"completed chapter(s), last completed chapter {last}")

    def _save_progress(self, status="running"):
        completed = sorted(self.completed_chapters)
        payload = {
            "book_name": self.book_name,
            "book_url": self.book_url,
            "requested_start_chapter": self.start_chapter,
            "requested_end_chapter": self.end_chapter,
            "chapters_per_part": self.chapters_per_part,
            "status": status,
            "completed_count": len(completed),
            "completed_chapters": completed,
            "last_completed_chapter": completed[-1] if completed else None,
            "last_completed_url": self.last_completed_url,
            "failed_chapters": {
                str(num): self.failed_chapters[num]
                for num in sorted(self.failed_chapters)
            },
            "updated_at": time.strftime("%Y-%m-%dT%H:%M:%S%z"),
        }

        tmp_path = self.progress_path + ".tmp"
        with open(tmp_path, "w", encoding="utf-8") as f:
            json.dump(payload, f, indent=2, ensure_ascii=False)
            f.write("\n")
        os.replace(tmp_path, self.progress_path)

    def _mark_completed(self, chap_num, url):
        self.completed_chapters.add(chap_num)
        self.last_completed_url = url
        self.failed_chapters.pop(chap_num, None)
        self._save_progress(status="running")

    def _mark_failed(self, chap_num, error):
        self.failed_chapters[chap_num] = str(error)
        self._save_progress(status="running")

    def _restart_driver(self):
        """Tear down and recreate the browser (recovers from a crashed tab)."""
        try:
            self.driver.quit()
        except Exception:
            pass
        time.sleep(2)
        self.driver = setup_driver()

    def _scrape_one(self, chap_num, url):
        """Load and extract a single chapter, restarting the browser and
        retrying if the tab crashes. Returns (title, content) or raises."""
        for attempt in range(1, MAX_CHAPTER_RETRIES + 1):
            try:
                self.driver.get(url)
                time.sleep(1)
                title = extract_chapter_title(self.driver, chap_num)
                content = extract_chapter_content(self.driver, chap_num)
                return title, content
            except Exception as e:
                msg = str(e).lower()
                crashed = ("tab crashed" in msg or "session" in msg
                           or "chrome not reachable" in msg or "disconnected" in msg)
                if crashed and attempt < MAX_CHAPTER_RETRIES:
                    print(f"  Browser crashed on chapter {chap_num} "
                          f"(attempt {attempt}/{MAX_CHAPTER_RETRIES}) — restarting...")
                    self._restart_driver()
                    continue
                raise

    def _save_part(self, is_final=False, quiet=False):
        if self.chapters_in_part == 0:
            return
        start = self.part_start_chapter or self.start_chapter
        end = self.part_end_chapter or start + self.chapters_in_part - 1
        fname = f"{self.book_name}_chapters_{start}_to_{end}.docx"
        path = os.path.join(self.output_dir, fname)
        self.doc.save(path)
        if (self.current_part_checkpoint_path
                and self.current_part_checkpoint_path != path
                and os.path.exists(self.current_part_checkpoint_path)):
            try:
                os.remove(self.current_part_checkpoint_path)
            except OSError:
                pass
        self.current_part_checkpoint_path = path
        if not quiet:
            print(f"  Saved: {fname} ({self.chapters_in_part} chapters)")
        if not is_final:
            self.part += 1
            self.doc = Document()
            self.chapters_in_part = 0
            self.part_start_chapter = None
            self.part_end_chapter = None
            self.current_part_checkpoint_path = None

    def run(self):
        # Step 1: collect TOC
        toc = collect_toc_urls(self.driver, self.book_url)

        # Step 2: filter to requested range
        end = self.end_chapter if self.end_chapter else toc[-1][0]
        if self.end_chapter and self.end_chapter > toc[-1][0]:
            print(f"Requested end chapter {self.end_chapter}, but latest chapter "
                  f"found is {toc[-1][0]}; scraping through {toc[-1][0]}.")
        selected_chapters = [
            (num, url) for num, url in toc
            if self.start_chapter <= num <= end
        ]

        if not selected_chapters:
            print(f"No chapters found in range {self.start_chapter}–{end}. "
                  f"TOC covers chapters {toc[0][0]}–{toc[-1][0]}.")
            self.driver.quit()
            return

        completed_in_range = [
            num for num, _ in selected_chapters
            if num in self.completed_chapters
        ]
        if completed_in_range:
            print(f"Resuming from progress.json: skipping "
                  f"{len(completed_in_range)} completed chapter(s): "
                  f"{_format_number_ranges(completed_in_range)}")

        chapters_to_scrape = [
            (num, url) for num, url in selected_chapters
            if num not in self.completed_chapters
        ]

        if not chapters_to_scrape:
            print(f"All requested chapters {selected_chapters[0][0]}–"
                  f"{selected_chapters[-1][0]} are already complete.")
            self._save_progress(status="complete")
            self.driver.quit()
            return

        print(f"\nScraping chapters {chapters_to_scrape[0][0]}–{chapters_to_scrape[-1][0]} "
              f"({len(chapters_to_scrape)} pending of {len(selected_chapters)} selected)")
        print(f"Output dir: {self.output_dir}")
        print("-" * 50)

        try:
            for i, (chap_num, url) in enumerate(chapters_to_scrape):
                # Proactively recycle the browser to avoid renderer OOM.
                if self.total_scraped > 0 and self.total_scraped % RESTART_EVERY == 0:
                    print(f"  Recycling browser after {self.total_scraped} chapters "
                          f"to free memory...")
                    self._restart_driver()

                try:
                    title, content = self._scrape_one(chap_num, url)

                    if not content:
                        print(f"  Warning: no content for chapter {chap_num}")

                    if self.chapters_in_part == 0:
                        self.part_start_chapter = chap_num
                    self.part_end_chapter = chap_num

                    self.doc.add_heading(title, level=1)
                    for para in content:
                        self.doc.add_paragraph(para)

                    self.chapters_in_part += 1

                    if self.chapters_in_part >= self.chapters_per_part:
                        self._save_part()
                    else:
                        self._save_part(is_final=True, quiet=True)

                    self._mark_completed(chap_num, url)
                    self.total_scraped += 1
                    print(f"  [{self.total_scraped}/{len(chapters_to_scrape)}] {title}")

                    if i < len(chapters_to_scrape) - 1:
                        time.sleep(DELAY_BETWEEN_CHAPTERS)

                except Exception as e:
                    print(f"  Error on chapter {chap_num}: {e}")
                    self._mark_failed(chap_num, e)
                    continue

            self._save_part(is_final=True)
            missing = [
                num for num, _ in selected_chapters
                if num not in self.completed_chapters
            ]
            self._save_progress(status="complete" if not missing else "partial")

            print("\n" + "=" * 50)
            print(f"Done! Scraped {self.total_scraped} chapter(s) this run.")
            if missing:
                print(f"Still missing chapter(s): {_format_number_ranges(missing)}")
            print(f"Files saved in: {self.output_dir}")

        except KeyboardInterrupt:
            print("\nInterrupted by user.")
            self._save_part(is_final=True)
            self._save_progress(status="interrupted")
            print(f"Progress saved. Scraped {self.total_scraped} chapters.")

        finally:
            self.driver.quit()


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def parse_args():
    parser = argparse.ArgumentParser(
        description="Scrape chapters from freewebnovel.com using Table of Contents"
    )
    parser.add_argument("--url-file", default=None,
                        help=f"Text file with lines: \"Book Name\" <book-url> [start] [end] (default: {DEFAULT_URL_FILE} when present)")
    parser.add_argument("--book-url", default=None,
                        help="URL of the book's main page (e.g. https://freewebnovel.com/novel/goblin-dependency/)")
    parser.add_argument("--book-name", default=None,
                        help="Name used for output filenames")
    parser.add_argument("--start-chapter", type=int, default=1,
                        help="First chapter number to scrape (default: 1)")
    parser.add_argument("--end-chapter", type=int, default=None,
                        help="Last chapter number to scrape (default: all)")
    parser.add_argument("--chapters-per-part", type=int, default=CHAPTERS_PER_PART,
                        help=f"Chapters per .docx file (default: {CHAPTERS_PER_PART})")
    parser.add_argument("--output-dir", default=None,
                        help="Directory for output files (default: <script_dir>/ouput/<book_name>/)")
    args = parser.parse_args()

    if args.url_file:
        if args.book_url or args.book_name:
            parser.error("use either --url-file or --book-url/--book-name, not both")
        return args

    if not args.book_url and not args.book_name:
        default_url_file = resolve_path(DEFAULT_URL_FILE)
        if os.path.exists(default_url_file):
            args.url_file = DEFAULT_URL_FILE
            return args

    if not args.book_url or not args.book_name:
        parser.error("--book-url and --book-name are required unless --url-file is used")

    return args


def run_scraper(book_url, book_name, start_chapter, end_chapter,
                chapters_per_part, output_dir):
    scraper = None
    try:
        scraper = BookScraper(
            book_url=book_url,
            book_name=book_name,
            start_chapter=start_chapter,
            end_chapter=end_chapter,
            chapters_per_part=chapters_per_part,
            output_dir=output_dir,
        )
        scraper.run()
    finally:
        if scraper is not None:
            try:
                scraper.driver.quit()
            except Exception:
                pass


def main():
    args = parse_args()

    if args.url_file:
        entries = parse_url_file(args.url_file)
        failures = 0
        print(f"Loaded {len(entries)} book(s) from {resolve_path(args.url_file)}")

        for index, entry in enumerate(entries, start=1):
            print("\n" + "=" * 60)
            print(f"Book {index}/{len(entries)}: {entry['book_name']}")
            print("=" * 60)

            output_dir = args.output_dir
            if output_dir and len(entries) > 1:
                output_dir = os.path.join(output_dir, sanitize_filename(entry["book_name"]))

            try:
                run_scraper(
                    book_url=entry["book_url"],
                    book_name=entry["book_name"],
                    start_chapter=entry["start_chapter"],
                    end_chapter=entry["end_chapter"],
                    chapters_per_part=args.chapters_per_part,
                    output_dir=output_dir,
                )
            except KeyboardInterrupt:
                raise
            except Exception as e:
                failures += 1
                print(f"Error scraping '{entry['book_name']}': {e}")

        if failures:
            raise SystemExit(f"{failures} book(s) failed.")
        return

    run_scraper(
        book_url=args.book_url,
        book_name=args.book_name,
        start_chapter=args.start_chapter,
        end_chapter=args.end_chapter,
        chapters_per_part=args.chapters_per_part,
        output_dir=args.output_dir,
    )


if __name__ == "__main__":
    main()
