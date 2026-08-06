#!/usr/bin/env python3
"""Download Hengyan novels into resumable, chapter-batched DOCX files.

The default input is ``url.txt`` beside this script. Each non-comment line is:

    "Book name" http://www.hengyan.com/book/312 1 100

The two trailing chapter positions are optional. They are inclusive and refer
to the order in Hengyan's catalog, not to a number parsed from the title.
"""

from __future__ import annotations

import argparse
import json
import re
import shlex
import time
from dataclasses import asdict, dataclass
from http.cookiejar import CookieJar
from pathlib import Path
from typing import Any, Iterable, Optional
from urllib.error import HTTPError, URLError
from urllib.parse import urljoin, urlsplit, urlunsplit
from urllib.request import HTTPCookieProcessor, Request, build_opener

from bs4 import BeautifulSoup
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


SCRIPT_DIR = Path(__file__).resolve().parent
DEFAULT_URL_FILE = SCRIPT_DIR / "url.txt"
USER_AGENT = (
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
    "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/125.0 Safari/537.36"
)


@dataclass(frozen=True)
class Job:
    book_name: str
    book_url: str
    start_chapter: int = 1
    end_chapter: Optional[int] = None


@dataclass(frozen=True)
class ChapterRef:
    position: int
    title: str
    url: str


@dataclass(frozen=True)
class ChapterContent:
    position: int
    title: str
    url: str
    paragraphs: list[str]


def clean_space(text: str) -> str:
    return re.sub(r"[\t\r\f\v ]+", " ", text).strip()


def safe_name(name: str) -> str:
    name = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "", name).strip(" .")
    name = re.sub(r"\s+", "_", name)
    return name or "book"


def canonical_book_url(url: str) -> str:
    parsed = urlsplit(url.strip())
    scheme = parsed.scheme or "http"
    netloc = parsed.netloc.lower()
    path = parsed.path.rstrip("/")
    match = re.search(r"/book/(?:chapter/)?(\d+)$", path)
    if match:
        path = f"/book/{match.group(1)}"
    return urlunsplit((scheme, netloc, path, "", ""))


def book_id_from_url(url: str) -> str:
    match = re.search(r"/book/(?:chapter/|content/)?(\d+)(?:/|$)", urlsplit(url).path)
    if not match:
        raise ValueError(f"Could not find a Hengyan book ID in URL: {url}")
    return match.group(1)


def parse_optional_int(value: str, label: str, line_no: int) -> Optional[int]:
    if value.lower() in {"all", "none", "-"}:
        return None
    try:
        result = int(value)
    except ValueError as exc:
        raise ValueError(f"Line {line_no}: invalid {label} chapter {value!r}") from exc
    if result < 1:
        raise ValueError(f"Line {line_no}: {label} chapter must be at least 1")
    return result


def parse_url_file(path: Path) -> list[Job]:
    jobs: list[Job] = []
    with path.open("r", encoding="utf-8-sig") as handle:
        for line_no, raw in enumerate(handle, 1):
            line = raw.strip()
            if not line or line.startswith("#"):
                continue
            try:
                parts = shlex.split(line, comments=True)
            except ValueError as exc:
                raise ValueError(f"Line {line_no}: {exc}") from exc
            url_index = next(
                (i for i, part in enumerate(parts) if part.startswith(("http://", "https://"))),
                None,
            )
            if url_index is None or url_index == 0:
                raise ValueError(f"Line {line_no}: expected a book name followed by a URL")
            tail = parts[url_index + 1 :]
            if len(tail) > 2:
                raise ValueError(f"Line {line_no}: expected only start and end after the URL")
            start = parse_optional_int(tail[0], "start", line_no) if tail else 1
            end = parse_optional_int(tail[1], "end", line_no) if len(tail) == 2 else None
            if start is None:
                raise ValueError(f"Line {line_no}: start chapter cannot be open-ended")
            if end is not None and end < start:
                raise ValueError(f"Line {line_no}: end chapter cannot be before start chapter")
            jobs.append(Job(" ".join(parts[:url_index]), canonical_book_url(parts[url_index]), start, end))
    if not jobs:
        raise ValueError(f"No jobs found in {path}")
    return jobs


class Fetcher:
    def __init__(self, timeout: float, retries: int) -> None:
        self.timeout = timeout
        self.retries = retries
        self.opener = build_opener(HTTPCookieProcessor(CookieJar()))

    def get(self, url: str, referer: Optional[str] = None) -> tuple[str, str]:
        headers = {
            "User-Agent": USER_AGENT,
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
            "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.7",
            "Connection": "close",
        }
        if referer:
            headers["Referer"] = referer
        last_error: Optional[Exception] = None
        for attempt in range(1, self.retries + 1):
            try:
                with self.opener.open(Request(url, headers=headers), timeout=self.timeout) as response:
                    data = response.read()
                    charset = response.headers.get_content_charset() or "utf-8"
                    return data.decode(charset, "replace"), response.geturl()
            except (HTTPError, URLError, TimeoutError, OSError) as exc:
                last_error = exc
                if attempt < self.retries:
                    wait = min(2 ** attempt, 8)
                    print(f"  Request failed ({attempt}/{self.retries}): {exc}; retrying in {wait}s")
                    time.sleep(wait)
        raise RuntimeError(f"Failed to fetch {url}: {last_error}")


def catalog_url(book_url: str) -> str:
    book_id = book_id_from_url(book_url)
    parsed = urlsplit(book_url)
    return urlunsplit((parsed.scheme, parsed.netloc, f"/book/chapter/{book_id}", "", ""))


def parse_catalog(html: str, page_url: str, expected_book_id: str) -> list[ChapterRef]:
    soup = BeautifulSoup(html, "html.parser")
    candidates = soup.select(".article .list .txt a[href], .list a[href]")
    if not candidates:
        candidates = soup.select("a[href*='/book/content/']")

    pattern = re.compile(rf"/book/content/{re.escape(expected_book_id)}/\d+/?$")
    seen: set[str] = set()
    chapters: list[ChapterRef] = []
    for anchor in candidates:
        url = urljoin(page_url, str(anchor.get("href", "")))
        if not pattern.search(urlsplit(url).path) or url in seen:
            continue
        title = clean_space(anchor.get_text(" ", strip=True))
        if not title:
            continue
        seen.add(url)
        chapters.append(ChapterRef(len(chapters) + 1, title, url))
    if not chapters:
        raise RuntimeError("No chapter links found on Hengyan's catalog page")
    return chapters


def parse_chapter(html: str, page_url: str, position: int, title_hint: str) -> ChapterContent:
    soup = BeautifulSoup(html, "html.parser")
    reading = soup.select_one(".reading")
    title_node = (reading.select_one(".title h3") if reading else None) or soup.select_one("h3")
    title = clean_space(title_node.get_text(" ", strip=True)) if title_node else title_hint
    content = (reading.select_one(".cont .text") if reading else None) or soup.select_one(".cont .text")
    if content is None:
        raise RuntimeError(f"Chapter content container was not found at {page_url}")

    for unwanted in content.select("script, style, .pageings, ins, iframe"):
        unwanted.decompose()
    paragraphs = [clean_space(node.get_text(" ", strip=True)) for node in content.select("p")]
    paragraphs = [paragraph for paragraph in paragraphs if paragraph]
    if not paragraphs:
        paragraphs = [clean_space(line) for line in content.get_text("\n").splitlines()]
        paragraphs = [paragraph for paragraph in paragraphs if paragraph]
    if not paragraphs or sum(map(len, paragraphs)) < 20:
        raise RuntimeError(f"Chapter content was empty or unexpectedly short at {page_url}")
    return ChapterContent(position, title or title_hint, page_url, paragraphs)


def read_json(path: Path, default: Any) -> Any:
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return default


def write_json(path: Path, payload: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_suffix(path.suffix + ".tmp")
    temporary.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    temporary.replace(path)


def set_east_asian_font(run, font_name: str = "SimSun", size: int = 12) -> None:
    run.font.name = font_name
    run.font.size = Pt(size)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font_name)


def enable_update_fields(document: Document) -> None:
    settings = document.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def save_docx(path: Path, book_name: str, chapters: Iterable[ChapterContent]) -> None:
    chapter_list = list(chapters)
    document = Document()
    enable_update_fields(document)
    normal = document.styles["Normal"]
    normal.font.name = "SimSun"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

    heading = document.add_heading(book_name, level=0)
    for run in heading.runs:
        set_east_asian_font(run, "SimHei", 18)
    document.add_paragraph(
        f"Hengyan catalog chapters {chapter_list[0].position}-{chapter_list[-1].position}"
    )

    for index, chapter in enumerate(chapter_list):
        if index:
            document.add_page_break()
        heading = document.add_heading(chapter.title, level=1)
        for run in heading.runs:
            set_east_asian_font(run, "SimHei", 16)
        for text in chapter.paragraphs:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.first_line_indent = Pt(24)
            paragraph.paragraph_format.space_after = Pt(3)
            set_east_asian_font(paragraph.add_run(text))

    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_suffix(".tmp.docx")
    document.save(temporary)
    temporary.replace(path)


def chunked(items: list[ChapterRef], size: int) -> Iterable[list[ChapterRef]]:
    for offset in range(0, len(items), size):
        yield items[offset : offset + size]


def load_or_fetch_catalog(
    job: Job, book_dir: Path, fetcher: Fetcher, refresh: bool
) -> list[ChapterRef]:
    cache_path = book_dir / "chapter_urls.json"
    cached = read_json(cache_path, {}) if not refresh else {}
    if cached.get("book_url") == job.book_url and isinstance(cached.get("chapters"), list):
        try:
            chapters = [ChapterRef(**item) for item in cached["chapters"]]
        except (TypeError, KeyError):
            chapters = []
        if chapters:
            print(f"Loaded {len(chapters)} cached chapter links")
            return chapters

    url = catalog_url(job.book_url)
    print(f"Fetching catalog: {url}")
    html, final_url = fetcher.get(url, referer=job.book_url)
    chapters = parse_catalog(html, final_url, book_id_from_url(job.book_url))
    write_json(cache_path, {"book_url": job.book_url, "chapters": [asdict(c) for c in chapters]})
    print(f"Found {len(chapters)} chapters")
    return chapters


def run_job(job: Job, args: argparse.Namespace, fetcher: Fetcher) -> None:
    book_dir = args.output_dir / safe_name(job.book_name)
    book_dir.mkdir(parents=True, exist_ok=True)
    catalog = load_or_fetch_catalog(job, book_dir, fetcher, args.refresh_toc)
    end = min(job.end_chapter or len(catalog), len(catalog))
    if job.start_chapter > len(catalog):
        raise ValueError(
            f"Start chapter {job.start_chapter} exceeds catalog length {len(catalog)} for {job.book_name}"
        )
    selected = catalog[job.start_chapter - 1 : end]
    if not selected:
        raise ValueError(f"The selected chapter range is empty for {job.book_name}")
    print(f"Selected catalog positions {selected[0].position}-{selected[-1].position}")

    progress_path = book_dir / "progress.json"
    pending_path = book_dir / ".pending_batch.json"
    progress = read_json(progress_path, {})
    if progress and progress.get("book_url") not in {None, job.book_url}:
        raise RuntimeError(f"{progress_path} belongs to a different book URL")

    for refs in chunked(selected, args.chapters_per_file):
        first, last = refs[0].position, refs[-1].position
        output_path = book_dir / f"{safe_name(job.book_name)}_chapters_{first:04d}-{last:04d}.docx"
        if output_path.exists() and not args.overwrite:
            print(f"Skipping existing {output_path.name}")
            continue

        pending = read_json(pending_path, {}) if args.resume else {}
        if (
            pending.get("book_url") == job.book_url
            and pending.get("first") == first
            and pending.get("last") == last
        ):
            saved_items = pending.get("chapters", [])
            try:
                collected = [ChapterContent(**item) for item in saved_items]
            except (TypeError, KeyError):
                collected = []
            if collected:
                print(f"Resuming batch {first}-{last} after {len(collected)} saved chapter(s)")
        else:
            collected = []

        for ref in refs[len(collected) :]:
            print(f"[{ref.position}/{end}] {ref.title}")
            html, final_url = fetcher.get(ref.url, referer=job.book_url)
            chapter = parse_chapter(html, final_url, ref.position, ref.title)
            collected.append(chapter)
            write_json(
                pending_path,
                {
                    "book_url": job.book_url,
                    "first": first,
                    "last": last,
                    "chapters": [asdict(item) for item in collected],
                },
            )
            if args.delay:
                time.sleep(args.delay)

        save_docx(output_path, job.book_name, collected)
        pending_path.unlink(missing_ok=True)
        write_json(
            progress_path,
            {
                "book_name": job.book_name,
                "book_url": job.book_url,
                "completed_through": last,
                "last_output": output_path.name,
            },
        )
        print(f"Saved {output_path}")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Scrape Hengyan chapter content into DOCX files")
    parser.add_argument("--url-file", type=Path, help="Defaults to url.txt beside this script")
    parser.add_argument("--book-url", help="Run one book directly instead of reading url.txt")
    parser.add_argument("--book-name", help="Required with --book-url")
    parser.add_argument("--start-chapter", type=int, default=1)
    parser.add_argument("--end-chapter", type=int)
    parser.add_argument("--output-dir", type=Path, default=SCRIPT_DIR)
    parser.add_argument("--chapters-per-file", type=int, default=100)
    parser.add_argument("--delay", type=float, default=1.0, help="Seconds between chapter requests")
    parser.add_argument("--timeout", type=float, default=30.0)
    parser.add_argument("--retries", type=int, default=3)
    parser.add_argument("--refresh-toc", action="store_true")
    parser.add_argument("--no-resume", dest="resume", action="store_false")
    parser.add_argument("--overwrite", action="store_true", help="Replace existing DOCX batches")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    if args.chapters_per_file < 1 or args.retries < 1 or args.timeout <= 0 or args.delay < 0:
        raise SystemExit("chapters-per-file/retries must be positive; timeout positive; delay non-negative")
    args.output_dir = args.output_dir.expanduser().resolve()

    if args.book_url:
        if not args.book_name:
            raise SystemExit("--book-name is required with --book-url")
        if args.start_chapter < 1 or (args.end_chapter and args.end_chapter < args.start_chapter):
            raise SystemExit("Invalid direct chapter range")
        jobs = [
            Job(
                args.book_name,
                canonical_book_url(args.book_url),
                args.start_chapter,
                args.end_chapter,
            )
        ]
    else:
        url_file = (args.url_file or DEFAULT_URL_FILE).expanduser().resolve()
        jobs = parse_url_file(url_file)

    fetcher = Fetcher(args.timeout, args.retries)
    for number, job in enumerate(jobs, 1):
        if len(jobs) > 1:
            print(f"\n=== Job {number}/{len(jobs)}: {job.book_name} ===")
        run_job(job, args, fetcher)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
