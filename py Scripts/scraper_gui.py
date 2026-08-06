#!/usr/bin/env python3
"""
scraper_gui.py — Local GUI for the Novel Scraper pipeline.

No N8N required. Paste a URL, detect the platform, configure chapter range,
and hit Scrape. Supports all 17 platforms from layer1_router.

Run from the py Scripts/ directory:
    python scraper_gui.py
"""

import sys
import os
import re
import time
import json
import queue
import threading
import subprocess
from pathlib import Path

# Ensure this script's directory is on sys.path (for layer1_router import)
SCRIPT_DIR = Path(__file__).parent.resolve()
sys.path.insert(0, str(SCRIPT_DIR))

import tkinter as tk
from tkinter import ttk, filedialog, messagebox

# ── Dependency checks ────────────────────────────────────────────────────────

try:
    import requests
    from bs4 import BeautifulSoup
    HTTP_DEPS_OK = True
    HTTP_DEPS_ERROR = ""
except ImportError as e:
    HTTP_DEPS_OK = False
    HTTP_DEPS_ERROR = str(e)

try:
    from docx import Document
    DOCX_OK = True
    DOCX_ERROR = ""
except ImportError as e:
    DOCX_OK = False
    DOCX_ERROR = str(e)

try:
    from layer1_router import route as layer1_route
    ROUTER_OK = True
    ROUTER_ERROR = ""
except Exception as e:
    ROUTER_OK = False
    ROUTER_ERROR = str(e)


# ── HTTP headers ─────────────────────────────────────────────────────────────

HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/124.0.0.0 Safari/537.36"
    ),
    "Referer": "https://www.google.com/",
    "Accept-Language": "en-US,en;q=0.9",
}

CHAPTERS_PER_FILE = 100


# ── Inline scrapers (requests-based, no browser) ─────────────────────────────

def safe_book_dir_name(name: str) -> str:
    return re.sub(r'[<>:"/\\|?*\x00-\x1f]', "", name).strip() or "Novel"


def scrape_novelbin(start_url, book_name, start_ch, end_ch, out_dir, log, stop):
    """Full novelbin scraper — requests + BeautifulSoup."""
    safe_name = safe_book_dir_name(book_name)
    book_dir = os.path.join(out_dir, safe_name)
    os.makedirs(book_dir, exist_ok=True)
    log(f"Output: {book_dir}")

    def file_path(idx):
        s = (idx - 1) * CHAPTERS_PER_FILE + 1
        e = s + CHAPTERS_PER_FILE - 1
        return os.path.join(book_dir, f"{safe_name} {s}-{e}.docx")

    ch_num = start_ch
    file_idx = (ch_num - 1) // CHAPTERS_PER_FILE + 1
    ch_in_file = (ch_num - 1) % CHAPTERS_PER_FILE
    out_path = file_path(file_idx)
    doc = Document(out_path) if os.path.exists(out_path) else Document()
    url = start_url

    while url and ch_num <= end_ch:
        if stop.is_set():
            doc.save(out_path)
            log("Stopped.")
            return

        log(f"Ch {ch_num}: {url}")

        for attempt in range(1, 4):
            try:
                resp = requests.get(url, headers=HEADERS, timeout=15)
                if resp.status_code != 200:
                    if attempt < 3:
                        log(f"  HTTP {resp.status_code}, retry {attempt}/3...")
                        time.sleep(3 * attempt)
                        continue
                    log(f"  HTTP {resp.status_code}, skipping.")
                    ch_num += 1
                    break

                soup = BeautifulSoup(resp.content, "html.parser")

                title_tag = soup.find("span", class_="chr-text")
                title = title_tag.get_text(strip=True) if title_tag else f"Chapter {ch_num}"
                doc.add_heading(title, level=1)

                content_div = soup.find("div", id="chr-content")
                if not content_div:
                    if attempt < 3:
                        log(f"  No content, retry {attempt}/3...")
                        time.sleep(3 * attempt)
                        continue
                    log("  No content found, skipping.")
                else:
                    count = 0
                    for p in content_div.find_all("p"):
                        for br in p.find_all("br"):
                            br.replace_with("\n")
                        text = p.get_text(separator="", strip=True)
                        if text:
                            doc.add_paragraph(text)
                            count += 1
                    log(f"  Saved: {title} ({count} paragraphs)")

                doc.save(out_path)
                ch_num += 1
                ch_in_file += 1

                next_link = soup.find(
                    "a", class_="js-chapter-nav",
                    attrs={"data-chapter-nav": "next"}
                )
                if (not next_link
                        or next_link.get("disabled") is not None
                        or "disabled" in (next_link.get("class") or [])):
                    log("No more chapters.")
                    url = None
                    break

                url = next_link.get("data-chapter-url") or next_link.get("href")
                time.sleep(1)

                if ch_in_file >= CHAPTERS_PER_FILE:
                    log(f"100 chapters done, starting new file...")
                    file_idx += 1
                    ch_in_file = 0
                    out_path = file_path(file_idx)
                    doc = Document(out_path) if os.path.exists(out_path) else Document()

                break

            except Exception as exc:
                if attempt < 3:
                    log(f"  Error (attempt {attempt}/3): {exc}, retrying...")
                    time.sleep(5 * attempt)
                else:
                    log(f"  Skipping chapter {ch_num}: {exc}")
                    ch_num += 1
                    break

    log(f"Done. Files in: {book_dir}")


def scrape_royalroad(start_url, book_name, start_ch, end_ch, out_dir, log, stop):
    """RoyalRoad via requests — follows next-chapter links."""
    safe_name = safe_book_dir_name(book_name)
    book_dir = os.path.join(out_dir, safe_name)
    os.makedirs(book_dir, exist_ok=True)
    log(f"Output: {book_dir}")

    def file_path(idx):
        s = (idx - 1) * CHAPTERS_PER_FILE + 1
        e = s + CHAPTERS_PER_FILE - 1
        return os.path.join(book_dir, f"{safe_name} {s}-{e}.docx")

    ch_num = start_ch
    file_idx = (ch_num - 1) // CHAPTERS_PER_FILE + 1
    ch_in_file = (ch_num - 1) % CHAPTERS_PER_FILE
    out_path = file_path(file_idx)
    doc = Document(out_path) if os.path.exists(out_path) else Document()
    url = start_url

    while url and ch_num <= end_ch:
        if stop.is_set():
            doc.save(out_path)
            log("Stopped.")
            return

        log(f"Ch {ch_num}: {url}")

        for attempt in range(1, 4):
            try:
                resp = requests.get(url, headers=HEADERS, timeout=15)
                if resp.status_code != 200:
                    if attempt < 3:
                        log(f"  HTTP {resp.status_code}, retry {attempt}/3...")
                        time.sleep(3 * attempt)
                        continue
                    log(f"  HTTP {resp.status_code}, skipping.")
                    ch_num += 1
                    break

                soup = BeautifulSoup(resp.content, "html.parser")

                # Title
                h1 = soup.find("h1") or soup.find("h2")
                title = h1.get_text(strip=True) if h1 else f"Chapter {ch_num}"
                doc.add_heading(title, level=1)

                # Content
                content_div = soup.find("div", class_=lambda c: c and "chapter-inner" in c)
                if not content_div:
                    content_div = soup.find("div", class_="chapter-content")
                if not content_div:
                    if attempt < 3:
                        log(f"  No content, retry {attempt}/3...")
                        time.sleep(3 * attempt)
                        continue
                    log("  No content found, skipping.")
                else:
                    count = 0
                    for p in content_div.find_all("p"):
                        text = p.get_text(strip=True)
                        if text:
                            doc.add_paragraph(text)
                            count += 1
                    log(f"  Saved: {title} ({count} paragraphs)")

                doc.save(out_path)
                ch_num += 1
                ch_in_file += 1

                # Next chapter
                next_btn = soup.find(
                    "a",
                    class_=lambda c: c and "btn-primary" in c,
                    string=re.compile(r"next", re.I)
                )
                if not next_btn:
                    next_btn = soup.find(
                        "a",
                        href=re.compile(r"/chapter/"),
                        string=re.compile(r"next", re.I)
                    )
                if not next_btn:
                    log("No more chapters.")
                    url = None
                    break

                href = next_btn.get("href", "")
                url = href if href.startswith("http") else f"https://www.royalroad.com{href}"
                time.sleep(1)

                if ch_in_file >= CHAPTERS_PER_FILE:
                    log("100 chapters done, starting new file...")
                    file_idx += 1
                    ch_in_file = 0
                    out_path = file_path(file_idx)
                    doc = Document(out_path) if os.path.exists(out_path) else Document()

                break

            except Exception as exc:
                if attempt < 3:
                    log(f"  Error (attempt {attempt}/3): {exc}, retrying...")
                    time.sleep(5 * attempt)
                else:
                    log(f"  Skipping chapter {ch_num}: {exc}")
                    ch_num += 1
                    break

    log(f"Done. Files in: {book_dir}")


# ── Subprocess runner for browser-based scrapers ─────────────────────────────

def run_subprocess(cmd, log, stop, cwd=None):
    """Run a command, stream output to log(), check stop event."""
    log(f"Running: {' '.join(str(c) for c in cmd)}")
    try:
        proc = subprocess.Popen(
            [str(c) for c in cmd],
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
            text=True,
            cwd=str(cwd or SCRIPT_DIR),
            bufsize=1,
        )
        for line in proc.stdout:
            if stop.is_set():
                proc.terminate()
                log("Stopped.")
                return -1
            log(line.rstrip())
        proc.wait()
        return proc.returncode
    except Exception as exc:
        log(f"Failed to run script: {exc}")
        return -1


# ── Platform dispatch table ───────────────────────────────────────────────────
# Maps platform name -> how to run it.
# "inline"    : handled in Python directly (requests-based)
# "argparse"  : script accepts CLI args
# "hardcoded" : script has config vars at top — needs manual config

PLATFORM_MODE = {
    "novelbin":     "inline",
    "royalroad":    "inline",
    "freewebnovel": "argparse",
    "webnovel":     "hardcoded",
    "69shuba":      "hardcoded",
    "babelnovel":   "hardcoded",
    "tapas":        "hardcoded",
    "wuxiaworld":   "hardcoded",
    "wattpad":      "hardcoded",
    "kakao":        "hardcoded",
    "qidian":       "hardcoded",
    "qdmm":         "hardcoded",
    "hengyan":      "hardcoded",
}


# ── Main Application ──────────────────────────────────────────────────────────

class ScraperApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Novel Scraper")
        self.geometry("980x760")
        self.minsize(800, 600)
        self.configure(bg="#1e1e2e")

        self._result = None
        self._stop_event = threading.Event()
        self._log_queue = queue.Queue()

        self._setup_styles()
        self._build_ui()
        self._poll_log()

        # Startup messages
        if not HTTP_DEPS_OK:
            self._log(f"Missing deps: {HTTP_DEPS_ERROR}")
            self._log("Run: pip install requests beautifulsoup4 python-docx")
        if not DOCX_OK:
            self._log(f"Missing python-docx: {DOCX_ERROR}")
        if not ROUTER_OK:
            self._log(f"layer1_router unavailable: {ROUTER_ERROR}")
            self._log("Ensure layer1_router.py is in the same directory.")
        if HTTP_DEPS_OK and DOCX_OK and ROUTER_OK:
            self._log("Ready. Paste a novel URL above and click Detect.")

    # ── Styles ────────────────────────────────────────────────────────────────

    def _setup_styles(self):
        s = ttk.Style(self)
        s.theme_use("clam")

        BG    = "#1e1e2e"
        FG    = "#cdd6f4"
        MUTED = "#6c7086"
        PANEL = "#313244"
        BLUE  = "#89b4fa"
        GREEN = "#a6e3a1"
        RED   = "#f38ba8"

        s.configure(".",         background=BG, foreground=FG,  font=("Helvetica", 12))
        s.configure("TFrame",    background=BG)
        s.configure("TLabel",    background=BG, foreground=FG)
        s.configure("TButton",   background=PANEL, foreground=FG, padding=(10, 5), relief="flat")
        s.map("TButton",         background=[("active", "#45475a"), ("disabled", "#181825")])
        s.configure("Go.TButton",    background=BLUE,  foreground="#1e1e2e", padding=(12, 6), font=("Helvetica", 12, "bold"))
        s.map("Go.TButton",          background=[("active", "#74c7ec"), ("disabled", "#313244")])
        s.configure("Stop.TButton",  background=RED,   foreground="#1e1e2e", padding=(10, 5))
        s.map("Stop.TButton",        background=[("active", "#eba0ac"), ("disabled", "#313244")])
        s.configure("TEntry",        fieldbackground=PANEL, foreground=FG,  insertcolor=FG, relief="flat")
        s.configure("TSpinbox",      fieldbackground=PANEL, foreground=FG,  relief="flat")
        s.configure("TLabelframe",         background=BG,    foreground=BLUE)
        s.configure("TLabelframe.Label",   background=BG,    foreground=BLUE, font=("Helvetica", 10, "bold"))

        self._colors = {"bg": BG, "fg": FG, "panel": PANEL,
                        "blue": BLUE, "green": GREEN, "red": RED, "muted": MUTED}

    # ── UI Layout ─────────────────────────────────────────────────────────────

    def _build_ui(self):
        C = self._colors

        # ── URL row ──────────────────────────────────────────────────────────
        url_row = tk.Frame(self, bg=C["bg"], padx=14, pady=10)
        url_row.pack(fill=tk.X)

        tk.Label(url_row, text="URL", bg=C["bg"], fg=C["muted"],
                 font=("Helvetica", 11, "bold")).pack(side=tk.LEFT, padx=(0, 8))

        self._url_var = tk.StringVar()
        url_entry = tk.Entry(url_row, textvariable=self._url_var,
                             bg=C["panel"], fg=C["fg"], insertbackground=C["fg"],
                             relief=tk.FLAT, font=("Helvetica", 12),
                             highlightthickness=1, highlightbackground="#45475a",
                             highlightcolor=C["blue"])
        url_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, ipady=5)
        url_entry.bind("<Return>", lambda _: self._detect())

        self._detect_btn = ttk.Button(url_row, text="Detect", command=self._detect)
        self._detect_btn.pack(side=tk.LEFT, padx=(8, 0))

        # ── Middle section ────────────────────────────────────────────────────
        mid = tk.Frame(self, bg=C["bg"])
        mid.pack(fill=tk.BOTH, expand=False, padx=14, pady=(0, 6))

        # Left: Book info
        info_outer = tk.Frame(mid, bg=C["panel"], padx=14, pady=12)
        info_outer.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 6))

        self._badge_lbl = tk.Label(
            info_outer, text="", bg="#313244", fg=C["blue"],
            font=("Helvetica", 10, "bold"), padx=8, pady=3, anchor=tk.W
        )
        self._badge_lbl.pack(fill=tk.X, pady=(0, 8))

        self._title_lbl = tk.Label(
            info_outer, text="Paste a URL and click Detect",
            bg=C["panel"], fg=C["fg"],
            font=("Helvetica", 14, "bold"),
            wraplength=420, justify=tk.LEFT, anchor=tk.W
        )
        self._title_lbl.pack(fill=tk.X)

        self._author_lbl = tk.Label(
            info_outer, text="", bg=C["panel"], fg=C["green"],
            font=("Helvetica", 11), anchor=tk.W
        )
        self._author_lbl.pack(fill=tk.X, pady=(3, 0))

        self._ch_count_lbl = tk.Label(
            info_outer, text="", bg=C["panel"], fg=C["muted"],
            font=("Helvetica", 10), anchor=tk.W
        )
        self._ch_count_lbl.pack(fill=tk.X, pady=(2, 0))

        self._desc_lbl = tk.Label(
            info_outer, text="", bg=C["panel"], fg="#bac2de",
            font=("Helvetica", 10),
            wraplength=420, justify=tk.LEFT, anchor=tk.W
        )
        self._desc_lbl.pack(fill=tk.X, pady=(8, 0))

        # Translation candidates (shown when needed)
        self._trans_frame = tk.Frame(info_outer, bg=C["panel"])
        # packed dynamically

        # Right: Config
        cfg_outer = tk.Frame(mid, bg=C["panel"], padx=14, pady=12)
        cfg_outer.pack(side=tk.LEFT, fill=tk.Y, padx=(6, 0))

        tk.Label(cfg_outer, text="SCRAPE CONFIG", bg=C["panel"],
                 fg=C["muted"], font=("Helvetica", 9, "bold")).grid(
            row=0, column=0, columnspan=2, sticky=tk.W, pady=(0, 10))

        def cfg_row(row, label, widget_fn):
            tk.Label(cfg_outer, text=label, bg=C["panel"], fg=C["fg"],
                     font=("Helvetica", 11)).grid(row=row, column=0, sticky=tk.W, pady=5, padx=(0, 10))
            w = widget_fn(cfg_outer)
            w.grid(row=row, column=1, sticky=tk.EW, pady=5)
            return w

        self._start_var = tk.IntVar(value=1)
        self._end_var   = tk.IntVar(value=100)

        cfg_row(1, "Start Ch.", lambda p: ttk.Spinbox(
            p, from_=1, to=99999, textvariable=self._start_var, width=9))
        cfg_row(2, "End Ch.",   lambda p: ttk.Spinbox(
            p, from_=1, to=99999, textvariable=self._end_var,   width=9))

        # Output dir row (spans full width)
        tk.Label(cfg_outer, text="Output", bg=C["panel"], fg=C["fg"],
                 font=("Helvetica", 11)).grid(row=3, column=0, sticky=tk.W, pady=5, padx=(0, 10))
        out_sub = tk.Frame(cfg_outer, bg=C["panel"])
        out_sub.grid(row=3, column=1, sticky=tk.EW, pady=5)
        self._outdir_var = tk.StringVar(value=str(Path.home() / "Downloads" / "Novels"))
        tk.Entry(out_sub, textvariable=self._outdir_var,
                 bg=C["panel"], fg=C["fg"], insertbackground=C["fg"],
                 relief=tk.FLAT, font=("Helvetica", 10),
                 highlightthickness=1, highlightbackground="#45475a",
                 width=22).pack(side=tk.LEFT, ipady=3)
        ttk.Button(out_sub, text="…", width=3,
                   command=self._browse_out).pack(side=tk.LEFT, padx=(4, 0))

        # Scraper info label
        self._scraper_info = tk.Label(
            cfg_outer, text="", bg=C["panel"], fg=C["muted"],
            font=("Helvetica", 9), wraplength=220, justify=tk.LEFT
        )
        self._scraper_info.grid(row=4, column=0, columnspan=2, pady=(10, 0), sticky=tk.W)

        # Buttons
        btn_frame = tk.Frame(cfg_outer, bg=C["panel"])
        btn_frame.grid(row=5, column=0, columnspan=2, pady=(14, 0), sticky=tk.EW)

        self._run_btn = ttk.Button(
            btn_frame, text="  Scrape", style="Go.TButton",
            command=self._start_scrape, state=tk.DISABLED
        )
        self._run_btn.pack(fill=tk.X, pady=(0, 6))

        self._stop_btn = ttk.Button(
            btn_frame, text="  Stop", style="Stop.TButton",
            command=self._stop_scrape, state=tk.DISABLED
        )
        self._stop_btn.pack(fill=tk.X)

        # Open output folder button
        self._open_btn = ttk.Button(
            btn_frame, text="Open Output Folder",
            command=self._open_output_folder
        )
        self._open_btn.pack(fill=tk.X, pady=(6, 0))

        # ── Log ──────────────────────────────────────────────────────────────
        log_frame = ttk.LabelFrame(self, text="Log", padding=(8, 4))
        log_frame.pack(fill=tk.BOTH, expand=True, padx=14, pady=(0, 8))

        self._log_text = tk.Text(
            log_frame,
            state=tk.DISABLED,
            bg="#11111b", fg="#cdd6f4",
            font=("Menlo", 10),
            wrap=tk.WORD, relief=tk.FLAT,
            padx=8, pady=6,
            selectbackground="#313244",
        )
        sb = ttk.Scrollbar(log_frame, command=self._log_text.yview)
        self._log_text.configure(yscrollcommand=sb.set)
        sb.pack(side=tk.RIGHT, fill=tk.Y)
        self._log_text.pack(fill=tk.BOTH, expand=True)

        # ── Status bar ───────────────────────────────────────────────────────
        self._status_var = tk.StringVar(value="Ready")
        tk.Label(self, textvariable=self._status_var, bg=C["bg"], fg=C["muted"],
                 font=("Helvetica", 10), anchor=tk.W,
                 padx=14, pady=4).pack(fill=tk.X, anchor=tk.W)

    # ── Log helpers ───────────────────────────────────────────────────────────

    def _log(self, msg: str):
        """Thread-safe log append."""
        self._log_queue.put(str(msg) + "\n")

    def _poll_log(self):
        try:
            while True:
                msg = self._log_queue.get_nowait()
                self._log_text.configure(state=tk.NORMAL)
                self._log_text.insert(tk.END, msg)
                self._log_text.see(tk.END)
                self._log_text.configure(state=tk.DISABLED)
        except queue.Empty:
            pass
        self.after(80, self._poll_log)

    def _status(self, msg: str):
        self._status_var.set(msg)

    # ── Utility ───────────────────────────────────────────────────────────────

    def _browse_out(self):
        d = filedialog.askdirectory(initialdir=self._outdir_var.get())
        if d:
            self._outdir_var.set(d)

    def _open_output_folder(self):
        folder = self._outdir_var.get()
        if os.path.isdir(folder):
            subprocess.Popen(["open", folder])
        else:
            messagebox.showinfo("Not found", f"Folder does not exist yet:\n{folder}")

    # ── Detection ─────────────────────────────────────────────────────────────

    def _detect(self):
        url = self._url_var.get().strip()
        if not url:
            messagebox.showwarning("No URL", "Enter a novel URL first.")
            return
        if not ROUTER_OK:
            self._log(f"layer1_router not available: {ROUTER_ERROR}")
            return

        self._detect_btn.configure(state=tk.DISABLED)
        self._run_btn.configure(state=tk.DISABLED)
        self._status("Detecting...")
        self._log(f"\nDetecting: {url}")

        def _worker():
            try:
                res = layer1_route(url, search_translations=True)
                self.after(0, lambda: self._on_detected(res))
            except Exception as exc:
                self.after(0, lambda: self._on_detect_error(str(exc)))

        threading.Thread(target=_worker, daemon=True).start()

    def _on_detected(self, result: dict):
        self._detect_btn.configure(state=tk.NORMAL)
        self._result = result
        C = self._colors

        if result.get("status") != "ok":
            err = result.get("error", "Unknown error")
            self._title_lbl.configure(text=f"Error: {err}")
            self._badge_lbl.configure(text="UNKNOWN PLATFORM")
            self._status("Detection failed")
            self._log(f"Error: {err}")
            return

        platform    = result.get("platform", "unknown").upper()
        lang        = result.get("source_language", "?").upper()
        needs_trans = result.get("needs_translation", False)
        needs_login = result.get("needs_login", False)
        needs_br    = result.get("needs_browser", True)
        ch_count    = result.get("chapter_count", 0)
        book_name   = result.get("book_name") or ""
        author      = result.get("author") or ""
        desc        = (result.get("description") or "")[:250]
        if len(result.get("description") or "") > 250:
            desc += "…"
        scraper     = result.get("scraper_script") or "?"

        # Badge
        badge_parts = [platform, lang]
        if needs_trans:
            badge_parts.append("needs EN translation")
        self._badge_lbl.configure(text="  ".join(badge_parts))

        self._title_lbl.configure(
            text=book_name or "(title not found — check URL / site may block bots)")
        self._author_lbl.configure(text=author)
        self._ch_count_lbl.configure(
            text=f"{ch_count:,} chapters detected" if ch_count else "chapter count unknown")
        self._desc_lbl.configure(text=desc)

        if ch_count:
            self._end_var.set(ch_count)

        # Scraper info
        mode = PLATFORM_MODE.get(result.get("platform", ""), "hardcoded")
        mode_label = {
            "inline":    "built-in (requests, fast)",
            "argparse":  "subprocess with CLI args",
            "hardcoded": "browser script (Chrome required)",
        }[mode]
        info_lines = [
            f"Script: {scraper}",
            f"Method: {mode_label}",
        ]
        if needs_login:
            info_lines.append("Requires login (see script)")
        self._scraper_info.configure(text="\n".join(info_lines))

        # Translation candidates
        # Clear previous
        for w in self._trans_frame.winfo_children():
            w.destroy()
        candidates = result.get("translation_candidates", [])
        if candidates:
            self._trans_frame.pack(fill=tk.X, pady=(12, 0))
            tk.Label(self._trans_frame, text="English translation candidates:",
                     bg=C["panel"], fg=C["muted"],
                     font=("Helvetica", 9, "bold")).pack(anchor=tk.W)
            for c in candidates[:6]:
                src  = c.get("source", "?")
                ttl  = c.get("title", "")
                curl = c.get("url", "")
                row  = tk.Frame(self._trans_frame, bg=C["panel"])
                row.pack(fill=tk.X, pady=1)
                tk.Label(row, text=f"[{src}] {ttl}",
                         bg=C["panel"], fg="#89dceb",
                         font=("Helvetica", 9), cursor="hand2").pack(side=tk.LEFT)
                use_btn = ttk.Button(row, text="Use this URL",
                                     command=lambda u=curl: self._use_translation(u))
                use_btn.pack(side=tk.RIGHT)
        else:
            self._trans_frame.pack_forget()

        self._run_btn.configure(state=tk.NORMAL)
        self._status(f"Detected: {platform} — {book_name}")
        self._log(f"Platform: {platform} | Lang: {lang} | Chapters: {ch_count}")
        if needs_trans and candidates:
            self._log(f"Found {len(candidates)} EN translation candidates (shown above).")
        elif needs_trans:
            self._log("No EN translations found — would need direct scrape + translation.")

    def _on_detect_error(self, err: str):
        self._detect_btn.configure(state=tk.NORMAL)
        self._title_lbl.configure(text="Detection failed")
        self._status("Error")
        self._log(f"Detection error: {err}")

    def _use_translation(self, url: str):
        """Load a translation candidate URL into the URL bar and re-detect."""
        self._url_var.set(url)
        self._log(f"\nSwitching to translation URL: {url}")
        self._detect()

    # ── Scraping ──────────────────────────────────────────────────────────────

    def _start_scrape(self):
        if not self._result:
            return
        out_dir = self._outdir_var.get().strip()
        if not out_dir:
            messagebox.showwarning("No output", "Set an output directory first.")
            return

        start_ch = self._start_var.get()
        end_ch   = self._end_var.get()
        if start_ch > end_ch:
            messagebox.showwarning("Invalid range", "Start chapter must be <= end chapter.")
            return

        os.makedirs(out_dir, exist_ok=True)

        self._stop_event.clear()
        self._run_btn.configure(state=tk.DISABLED)
        self._stop_btn.configure(state=tk.NORMAL)
        self._status("Scraping...")

        t = threading.Thread(
            target=self._run_scrape,
            args=(self._result, start_ch, end_ch, out_dir),
            daemon=True,
        )
        t.start()

    def _stop_scrape(self):
        self._stop_event.set()
        self._stop_btn.configure(state=tk.DISABLED)
        self._status("Stopping...")

    def _run_scrape(self, result: dict, start_ch: int, end_ch: int, out_dir: str):
        platform  = result.get("platform", "unknown")
        book_name = result.get("book_name") or "Novel"
        args      = result.get("scraper_args", {})

        self._log(f"\n{'='*50}")
        self._log(f"Scraping: {book_name}")
        self._log(f"Platform: {platform} | Chapters: {start_ch}-{end_ch}")
        self._log(f"Output:   {out_dir}")
        self._log(f"{'='*50}")

        mode = PLATFORM_MODE.get(platform, "hardcoded")

        try:
            if mode == "inline" and platform == "novelbin":
                start_url = args.get("start_url") or result.get("canonical_url")
                if not start_url:
                    self._log("No start URL found in metadata.")
                else:
                    scrape_novelbin(start_url, book_name, start_ch, end_ch,
                                    out_dir, self._log, self._stop_event)

            elif mode == "inline" and platform == "royalroad":
                start_url = args.get("url") or result.get("canonical_url")
                if not start_url:
                    self._log("No start URL found in metadata.")
                else:
                    scrape_royalroad(start_url, book_name, start_ch, end_ch,
                                     out_dir, self._log, self._stop_event)

            elif mode == "argparse" and platform == "freewebnovel":
                book_url  = args.get("book_url") or result.get("canonical_url")
                script    = SCRIPT_DIR / "freewebnovel" / "script_new.py"
                if not script.exists():
                    self._log(f"Script not found: {script}")
                else:
                    run_subprocess(
                        [sys.executable, script,
                         "--book-url", book_url,
                         "--start-chapter", str(start_ch),
                         "--end-chapter", str(end_ch),
                         "--book-name", book_name],
                        self._log, self._stop_event
                    )

            else:
                # Browser-based or unknown — guide the user
                self._run_hardcoded_guide(result, start_ch, end_ch, out_dir)

        except Exception as exc:
            self._log(f"\nScrape error: {exc}")
        finally:
            self.after(0, self._on_done)

    def _run_hardcoded_guide(self, result: dict, start_ch: int, end_ch: int, out_dir: str):
        """
        For browser-based scripts (UC Chrome), show configuration instructions
        and optionally open the script in the system editor.
        """
        platform  = result.get("platform", "?")
        scraper   = result.get("scraper_script", "")
        book_name = result.get("book_name", "")
        canonical = result.get("canonical_url", "")
        args      = result.get("scraper_args", {})

        self._log(f"\nPlatform '{platform}' uses a browser-based scraper.")
        self._log(f"Script:  py Scripts/{scraper}")
        self._log(f"\nOpen the script and set these config values at the top:")
        self._log("-" * 40)

        # Show suggested config per platform
        if platform == "webnovel":
            self._log(f'  book_url    = "{canonical}"')
            self._log(f'  start_chapter = {start_ch}')
            self._log(f'  end_chapter   = {end_ch}')
        elif platform in ("69shuba", "hengyan"):
            self._log(f'  # Starting chapter URL:')
            self._log(f'  start_url     = "{args.get("start_url", canonical)}"')
            self._log(f'  start_chapter = {start_ch}')
            self._log(f'  end_chapter   = {end_ch}')
            self._log(f'  book_name     = "{book_name}"')
        elif platform == "babelnovel":
            self._log(f'  # Open the book page and log in manually, then set:')
            self._log(f'  book_url = "{canonical}"')
        elif platform == "tapas":
            self._log(f'  # Set in the script config:')
            self._log(f'  start_url    = "{args.get("start_url", canonical)}"')
            self._log(f'  book_name    = "{book_name}"')
            self._log(f'  max_chapters = {end_ch - start_ch + 1}')
        elif platform == "wuxiaworld":
            self._log(f'  url = "{canonical}"')
        elif platform == "wattpad":
            self._log(f'  url = "{canonical}"')
        elif platform in ("qidian", "qdmm"):
            self._log(f'  book_url      = "{canonical}"')
            self._log(f'  book_name     = "{book_name}"')
            self._log(f'  start_chapter = {start_ch}')
            self._log(f'  end_chapter   = {end_ch}')
        elif platform == "kakao":
            self._log(f'  # Requires Korean account login.')
            self._log(f'  book_url = "{canonical}"')
        else:
            for k, v in args.items():
                self._log(f"  {k} = {repr(v)}")

        self._log("-" * 40)
        self._log(f"\nThen run the script from the py Scripts/ directory:")
        self._log(f"  python {scraper}")
        self._log(f"\nOutput will go wherever the script saves it (check its OUT_DIR).")

        # Try to open the script
        script_path = SCRIPT_DIR / scraper if scraper else None
        if script_path and script_path.exists():
            self.after(0, lambda: self._offer_open(script_path))

    def _offer_open(self, path: Path):
        if messagebox.askyesno(
            "Open Script",
            f"Open {path.name} in your editor to configure it?\n\n{path}"
        ):
            try:
                subprocess.Popen(["open", str(path)])
            except Exception as e:
                self._log(f"Could not open file: {e}")

    def _on_done(self):
        self._run_btn.configure(state=tk.NORMAL)
        self._stop_btn.configure(state=tk.DISABLED)
        self._status("Done")
        self._log("\nFinished.")


# ── Entry point ───────────────────────────────────────────────────────────────

if __name__ == "__main__":
    app = ScraperApp()
    app.mainloop()
