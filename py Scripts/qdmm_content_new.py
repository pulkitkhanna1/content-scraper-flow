#!/usr/bin/env python3
"""
QDMM Chapter Content Scraper — Save to DOCX
=============================================
Extracts chapter text content from QDMM book pages and saves to Word documents.
Each .docx file contains 100 chapters. Chapter number is used as the heading.

Features:
- Aggressive driver restart: every N chapters + per-chapter timeout watchdog
- If a chapter takes longer than CHAPTER_TIMEOUT seconds, the driver is killed
  and recreated automatically
- Fresh driver + cookies for every batch to avoid session staleness
- 2Captcha Tencent solver for WAF/captcha bypass
- Checkpoint/resume support

SETUP:
  pip install undetected-chromedriver requests beautifulsoup4 lxml python-docx
"""

import os
import re
import json
import time
import random
import traceback
import threading

import requests as http_requests
import undetected_chromedriver as uc
from bs4 import BeautifulSoup as bs
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ──────────────────────────────────────────────
# CONFIG
# ──────────────────────────────────────────────

BOOK_URL = "https://www.qidian.com/book/1015648531/"
BOOK_NAME = "I really didn’t want to be reborn at all"
START_CHAPTER = 201
END_CHAPTER = 400

CHAPTERS_PER_DOC = 100  # chapters per .docx file

TWOCAPTCHA_API_KEY = os.environ.get("TWOCAPTCHA_API_KEY", "7c802b56a424742db63aa7d8640377eb")
MAX_CAPTCHA_RETRIES = 3

# Output
BOOK_ID = BOOK_URL.rstrip("/").split("/")[-1]
OUTPUT_FOLDER = f"qdmm_docs_{BOOK_ID}"
CHECKPOINT_FILE = f"qdmm_chapter_checkpoint_{BOOK_ID}.json"

# Delays
CHAPTER_DELAY_MIN = 3.0
CHAPTER_DELAY_MAX = 7.0

# ── DRIVER RESTART SETTINGS ──
BATCH_SIZE = 15               # restart driver every N chapters (lower = safer)
MAX_CONSECUTIVE_FAILS = 3     # restart after N consecutive failures
CHAPTER_TIMEOUT = 90          # kill driver if a single chapter takes longer than this (seconds)
COOLDOWN_MIN = 8.0            # min cooldown between driver restarts
COOLDOWN_MAX = 18.0           # max cooldown between driver restarts

# Cookies exported from browser (JSON list of dicts, qidian.com domain)
COOKIES_JSON = r"""[{"domain":".qidian.com","expirationDate":1786506457,"hostOnly":false,"httpOnly":false,"name":"e1","path":"/book/1037090668","sameSite":"unspecified","secure":false,"session":false,"storeId":"0","value":"%7B%22l6%22%3A%221%22%2C%22l7%22%3A%22%22%2C%22l1%22%3A3%2C%22l3%22%3A%22%22%2C%22pid%22%3A%22qd_P_xiangqing%22%2C%22eid%22%3A%22qd_G1001%22%7D"},{"domain":".qidian.com","expirationDate":1786506457,"hostOnly":false,"httpOnly":false,"name":"e2","path":"/book/1037090668","sameSite":"unspecified","secure":false,"session":false,"storeId":"0","value":"%7B%22l6%22%3A%221%22%2C%22l7%22%3A%22%22%2C%22l1%22%3A3%2C%22l3%22%3A%22%22%2C%22pid%22%3A%22qd_P_xiangqing%22%2C%22eid%22%3A%22qd_G1000%22%7D"},{"domain":".qidian.com","expirationDate":1798954185.425811,"hostOnly":false,"httpOnly":false,"name":"supportWebp","path":"/","sameSite":"unspecified","secure":false,"session":false,"storeId":"0","value":"true"},{"domain":".qidian.com","expirationDate":1799466385.487356,"hostOnly":false,"httpOnly":false,"name":"supportwebp","path":"/","sameSite":"unspecified","secure":true,"session":false,"storeId":"0","value":"true"},{"domain":".qidian.com","expirationDate":1784112256.525724,"hostOnly":false,"httpOnly":false,"name":"fu","path":"/","sameSite":"unspecified","secure":false,"session":false,"storeId":"0","value":"258984246"},{"domain":".qidian.com","expirationDate":1791691519.992975,"hostOnly":false,"httpOnly":false,"name":"_csrfToken","path":"/","sameSite":"unspecified","secure":false,"session":false,"storeId":"0","value":"94c4e9ae-735c-47c6-9caf-0cef4d542924"},{"domain":".qidian.com","expirationDate":1794978256.771234,"hostOnly":false,"httpOnly":false,"name":"newstatisticUUID","path":"/","sameSite":"unspecified","secure":false,"session":false,"storeId":"0","value":"1779426256_1143150479"},{"domain":".qidian.com","expirationDate":1799466390.092857,"hostOnly":false,"httpOnly":false,"name":"Hm_lvt_f00f67093ce2f38f215010b699629083","path":"/","sameSite":"unspecified","secure":false,"session":false,"storeId":"0","value":"1781782422"},{"domain":".qidian.com","hostOnly":false,"httpOnly":false,"name":"HMACCOUNT","path":"/","sameSite":"unspecified","secure":false,"session":true,"storeId":"0","value":"BA3EE2CC7EE8EEB9"},{"domain":".qidian.com","hostOnly":false,"httpOnly":false,"name":"ywguid","path":"/","sameSite":"no_restriction","secure":true,"session":true,"storeId":"0","value":"114200670318"},{"domain":".qidian.com","hostOnly":false,"httpOnly":false,"name":"ywopenid","path":"/","sameSite":"no_restriction","secure":true,"session":true,"storeId":"0","value":"3E553402386F1E4B20890FEA8049968C"},{"domain":".qidian.com","expirationDate":1784006985.425953,"hostOnly":false,"httpOnly":false,"name":"abPolicies","path":"/","sameSite":"unspecified","secure":false,"session":false,"storeId":"0","value":"%7B%22g17%22%3A1%2C%22g20%22%3A1%2C%22g16%22%3A0%2C%22g18%22%3A1%2C%22g19%22%3A0%2C%22g14%22%3A1%7D"},{"domain":".qidian.com","expirationDate":1784033136,"hostOnly":false,"httpOnly":false,"name":"traffic_search_engine","path":"/","sameSite":"unspecified","secure":false,"session":false,"storeId":"0","value":""},{"domain":".qidian.com","hostOnly":false,"httpOnly":false,"name":"se_ref","path":"/","sameSite":"unspecified","secure":false,"session":true,"storeId":"0","value":""},{"domain":".qidian.com","expirationDate":1785626095,"hostOnly":false,"httpOnly":false,"name":"e1","path":"/","sameSite":"unspecified","secure":false,"session":false,"storeId":"0","value":"%7B%22pid%22%3A%22qd_P_mycenter%22%2C%22eid%22%3A%22qd_H_mall_bottomaddownload%22%2C%22l7%22%3A%22hddl%22%7D"},{"domain":".qidian.com","expirationDate":1785626095,"hostOnly":false,"httpOnly":false,"name":"e2","path":"/","sameSite":"unspecified","secure":false,"session":false,"storeId":"0","value":"%7B%22pid%22%3A%22qd_P_mycenter%22%2C%22eid%22%3A%22qd_H_mall_bottomaddownload%22%2C%22l7%22%3A%22hddl%22%7D"},{"domain":".qidian.com","hostOnly":false,"httpOnly":false,"name":"traffic_utm_referer","path":"/","sameSite":"unspecified","secure":false,"session":true,"storeId":"0","value":""},{"domain":".qidian.com","hostOnly":false,"httpOnly":false,"name":"Hm_lpvt_f00f67093ce2f38f215010b699629083","path":"/","sameSite":"unspecified","secure":false,"session":true,"storeId":"0","value":"1783914390"},{"domain":".qidian.com","hostOnly":false,"httpOnly":false,"name":"ywkey","path":"/","sameSite":"no_restriction","secure":true,"session":true,"storeId":"0","value":"ywVenSkxOVYS"},{"domain":"www.qidian.com","expirationDate":1783914527,"hostOnly":true,"httpOnly":false,"name":"w_tsfp","path":"/","sameSite":"unspecified","secure":false,"session":false,"storeId":"0","value":"ltvuV0MF2utBvS0Q4anglU+rFzkmdTg4h0wpEaR0f5thQLErU5mN0o97us72MnXW68xnvd7DsZoyJTLYCJI3dwMdQ5qRe41DiQuSkdcs3ooSBRUyEs3VUVUccu5yvjhCL3hCNxS00jA8eIUd379yilkMsyN1zap3TO14fstJ019E6KDQmI5uDW3HlFWQRzaLbjcMcuqPr6g18L5a5TvasVr5eA0nAOtK2E3D3ShOC3B14ha9deEIPR6rJMqpSqA="}]"""

# ──────────────────────────────────────────────
# TIMEOUT EXCEPTION
# ──────────────────────────────────────────────

class ChapterTimeoutError(Exception):
    """Raised when a single chapter scrape exceeds CHAPTER_TIMEOUT."""
    pass


# ──────────────────────────────────────────────
# THREADED TIMEOUT WRAPPER
# ──────────────────────────────────────────────

def run_with_timeout(func, args=(), kwargs=None, timeout=CHAPTER_TIMEOUT):
    """
    Run func(*args, **kwargs) in a thread. If it doesn't finish within
    `timeout` seconds, return None (the thread is abandoned and the caller
    should kill the driver).
    Works on all platforms (no SIGALRM needed).
    """
    if kwargs is None:
        kwargs = {}

    result = [None]
    error = [None]
    finished = threading.Event()

    def wrapper():
        try:
            result[0] = func(*args, **kwargs)
        except Exception as e:
            error[0] = e
        finally:
            finished.set()

    t = threading.Thread(target=wrapper, daemon=True)
    t.start()
    finished.wait(timeout=timeout)

    if not finished.is_set():
        raise ChapterTimeoutError(
            f"Operation timed out after {timeout}s — driver is likely stuck"
        )

    if error[0] is not None:
        raise error[0]

    return result[0]


# ──────────────────────────────────────────────
# COOKIE PARSER
# ──────────────────────────────────────────────

def parse_cookie_string(cookie_str):
    """Parse either browser-exported JSON cookie array or name=value cookie string into Selenium cookie dicts."""
    cookie_str = cookie_str.strip()
    if not cookie_str:
        return []

    # Check if it looks like a JSON array
    if cookie_str.startswith("[") and cookie_str.endswith("]"):
        try:
            raw = json.loads(cookie_str)
            cookies = []
            for c in raw:
                cookie = {
                    "name": c["name"],
                    "value": c["value"],
                    "domain": c.get("domain", ".qidian.com"),
                    "path": c.get("path", "/"),
                    "secure": bool(c.get("secure", False)),
                    "httpOnly": bool(c.get("httpOnly", False)),
                }
                exp = c.get("expirationDate")
                if exp is not None:
                    cookie["expiry"] = int(exp)
                cookies.append(cookie)
            return cookies
        except Exception as e:
            print(f"Error parsing JSON cookies: {e}")

    # Fallback: semicolon-separated name=value pairs
    cookies = []
    for part in cookie_str.split(";"):
        part = part.strip()
        if "=" in part:
            k, v = part.split("=", 1)
            cookies.append({
                "name": k.strip(),
                "value": v.strip(),
                "domain": ".qidian.com",
                "path": "/"
            })
    return cookies


# ──────────────────────────────────────────────
# 2CAPTCHA TENCENT SOLVER
# ──────────────────────────────────────────────

def solve_tencent_captcha(website_url: str, app_id: str) -> dict | None:
    if TWOCAPTCHA_API_KEY == "YOUR_API_KEY_HERE":
        print("  ⚠ 2Captcha API key not set!")
        return None

    create_payload = {
        "clientKey": TWOCAPTCHA_API_KEY,
        "task": {
            "type": "TencentTaskProxyless",
            "websiteURL": website_url,
            "appId": app_id,
        }
    }

    print(f"  🔑 Sending captcha to 2Captcha (appId={app_id})...")
    try:
        resp = http_requests.post("https://api.2captcha.com/createTask", json=create_payload, timeout=30)
        result = resp.json()
    except Exception as e:
        print(f"  ✗ 2Captcha createTask failed: {e}")
        return None

    if result.get("errorId", 1) != 0:
        print(f"  ✗ 2Captcha error: {result.get('errorDescription', result)}")
        return None

    task_id = result["taskId"]
    print(f"  ⏳ Task {task_id} — polling...")

    for poll in range(60):
        time.sleep(3)
        try:
            poll_resp = http_requests.post(
                "https://api.2captcha.com/getTaskResult",
                json={"clientKey": TWOCAPTCHA_API_KEY, "taskId": task_id},
                timeout=30,
            )
            poll_result = poll_resp.json()
        except Exception as e:
            print(f"  ⚠ Poll error: {e}")
            continue

        if poll_result.get("errorId", 0) != 0:
            print(f"  ✗ Poll error: {poll_result.get('errorDescription', poll_result)}")
            return None

        status = poll_result.get("status")
        if status == "ready":
            solution = poll_result.get("solution", {})
            print(f"  ✓ Captcha solved! ticket={str(solution.get('ticket', '?'))[:50]}...")
            return solution
        elif status == "processing":
            if poll % 5 == 0:
                print(f"    Still solving... ({poll * 3}s)")

    print("  ✗ 2Captcha timed out after 180s")
    return None


def detect_tencent_captcha(html: str) -> tuple[bool, str | None, str | None]:
    if "TencentCaptcha" not in html or "captcha.qq.com" not in html:
        return False, None, None
    app_match = re.search(r"new\s+TencentCaptcha\s*\(\s*['\"](\d+)['\"]", html)
    app_id = app_match.group(1) if app_match else None
    seq_match = re.search(r'var\s+seqid\s*=\s*["\']([^"\']+)["\']', html)
    seqid = seq_match.group(1) if seq_match else None
    return True, app_id, seqid


def detect_waf_probe(html: str) -> bool:
    return "probev3.js" in html and len(html) < 2000


def submit_captcha_solution(driver, solution: dict, seqid: str | None) -> bool:
    ticket = solution.get("ticket", "")
    randstr = solution.get("randstr", "")
    ret = solution.get("ret", 0)

    js_code = f"""
    return new Promise((resolve, reject) => {{
        const seqid = "{seqid or ''}";
        const captchaResult = [];
        captchaResult.push({ret});
        if ({ret} === 0) {{
            captchaResult.push("{ticket}");
            captchaResult.push("{randstr}");
            captchaResult.push(seqid);
        }}
        const content = captchaResult.join('\\n');
        const xmlhttp = new XMLHttpRequest();
        xmlhttp.open("POST", "/WafCaptcha", true);
        xmlhttp.onreadystatechange = function() {{
            if (xmlhttp.readyState === 4) {{
                resolve(xmlhttp.status);
            }}
        }};
        xmlhttp.onerror = function() {{ reject("network error"); }};
        xmlhttp.send(content);
    }});
    """
    try:
        print("  📤 Submitting captcha token to /WafCaptcha...")
        driver.execute_script(js_code)
        time.sleep(1)
        driver.refresh()
        time.sleep(2)
        print("  ✓ Page reloaded after captcha submission")
        return True
    except Exception as e:
        print(f"  ⚠ Captcha submission error: {e}")
        try:
            driver.refresh()
            time.sleep(2)
            return True
        except Exception:
            return False


def wait_for_waf_probe(driver, timeout=30):
    print("  🔍 WAF probe detected — waiting for fingerprint check...")
    for i in range(timeout // 2):
        time.sleep(2)
        html = driver.page_source
        if not detect_waf_probe(html):
            print(f"  ✓ Probe completed after ~{(i + 1) * 2}s")
            return html
        try:
            driver.execute_script("""
                var evt = new MouseEvent('mousemove', {
                    clientX: Math.random() * window.innerWidth,
                    clientY: Math.random() * window.innerHeight
                });
                document.dispatchEvent(evt);
            """)
        except Exception:
            pass
    print(f"  ⚠ Probe still showing after {timeout}s")
    return driver.page_source


def handle_captcha(driver, url, max_retries=MAX_CAPTCHA_RETRIES) -> bool:
    html = driver.page_source

    if detect_waf_probe(html):
        html = wait_for_waf_probe(driver)
        if detect_waf_probe(html):
            print("  ⚠ Probe didn't resolve — reloading...")
            try:
                driver.refresh()
                time.sleep(3)
                html = driver.page_source
            except Exception:
                return False

    for attempt in range(max_retries):
        is_captcha, app_id, seqid = detect_tencent_captcha(html)
        if not is_captcha:
            return True

        print(f"  🛡️ Tencent Captcha detected (attempt {attempt + 1}/{max_retries})")
        if not app_id:
            print("  ✗ Could not extract appId")
            return False

        solution = solve_tencent_captcha(website_url=url, app_id=app_id)
        if not solution or solution.get("ret") != 0:
            print("  ✗ Captcha solve failed")
            time.sleep(3)
            continue

        submitted = submit_captcha_solution(driver, solution, seqid)
        if not submitted:
            continue

        time.sleep(2)
        html = driver.page_source

        is_still, _, _ = detect_tencent_captcha(html)
        if not is_still:
            print("  ✓ Captcha bypassed!")
            return True

    print(f"  ✗ Could not bypass captcha after {max_retries} attempts")
    return False


# ──────────────────────────────────────────────
# DRIVER MANAGEMENT
# ──────────────────────────────────────────────

def kill_driver(driver):
    """Force-kill the driver. Ignores all errors."""
    if driver is None:
        return
    try:
        driver.quit()
    except Exception:
        pass
    try:
        import psutil
        current_pid = os.getpid()
        for proc in psutil.process_iter(['pid', 'name', 'ppid']):
            try:
                if proc.info['ppid'] == current_pid and 'chrome' in proc.info['name'].lower():
                    proc.kill()
            except Exception:
                pass
    except ImportError:
        pass


def create_driver(headless=False):
    print("\n🚀 Launching fresh Chrome...")
    options = uc.ChromeOptions()
    options.add_argument("--start-maximized")
    options.add_argument("--window-size=1920,1080")
    options.add_argument("--disable-web-security")
    options.add_argument("--log-level=3")
    options.add_argument("--dns-prefetch-disable")

    driver = uc.Chrome(
        options=options,
        version_main=150,
        use_subprocess=True,
        headless=headless,
    )

    driver.set_page_load_timeout(45)
    driver.set_script_timeout(30)
    driver.implicitly_wait(5)

    print("✅ Chrome started")
    return driver


def create_driver_with_cookies(headless=False):
    """Create a fresh driver and inject cookies. Returns the driver."""
    driver = create_driver(headless=headless)

    print("🌐 Loading qidian.com to set cookies...")
    try:
        driver.get("https://www.qidian.com/")
        time.sleep(2)
        handle_captcha(driver, "https://www.qidian.com/")
    except Exception as e:
        print(f"  ⚠ Homepage load issue: {e}")

    cookies = parse_cookie_string(COOKIES_JSON)
    print(f"🍪 Injecting {len(cookies)} cookies...")
    added = 0
    for cookie in cookies:
        try:
            driver.add_cookie(cookie)
            added += 1
        except Exception as e:
            print(f"  ⚠ Skipped cookie {cookie.get('name')}: {e}")
    print(f"✅ Cookies injected ({added}/{len(cookies)})\n")
    return driver


def safe_navigate(driver, url, retries=2):
    for attempt in range(1, retries + 1):
        try:
            driver.get(url)
            time.sleep(random.uniform(1.5, 3.0))
            if not handle_captcha(driver, url):
                print(f"  ⚠ Captcha unresolved (attempt {attempt})")
                if attempt < retries:
                    time.sleep(3)
                    continue
                return False
            return True
        except Exception as e:
            print(f"  ⚠ Navigation error (attempt {attempt}): {e}")
            if attempt < retries:
                time.sleep(3)
    return False


# ──────────────────────────────────────────────
# CHAPTER URL EXTRACTION
# ──────────────────────────────────────────────

def extract_chapter_links(driver, book_url, start_ch, end_ch):
    print(f"\n📚 Extracting chapter links from: {book_url}")
    print(f"   Range: chapter {start_ch} to {end_ch}")

    if not safe_navigate(driver, book_url):
        print("❌ Could not load book page!")
        return []

    driver.execute_script("window.scrollTo(0, document.body.scrollHeight/3);")
    time.sleep(1)
    driver.execute_script("window.scrollTo(0, document.body.scrollHeight*2/3);")
    time.sleep(1)
    driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
    time.sleep(1)

    html = driver.page_source
    soup = bs(html, "lxml")

    atags = soup.select("ul.volume-chapters li.chapter-item a.chapter-name")
    if not atags:
        atags = soup.select("div.catalog-content-wrap a[href*='/chapter/']")
    if not atags:
        atags = soup.select("a[href*='/chapter/']")

    print(f"   Found {len(atags)} total chapter links on page")

    if len(atags) == 0:
        print("   ⚠ No chapter links found! Saving debug HTML...")
        with open("debug_book_page.html", "w", encoding="utf-8") as f:
            f.write(html)
        return []

    actual_end = min(end_ch, len(atags))
    if start_ch > len(atags):
        print(f"   ⚠ Start chapter {start_ch} exceeds available {len(atags)} chapters")
        return []

    selected = atags[start_ch - 1:actual_end]
    print(f"   Selected chapters {start_ch} to {actual_end} ({len(selected)} chapters)")

    chapters = []
    for idx, a in enumerate(selected, start=start_ch):
        href = a.get("href", "")
        if href.startswith("//"):
            href = "https:" + href
        elif href.startswith("/"):
            href = "https://www.qidian.com" + href

        name = a.get_text(strip=True)

        chapter_id = ""
        parts = href.rstrip("/").split("/")
        if len(parts) >= 2:
            chapter_id = parts[-1]

        chapters.append({
            "chapter_no": idx,
            "chapter_name": name,
            "chapter_url": href,
            "chapter_id": chapter_id,
        })

    print(f"   ✅ Prepared {len(chapters)} chapters for scraping\n")
    return chapters


# ──────────────────────────────────────────────
# CONTENT EXTRACTION
# ──────────────────────────────────────────────

def _extract_chapter_inner(driver, chapter_url):
    """
    Inner function that does the actual navigation + extraction.
    This runs inside run_with_timeout() so it can be killed if stuck.
    """
    if not safe_navigate(driver, chapter_url):
        return None, None

    html = driver.page_source
    soup = bs(html, "lxml")

    chapter_title = ""
    title_el = soup.select_one("h1.chapter-title, h3.chapter-title, span.chapter-title")
    if title_el:
        chapter_title = title_el.get_text(strip=True)

    content_el = soup.select_one("main.content")
    if not content_el:
        content_el = (
            soup.select_one("div.read-content")
            or soup.select_one("div#chapterContent")
            or soup.select_one("div.chapter-content")
            or soup.select_one("div.text-content")
        )

    if not content_el:
        print("  ⚠ Content element not found!")
        debug_name = f"debug_chapter_{int(time.time())}.html"
        with open(debug_name, "w", encoding="utf-8") as f:
            f.write(html)
        print(f"  💾 Debug saved: {debug_name}")
        return None, chapter_title

    content_spans = content_el.select("span.content-text")

    if content_spans:
        paragraphs = []
        for span in content_spans:
            text = span.get_text(strip=True)
            if text:
                paragraphs.append(text)
    else:
        p_tags = content_el.find_all("p")
        if p_tags:
            paragraphs = []
            for p in p_tags:
                text = p.get_text(strip=True)
                if text:
                    paragraphs.append(text)
        else:
            paragraphs = [content_el.get_text("\n", strip=True)]

    return paragraphs, chapter_title


def extract_chapter_content(driver, chapter_url):
    """
    Wrapper that runs _extract_chapter_inner with a timeout.
    Raises ChapterTimeoutError if the driver is stuck.
    """
    return run_with_timeout(
        _extract_chapter_inner,
        args=(driver, chapter_url),
        timeout=CHAPTER_TIMEOUT,
    )


# ──────────────────────────────────────────────
# CHECKPOINT
# ──────────────────────────────────────────────

def load_checkpoint():
    if os.path.exists(CHECKPOINT_FILE):
        try:
            with open(CHECKPOINT_FILE, "r", encoding="utf-8") as f:
                data = json.load(f)
            print(f"📋 Checkpoint loaded: {len(data.get('completed', []))} chapters done")
            return data
        except Exception:
            pass
    return {"completed": []}


def save_checkpoint(completed_list):
    data = {
        "completed": completed_list,
        "timestamp": datetime.now().isoformat(),
    }
    with open(CHECKPOINT_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2)


# ──────────────────────────────────────────────
# DOCX WRITER
# ──────────────────────────────────────────────

class DocxBatchWriter:
    def __init__(self, output_folder, book_name, chapters_per_doc=100):
        self.output_folder = output_folder
        self.book_name = book_name
        self.chapters_per_doc = chapters_per_doc
        os.makedirs(output_folder, exist_ok=True)
        self._doc_cache = {}

    def _get_doc_key(self, chapter_no):
        doc_start = ((chapter_no - 1) // self.chapters_per_doc) * self.chapters_per_doc + 1
        doc_end = doc_start + self.chapters_per_doc - 1
        return (doc_start, doc_end)

    def _get_or_create_doc(self, doc_key):
        if doc_key in self._doc_cache:
            return self._doc_cache[doc_key]

        doc_start, doc_end = doc_key
        safe_name = re.sub(r'[\\/:\"\*?<>|]+', '_', self.book_name)
        filename = f"{safe_name}_ch{doc_start}-{doc_end}.docx"
        filepath = os.path.join(self.output_folder, filename)

        if os.path.exists(filepath):
            doc = Document(filepath)
            print(f"  📄 Loaded existing doc: {filename}")
        else:
            doc = Document()

            style = doc.styles['Normal']
            font = style.font
            font.name = 'SimSun'
            font.size = Pt(11)

            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.add_run(self.book_name)
            title_run.bold = True
            title_run.font.size = Pt(18)

            subtitle_para = doc.add_paragraph()
            subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_run = subtitle_para.add_run(f"Chapters {doc_start} – {doc_end}")
            subtitle_run.font.size = Pt(14)
            subtitle_run.font.color.rgb = RGBColor(100, 100, 100)

            doc.add_page_break()
            print(f"  📄 Created new doc: {filename}")

        entry = {"doc": doc, "path": filepath, "count": 0}
        self._doc_cache[doc_key] = entry
        return entry

    def add_chapter(self, chapter_no, chapter_name, paragraphs):
        doc_key = self._get_doc_key(chapter_no)
        entry = self._get_or_create_doc(doc_key)
        doc = entry["doc"]

        heading_text = f"Chapter {chapter_no}: {chapter_name}"
        doc.add_heading(heading_text, level=1)

        for para_text in paragraphs:
            para_text = para_text.strip()
            if para_text:
                p = doc.add_paragraph(para_text)
                p.paragraph_format.first_line_indent = Pt(24)
                p.paragraph_format.space_after = Pt(4)

        doc.add_paragraph("")
        entry["count"] += 1

    def save_all(self):
        for doc_key, entry in self._doc_cache.items():
            try:
                entry["doc"].save(entry["path"])
                print(f"  💾 Saved: {os.path.basename(entry['path'])} ({entry['count']} chapters added)")
            except Exception as e:
                print(f"  ❌ Error saving {entry['path']}: {e}")

    def save_doc_for_chapter(self, chapter_no):
        doc_key = self._get_doc_key(chapter_no)
        if doc_key in self._doc_cache:
            entry = self._doc_cache[doc_key]
            try:
                entry["doc"].save(entry["path"])
            except Exception as e:
                print(f"  ❌ Error saving {entry['path']}: {e}")


# ──────────────────────────────────────────────
# MAIN SCRAPER
# ──────────────────────────────────────────────

def main():
    import argparse
    parser = argparse.ArgumentParser(description="QDMM/Qidian Scraper")
    parser.add_argument("book_url", nargs="?", default=None, help="Qidian book page URL")
    parser.add_argument("--book-name", default=None, help="Name of the book (for docx output)")
    parser.add_argument("--start-chapter", "-s", type=int, default=None, help="Start chapter number")
    parser.add_argument("--end-chapter", "-e", type=int, default=None, help="End chapter number")
    parser.add_argument("--output-folder", "-o", default=None, help="Output folder path")
    parser.add_argument("--cookies", default=None, help="Pasted cookies string or file path")
    parser.add_argument("--headless", action="store_true", help="Run in headless mode")
    args = parser.parse_args()

    global BOOK_URL, BOOK_NAME, START_CHAPTER, END_CHAPTER, OUTPUT_FOLDER, CHECKPOINT_FILE, BOOK_ID, COOKIES_JSON

    if args.book_url:
        BOOK_URL = args.book_url
        BOOK_ID = BOOK_URL.rstrip("/").split("/")[-1]
    if args.book_name:
        BOOK_NAME = args.book_name
    if args.start_chapter is not None:
        START_CHAPTER = args.start_chapter
    if args.end_chapter is not None:
        END_CHAPTER = args.end_chapter
    if args.output_folder:
        OUTPUT_FOLDER = args.output_folder
    else:
        OUTPUT_FOLDER = f"qdmm_docs_{BOOK_ID}"
        
    CHECKPOINT_FILE = f"qdmm_chapter_checkpoint_{BOOK_ID}.json"

    if args.cookies:
        if os.path.exists(args.cookies):
            try:
                with open(args.cookies, "r", encoding="utf-8") as f:
                    COOKIES_JSON = f.read()
            except Exception as e:
                print(f"Error reading cookies file: {e}")
        else:
            COOKIES_JSON = args.cookies

    headless_mode = False
    if args.headless:
        headless_mode = True

    print("=" * 70)
    print("QDMM Chapter Content Scraper — DOCX + Watchdog Timeout")
    print("=" * 70)
    print(f"📚 Book: {BOOK_NAME}")
    print(f"🔗 URL:  {BOOK_URL}")
    print(f"📖 Chapters: {START_CHAPTER} to {END_CHAPTER}")
    print(f"📄 {CHAPTERS_PER_DOC} chapters per .docx file")
    print(f"📂 Output folder: {OUTPUT_FOLDER}")
    print(f"⏱  Chapter timeout: {CHAPTER_TIMEOUT}s")
    print(f"🔄 Restart every {BATCH_SIZE} chapters")
    print(f"👻 Headless: {headless_mode}")
    print()

    # Load checkpoint
    checkpoint = load_checkpoint()
    completed = set(checkpoint.get("completed", []))
    is_resuming = len(completed) > 0

    if is_resuming:
        print(f"🔄 Resuming — {len(completed)} chapters already done\n")

    # ── Phase 1: Extract chapter links (needs a driver) ──
    driver = create_driver_with_cookies(headless=headless_mode)
    chapters = extract_chapter_links(driver, BOOK_URL, START_CHAPTER, END_CHAPTER)

    if not chapters:
        print("❌ No chapters to scrape. Exiting.")
        kill_driver(driver)
        return

    # Kill this driver — we'll create a fresh one for scraping
    kill_driver(driver)
    driver = None

    # ── Phase 2: Scrape chapters ──
    doc_writer = DocxBatchWriter(OUTPUT_FOLDER, BOOK_NAME, CHAPTERS_PER_DOC)

    batch_count = 0
    consecutive_fails = 0
    total_scraped = len(completed)
    total_to_scrape = len(chapters)
    failed_chapters = []

    # Filter to only pending chapters
    pending = [ch for ch in chapters if ch["chapter_no"] not in completed]
    print(f"📝 {len(pending)} chapters remaining to scrape\n")

    if not pending:
        print("✅ All chapters already scraped!")
        return

    try:
        for ch in pending:
            ch_no = ch["chapter_no"]
            ch_name = ch["chapter_name"]
            ch_url = ch["chapter_url"]

            # ── Proactive restart every BATCH_SIZE chapters ──
            if batch_count >= BATCH_SIZE:
                print(f"\n🔄 Batch limit ({BATCH_SIZE}) — recycling driver...")
                doc_writer.save_all()
                kill_driver(driver)
                driver = None

            # ── Restart on consecutive failures ──
            if consecutive_fails >= MAX_CONSECUTIVE_FAILS:
                print(f"\n🔄 {MAX_CONSECUTIVE_FAILS} consecutive failures — recycling driver...")
                doc_writer.save_all()
                kill_driver(driver)
                driver = None
                consecutive_fails = 0

            # ── Create driver if we don't have one ──
            if driver is None:
                cooldown = random.uniform(COOLDOWN_MIN, COOLDOWN_MAX)
                print(f"   ⏳ Cooling down {cooldown:.1f}s before new driver...")
                time.sleep(cooldown)
                driver = create_driver_with_cookies(headless=headless_mode)
                batch_count = 0

            # ── Delay between chapters ──
            delay = random.uniform(CHAPTER_DELAY_MIN, CHAPTER_DELAY_MAX)
            remaining = total_to_scrape - total_scraped
            print(f"\n[Ch {ch_no} | {remaining} left] "
                  f"Waiting {delay:.1f}s → {ch_name[:50]}")
            time.sleep(delay)

            # ── Scrape with timeout + retries ──
            success = False
            for attempt in range(1, 4):
                if attempt > 1:
                    retry_delay = random.uniform(5, 12)
                    print(f"  Retry {attempt}/3 — restarting driver, waiting {retry_delay:.1f}s")
                    kill_driver(driver)
                    driver = None
                    time.sleep(retry_delay)

                if driver is None:
                    cooldown = random.uniform(COOLDOWN_MIN, COOLDOWN_MAX)
                    print(f"   ⏳ Cooling down {cooldown:.1f}s before new driver...")
                    time.sleep(cooldown)
                    driver = create_driver_with_cookies(headless=headless_mode)
                    batch_count = 0

                try:
                    start_time = time.time()
                    paragraphs, page_title = extract_chapter_content(driver, ch_url)
                    elapsed = time.time() - start_time
                    print(f"  ⏱ Took {elapsed:.1f}s")

                    if paragraphs:
                        final_name = ch_name or page_title or f"Chapter {ch_no}"
                        char_count = sum(len(p) for p in paragraphs)

                        doc_writer.add_chapter(ch_no, final_name, paragraphs)
                        doc_writer.save_doc_for_chapter(ch_no)

                        completed.add(ch_no)
                        save_checkpoint(list(completed))
                        total_scraped += 1
                        batch_count += 1
                        consecutive_fails = 0

                        print(f"  ✅ {final_name[:60]}")
                        print(f"     {char_count} chars, {len(paragraphs)} paras | ")
                        print(f"Progress: {total_scraped}/{total_to_scrape}")

                        success = True
                        break
                    else:
                        print(f"  ✗ No content (attempt {attempt})")

                except ChapterTimeoutError:
                    print(f"  ⏰ TIMEOUT after {CHAPTER_TIMEOUT}s — killing driver!")
                    kill_driver(driver)
                    driver = None
                    continue

                except Exception as e:
                    print(f"  ❌ Error (attempt {attempt}): {e}")
                    try:
                        _ = driver.title
                    except Exception:
                        print("  💀 Driver is dead — will recreate")
                        kill_driver(driver)
                        driver = None

            if not success:
                consecutive_fails += 1
                failed_chapters.append(ch_no)
                print(f"  ❌ FAILED Ch {ch_no} after 3 attempts "
                      f"(consecutive fails: {consecutive_fails})")

    except KeyboardInterrupt:
        print("\n\n⚠️ Interrupted by user!")

    except Exception as e:
        print(f"\n❌ Fatal error: {e}")
        traceback.print_exc()

    finally:
        doc_writer.save_all()

        print(f"\n{'=' * 70}")
        print(f"📊 SUMMARY")
        print(f"{'=' * 70}")
        print(f"   Total scraped:   {total_scraped}/{total_to_scrape}")
        print(f"   Failed chapters: {len(failed_chapters)}")
        if failed_chapters:
            print(f"   Failed list: {failed_chapters}")
        print(f"   Output folder: {OUTPUT_FOLDER}/")

        if os.path.exists(OUTPUT_FOLDER):
            docx_files = sorted(f for f in os.listdir(OUTPUT_FOLDER) if f.endswith(".docx"))
            for df in docx_files:
                fpath = os.path.join(OUTPUT_FOLDER, df)
                size_kb = os.path.getsize(fpath) / 1024
                print(f"   📄 {df} ({size_kb:.1f} KB)")

        print(f"   Checkpoint: {CHECKPOINT_FILE}")
        print(f"{'=' * 70}")

        kill_driver(driver)
        print("🔒 Browser closed")

        if total_scraped >= total_to_scrape and not failed_chapters:
            try:
                os.remove(CHECKPOINT_FILE)
                print("🎉 All chapters scraped! Checkpoint cleaned up.")
            except Exception:
                pass


if __name__ == "__main__":
    main()